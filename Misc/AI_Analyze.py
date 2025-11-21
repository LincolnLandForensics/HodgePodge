#!/usr/bin/env python3
# coding: utf-8
"""
Read .docx, .txt, or .pdf files and analyze if they were written by AI
Author: LincolnLandForensics
Version: 0.0.3
"""
 
# <<<<<<<<<<<<<<<<<<<<<<<<<<      Imports        >>>>>>>>>>>>>>>>>>>>>>>>>> #

import os
import re
import sys
import math
import hashlib
import argparse
import tkinter as tk
from tkinter import filedialog

import unicodedata

from collections import Counter
from typing import Dict, Any, List
from docx import Document
from PyPDF2 import PdfReader
from docx.opc import coreprops
import threading



# colors
color_red = color_yellow = color_green = color_blue = color_purple = color_reset = ''
from colorama import Fore, Back, Style
print(Back.BLACK)
color_red, color_yellow, color_green = Fore.RED, Fore.YELLOW, Fore.GREEN
color_blue, color_purple, color_reset = Fore.BLUE, Fore.MAGENTA, Style.RESET_ALL


DEFAULT_OUTPUT = "sample_AI_Cleaned.docx"

# Global log function, defaults to print
log_func = print


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Main           >>>>>>>>>>>>>>>>>>>>>>>>>> #

def main():
    global args

    parser = argparse.ArgumentParser(description="Analyze DOCX, TXT, or PDF for AI-authorship indicators")

    parser.add_argument('-I', '--input', help='Input DOCX, TXT, or PDF file', required=False)
    parser.add_argument('-O', '--output', help='Write cleaned text to this file', required=False)
    parser.add_argument('-a', '--analyzeAI', help='Analyze the input file for AI-like characteristics', action='store_true')
    parser.add_argument('-g', '--gui', help='Force GUI mode', action='store_true')

    args = parser.parse_args()

    # ⬅ DEFAULT BEHAVIOR: start GUI if no CLI action provided
    if len(sys.argv) == 1 or args.gui:
        build_gui()
        return

    # CLI MODE
    input_file = args.input if args.input else "sample.docx"
    parts = input_file.rsplit('.', 1)
    suggested_output = f"{parts[0]}_cleaned.{parts[1]}" if len(parts) == 2 else input_file + "_cleaned"

    output_file = args.output if args.output else suggested_output

    if args.analyzeAI:
        analyze_AI(input_file, output_file)
    else:
        usage()


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Core Logic     >>>>>>>>>>>>>>>>>>>>>>>>>> #

def analyze_AI(input_file, output_file):
    if not os.path.exists(input_file):
        msg_blurb_square(f"{input_file} does not exist", color_red)
        sys.exit(1)

    msg_blurb_square(f"Reading {input_file}", color_green)

    result = analyze_document(input_file, highlight=True, clean=True)

    if "highlighted_text" in result:
        print("\n--- Highlighted Text Preview ---\n")
        print(result["highlighted_text"][:500])

    if "cleaned_text" in result:
        write_to_docx(output_file, result["cleaned_text"])
        msg_blurb_square(f"Cleaned output written to {output_file}", color_green)

    print_pretty_summary(result)




def analyze_document(q: str, highlight: bool = False, clean: bool = False) -> Dict[str, Any]:
    if q.lower().endswith(".docx") or q.lower().endswith(".doc"):
        text = extract_text_from_docx(q)
        metadata = extract_docx_metadata(q)

    elif q.lower().endswith(".pdf"):
        text = extract_text_from_pdf(q)
        metadata = extract_pdf_metadata(q)

    elif q.lower().endswith((
    ".txt", ".text", ".asc", ".nfo", ".log", ".csv", ".md",
    ".xml", ".json", ".ini", ".yaml", ".yml", ".tex", ".cfg", ".dat"
    )):

        with open(q, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        metadata = {
            "filename": q,
            "length": len(text),
            "lines": text.count("\n") + 1,
            "encoding": "utf-8"
        }

    else:
        raise ValueError("Unsupported file type. Only DOCX, DOC, PDF, or TXT files are allowed.")

    hidden_chars = detect_hidden_characters(text)
    score = compute_ai_likelihood(text, hidden_chars)
    # filehash = file_hash(q)

    result = {
        "file": q,
        "hidden_characters": hidden_chars,
        "ai_likelihood_score": score,
        # "file_hash": filehash,
        "metadata": metadata,
    }

    if highlight:
        result["highlighted_text"] = highlight_hidden_characters(text, hidden_chars)
    if clean:
        result["cleaned_text"] = clean_text(text, hidden_chars)

    return result

# <<<<<<<<<<<<<<<<<<<<<<<<<<      Utilities      >>>>>>>>>>>>>>>>>>>>>>>>>> #

def build_gui():
    global input_entry, output_entry, message_box

    root = tk.Tk()
    root.title("AI Document Analyzer")

    tk.Label(root, text="🧠 AI Authorship Analyzer",
             font=("Arial", 12, "bold")).pack(pady=(5, 0))

    frame = tk.Frame(root)
    frame.pack(pady=5)

    # Input
    tk.Label(frame, text="Input File:").grid(row=0, column=0, sticky="e")
    input_entry = tk.Entry(frame, width=50)
    input_entry.insert(0, os.path.join(os.getcwd(), "sample.docx"))
    input_entry.grid(row=0, column=1)
    tk.Button(
        frame, text="Browse",
        command=lambda: input_entry.delete(0, tk.END) or
                        input_entry.insert(0, filedialog.askopenfilename(
                            filetypes=[("Documents", "*.docx *.doc *.pdf *.txt")]))
    ).grid(row=0, column=2)

    # Output
    tk.Label(frame, text="Output File:").grid(row=1, column=0, sticky="e")
    output_entry = tk.Entry(frame, width=50)
    output_entry.insert(0, "sample_Cleaned.docx")
    output_entry.grid(row=1, column=1)
    tk.Button(
        frame, text="Browse",
        command=lambda: output_entry.delete(0, tk.END) or
                        output_entry.insert(0, filedialog.asksaveasfilename(
                            defaultextension=".docx",
                            filetypes=[("DOCX File", "*.docx")]))
    ).grid(row=1, column=2)

    # Button
    tk.Button(root, text="Analyze",
              command=start_gui_analysis,
              bg="lightblue").pack(pady=10)

    # Colored output window
    message_box = tk.Text(root, width=90, height=12, state="disabled")
    message_box.tag_config("red", foreground="red")
    message_box.tag_config("green", foreground="green")
    message_box.tag_config("blue", foreground="blue")
    message_box.tag_config("black", foreground="black")
    message_box.pack(pady=5)

    root.mainloop()

    
    
def entropy_score(text: str) -> float:
    if not text:
        return 0.0
    normalized = text.lower().replace(" ", "")
    freq = Counter(normalized)
    total = sum(freq.values())
    entropy = -sum((count / total) * math.log2(count / total) for count in freq.values())
    log_func(f'Entropy score: {round(entropy, 4)}', "blue")
    return round(entropy, 4)



def extract_text_from_docx(input_file: str) -> str:
    ext = os.path.splitext(input_file)[1].lower()

    if ext == ".docx":
        try:
            doc = Document(input_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            return (f"{color_red}Error reading .docx file: {e}{color_reset}")
            try:
                import textract
                text = textract.process(input_file).decode("utf-8")
                return "\n".join([line.strip() for line in text.splitlines() if line.strip()])
            except Exception as e:
                return f"Error reading .doc file: {e}"


    elif ext == ".doc":
        print(f'I havent written the .doc reader yet')
        # try:
            # import textract
            # text = textract.process(input_file).decode("utf-8")
            # return "\n".join([line.strip() for line in text.splitlines() if line.strip()])
        # except Exception as e:
            # return f"Error reading .doc file: {e}"

    else:
        return "Unsupported file format. Please provide a .docx or .pdf"


def extract_text_from_docxOLD(input_file: str) -> str:
    doc = Document(input_file)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_pdf(input_file: str) -> str:
    reader = PdfReader(input_file)
    return "\n".join([page.extract_text() or "" for page in reader.pages])

def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    doc = Document(file_path)
    props = doc.core_properties

    metadata = {
        "title": props.title,
        "subject": props.subject,
        "author": props.author,
        "category": props.category,
        "comments": props.comments,
        "keywords": props.keywords,
        "content_status": props.content_status,
        "identifier": props.identifier,
        "language": props.language,
        "revision": props.revision,
        "version": props.version,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_printed": str(props.last_printed) if props.last_printed else None,
        "last_modified_by": props.last_modified_by,
    }

    # Filter out empty or None values
    return {k: v for k, v in metadata.items() if v not in [None, ""]}

def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    reader = PdfReader(file_path)
    return dict(reader.metadata or {})

def file_hash(path: str) -> str:
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()

def detect_hidden_characters(text: str) -> Dict[str, Any]:
    suspicious_chars = {
        "ZERO_WIDTH_SPACE": "\u200b",
        "ZERO_WIDTH_NON_JOINER": "\u200c",
        "ZERO_WIDTH_JOINER": "\u200d",
        "NON_BREAKING_SPACE": "\u00a0",
        "SOFT_HYPHEN": "\u00ad",
        "LEFT_TO_RIGHT_MARK": "\u200e",
        "RIGHT_TO_LEFT_MARK": "\u200f",
    }

    findings = {}
    for name, char in suspicious_chars.items():
        count = text.count(char)
        if count > 0:
            findings[name] = {"char": char, "count": count}

    unusual = [ch for ch in text if ord(ch) > 126 and not unicodedata.category(ch).startswith("Z")]
    if unusual:
        findings["UNUSUAL_UNICODE"] = {
            "char": None,
            "count": len(unusual),
            "examples": Counter(unusual).most_common(5)
        }

    return findings

def highlight_hidden_characters(text: str, findings: Dict[str, Any]) -> str:
    highlighted = text
    for name, data in findings.items():
        if name == "UNUSUAL_UNICODE":
            continue
        char = data["char"]
        highlighted = highlighted.replace(char, f"[{name}]")
    return highlighted

def log_gui(text, tag="black"):
    message_box.config(state="normal")
    message_box.insert(tk.END, text + "\n", tag)
    message_box.config(state="disabled")
    message_box.see(tk.END)


def clean_text(text: str, findings: Dict[str, Any]) -> str:
    cleaned = (
        text.replace("’", "'")
             .replace("—", ", ")
             .replace("“", "\"")
             .replace("”", "\"")
             .replace("️", "")
             .replace("️", "")
             .replace("⚠", "")
             .replace("🧠", "")
             .replace("☠", "")
             .replace("🧭", "")
)
    
    for name, data in findings.items():
        if name == "UNUSUAL_UNICODE":
            continue
        char = data["char"]
        cleaned = cleaned.replace(char, "")
    return cleaned

def compute_ai_likelihood(text: str, hidden_char_findings: Dict[str, Any]) -> float:
    score = 0.0

    if hidden_char_findings:
        score += min(30, sum(v["count"] for v in hidden_char_findings.values() if "count" in v) * 5)

    sentences = re.split(r"[.!?]", text)
    sentences = [s.strip() for s in sentences if s.strip()]
    avg_len = sum(len(s.split()) for s in sentences) / max(1, len(sentences))
    if avg_len > 25:
        score += 20

    words = re.findall(r"\w+", text.lower())
    freq = Counter(words)
    repeated = [w for w, c in freq.items() if c > 10 and len(w) > 6]
    if repeated:
        score += 15

    # Passive voice detection (naive)
    passive_hits = len(re.findall(r"\b(is|was|were|been|being|be)\b\s+\w+ed\b", text))
    if passive_hits > 5:
        score += 10

    # Placeholder for entropy/burstiness
    score += entropy_score(text)

    return min(100.0, score)

def generate_summary(result: Dict[str, Any]) -> str:
    lines = [f"# AI Authorship Analysis for `{result['file']}`"]
    lines.append(f"**AI-Likelihood Score**: `{result['ai_likelihood_score']}`")
    lines.append(f"**File Hash**: `{result['file_hash']}`")
    lines.append("## Metadata:")
    for k, v in result["metadata"].items():
        lines.append(f"- **{k}**: {v}")
    lines.append("## Hidden Characters Detected:")
    for k, v in result["hidden_characters"].items():
        lines.append(f"- **{k}**: {v['count']}")
    return "\n".join(lines)

def msg_blurb_square(msg, color):
    border = f"+{'-' * (len(msg) + 2)}+"
    log_func(f"{border}\n| {msg} |\n{border}", color)

def print_pretty_summary(result: Dict[str, Any]):
    print("\n" + "=" * 60)
    print("📄 AI Authorship Analysis Summary")
    print("=" * 60)

    print(f"\n📊 AI-Likelihood Score:\n  {result['ai_likelihood_score']} (Moderate indicators based on hidden characters, entropy and repetition)")

    print("\n🕵️ Hidden Characters Detected:")
    print(f"{'Type':<25} {'Unicode':<10} {'Count':<6} {'Examples / Notes'}")
    print("-" * 60)
    for k, v in result["hidden_characters"].items():
        char = repr(v["char"]) if v["char"] else "—"
        count = v["count"]
        if k == "UNUSUAL_UNICODE":
            examples = ", ".join([f"{repr(ch)}×{n}" for ch, n in v["examples"]])
        else:
            examples = ""
        print(f"{k:<25} {char:<10} {str(count):<6} {examples}")

    print("\n🧾 Document Metadata:")
    print(f"{'Field':<20} {'Value'}")
    print("-" * 40)
    for k, v in result["metadata"].items():
        print(f"{k:<20} {v}")
    # print(f"\n🔐 File Hash (SHA256):\n  {result['file_hash']}")

        
    print("\n" + "=" * 60 + "\n")

def start_gui_analysis():
    input_file = input_entry.get().strip()
    output_file = output_entry.get().strip()

    log_gui(f"Analyzing: {input_file}", "blue")

    # Run analysis in separate thread to avoid freezing GUI
    threading.Thread(target=gui_analysis_thread, args=(input_file, output_file), daemon=True).start()

def gui_analysis_thread(input_file, output_file):
    global log_func
    old_log_func = log_func
    log_func = lambda msg, color="black": log_gui(msg, color)

    try:
        analyze_AI(input_file, output_file)
        log_gui(f"Saved cleaned document → {output_file}", "green")
    except Exception as e:
        log_gui(f"Error: {e}", "red")
    finally:
        log_func = old_log_func  # restore original log function

def write_to_docx(output_path: str, text: str) -> None:
    doc = Document()
    for line in text.splitlines():
        if line.strip():  # Avoid empty lines
            doc.add_paragraph(line.strip())
    doc.save(output_path)

def usage():
    print(f"Usage: {sys.argv[0]} -a [-I sample.docx] [-O sample_updated.docx]")
    print("Example:")
    print(f"    {sys.argv[0]} -a")
    print(f"    {sys.argv[0]} -a -I sample.docx -O sample_cleaned.docx")
    print(f"    {sys.argv[0]} -a -I sample_AI_written.txt  -O sample_AI_Cleaned.docx")
    print(f"    {sys.argv[0]} -a -I sample_AI_written.txt  -a -O sample_AI_Cleaned.docx")


if __name__ == '__main__':
    main()

# <<<<<<<<<<<<<<<<<<<<<<<<<< Revision History >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
0.1.2 - original version

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<< Future Wishlist  >>>>>>>>>>>>>>>>>>>>>>>>>>

"""


"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<      notes            >>>>>>>>>>>>>>>>>>>>>>>>>>

"""



"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<      The End        >>>>>>>>>>>>>>>>>>>>>>>>>>


'''

'''

