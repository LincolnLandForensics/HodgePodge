#!/usr/bin/python
# coding: utf-8

# <<<<<<<<<<<<<<<<<<<<<<<<<<     Imports        >>>>>>>>>>>>>>>>>>>>>>>>>>
try:
    from bs4 import BeautifulSoup
    from playwright.async_api import async_playwright   # pip install playwright
    from playwright_stealth import Stealth  # pip install playwright_stealth    
except:
    print(f'install missing modules:    pip install -r requirements_identity_hunt.txt')
    exit()

import os
import re
import sys
import json
import time
from datetime import datetime
import random
import openpyxl
import requests # pip install requests

from docx import Document   # pip install python-docx

from openpyxl import load_workbook, Workbook    # pip install openpyxl
from openpyxl.styles import Font, Alignment, PatternFill

from tkinter import * 
from tkinter import messagebox

import socket
import argparse  # for menu system

import importlib.util
import platform
from pathlib import Path
import asyncio
import threading
from tkinter import ttk, filedialog, scrolledtext


# <<<<<<<<<<<<<<<<<<<<<<<<<<     Pre-Sets       >>>>>>>>>>>>>>>>>>>>>>>>>>

author = 'LincolnLandForensics'
description2 = "OSINT: track people down by username, email, ip, phone and website"
version = '3.5.4'

headers_intel = [
    "query", "ranking", "fullname", "url", "email", "user", "phone",
    "business", "fulladdress", "city", "state", "country", "note", "AKA",
    "DOB", "SEX", "info", "misc", "firstname", "middlename", "lastname",
    "associates", "case", "sosfilenumber", "owner", "president", "sosagent",
    "managers", "Time", "Latitude", "Longitude", "Coordinate",
    "original_file", "Source", "Source file information", "Plate", "VIS", "VIN",
    "VYR", "VMA", "LIC", "LIY", "DLN", "DLS", "content", "referer", "osurl",
    "titleurl", "pagestatus", "ip", "dnsdomain", "Tag", "Icon", "Type"
]

headers_locations = [
    "#", "Time", "Latitude", "Longitude", "Address", "Group", "Subgroup"
    , "Description", "Type", "Source", "Deleted", "Tag", "Source file information"
    , "Service Identifier", "Carved", "Name", "business", "number", "street"
    , "city", "county", "state", "zipcode", "country", "fulladdress", "query"
    , "Sighting State", "Plate", "Capture Time", "Capture Network", "Highway Name"
    , "Coordinate", "Capture Location Latitude", "Capture Location Longitude"
    , "Container", "Sighting Location", "Direction", "Time Local", "End time"
    , "Category", "Manually decoded", "Account", "PlusCode", "Time Original", "Timezone"
    , "Icon", "original_file", "case", "Index"
    ]

# --- Global Configuration ---
# User Agent (mimicking the Perl script's UA)
USER_AGENT = "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.31 (KHTML, like Gecko) Chrome/26.0.1410.63 Safari/537.31"
HEADERS = {'User-Agent': USER_AGENT}
    

# Regex section

regex_host = re.compile(
    r'(?i)\b((?:(?!-)[a-zA-Z0-9-]{1,63}(?<!-)\.)+'
    '(?!exe|php|dll|doc|docx|txt|rtf|odt|xls|xlsx|ppt|pptx|bin|pcap|ioc|pdf|mdb|asp|html|xml|jpg|gif$|png'
    '|lnk|log|vbs|lco|bat|shell|quit|pdb|vbp|bdoda|bsspx|save|cpl|wav|tmp|close|ico|ini'
    '|sleep|run|dat$|scr|jar|jxr|apt|w32|css|js|xpi|class|apk|rar|zip|hlp|cpp|crl'
    '|cfg|cer|plg|lxdns|cgi|xn$)'
    '(?:xn--[a-zA-Z0-9]{2,22}|[a-zA-Z]{2,13}))(?:\s|$)')

regex_md5 = re.compile(r'^([a-fA-F\d]{32})$')  # regex_md5        [a-f0-9]{32}$/gm
regex_sha1 = re.compile(r'^([a-fA-F\d]{40})$')  # regex_sha1
regex_sha256 = re.compile(r'^([a-fA-F\d]{64})$')  # regex_sha256
regex_sha512 = re.compile(r'^([a-fA-F\d]{128})$')  # regex_sha512

regex_number = re.compile(r'^(^\d)$')  # regex_number    #Beta
regex_number_fb = re.compile(r'^\d{9,15}$')  # regex_number    #to match facebook user id

regex_ipv4 = re.compile('([1-2]?[0-9]?[0-9]\.[1-2]?[0-9]?[0-9]\.[1-2]?[0-9]?[0-9]\.[1-2]?[0-9]?[0-9])') # test

regex_ipv6 = re.compile('(([0-9a-fA-F]{1,4}:){7,7}[0-9a-fA-F]{1,4}|([0-9a-fA-F]{1,4}:){1,7}:|([0-9a-fA-F]{1,4}:){1,6}:[0-9a-fA-F]{1,4}|([0-9a-fA-F]{1,4}:){1,5}(:[0-9a-fA-F]{1,4}){1,2}|([0-9a-fA-F]{1,4}:){1,4}(:[0-9a-fA-F]{1,4}){1,3}|([0-9a-fA-F]{1,4}:){1,3}(:[0-9a-fA-F]{1,4}){1,4}|([0-9a-fA-F]{1,4}:){1,2}(:[0-9a-fA-F]{1,4}){1,5}|[0-9a-fA-F]{1,4}:((:[0-9a-fA-F]{1,4}){1,6})|:((:[0-9a-fA-F]{1,4}){1,7}|:)|fe80:(:[0-9a-fA-F]{0,4}){0,4}%[0-9a-zA-Z]{1,}|::(ffff(:0{1,4}){0,1}:){0,1}((25[0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9])\.){3,3}(25[0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9])|([0-9a-fA-F]{1,4}:){1,4}:((25[0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9])\.){3,3}(25[0-5]|(2[0-4]|1{0,1}[0-9]){0,1}[0-9]))')    # test

regex_phone = re.compile(r'^(\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}$')  # test

regex_phone11 = re.compile(r'^1\d{10}$')
regex_phone2 = re.compile(r'(\d{3}) \W* (\d{3}) \W* (\d{4}) \W* (\d*)$')


# <<<<<<<<<<<<<<<<<<<<<<<<<<     Menu           >>>>>>>>>>>>>>>>>>>>>>>>>>

# <<<<<<<<<<<<<<<<<<<<<<<<<<     GUI            >>>>>>>>>>>>>>>>>>>>>>>>>>

class TextRedirector(object):
    def __init__(self, widget, tag="stdout"):
        self.widget = widget
        self.tag = tag
        self.terminal = sys.stdout

    def write(self, str):
        self.terminal.write(str)
        try:
            self.widget.configure(state='normal')
            self.widget.insert(END, str, (self.tag,))
            self.widget.see(END)
            self.widget.configure(state='disabled')
        except:
            pass

    def flush(self):
        self.terminal.flush()

def launch_gui():
    """Launch the Tkinter GUI interface."""
    global root, input_file_entry, output_file_entry, progress_bar, status_text, process_button
    global mode_var
    
    root = Tk()
    script_name = os.path.basename(sys.argv[0])
    root.title(f"{script_name} {version}")
    root.geometry("800x650")
    
    # Apply vista theme
    try:
        style = ttk.Style()
        style.theme_use('vista')
    except:
        pass
    
    # Description
    desc_label = Label(root, text=description2, font=("Arial", 10, "bold"), wraplength=750, justify="center")
    desc_label.pack(pady=10)
    
    # Mode Selection (Radio Buttons)
    mode_frame = LabelFrame(root, text="Processing Mode", padx=10, pady=10)
    mode_frame.pack(pady=10, padx=20, fill=X)
    
    mode_var = StringVar(value="hunt")
    
    def update_defaults(*args):
        mode = mode_var.get()
        input_file_entry.delete(0, END)
        output_file_entry.delete(0, END)
        if mode == "hunt":
            input_file_entry.insert(0, "input.txt")
            output_file_entry.insert(0, "Intel__DRAFT_V1.xlsx")
        elif mode == "blurb":
            input_file_entry.insert(0, "Intel__DRAFT_V1.xlsx")
            output_file_entry.insert(0, "OSINT__DRAFT_V1.docx")
        elif mode == "blank":
            input_file_entry.insert(0, "") # No input
            output_file_entry.insert(0, "Intel__DRAFT_V1.xlsx")
        elif mode == "samples":
            input_file_entry.insert(0, "") # No input
            output_file_entry.insert(0, "") # No output
            
    Radiobutton(mode_frame, text="Hunt", variable=mode_var, value="hunt").pack(side=LEFT, padx=20)
    Radiobutton(mode_frame, text="Blurb", variable=mode_var, value="blurb").pack(side=LEFT, padx=20)
    Radiobutton(mode_frame, text="Blank", variable=mode_var, value="blank").pack(side=LEFT, padx=20)
    Radiobutton(mode_frame, text="Samples", variable=mode_var, value="samples").pack(side=LEFT, padx=20)

    # Input File Row
    input_frame = Frame(root)
    input_frame.pack(pady=5, padx=20, fill=X)
    Label(input_frame, text="Input File:", width=15, anchor='w').pack(side=LEFT)
    input_file_entry = Entry(input_frame)
    input_file_entry.insert(0, "input.txt")
    input_file_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
    
    def browse_input():
        file = filedialog.askopenfilename(title="Select Input File")
        if file:
            input_file_entry.delete(0, END)
            input_file_entry.insert(0, file)
                
    Button(input_frame, text="Browse", command=browse_input).pack(side=LEFT)
    
    # Output File Row
    output_frame = Frame(root)
    output_frame.pack(pady=5, padx=20, fill=X)
    Label(output_frame, text="Output File:", width=15, anchor='w').pack(side=LEFT)
    output_file_entry = Entry(output_frame)
    output_file_entry.insert(0, "Intel__DRAFT_V1.xlsx")
    output_file_entry.pack(side=LEFT, fill=X, expand=True, padx=5)
    
    def browse_output():
        file = filedialog.asksaveasfilename(defaultextension=".xlsx", title="Select Output File")
        if file:
            output_file_entry.delete(0, END)
            output_file_entry.insert(0, file)
            
    Button(output_frame, text="Browse", command=browse_output).pack(side=LEFT)
    
    # Add trace after entry widgets are created
    mode_var.trace_add("write", update_defaults)
    
    # Progress Bar
    progress_bar = ttk.Progressbar(root, mode='indeterminate')
    progress_bar.pack(fill=X, padx=22, pady=5)
    
    # Status ScrolledText
    status_label = Label(root, text="Status:")
    status_label.pack(anchor=W, padx=22)
    status_text = scrolledtext.ScrolledText(root, height=15, width=80)
    status_text.pack(pady=5, padx=20, fill=BOTH, expand=True)
    status_text.configure(state='disabled')
    
    # Hunt Button
    def start_processing():
        process_button.config(state=DISABLED)
        progress_bar.start()
        status_text.configure(state='normal')
        status_text.delete(1.0, END)
        status_text.configure(state='disabled')
        
        # Capture settings before starting thread
        current_input = input_file_entry.get()
        current_output = output_file_entry.get()
        current_mode = mode_var.get()
        
        t = threading.Thread(target=run_hunt_task, args=(current_input, current_output, current_mode))
        t.daemon = True
        t.start()
        
    process_button = Button(root, text="Hunt", command=start_processing, 
                            bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), padx=20, pady=5)
    process_button.pack(pady=10)
    
    root.mainloop()

def run_hunt_task(input_path, output_path, mode):
    """Run the hunting process in a background thread."""
    global filename, input_xlsx, output_xlsx, data
    
    # Redirect stdout
    old_stdout = sys.stdout
    sys.stdout = TextRedirector(status_text)
    
    try:
        print(f"Starting {mode} process...")
        # if input_path:
            # print(f"Input file: {input_path}")
            
        # Set up global variables as if we were in CLI
        filename = input_path
        output_xlsx = output_path
        
        # Mocking arguments for the internal logic
        class Args:
            pass
        args = Args()
        args.input = input_path
        args.output = output_path
        args.blurb = (mode == "blurb")
        args.blank = (mode == "blank")
        args.samples = (mode == "samples")
        args.search = (mode == "hunt")
        args.convert = False
        args.locations = False
        args.test = False

        # Run the core logic from main()
        # Note: We need to handle the case where main() would exit
        run_core_logic(args)
        
        # print(f"\nDone! Output saved to: {output_xlsx} line 281")
        print(f"\nDone!")        
        
        # root.after(0, lambda: messagebox.showinfo("Complete", f"Identity Hunt finished successfully!\nOutput: {output_xlsx}"))
        
    except Exception as e:
        print(f"\nError: {str(e)}")
        # root.after(0, lambda: messagebox.showerror("Error", f"An error occurred: {str(e)}"))
    finally:
        sys.stdout = old_stdout
        root.after(0, progress_bar.stop)
        root.after(0, lambda: process_button.config(state=NORMAL))

def run_core_logic(args):
    """Refactored core logic from main() to be callable by CLI and GUI."""
    global data, filename, input_xlsx, output_xlsx, row, emails, ips, phones, users, dnsdomains, websites
    
    # check internet status
    status = internet()
    status2 = is_running_in_virtual_machine()

    if status2 == True:
        print(f'This is a virtual machine. Not checking for internet connectivity')
    elif status == False:
        print(f'CONNECT TO THE INTERNET FIRST.')
        return
    else:
        print('\nINTERNET IS CONNECTED\n')

    filename = 'input.txt'
    input_xlsx = 'Intel2.xlsx'
    row = 1
    data = []
    emails = []
    ips = []
    phones = []
    users = []
    dnsdomains = []
    websites = [] 
    input_file_type = ''

    if args.samples:  
        samples()
        return 

    # default input
    if not args.input: 
        if args.blurb:
            input_file_type = 'xlsx'
            input_xlsx = 'Intel_.xlsx'
        else:
            input_file_type = 'txt'
            filename = 'input.txt'
    elif '.txt' in args.input:
        input_file_type = 'txt'
        filename = args.input
        if os.path.exists(filename) and os.path.getsize(filename) > 0:
            emails,dnsdomains,ips,users,phones,websites = read_text(filename)
    elif '.xlsx' in args.input:
        input_file_type = 'xlsx'
        input_xlsx = args.input
    else:
        input_file_type = 'xlsx'
        input_xlsx = args.input 
   
    # output xlsx
    if not args.output:
        output_xlsx = "Intel__DRAFT_V1.xlsx"          
    else:
        output_xlsx = args.output

    if args.blank:  
        write_intel(data)
        return 

    # if text file input
    if input_file_type == 'txt':
        if not os.path.exists(filename):
            print(f"{filename} doesnt exist.")
            return
        elif os.path.getsize(filename) == 0:
            print(f'{filename} is empty. Fill it with username, email, ip, phone and/or websites.')
            return
        elif os.path.isfile(filename):
            emails,dnsdomains,ips,users,phones,websites = read_text(filename)
            
    # if xlsx input
    elif input_file_type == 'xlsx':
        if not os.path.exists(input_xlsx):
            print(f"{input_xlsx} doesnt exist.")
            return
        elif os.path.getsize(input_xlsx) == 0:
            print(f'{input_xlsx} is empty. Fill it with username, email, ip, phone and/or websites.')
            return
        elif os.path.isfile(input_xlsx):
            data = read_xlsx(input_xlsx)
            if args.convert:
                # print(f'data = {data}') # temp
                print(f' converting {input_xlsx}')
                write_intel_basic(data, output_xlsx)
                return
            if args.blurb:
                write_blurb()
                return
            if args.locations:
                data = read_xlsx_basic(input_xlsx)
                write_locations(data)
                return
            data = read_xlsx(input_xlsx)
  
    # If no modules selected, default to all
    if not any([args.samples, args.blank, args.convert, args.blurb, args.locations]):
        # Default behavior: run all major modules if items are present
        if len(emails) > 0:  
            print(f'Emails = {emails}')
            main_email()
            breachbase()
            carrot_email()
            cyberbackground_email()
            emailosint()
            epios_email()
            etsy_email()
            ghunt()
            google_calendar()
            have_i_been_pwned()
            lookups_io_email()
            holehe_email()
            osintIndustries_email()
            thatsthememail()
            truepeople_email()

        if len(ips) > 0:     
            print(f'IPs = {ips}')
            arinip()
            geoiptool()
            ipinfo()
            main_ip()
            resolverRS()
            whoisip()
            whatismyip()
            
        if len(phones) > 0:
            print(f'phones = {phones}')
            main_phone()
            familytreephone()
            thatsthemphone()
            reversephonecheck()
            validnumber()
            whitepagesphone()
            whocalld()
            zabasearch()  # cloudflare
            
        if args.test:  
            print(f' using test module')
            etsy_email()

        if len(users) > 0:  
            print(f'users = {users}')    
            main_user()
            about()
            allmylinks()
            behance()
            bitbucket()
            blogspot_users()
            bsky()
            cashapp()
            calendly()
            disqus()
            ebay()
            etsy()
            linktree()
            facebook()
            familytree()
            flickr()  # errors 
            freelancer()
            garmin()
            github()
            go()
            goodread()
            # goodread2()   # read error
            ham_radio()
            heylink()
            hodgepodge()    # test
            imageshack()
            instagram()
            instantusername()
            instructables() # # errors 
            inteltechniques() 
            gab()
            keybase()
            kick()
            mastadon()
            medium()
            myshopify()
            myspace_users()
            osint_rocks()
            paypal()
            patreon()
            pinterest()
            poshmark()
            reddit()
            rumble()
            roblox()
            sherlock()
            slack()
            snapchat()
            sportstracker()
            spotify()   # error
            substack()
            threads()
            tiktok()
            tinder()
            truthSocial()
            tumblr()
            twitch()
            twitter()
            venmo()
            vimeo()
            whatnot()
            whatsmyname()
            wordpress()
            wordpress_profiles()
            youtube()
            
            # gravatar()
            
            # linkedin() # auth issue?

        if len(websites) > 0:          
            print(f'websites = {websites}')    
            centralops()
            main_website()
            robtex()
            titles()
            viewdnsdomain()
            whoiswebsite()
            titles()
        
        write_intel(data)

def main():
    parser = argparse.ArgumentParser(description=description2)
    parser.add_argument('-I', '--input', help='', required=False)
    parser.add_argument('-O', '--output', help='', required=False)
    parser.add_argument('-c','--convert', help='convert from old headers', required=False, action='store_true')
    parser.add_argument('-b','--blurb', help='write ossint blurb', required=False, action='store_true')
    parser.add_argument('-B','--blank', help='create blank intel sheet', required=False, action='store_true')
    parser.add_argument('-l','--locations', help='convert intel 2 locations format', required=False, action='store_true')
    parser.add_argument('-s','--samples', help='print sample inputs', required=False, action='store_true')
    parser.add_argument('-S','--search', help='search for OSINT', required=False, action='store_true')    
    parser.add_argument('-t','--test', help='testing individual modules', required=False, action='store_true')

    args = parser.parse_args()

    # If no arguments provided, launch GUI
    if len(sys.argv) == 1:
        launch_gui()
        return 0

    cls()
    print_logo()
    
    # Run the core logic
    run_core_logic(args)

    if not args.blurb:
        input(f"See '{output_xlsx}' for output. Hit Enter to exit...")

    return 0


# <<<<<<<<<<<<<<<<<<<<<<<<<<  Sub-Routines   >>>>>>>>>>>>>>>>>>>>>>>>>>

# <<<<<<<<<<<<<<<<<<<<<<<<<<  Sub-Routines   >>>>>>>>>>>>>>>>>>>>>>>>>>


def about(): # testuser = kevinrose
    """
        parses each user
        creates a url based on the username
        if the webpage exists,it writes it to the output sheet
    """
    
    print(f'\n\t<<<<< about.me users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - about.me')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note, city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '', '')
        (lastname, firstname) = ('', '')
        user = user.rstrip()
        url = (f'https://about.me/{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        for eachline in content.split("\n"):
            if " | about.me" in eachline:        
                fullname = eachline.strip().replace(" | about.me","") # .split("-")(0)

                if ' - ' in fullname:
                    note = fullname.split(' - ')[1]
                    fullname = fullname.split(' - ')[0]  
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''




        if '404' not in str(pagestatus):
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["note"] = note
            row_data["user"] = user
            row_data["city"] = city
            row_data["country"] = country
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname

            data.append(row_data)


def allmylinks():   # https://allmylinks.com/terminator
    print(f'\n\t<<<<< allmylinks users >>>>>')



    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:
        row_data = {}
        (query, ranking) = (user, '4 - allmylinks')
        url = f"https://allmylinks.com/{user}"
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ("")
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        
        # if 1==1:
        if pagestatus == 200:

            fullname = str(titleurl)
            if ' (' in fullname:
                fullname = fullname.split(' (')[0]

            if ' ' in fullname:
                (fullname, firstname, lastname, middlename) = fullname_parse(fullname)

            print(f'{url}	   {fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl             
            data.append(row_data)
            
            
def arinip():    # testuser=    77.15.67.232
    from subprocess import call, Popen, PIPE
    print(f"\n\t<<<<< arin IP's >>>>>")
    for ip in ips:    
        row_data = {}
        (query, note) = (ip, '')
        (city, business, country, zipcode, state) = ('', '', '', '', '')
        (Latitude, Longitude) = ('', '')

        (content, titleurl, pagestatus) = ('', '', '')
        (email, phone, fullname, entity, fulladdress) = ('', '', '', '', '') 
        url = (f'https://search.arin.net/rdap/?query={ip}')        
        
        if sys.platform == 'win32' or sys.platform == 'win64':    
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            
            
            if '403 Forbidden' in content:
                pagestatus = '403 Forbidden'
                content = ''
            for eachline in content.split("\n"):
                
                if "www.ip-adress.com/legal-notice" in eachline:
                    for eachline in content.split("<tr><th>"):
                        if "Country<td>" in eachline:
                            country = eachline.strip().split("Country<td>")[1]
                        elif "City<td>" in eachline:
                            city = eachline.strip().split("City<td>")[1]
                        elif "ISP</abbr><td>" in eachline:
                            business = eachline.strip().split("ISP</abbr><td>")[1]
                        elif "Postal Code<td>" in eachline:
                            zipcode = eachline.strip().split("Postal Code<td>")[1]
                        elif "State<td>" in eachline:
                            state = eachline.strip().split("State<td>")[1]
               
                elif "<tr><th>Country</th><td>" in eachline:
                    country = eachline.strip().split("<tr><th>Country</th><td>")[1]
                    country = country.split("<")[0]
                elif "City: " in eachline:
                    city = eachline.strip().split("<tr><th>City</th><td>")[1]
                    city = city.split("<")[0]            
                elif "<tr><th>Postal Code</th><td>" in eachline:
                    zipcode = eachline.strip().split("<tr><th>Postal Code</th><td>")[1]
                    zipcode = zipcode.split("<")[0] 
            # time.sleep(3) #will sleep for 30 seconds
            if "Fail" in pagestatus:
                pagestatus = 'fail'

        print(f'{ip}	{country}	{city}	{zipcode}')

        if '.' in ip:
            ranking = '9 - ARIN'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["ip"] = ip
            row_data["note"] = note
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["email"] = email
            row_data["phone"] = phone
            row_data["note"] = entity
            row_data["fulladdress"] = fulladdress
            row_data["query"] = query
            
            row_data["city"] = city
            row_data["country"] = country
            row_data["state"] = state
            row_data["zipcode"] = zipcode            
            row_data["Latitude"] = Latitude
            row_data["Longitude"] = Longitude
          
           
            data.append(row_data)

def behance(): # testuser = kevinrose https://www.behance.net/kevinrose
    print(f'\n\t<<<<< behance users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)


# https://www.behance.net/soulja  Robert Dabi - Designer in Nuremberg, Germany

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - behance')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.behance.net/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 


        if str(pagestatus) != '404':

            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - behance",'')
            titleurl = titleurl.replace(" :: Behance",'')
            if ' - Designer in ' in titleurl:
                titletemp = titleurl.split(' - Designer in ')
                try:
                    fullname = titletemp[0]
                    city = titletemp[1]
                except:
                    fullname = str(titleurl)
                
            
            
            if ' ' in str(fullname):
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if fullname == 'behance.com':
                ranking = '9 - behance'
                fullname = ''

            if 1 == 1:
            # if ranking == '4 - behance':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds
        
        
def bitbucket(): # testuser = rick
    print(f'\n\t<<<<< bitbucket users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - bitbucket')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://bitbucket.org/{user}/')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "display_name" in eachline:
                pattern = r'"display_name":\s*"([^"]+)"'
                match = re.search(pattern, eachline)

                if match:
                    fullname = match.group(1)

        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        else:
            fullname = ''

        if '404' not in str(pagestatus):
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname            
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["user"] = user
            row_data["city"] = city
            row_data["country"] = country

            data.append(row_data)

def blogspot_users(): # testuser = kevinrose
    print(f'\n\t<<<<< blogspot users >>>>>')

    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return     
    for user in users:
        row_data = {}
        (query, ranking) = (user, '4 - blogspot')
    
        url = f"https://{user}.blogspot.com"


        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}')

        for eachline in content.split("<"):
            if "og:title" in eachline:
                fullname = eachline
                # print(f'fullname = {fullname}') # temp
                fullname = eachline.strip().split("\"")[1]
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]


        if pagestatus == 200:    
        # if 'Success' in pagestatus:
            titleurl = titleurl_og(content)

            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''
            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname 
            row_data["url"] = url
            row_data["note"] = note
            row_data["user"] = user
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl             
            data.append(row_data)


def bsky(): # testuser = kevinrose
    print(f'\n\t<<<<< bsky users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '5 - bsky')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://bsky.app/profile/{user}.bsky.social')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("  <"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1]
                fullname = fullname.split(" (")[0]
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]
                print(f'note = {note}')

        if '@' in titleurl:
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user  
            row_data["note"] = note 

            data.append(row_data)  


def calendly(): # testuser = kevinrose or jeff
    print(f'\n\t<<<<< calendly users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - calendly')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://calendly.com/{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("  <"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1]
                fullname = str(fullname).split(" (")[0]
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]
                print(f'note = {note}')

        if fullname != '':
            ranking = '5 - calendly'
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user  
            row_data["note"] = note 
            row_data["pagestatus"] = pagestatus 
            # row_data["titleurl"] = titleurl 
            data.append(row_data)   



def carrot_email(): 
    print(f'\n\t<<<<< carrot2 emails >>>>>')
    
    for email in emails:
        row_data = {}
        (query, content, note) = (email, '', '')
        url = f'https://search.carrot2.org/#/search/web/{email}/folders'
        
        # Uncomment this line when request() function is implemented
        # (content, referer, osurl, titleurl, pagestatus) = request(url)

        if '@' in email.lower():
            if 'Enable JavaScript to run this app' in content:
                ranking = '9 - carrot'
            elif 'All retrieved results' in content:
                note = 'Events shown in time zone'
                ranking = '4 - carrot2'
                # print(f' {email}')                
            else:
                ranking = '9 - carrot2'
                # print(f' {email}  ')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email
            row_data["note"] = note
            data.append(row_data)
            # print(f'row_data = {row_data}') # temp

def cashapp(): # testuser = kevinrose
    print(f'\n\t<<<<< cash.app users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - cashapp')
        (fullname, firstname, lastname, middlename, country) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://cash.app/${user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        if '404' not in str(pagestatus):
            pattern = r'"display_name":"([^"]+)"'
            match = re.search(pattern, content)
            if match:
                fullname = match.group(1)

            if fullname.count(' ') == 1:
                firstname, lastname = fullname.split(' ')
                firstname = firstname.title()
                lastname = lastname.upper()
                fullname = (f'{firstname} {lastname}')
                
            elif fullname.count(' ') == 2:
                firstname, middlename, lastname = fullname.split(' ')
                firstname = firstname.title()
                lastname = lastname.upper()
                middlename = middlename.title()
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            
            
            pattern2 = r'"country_code":"([^"]+)"'
            match2 = re.search(pattern2, content)
            if match:
                country = match2.group(1)        
            
            print(f'{url}\t{fullname}\t{country}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["country"] = country
                      
            data.append(row_data)
     

def centralops(): 
    if len(websites) > 0:
        row_data = {}
        ranking = '9 - manual'
        url = ('https://centralops.net/co/')

        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)    

def cls():
    """
        Clears the screen    
    """
    
    linux = 'clear'
    windows = 'cls'
    os.system([linux, windows][os.name == 'nt'])

def convert_timestamp(timestamp, time_orig, timezone):
    timezone = timezone or ''
    time_orig = time_orig or ''

    timestamp = str(timestamp)

    if re.match(r'\d{1,2}/\d{1,2}/\d{4} \d{1,2}:\d{2}:\d{2}\.\d{3} (AM|PM)', timestamp):
        # Define the expected format
        expected_format = "%m/%d/%Y %I:%M:%S.%f %p"

        # Parse the string into a datetime object
        dt_obj = datetime.strptime(timestamp, expected_format)

        # Remove microseconds
        dt_obj = dt_obj.replace(microsecond=0)

        # Format the datetime object back into a string with the specified format
        timestamp = dt_obj.strftime("%Y/%m/%d %I:%M:%S %p")


    # Regular expression to find all timezones
    timezone_pattern = r"([A-Za-z ]+)$"
    matches = re.findall(timezone_pattern, timestamp)

    # Extract the last timezone match
    if matches:
        timezone = matches[-1]
        timestamp = timestamp.replace(timezone, "").strip()
    else:
        timezone = ''
        
    if time_orig == "":
        time_orig = timestamp
    else:
        timezone = ''

    # timestamp = timestamp.replace(' at ', ' ')
    if "(" in timestamp:
        timestamp = timestamp.split('(')
        timezone = timestamp[1].replace(")", '')
        timestamp = timestamp[0]
    elif " CDT" in timestamp:
        timezone = "CDT"
        timestamp = timestamp.replace(" CDT", "")
    elif " CST" in timestamp:
        timezone = "CST"
        timestamp = timestamp.replace(" CST", "")

    formats = [
        "%B %d, %Y, %I:%M:%S %p %Z",    # June 13, 2022, 9:41:33 PM CDT (Flock)
        "%Y:%m:%d %H:%M:%S",
        "%Y-%m-%d %H:%M:%S",
        "%m/%d/%Y %I:%M:%S %p",
        "%m/%d/%Y %I:%M:%S.%f %p", # '3/8/2024 11:06:47.358 AM  # task
        "%m/%d/%Y %I:%M %p",  # timestamps without seconds
        "%m/%d/%Y %H:%M:%S",  # timestamps in military time without seconds
        "%m-%d-%y at %I:%M:%S %p %Z", # test 09-10-23 at 4:29:12 PM CDT
        "%m-%d-%y %I:%M:%S %p",
        "%B %d, %Y at %I:%M:%S %p %Z",
        "%B %d, %Y at %I:%M:%S %p",
        "%B %d, %Y %I:%M:%S %p %Z",
        "%B %d, %Y %I:%M:%S %p",
        "%B %d, %Y, %I:%M:%S %p %Z",
        "%Y-%m-%dT%H:%M:%SZ",  # ISO 8601 format with UTC timezone
        "%Y/%m/%d %H:%M:%S",  # 2022/06/13 21:41:33
        "%d-%m-%Y %I:%M:%S %p",  # 13-06-2022 9:41:33 PM
        "%d/%m/%Y %H:%M:%S",  # 13/06/2022 21:41:33
        "%Y-%m-%d %I:%M:%S %p",  # 2022-06-13 9:41:33 PM
        "%Y%m%d%H%M%S",  # 20220613214133
        "%Y%m%d %H%M%S",  # 20220613 214133
        "%m/%d/%y %H:%M:%S",  # 06/13/22 21:41:33
        "%d-%b-%Y %I:%M:%S %p",  # 13-Jun-2022 9:41:33 PM
        "%d/%b/%Y %H:%M:%S",  # 13/Jun/2022 21:41:33
        "%Y/%b/%d %I:%M:%S %p",  # 2022/Jun/13 9:41:33 PM
        "%d %b %Y %H:%M:%S",  # 13 Jun 2022 21:41:33
        "%A, %B %d, %Y %I:%M:%S %p %Z",  # Monday, June 13, 2022 9:41:33 PM CDT ?
        "%A, %B %d, %Y %I:%M:%S %p"     # Monday, June 13, 2022 9:41:33 PM CDT
    ]

    for fmt in formats:
        try:
            dt_obj = datetime.strptime(timestamp.strip(), fmt)
            timestamp = dt_obj
            return timestamp, time_orig, timezone
                        
        except ValueError:
            pass

    # raise ValueError(f"{time_orig} Timestamp format not recognized")
    return timestamp, time_orig, timezone
    
def cyberbackground_email():# testEmail= kevinrose@gmail.com    
    print(f'\n\t<<<<< cyberbackground emails >>>>>')
    
    for email in emails:
        row_data = {}
        (query, content, note) = (email, '', '')
        url = (f'https://www.cyberbackgroundchecks.com/email/{email}')
        # (content, referer, osurl, titleurl, pagestatus) = request(url)        

        if 1==1:
            if ('results for') in content: 
                note = 'results for'
                ranking = '7 - cyberbackground'

                print(f' {email}   {url}')
          
            else:
            
                ranking = '9 - cyberbackground'

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email
            row_data["note"] = note
            # row_data["content"] = content
                        
            data.append(row_data)

            
def digitalfootprintcheckemail():    # testuser=    kevinrose@gmail.com 
    print(f'\n\t<<<<< digitalfootprintcheck emails >>>>>')    

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return    
    
    for email in emails:
        row_data = {}
        (query, content, note, ranking) = (email, '', '', '9 - digitalfootprintcheck')
        
        url = (f'https://www.digitalfootprintcheck.com/free-checker.html?q={email}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        

        if 'No results found' in content:
            
            print(f'{url} {email}')
        # else:
            # print(f'{url} {email}')
            ranking = '5 - digitalfootprintcheck'
                            
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email  
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl        
            data.append(row_data)
        time.sleep(5) #will sleep for 5 seconds    

def discoverprofile(): 
    if len(users) > 0:
        row_data = {}
        ranking = '7 - manual'
        url = ('https://discoverprofile.com/')
        # url = (f'https://discoverprofile.com/{user}')

        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)


def disqus(): # testuser = kevinrose
    # task: add City, State
    print(f'\n\t<<<<< discus users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking, note) = (user, '5 - discus', '')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        # url = (f'http://disqus.com/{user}')
        # url = (f'http://disqus.com/by/{user}')  
        url = (f'http://disqus.com/by/{user}/about/')  

        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        for eachline in content.split("<"):
            if "twitter:description" in eachline:
                note = eachline.strip().split("\"")[3]
                if 's community of communities now has one central hub.' in note:
                    note = ''
            elif "twitter:title" in eachline:
                fullname = eachline.split('"')[3]
                if ' · ' in fullname:
                    fullname = fullname.split(' · ')[0]
                if user.lower() == fullname.lower():
                    fullname = ''
                fullname = fullname.replace('&Quot;', '"').replace('&quot;', '"')

                if ' ' in fullname:
                    (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                
        if pagestatus == 200:
            print(f'{url}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["user"] = user
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname            
            row_data["note"] = note
            data.append(row_data)

def ebay(): # testuser = kevinrose
    print(f'\n\t<<<<< ebay users >>>>>')
    print(f'\n\tthis can take a while >>>>>')

    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - ebay')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (city, country, fullname, titleurl, pagestatus, note) = ('', '', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.ebay.com/str/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if pagestatus == 200:

            titleurl = titleurl.replace(' | eBay Stores', '')
            fullname = str(titleurl)
            if 'Security Measure | eBay' in fullname:
                fullname = ''
                ranking = '9 - ebay'
                print(f' Security Measure')
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                ranking = '7 - ebay'
            # else:
                # fullname = ''          
            
            
            print(f'{url} {fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["user"] = user
            row_data["note"] = note            
            row_data["lastname"] = lastname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl              

            data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds

def emailosint(): 
    if len(emails) > 0:
        row_data = {}
        ranking ='7 - manual'
        url = ('https://emailosint.org/')
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)

def epios_email(): 
    if len(emails) > 0:
        row_data = {}
        ranking = '9 - manual $$'
        url = (f'https://epieos.com/')
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)
            
            
def etsy(): # testuser = kevinrose https://www.etsy.com/people/kevinrose    # protected by a captcha
    print(f'\n\t<<<<< etsy users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)


    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - etsy')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.etsy.com/people/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            test = 'test'
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            # (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if 1 == 1:
        # if '404' not in str(pagestatus):
            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - Etsy",'')

            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if fullname == 'etsy.com':
                ranking = '9 - etsy'
                fullname = ''

            if 1 == 1:
            # if ranking == '4 - etsy':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds



def etsy_email():  # testemail = kevinrose@gmail.com
    print(f'\n\t<<<<< etsy email >>>>>')

    for email in emails:
        row_data = {}
        email = email.strip()
        query = email.rstrip()
        ranking = '7 - etsy'
        fullname = firstname = middlename = lastname = ''
        business = city = state = country = ''
        user = ''
        SEX = ''
        info = ''
        note = ''
        misc = ''
        associates = AKA = owner = president = ''
        otherurls = ''
        pagestatus = ''
        titleurl = ''

        # Build API URL
        url = f'https://www.etsy.com/api/v3/ajax/public/users/by-identity-optional?identity={email}'

        # Skip non-alpha emails (your original logic)
        if not any(char.isalpha() for char in query):
            continue

        # Request
        content, referer, osurl, titleurl, pagestatus = request(url)

        # Bail out if Etsy returns "fail"
        if 'ail' in pagestatus:
            continue

        # Parse JSON
        try:
            parsed_data = json.loads(content)
        except Exception as error:
            print(f'JSON error: {error}')
            continue

        # -----------------------------
        # USERNAME EXTRACTION
        # -----------------------------
        # Old API (your sample JSON)
        if not user:
            try:
                user = parsed_data.get('login_name', '')
            except:
                pass

        # fullname
        if not fullname:
            try:
                fullname = parsed_data.get('real_name', '')
            except:
                pass

        # firstname
        if not firstname:
            try:
                firstname = parsed_data.get('first_name', '')
            except:
                pass

        # lastname
        if not lastname:
            try:
                lastname = parsed_data.get('last_name', '')
            except:
                pass

        # city
        if not city:
            try:
                city = parsed_data.get('location', '')
            except:
                pass

        # info
        if not info:
            try:
                info = parsed_data.get('bio', '')
            except:
                pass
 
        # SEX
        if not SEX:
            try:
                SEX = parsed_data.get('gender', '')
            except:
                pass

        try:
            accounts = parsed_data['entry'][0]['accounts']
            if len(accounts) > 0:
                misc = accounts[0]['url']
            if len(accounts) > 1:
                associates = accounts[1]['url']
            if len(accounts) > 2:
                AKA = accounts[2]['url']
            if len(accounts) > 3:
                country = accounts[3]['url']
            if len(accounts) > 4:
                owner = accounts[4]['url']
            if len(accounts) > 5:
                president = accounts[5]['url']
        except:
            pass

        # -----------------------------
        # PROFILE URL (FIXED)
        # -----------------------------
        # profile_url = f'http://en.etsy.com/{user}' if user else ''
        profile_url = f'https://www.etsy.com/people/{user}' if user else ''

        print(f'{profile_url}\t{fullname}')

        # Ranking bump
        if fullname or otherurls or note:
            ranking = '3 - etsy'

        # -----------------------------
        # BUILD ROW DATA
        # -----------------------------
        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["fullname"] = fullname
        row_data["middlename"] = middlename
        row_data["firstname"] = firstname
        row_data["lastname"] = lastname
        row_data["url"] = profile_url
        row_data["email"] = email
        row_data["business"] = business
        row_data["city"] = city
        row_data["state"] = state
        row_data["country"] = country
        row_data["user"] = user
        row_data["SEX"] = SEX
        row_data["info"] = info
        row_data["misc"] = misc
        row_data["note"] = note
        row_data["titleurl"] = titleurl
        row_data["associates"] = associates
        row_data["owner"] = owner
        row_data["president"] = president
        row_data["AKA"] = AKA
        # row_data["content"] = content
        data.append(row_data)



def facebook(): # testuser = kevinrose
    print(f'\n\t<<<<< facebook users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - Facebook')

        (fullname,lastname,firstname) = ('','','')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        user = user.rstrip()
        url = (f'https://facebook.com/{user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
            
        fullname = str(titleurl).strip()
        if ' | Facebook' in titleurl:
            fullname = fullname.replace(' | Facebook', '')
            ranking = '4 - Facebook'

        if fullname == 'Facebook':
            fullname = ''            
            
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)        

        # if 1==1:
        if fullname != '':
            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname            
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user
            row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            row_data["titleurl"] = titleurl 
            data.append(row_data)        


def familytree(): 
    print(f'\n\t<<<<< Manually check familytreenow.com >>>>>')
    row_data = {}
    (query, ranking, note) = ('', '9 - manual', 'See FamilyTree link for Possible Relatives and Possible Associates')
    url = ('https://www.familytreenow.com/search/')

    # row_data["query"] = query
    row_data["ranking"] = ranking
    row_data["url"] = url
    row_data["note"] = note
    data.append(row_data)


def familytreephone():# DROP THE LEADING 1
    print(f'\n\t<<<<< familytree phone numbers >>>>>')
    for phone in phones:
        row_data = {}
        (query, ranking) = (phone, '7 - familytree')

        (country, city, state, zipcode, case, note) = ('', '', '', '', '', 'See FamilyTree link for Possible Relatives and Possible Associates')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', '')

        url = (f'https://www.familytreenow.com/search/genealogy/results?phoneno=%s' %(phone.lstrip('1')))
        # url = (f'https://www.familytreenow.com/search/genealogy/results?phoneno={phone.lstrip('1')}')

        state = phone_state_check(phone, state).replace('?', '')
        
        if 1==1:        
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone
            row_data["state"] = state
            row_data["note"] = note
            data.append(row_data)

def findwhocallsyou():# testPhone= 
    # note: need to add dishes if there are none
    print(f'\n\t<<<<< findwhocallsyou phone numbers >>>>>')
    for phone in phones:
        phone = phone.replace('-', '')
 
        row_data = {}
        (query, ranking) = (phone, '9 - findwhocallsyou')
  
        (country, city, zipcode, case, note) = ('', '', '', '', '')

        url = (f'https://findwhocallsyou.com/{phone}')    

        if 1==1:
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone
            row_data["note"] = note

            data.append(row_data)

def fiverr():    # testuser=    kevinrose
    print(f'\n\t<<<<< fiverr users >>>>>')

    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - fiverr')
        (fullname, firstname, lastname, middlename, note, DOB)  = ('','','','', '', '')
        (misc) = ('')

        url = (f'https://www.fiverr.com/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

            
        # time.sleep(1) # will sleep for 1 seconds
        if pagestatus == 200:
            ranking = '3 - fiverr'
        
        if 1==1:
        # if 'alternate' in content:
            print(f'{url}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["note"] = note
            row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            row_data["titleurl"] = titleurl  

            data.append(row_data)


def flickr(): # testuser = kevinrose
    print(f'\n\t<<<<< flickr users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - flickr')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.rstrip()
        url = (f'https://www.flickr.com/people/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))            
            for eachline in content.split("\n"):
                if "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1]
                elif "<meta content=" in eachline and "name=\"description\"" in eachline:
                    note = eachline.split('"')[1]
        except TypeError as error:
            print(f'{error}')
        pagestatus = str(pagestatus)


        if '404' not in pagestatus and 'ail' not in pagestatus:
            if fullname.lower() == user.lower():
                fullname = ''
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                fullname = ''

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user            
            row_data["note"] = note              
            data.append(row_data)           


def freelancer(): # testuser = kevinrose
    print(f'\n\t<<<<< freelancer users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '5 - freelancer')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        city = ''
        user = user.rstrip()
        url = (f'https://www.freelancer.com/u/{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        titleurl = titleurl.replace(' Profile | Freelancer','')
        if '404' not in str(pagestatus):
           
            if ' ' in titleurl:
                fullname = str(titleurl)
            
            if fullname.lower() == user.lower():
                fullname = ''

            if '' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if 'Sign Up Free' in fullname:
                ranking = '9 - freelancer'
                fullname = ''

            
            if 'Browser ' not in fullname:
                print(f'{url}	{fullname}') 
                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["city"] = city
                row_data["user"] = user            
                row_data["titleurl"] = titleurl            
                row_data["pagestatus"] = pagestatus                             
                data.append(row_data)  

def friendfinder():    # testuser=  kevinrose   # java math problem
    print(f'\n\t<<<<< friendfinder users >>>>>')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - friendfinder')
        (fullname, city, country, note, DOB, SEX) = ('', '', '', '', '', '')
        (firstname, middlename, lastname) = ('', '', '')
        url = (f'https://www.friendfinder-x.com/profile/{user}')

        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        # (fullname, info, note) = ('', '', '')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        if 1==1:
        # if 'Register to Find' not in titleurl:
            print(f'{url}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user 
            row_data["city"] = city            
            row_data["country"] = country            
            row_data["note"] = note            
            row_data["DOB"] = DOB            
            row_data["SEX"] = SEX            
            row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            row_data["titleurl"] = titleurl  

            data.append(row_data) 


def fullname_parse(fullname):
    (firstname, lastname, middlename) = ('', '', '')
    fullname = fullname.strip()
    if ' ' in fullname:
        fullname_parts = fullname.split(' ')
        firstname = fullname_parts[0].title()
        lastname = fullname_parts[-1].upper()
        
        if len(fullname_parts) == 2:
            fullname = (f'{firstname} {lastname}')
        elif len(fullname_parts) == 3:
            middlename = fullname_parts[1].title()
            fullname = (f'{firstname} {middlename} {lastname}')

    return fullname, firstname, middlename, lastname


def garmin(): # testuser = kevinrose
    print(f'\n\t<<<<< garmin users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - garmin')
        # (fullname, lastname, firstname, case, SEX) = ('','','','','')
        user = user.rstrip()

        url = (f'https://connect.garmin.com/modern/profile/{user}')

        (content, referer, osurl, titleurl, pagestatus) = request(url)

        if 'twitter:card' not in content:
        
            fullname = str(titleurl)
            fullname = fullname.split(" (")[0]
            fullname = fullname.replace("Garmin Connect","").strip()

            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user            
            # row_data["titleurl"] = titleurl            
            # row_data["city"] = city            
            # row_data["country"] = country            
            # row_data["note"] = note            
            # row_data["DOB"] = DOB            
            # row_data["SEX"] = SEX            
            # row_data["content"] = content            

            data.append(row_data) 


def geoiptool():
    print(f'\n\t<<<<< geoiptool IPs >>>>>')



    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for ip in ips:
        row_data = {}
        (query, ranking) = (ip, '9 - geoiptool')
        url = f"https://www.geodatatool.com/en/?ip={ip}"
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ("")
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 


        if pagestatus == 200:

            fullname = str(titleurl)
            if ' (' in fullname:
                fullname = fullname.split(' (')[0]

            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            print(f'{url}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["ip"] = ip
            row_data["url"] = url
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl    
            
            data.append(row_data)
            
            
def ghunt():    # testEmail= kevinrose@gmail.com
    for email in emails:
        row_data = {}
        (query, note) = (email, '')
        note = (f'cd C:\Forensics\scripts\python\git-repo\GHunt && ghunt email {email}')

        if '@' in email.lower():
            if email.endswith('gmail.com'):
                ranking = ('9 - ghunt')
                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["email"] = email
                row_data["note"] = note
                data.append(row_data)




def github(): # testuser = kevinrose
    print(f'\n\t<<<<< github users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - github')
        (city, country, fullname, titleurl, pagestatus, content, info) = ('', '', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://github.com/{user}')

        (content, referer, osurl, titleurl, pagestatus) = request(url)
        if "(" in titleurl:
            fullname = str(titleurl).strip()
            if "(" in fullname:
                fullname = fullname.split("(")[1]
                fullname = fullname.strip().split(")")[0]

        for eachline in content.split("  <"):

                
            if "og:description" in eachline:
                note = eachline.strip()
                note = note.strip().split("\"")[1]
            elif "twitter:title" in eachline:
                info = eachline.strip()
                info = info.strip().split("\"")[1].replace(' - Overview', '')
                info = (f'https://x.com/{info}')
  
        if 'GitHub is where ' in note:
            note = ''

        if '404' not in pagestatus and 'ail' not in pagestatus:
            if fullname.lower() == user.lower():
                fullname = ''
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                fullname = ''

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename            
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user 
            row_data["note"] = note             
            row_data["info"] = info
            data.append(row_data)  

def go(): 
    # https://breachdirectory.org  crack hashes with weakpass.com - cloudflare
    # gosearch.exe kevinrose --no-false-positives

    if len(users) > 0:
        row_data = {}
        ranking = '9 - manual'
        url = ('https://breachdirectory.org')

        row_data["ranking"] = ranking
        row_data["url"] = url
        row_data["note"] = 'gosearch.exe {user} --no-false-positives'
        data.append(row_data)     


def goodread2():
    from user_scanner.core.helpers import get_random_user_agent
    from user_scanner.core.orchestrator import Result, make_request


    print(f'\n\t<<<<< goodread users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    # if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        # return (content, referer, osurl, titleurl, pagestatus)


    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - goodread')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.goodreads.com/{user}')
            

        headers = {
            "User-Agent": get_random_user_agent(),
        }

        # Perform HTTP request
        resp = make_request(
            url,
            headers=headers,
            http2=True,
            follow_redirects=True
        )

        final_url = str(resp.url)

        # 404 → available
        if resp.status_code == 404:
            # return Result.available(url=final_url)
            url = Result.available(url=final_url)
            
        # 200 → taken
        elif resp.status_code == 200:
            extra = {}

            # Extract <title>
            titleurl = re.search(
                r'<title>(.*?)</title>',
                resp.text,
                re.IGNORECASE
            )
            # if title_match:
                # titleurl = title_match.group(1).strip()

            # Extract og:title
            name_match = re.search(
                r'<meta[^>]*property=["\']og:title["\'][^>]*content=["\']([^"\']+)["\']',
                resp.text,
                re.IGNORECASE
            )
            if name_match:
                extracted = name_match.group(1).strip()
                extracted = extracted.replace(' (', '')
                fullname = extracted

            # return Result.taken(extra=extra, url=final_url)

            # Unexpected status
            # return Result.error(f"Unexpected response status: {resp.status_code}")


            if 1 == 1:
            # if ranking == '4 - goodread':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds


def goodread(): # testuser = kevinrose https://www.goodread.net/kevinrose
    print(f'\n\t<<<<< goodread users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    # if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        # try:
            # return (content, referer, osurl, titleurl, pagestatus)
        # except:pass

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - goodread')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.goodreads.com/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if '404' not in str(pagestatus):
            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - goodread",'')

            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if fullname == 'goodread.com':
                ranking = '9 - goodread'
                fullname = ''

            if 1 == 1:
            # if ranking == '4 - goodread':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds


def gravatar(): # testuser = kevinrose      https://en.gravatar.com/kevinrose
    print(f'\n\t<<<<< gravatar users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking, SEX, city, business, email) = (user, '7 - gravatar', '', '', '', '')
        (city, country, fullname, titleurl, pagestatus, info) = ('', '', '', '', '', '')
        (info, lastname, firstname, note, otherurls, misc) = ('', '','', '', '', '')
        (associates, AKA, country, state, owner, president) = ('', '', '', '', '', '')
        (middlename) = ('')
        user = user.rstrip()
        url = (f' https://gravatar.com/{user}.json')        

        if any(char.isalpha() for char in user):
            (content, referer, osurl, titleurl, pagestatus) = request(url)

            (parsed_data) = []
            if 'ail' not in pagestatus:
                try:            
                    parsed_data = json.loads(content)
                except TypeError as error:
                    print(f'{error}')
                print(f'I got this far')    # temp

                # print(f'parsed_data = {parsed_data}')   # temp
                try:
                    info = parsed_data['entry'][0]['photos'][0]['value']
                except:pass  

                if 'familyName' in content:
                    try:
                        fullname = parsed_data['entry'][0]['name']['formatted']
                    except:pass  



                if 'preferredUsername' in content:
                    try:
                        user = parsed_data['entry'][0]['preferredUsername']
                    except:pass  

                if 'emails' in content:
                    try:
                        email = parsed_data['entry'][0]['emails'][0]['value']
                    except:pass  
                    
                if 'company' in content:
                    try:
                        business = parsed_data['entry'][0]['company']
                    except:pass  


                if 'pronouns' in content:
                    try:
                        SEX = parsed_data['entry'][0]['pronouns']
                    except:pass  
                    
                    if 'she/her' in SEX:
                        SEX = 'F'
                    elif 'he/him' in SEX:
                        SEX = 'M'                    

                if 'currentLocation' in content:
                    
                    try:
                        city = parsed_data['entry'][0]['currentLocation']
                    except:pass  
                    
                    if isinstance(city, str) and ', ' in city:
                        
                        try:
                            parts = city.split(', ', 1)
                            city, state = parts[0], parts[1]
                        except:pass
                    
                if 'aboutMe' in content:
                    try:
                        note = parsed_data['entry'][0]['aboutMe'].replace('&amp', '&')
                    except:pass  

                if 'displayName' in content:
                    
                    try:
                        fullname = parsed_data['entry'][0]['displayName']
                    except:pass  
                    fullname = fullname.replace('&amp', '&')
                    if fullname == user:
                        fullname = ''


                if 'accounts' in content:
                    misc = parsed_data['entry'][0]['accounts'][0]['url']
                    try:
                        associates = parsed_data['entry'][0]['accounts'][1]['url']
                    except:pass
                    try:
                        AKA = parsed_data['entry'][0]['accounts'][2]['url']
                    except:pass                    
                    try:
                        country = parsed_data['entry'][0]['accounts'][3]['url']
                    except:pass
                    try:
                        owner = parsed_data['entry'][0]['accounts'][4]['url']
                    except:pass                    
                    try:
                        president = parsed_data['entry'][0]['accounts'][5]['url']
                    except:pass  

                if ' ' in fullname:
                    (fullname, firstname, middlename, lastname) = fullname_parse(fullname)

                    

                url = (f'http://en.gravatar.com/{user}')
                print(f'{url}	{fullname}') 
                
                if fullname != '' or otherurls != '' or note != '': 
                    ranking = '3 - gravatar'

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["middlename"] = middlename                
                row_data["firstname"] = firstname
                row_data["lastname"] = lastname
                
                row_data["url"] = url
                row_data["email"] = email
                row_data["business"] = business
                row_data["city"] = city 
                row_data["user"] = user 
                row_data["SEX"] = SEX
                row_data["info"] = info 
                row_data["misc"] = misc                 
                row_data["note"] = note 
                row_data["titleurl"] = titleurl            
                row_data["city"] = city
                row_data["state"] = state                 
                row_data["country"] = country            
                row_data["note"] = note            
                row_data["associates"] = associates   
                row_data["owner"] = owner  
                row_data["president"] = president  
                row_data["AKA"] = AKA 
                # row_data["content"] = content                
                row_data["country"] = country           
                data.append(row_data)

def google_calendar():# testEmail= kevinrose@gmail.com    
    print(f'\n\t<<<<< google calendar emails >>>>>')
    
    for email in emails:
        row_data = {}
        (query, content, note) = (email, '', '')
        url = (f'https://calendar.google.com/calendar/u/0/embed?src={email}&pli=1')
        # (content, referer, osurl, titleurl, pagestatus) = request(url)        

        if 'gmail.com' in email.lower():
            if ('you do not have the permission to view') in content: 
                note = 'you do not have the permission to view'
                ranking = '7 - calendar'

                print(f' {email}   {url}')
            elif ('Events shown in time zone') in content:
                note = 'Events shown in time zone'
                ranking = '4 - calendar'             
            else:
                ranking = '9 - calendar'

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email
            row_data["note"] = note
            data.append(row_data)
            # print(f'row_data = {row_data}') # temp


def ham_radio(): # testuser = K9CYC
    '''
    make a list of all users that might also be FCC Ham radio call signs.
    Ham radio call signs can lead to verified full names, address, etc.
    '''
    print(f'\n\t<<<<< ham radio users >>>>>')
    callsigns = []
    call_sign_pattern = r'([BFGKIMNRW]|[0-9][A-Z]|[A-Z][0-9]|[A-Z][A-Z])[0-9][0-9A-Z]{0,2}[A-Z]?'

    for user in users:    
        row_data = {}
        (query, ranking, info) = (user, '5 - ham radio', '')
        (note) = ('https://www.arrl.org/advanced-call-sign-search')
        user = user.rstrip().upper()
        url = (f'https://www.radioqth.net/lookup')        

        if len(user) <= 6 and re.match(call_sign_pattern, user):
            callsigns.append(user)
            
            
    info = f'{", ".join(callsigns)}'  # Join all call signs with commas and append to info string
    
    if callsigns:
            print(f'{url}	{info}') 
            row_data["query"] = info
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["note"] = note
            data.append(row_data)


def have_i_been_pwned(): 
    if len(emails) > 0:
        # hibp-api-key = 'YOUR_API_KEY'
        row_data = {}
        ranking ='7 - manual'
        url = ('https://haveibeenpwned.com')
        # url = ('https://haveibeenpwned.com/api/v3/breachedaccount/{email}')
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)

def hibp_email():# testEmail= kevinrose@gmail.com    
    print(f'\n\t<<<<< HaveIBeenPwned emails >>>>>')
    headers = {"User-Agent": "OSINTTool"}
    for email in emails:
        row_data = {}
        (query, ranking, content, note) = (email, '', '', '')
        url = f"https://haveibeenpwned.com/api/v3/breachedaccount/{email}"
        try:
            response = requests.get(url, headers=headers, timeout=10)
            if response.status_code == 200:
                ranking = ('3 - HaveIBeenPwned')

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["url"] = url
                row_data["email"] = email
                row_data["note"] = note
                data.append(row_data)

            if response.status_code == 404:
                print(f"No breach found for {email}")
        except Exception as e:
            print(f"HIBP check failed for {email}: {e}")
        return False


def heylink():  # kevin
    print(f'\n\t<<<<< heylink users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:
        row_data = {}
        (query, ranking) = (user, '5 - heylink')
        url = f"https://heylink.me/{user}"
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ("")
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        # if 1==1:
        if pagestatus == 200:

            fullname = str(titleurl)
            fullname = fullname.replace('HeyLink.me | ','')


            if ' ' in fullname:
                (fullname, firstname, lastname, middlename) = fullname_parse(fullname)


            print(f'{url}	   {fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl    
            
            data.append(row_data)
            
def hodgepodge():            
    SITES = [
        "http://dating.ru/{user}",
        "http://en.gravatar.com/{user}",
        "http://promodj.com/{user}",
        "http://uid.me/{user}",
        "https://{user}.blogspot.com/",
        "https://{user}.booth.pm/",
        "https://{user}.carbonmade.com",
        "https://{user}.carrd.co/",
        "https://{user}.contently.com/",
        "https://{user}.crevado.com",
        "https://{user}.empretienda.com.ar",
        "https://{user}.exposure.co/",
        "https://{user}.gitbook.io/",
        "https://{user}.itch.io/",
        "https://{user}.jimdosite.com",
        "https://{user}.livejournal.com",
        "https://{user}.mitiendanube.com/",
        "https://{user}.newgrounds.com",
        "https://{user}.omg.lol",
        "https://{user}.rajce.idnes.cz/",
        "https://{user}.smugmug.com",
        "https://{user}.tistory.com/",
        "https://{user}.tumblr.com",
        "https://{user}.webnode.cz/",
        "https://{user}.weebly.com/",
        "https://{user}.wix.com",
        "https://{user}.www.nn.ru/",
        "https://2Dimensions.com/a/{user}",
        "https://7dach.ru/profile/{user}",
        "https://addons.wago.io/user/{user}",
        "https://admireme.vip/{user}/",
        "https://airbit.com/{user}",
        "https://akniga.org/profile/{user}",
        "https://ameblo.jp/{user}",
        "https://apclips.com/{user}",
        "https://apex.tracker.gg/apex/profile/origin/{user}/overview",
        "https://api.boosty.to/v1/blog/{user}",
        "https://api.cropty.io/v1/auth/{user}",
        "https://api.dailymotion.com/user/{user}?fields=id,username,screenname,description,avatar_720_url,cover_250_url,followers_total,following_total,videos_total,country,created_time,verified,url",
        "https://api.destream.net/siteapi/v2/live/details/{user}",
        "https://api.discogs.com/users/{user}",
        "https://api.mojang.com/minecraft/profile/lookup/name/{user}",
        "https://api.mojang.com/users/profiles/minecraft/{user}",
        "https://api.niftygateway.com/user/profile-and-offchain-nifties-by-url/?profile_url={user}",
        "https://api.omg.lol/address/{user}/info",
        "https://api.scratch.mit.edu/users/{user}",
        "https://api.stats.fm/api/v1/users/{user}",
        "https://api.tracker.gg/api/v2/apex/standard/profile/origin/{user}",
        "https://api.warframe.market/v2/user/{user}",
        "https://api.zhihu.com/books/people/{user}/publications?offset=0&limit=5",
        "https://archiveofourown.org/users/{user}",
        "https://asciinema.org/{user}",
        "https://asciinema.org/~{user}",
        "https://ask.fedoraproject.org/u/{user}",
        "https://atcoder.jp/users/{user}",
        "https://audiojungle.net/user/{user}",
        "https://bandcamp.com/{user}",
        "https://bezuzyteczna.pl/uzytkownicy/{user}",
        "https://bit.ly/{user}",
        "https://bitbucket.org/{user}/",
        "https://blitztactics.com/{user}",
        "https://blog.naver.com/{user}",
        "https://boardgamegeek.com/user/{user}",
        "https://boosty.to/{user}",
        "https://bugcrowd.com/{user}",
        "https://buymeacoff.ee/{user}",
        "https://buzzfeed.com/{user}",
        "https://caddy.community/u/{user}/summary",
        "https://calendly.com/api/booking/profiles/{user}",
        "https://career.habr.com/{user}",
        "https://ch.tetr.io/u/{user}",
        "https://chaos.social/@{user}",
        "https://chaturbate.com/{user}",
        "https://choice.community/u/{user}/summary",
        "https://clapperapp.com/{user}",
        "https://client.warpcast.com/v2/user-by-username?username={user}",
        "https://codeberg.org/{user}",
        "https://codeberg.org/api/v1/users/{user}",
        "https://codeforces.com/api/user.info?handles={user}",
        "https://codeforces.com/profile/{user}",
        "https://codepen.io/{user}",
        "https://coderlegion.com/user/{user}",
        "https://coderwall.com/{user}.json",
        "https://coderwall.com/{user}",
        "https://codesandbox.io/u/{user}",
        "https://codesnippets.fandom.com/wiki/User:{user}",
        "https://coinvote.cc/profile/{user}",
        "https://community.bitwarden.com/u/{user}/summary",
        "https://community.brave.com/u/{user}/",
        "https://community.cartalk.com/u/{user}/summary",
        "https://community.cloudflare.com/u/{user}",
        "https://community.cryptomator.org/u/{user}",
        "https://community.eintracht.de/fans/{user}",
        "https://community.icons8.com/u/{user}/summary",
        "https://community.n8n.io/u/{user}/summary",
        "https://community.native-instruments.com/profile/{user}",
        "https://community.oracle.com/people/{user}",
        "https://community.signalusers.org/u/{user}",
        "https://community.windy.com/user/{user}",
        "https://community.wolfram.com/web/{user}/home",
        "https://cracked.ax/{user}",
        "https://crates.io/api/v1/users/{user}",
        "https://crates.io/users/{user}",
        "https://crowdin.com/profile/{user}",
        "https://ctan.org/author/{user}",
        "https://cults3d.com/en/users/{user}/creations",
        "https://cyber.harvard.edu/people/{user}",
        "https://d3.ru/user/{user}/posts",
        "https://dev.to/{user}",
        "https://dev.to/api/users/by_username?url={user}",
        "https://discourse.joplinapp.org/u/{user}",
        "https://discourse.jupyter.org/u/{user}.json",
        "https://discourse.jupyter.org/u/{user}",
        "https://discourse.jupyter.org/u/{user}/summary",
        "https://discourse.mozilla.org/u/{user}.json",
        "https://discourse.mozilla.org/u/{user}",
        "https://discourse.wicg.io/u/{user}/summary",
        "https://discuss.elastic.co/u/{user}",
        "https://discuss.hashicorp.com/u/{user}.json",
        "https://discuss.hashicorp.com/u/{user}",
        "https://discuss.kotlinlang.org/u/{user}.json",
        "https://discuss.kotlinlang.org/u/{user}",
        "https://discuss.python.org/u/{user}.json",
        "https://discuss.python.org/u/{user}",
        "https://discuss.python.org/u/{user}/summary",
        "https://dmoj.ca/user/{user}",
        "https://donatello.to/{user}",
        "https://dribbble.com/{user}",
        "https://dribbble.com/{user}/about",
        "https://egpu.io/forums/profile/{user}/",
        "https://en.liberapay.com/{user}",
        "https://en.wikipedia.org/wiki/Special:CentralAuth/{user}?uselang=qqx",
        "https://en.wikipedia.org/wiki/User:{user}  ",
        "https://f3.cool/{user}/",
        "https://fameswap.com/user/{user}",
        "https://flipboard.com/@{user}",
        "https://fortnitetracker.com/profile/all/{user}",
        "https://forum.arduino.cc/u/{user}.json",
        "https://forum.arduino.cc/u/{user}",
        "https://forum.arduino.cc/u/{user}/summary",
        "https://forum.cfx.re/u/{user}/summary",
        "https://forum.dangerousthings.com/u/{user}",
        "https://forum.elixirforum.com/u/{user}.json",
        "https://forum.elixirforum.com/u/{user}",
        "https://forum.f-droid.org/u/{user}.json",
        "https://forum.f-droid.org/u/{user}",
        "https://forum.ghost.org/u/{user}.json",
        "https://forum.ghost.org/u/{user}",
        "https://forum.hackersploit.org/u/{user}",
        "https://forum.hackthebox.com/u/{user}",
        "https://forum.ionicframework.com/u/{user}",
        "https://forum.leasehackr.com/u/{user}/summary/",
        "https://forum.rclone.org/u/{user}",
        "https://forum.spells8.com/u/{user}",
        "https://forum.sublimetext.com/u/{user}",
        "https://forums.digitalspy.com/profile/{user}",
        "https://forums.envato.com/u/{user}",
        "https://forums.mmorpg.com/profile/{user}",
        "https://forums.opera.com/api/user/{user}",
        "https://forums.opera.com/user/{user}",
        "https://forums.terraria.org/index.php?search/42798315/&c[users]={user}&o=relevance",
        "https://fosstodon.org/@{user}",
        "https://fotka.com/profil/{user}",
        "https://foursquare.com/{user}",
        "https://framapiaf.org/@{user}",
        "https://freesound.org/people/{user}/",
        "https://gallog.dcinside.com/{user}",
        "https://gamefaqs.gamespot.com/community/{user}",
        "https://genius.com/{user}",
        "https://genius.com/artists/{user}",
        "https://gitea.com/{user}",
        "https://gitea.com/api/v1/users/{user}",
        "https://gitee.com/{user}",
        "https://gitee.com/api/v5/users/{user}",
        "https://gitlab.com/{user}",
        "https://gpodder.net/user/{user}/",
        "https://habr.com/ru/users/{user}",
        "https://habr.com/ru/users/{user}/",
        "https://hackaday.io/{user}",
        "https://hackenproof.com/hackers/{user}",
        "https://hackerearth.com/@{user}",
        "https://hackerone.com/{user}",
        "https://hackerrank.com/{user}",
        "https://hackmd.io/@{user}",
        "https://help.nextcloud.com/u/{user}/summary",
        "https://hive.blog/@{user}",
        "https://hosted.weblate.org/user/{user}/",
        "https://hub.docker.com/u/{user}/",
        "https://hub.docker.com/v2/users/{user}/",
        "https://hubpages.com/@{user}",
        "https://huggingface.co/{user}",
        "https://ifttt.com/p/{user}",
        "https://ifunny.co/user/{user}",
        "https://imgup.cz/{user}",
        "https://imgur.com/user/{user}",
        "https://independent.academia.edu/{user}",
        "https://irc-galleria.net/user/{user}",
        "https://irecommend.ru/users/{user}",
        "https://issuu.com/{user}",
        "https://issuu.com/query?format=json&_=3210224608766&profileUsername={user}&action=issuu.user.get_anonymous",
        "https://itch.io/profile/{user}",
        "https://jbzd.com.pl/uzytkownik/{user}",
        "https://ko-fi.com/{user}",
        "https://kwork.ru/user/{user}",
        "https://laracasts.com/@{user}",
        "https://last.fm/user/{user}",
        "https://launchpad.net/~{user}",
        "https://leetcode.com/{user}",
        "https://leetcode.com/u/{user}/",
        "https://lemmy.world/api/v3/user?username={user}",
        "https://lemmy.world/u/{user}",
        "https://letterboxd.com/{user}",
        "https://lichess.org/@/{user}",
        "https://lichess.org/api/user/{user}",
        "https://linktr.ee/{user}",
        "https://linuxfr.org/users/{user}",
        "https://listed.to/@{user}",
        "https://lobste.rs/u/{user}",
        "https://lottiefiles.com/{user}",
        "https://mamot.fr/@{user}",
        "https://mastodon.social/@{user}",
        "https://mastodon.xyz/@{user}",
        "https://meta.discourse.org/u/{user}.json",
        "https://meta.discourse.org/u/{user}",
        "https://meta.wikimedia.org/wiki/Special:CentralAuth/{user}",
        "https://mix.com/{user}",
        "https://moikrug.ru/{user}",
        "https://mstdn.io/@{user}",
        "https://mstdn.social/@{user}",
        "https://musescore.com/{user}",
        "https://music.yandex.ru/handlers/library.jsx?owner={user}",
        "https://music.yandex.ru/users/{user}",
        "https://my.flightradar24.com/{user}",
        "https://myanimelist.net/profile/{user}",
        "https://namemc.com/profile/{user}",
        "https://namu.wiki/w/%EC%82%AC%EC%9A%A9%EC%9E%90:{user}",
        "https://nationstates.net/nation={user}",
        "https://nationstates.net/region={user}",
        "https://news.ycombinator.com/user?id={user}",
        "https://nl.pepper.com/profile/{user}",
        "https://notabug.org/{user}",
        "https://note.com/{user}",
        "https://nothing.community/u/{user}",
        "https://nyaa.si/user/{user}",
        "https://observablehq.com/@{user}",
        "https://omg.lol/{user}",
        "https://opencollective.com/{user}",
        "https://opengameart.org/users/{user}",
        "https://opengovus.com/search?q={user}%2C+LLC",
        "https://opensource.com/users/{user}",
        "https://osu.ppy.sh/users/{user}",
        "https://packagist.org/users/{user}/",
        "https://paragraph.com/@{user}",
        "https://paragraph.com/api/blogs/@{user}",
        "https://pastebin.com/u/{user}",
        "https://pentesterlab.com/profile/{user}",
        "https://pikabu.ru/@{user}",
        "https://pixelfed.social/{user}/",
        "https://platzi.com/p/{user}/",
        "https://play.google.com/store/apps/developer?id={user}",
        "https://playerdb.co/api/player/minecraft/{user}",
        "https://playstrategy.org/@/{user}",
        "https://plugins.gradle.org/u/{user}",
        "https://pokemonshowdown.com/users/{user}",
        "https://polymart.org/user/{user}",
        "https://pornhub.com/users/{user}",
        "https://pr0gramm.com/api/profile/info?name={user}",
        "https://profil.chatujme.cz/{user}",
        "https://prog.hu/azonosito/info/{user}",
        "https://programming.dev/u/{user}",
        "https://pronouns.page/@{user}",
        "https://psnprofiles.com/{user}",
        "https://rateyourmusic.com/~{user}",
        "https://replit.com/@{user}",
        "https://robertsspaceindustries.com/citizens/{user}",
        "https://ruby-forum.com/u/{user}/summary",
        "https://rubygems.org/profiles/{user}",
        "https://rumble.com/user/{user}",
        "https://satsis.info/user/{user}",
        "https://scholar.harvard.edu/{user}",
        "https://scratch.mit.edu/users/{user}",
        "https://seoforum.com/@{user}",
        "https://sessionize.com/{user}",
        "https://sketchfab.com/{user}",
        "https://slashdot.org/~{user}",
        "https://slides.com/{user}",
        "https://social.tchncs.de/@{user}",
        "https://soundcloud.com/{user}",
        "https://sourceforge.net/u/{user}",
        "https://sourceforge.net/u/{user}/",
        "https://spacehey.com/{user}",
        "https://speakerdeck.com/{user}",
        "https://speedrun.com/users/{user}",
        "https://spletnik.ru/user/{user}",
        "https://splits.io/users/{user}",
        "https://stats.fm/{user}",
        "https://status.cafe/users/{user}",
        "https://steemit.com/@{user}",
        "https://swapd.co/u/{user}",
        "https://tagged.com/profile.html?uid={user}",
        "https://tellonym.me/{user}",
        "https://tenor.com/users/{user}",
        "https://themeforest.net/user/{user}",
        "https://topmate.io/{user}",
        "https://traewelling.de/@{user}",
        "https://traktrain.com/{user}",
        "https://translate.jellyfin.org/user/{user}/",
        "https://trashbox.ru/users/{user}",
        "https://tuna.voicemod.net/user/{user}",
        "https://tweakers.net/gallery/{user}",
        "https://twitchtracker.com/{user}",
        "https://ubuntu-mate.community/u/{user}.json",
        "https://ubuntu-mate.community/u/{user}",
        "https://uk.advfn.com/forum/profile/{user}",
        "https://ultimate-guitar.com/u/{user}",
        "https://unsplash.com/@{user}",
        "https://untappd.com/user/{user}",
        "https://users.rust-lang.org/u/{user}.json",
        "https://users.rust-lang.org/u/{user}",
        "https://valorantforums.com/u/{user}",
        "https://velog.io/@{user}/posts",
        "https://vero.co/{user}",
        "https://vimeo.com/api/v2/{user}/info.json",
        "https://virgool.io/@{user}",
        "https://VJudge.net/user/{user}",
        "https://vk.com/{user}",
        "https://vsco.co/{user}",
        "https://wakatime.com/@{user}",
        "https://warframe.market/profile/{user}",
        "https://wiki.archlinux.org/title/User:{user}",
        "https://wowhead.com/user={user}",
        "https://write.as/{user}",
        "https://www.1337x.to/user/{user}/",
        "https://www.7cups.com/@{user}",
        "https://www.9gag.com/u/{user}",
        "https://www.adultism.com/profile/{user}",
        "https://www.adultism.com/profile/{user}/friends",
        "https://www.airliners.net/user/{user}/profile",
        "https://www.americanthinker.com/author/{user}/",
        "https://www.artstation.com/{user}",
        "https://www.autofrage.net/nutzer/{user}",
        "https://www.babepedia.com/user/{user}",
        "https://www.baby.ru/u/{user}",
        "https://www.babyblog.ru/user/{user}",
        "https://www.bandcamp.com/{user}",
        "https://www.bandlab.com/api/v1.3/users/{user}",
        "https://www.bazar.cz/{user}/",
        "https://www.behance.net/{user}/appreciated",
        "https://www.biggerpockets.com/users/{user}",
        "https://www.blipfoto.com/{user}",
        "https://www.bookcrossing.com/mybookshelf/{user}/",
        "https://www.buymeacoffee.com/{user}",
        "https://www.cgtrader.com/{user}",
        "https://www.championat.com/user/{user}",
        "https://www.chess.com/member/{user}",
        "https://www.chollometro.com/profile/{user}",
        "https://www.clozemaster.com/players/{user}",
        "https://www.clubhouse.com/@{user}",
        "https://www.cnet.com/profiles/{user}/",
        "https://www.codewars.com/users/{user}",
        "https://www.colourlovers.com/lover/{user}",
        "https://www.coroflot.com/{user}",
        "https://www.couchsurfing.com/people/{user}",
        "https://www.cracked.com/members/{user}/",
        "https://www.credly.com/users/{user}",
        "https://www.curseforge.com/members/{user}/projects",
        "https://www.dealabs.com/profile/{user}",
        "https://www.defensivecarry.com/members/?username={user}",
        "https://www.discogs.com/user/{user}",
        # "https://www.donationalerts.com/api/v1/user/{user}/donationpagesettings",
        "https://www.donationalerts.com/r/{user}",
        "https://www.drive2.ru/users/{user}",
        "https://www.erome.com/{user}",
        "https://www.exophase.com/user/{user}/",
        "https://www.eyeem.com/u/{user}",
        "https://www.fandom.com/u/{user}",
        "https://www.fanpop.com/fans/{user}",
        "https://www.figma.com/@{user}",
        "https://www.finanzfrage.net/nutzer/{user}",
        "https://www.fixya.com/users/{user}",
        "https://www.fl.ru/users/{user}",
        "https://www.freecodecamp.org/{user}",
        "https://www.freepik.com/author/{user}",
        "https://www.furaffinity.net/user/{user}",
        "https://www.gamespot.com/profile/{user}/",
        "https://www.geocaching.com/p/default.aspx?u={user}",
        "https://www.gesundheitsfrage.net/nutzer/{user}",
        "https://www.getmyuni.com/user/{user}",
        "https://www.giantbomb.com/profile/{user}/",
        "https://www.grailed.com/{user}",
        "https://www.gumroad.com/{user}",
        "https://www.gutefrage.net/nutzer/{user}",
        "https://www.hackster.io/{user}",
        "https://www.hotukdeals.com/profile/{user}",
        "https://www.iconfinder.com/{user}",
        "https://www.ifttt.com/p/{user}",
        "https://www.imagefap.com/profile/{user}",
        "https://www.imood.com/users/{user}",
        "https://www.instapaper.com/p/{user}",
        "https://www.instructables.com/member/{user}/",
        "https://www.interpals.net/{user}",
        "https://www.jeuxvideo.com/profil/{user}",
        "https://www.kaggle.com/{user}",
        "https://www.kongregate.com/accounts/{user}",
        "https://www.last.fm/user/{user}",
        "https://www.librarything.com/profile/{user}",
        "https://www.linux.org.ru/people/{user}/profile",
        "https://www.livelib.ru/reader/{user}",
        "https://www.mapmytracks.com/{user}",
        "https://www.memrise.com/user/{user}/",
        "https://www.mercadolivre.com.br/perfil/{user}",
        "https://www.mixcloud.com/{user}/",
        "https://www.motorradfrage.net/nutzer/{user}",
        "https://www.mydealz.de/profile/{user}",
        "https://www.mydramalist.com/profile/{user}",
        "https://www.myminifactory.com/users/{user}",
        "https://www.nairaland.com/{user}",
        "https://www.needrom.com/author/{user}/",
        "https://www.newamerica.org/our-people/{user}/",
        "https://www.nintendolife.com/users/{user}",
        "https://www.nitrotype.com/racer/{user}",
        "https://www.npmjs.com/~{user}",
        "https://www.openstreetmap.org/user/{user}",
        "https://www.pepper.pl/profile/{user}",
        "https://www.pepperdeals.com/profile/{user}",
        "https://www.pepperdeals.se/profile/{user}",
        "https://www.periscope.tv/{user}/",
        "https://www.pinkbike.com/u/{user}/",
        "https://www.polygon.com/users/{user}",
        "https://www.preisjaeger.at/profile/{user}",
        "https://www.producthunt.com/@{user}",
        "https://www.promodescuentos.com/profile/{user}",
        "https://www.redbubble.com/people/{user}",
        "https://www.reisefrage.net/nutzer/{user}",
        "https://www.reverbnation.com/{user}",
        "https://www.rusfootball.info/user/{user}/",
        "https://www.sbazar.cz/{user}",
        "https://www.scribd.com/{user}",
        "https://www.shitpostbot.com/user/{user}",
        "https://www.shpock.com/shop/{user}/items",
        "https://www.slant.co/users/{user}",
        "https://www.smule.com/{user}",
        "https://www.speedrun.com/api/v1/users/{user}",
        "https://www.speedrun.com/users/{user}",
        "https://www.sporcle.com/user/{user}/people",
        "https://www.sportlerfrage.net/nutzer/{user}",
        "https://www.sports.ru/profile/{user}/",
        "https://www.strava.com/athletes/{user}",
        "https://www.thefirearmsforum.com/members/?username={user}",
        "https://www.themoviedb.org/u/{user}",
        "https://www.threads.net/api/v1/users/web_profile_info/?username={user}",
        "https://www.toster.ru/user/{user}/answers",
        "https://www.tradingview.com/u/{user}",
        "https://www.vinted.pt/member/general/search?search_text={user}",
        "https://www.vlr.gg/user/{user}",
        "https://www.warriorforum.com/members/{user}.html",
        "https://www.wattpad.com/user/{user}",
        "https://www.weforum.org/people/{user}",
        "https://www.wordnik.com/users/{user}",
        "https://www.wykop.pl/ludzie/{user}",
        "https://www.zhihu.com/people/{user}",
        "https://www.znanylekarz.pl/{user}",
        "https://www.zomato.com/{user}/reviews",
        "https://xboxgamertag.com/search/{user}",
        "https://youpic.com/photographer/{user}/",
        "https://zmarsa.com/uzytkownik/{user}",
        ]

    print("\n\t<<<<< hodgepodge of user sites >>>>>")

    for username in users:
        row_data = {}
        query = username
        ranking = "8"
        case = 'test'
        for template in SITES:
            url = template.format(user=username)

            try:
                content, referer, osurl, titleurl, pagestatus = request(url)
            except Exception:
                continue
            pagestatus = str(pagestatus)
            
            if pagestatus and "success" in str(pagestatus).lower():
                print(url)
                m = re.search(r'//(?:[A-Za-z0-9_-]+\.)*([A-Za-z0-9_-]+)\.(?:com|org|net)', url, re.IGNORECASE)
                
                if m:
                    ranking = f'8 - {m.group(1)}'
                else:
                    # Match domain minus last part (e.g., carrd.co → carrd)
                    m2 = re.search(
                        r'//(?:[A-Za-z0-9_-]+\.)*([A-Za-z0-9_-]+)\.[A-Za-z]{2,6}(?:/|$)',
                        url,
                        re.IGNORECASE
                    )

                    if m2:
                        ranking = f'8 - {m2.group(1)}'
                    else:
                        ranking = '8'

                row_data = {
                    "query": query,
                    "ranking": ranking,
                    "user": username,
                    "url": url,
                    "titleurl": titleurl,
                    # "case": case,                    
                    "pagestatus": pagestatus

                }

                data.append(row_data) 
        time.sleep(2) #will sleep for 2 seconds        
            
def holehe_email(): # testEmail= kevinrose@gmail.com
    print(f'\n\t<<<<< holehe emails >>>>>')    # temp
    for email in emails:
        row_data = {}
        (query, ranking) = (email, '7 - manual')
        note = (f'cd C:\Forensics\scripts\python\git-repo\holehe && holehe -NP --no-color --no-clear --only-used {email}')

        if '@' in email.lower():
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["email"] = email
            row_data["note"] = note
            data.append(row_data)

def imageshack(): # testuser = ToddGilbert

    print(f'\n\t<<<<< imageshack users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - imageshack')
        try:
            user = user.strip()
        except:
            pass        
        url = (f'https://imageshack.com/user/{user}')

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 's Images' in titleurl:
            # fullname = str(titleurl)
            print(f'{url}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            data.append(row_data)
            
def instagram():    # testuser=    kevinrose     # add info
    print(f'\n\t<<<<< instagram users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - instagram')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        url = (f'https://instagram.com/{user}/')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)

            # Regular expression to find the profile_id
            pattern = r'"profile_id":"(\d+)"'

            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # Load the JSON data
                    datatemp = json.loads(eachline)

                    # Extract the description value and print it
                    note = datatemp['description']
                elif 'profile_id' in eachline:

                    # Search for the pattern in the content
                    match = re.search(pattern, content)

                    # Extract and print the profile_id if found
                    if match:
                        misc = match.group(1)   # user id
                    else:
                        misc = ''
        except:
            pass

        if '@' in titleurl:
            fullname = str(titleurl).split(" (")[0]
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                ranking = '3 - instagram'
            else:
                fullname = ''
            if "@" in fullname:
                (fullname, firstname, lastname, middlename) = ('', '', '', '')
        if '@' in titleurl:
            print(f'{url}	{fullname} {misc}')   
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            # row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            row_data["note"] = note
            row_data["misc"] = misc

            data.append(row_data)

def instagramtwo(): #alpha
    # from lib.colors import red,white,green,reset
    self = 'kevinrose'

    response = self.session.get(self.url)
    if response.status_code != 200:
        exit(f'[-] instagram: user not found')
    response = response.json()
    user = response['graphql']['user']
           
    data = {'Profile photo': user['profile_pic_url_hd'],
                 'Username': user['username'],
                 'User ID': user['id'],
                 'External URL': user['external_url'],
                 'Bio': user['biography'],
                 'Followers': user['edge_followed_by']['count'],
                 'Following': user['edge_follow']['count'],
                 'Pronouns': user['pronouns'],
                 'Images': user['edge_owner_to_timeline_media']['count'],
                 'Videos': user['edge_felix_video_timeline']['count'],
                 'Reels': user['highlight_reel_count'],
                 'Is private?': user['is_private'],
                 'Is verified?': user['is_verified'],
                 'Is business account?': user['is_business_account'],
                 'Is professional account?': user['is_professional_account'],
                 'Is recently joined?': user['is_joined_recently'],
                 'Business category': user['business_category_name'],
                 'Category': user['category_enum'],
                 'Has guides?': user['has_guides']
    }
    print
    
    print(f"\n{user['full_name']} | Instagram{reset}")
    for key, value in data.items():
       print(f"├─ {key}: {value}")


def instantusername(): # testuser = kevinrose
    print(f'\n\t<<<<< instantusername users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking, note) = (user, '9 - instantusername', '')
        
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        user = user.rstrip()
        url = (f'https://instantusername.com/?q={user}')
        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["fullname"] = fullname                
            row_data["url"] = url
            row_data["user"] = user
            # row_data["note"] = note
            data.append(row_data)

        except TypeError as error:
            print(f'{error}')


def instructables(): # testuser = kevinrose
    print(f'\n\t<<<<< instructables users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - instructables')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.instructables.com/member/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))            
        except TypeError as error:
            print(f'{error}')            
        pagestatus = str(pagestatus)    
        if '404' not in str(pagestatus):
            if "'" in titleurl:
                titleurl = titleurl.split("'")[0]
            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''           
            
            
            print(f'{url}	{fullname}') 
            
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            data.append(row_data)
            
def inteltechniques(): 
    if len(users) > 0:
        row_data = {}
        ranking = '9 - manual'
        url = ('https://inteltechniques.com/tools/')

        row_data["ranking"] = ranking
        row_data["url"] = url
        # row_data["user"] = user
        data.append(row_data)                

def internet(host="8.8.8.8", port=53, timeout=3):
    """
    Host: 8.8.8.8 (google-public-dns-a.google.com)
    OpenPort: 53/tcp
    Service: domain (DNS/TCP)
    """
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error as ex:
        print(f'this is an error')  # temp
        print(ex)
        return False

def ip_address(dnsdomain):
    (ip) = ('')
    """
    Ping the URL and return the IP address
    """
    try:
        ip = socket.gethostbyname(dnsdomain)
    except socket.gaierror:
        ip = ''
    return ip


def is_running_in_virtual_machine():
    # Check for common virtualization artifacts
    virtualization_artifacts = [
        "/dev/virtio-ports",
        "/dev/vboxguest",
        "/dev/vmware",
        "/dev/qemu",
        "/sys/class/dmi/id/product_name",
        "/proc/scsi/scsi",
    ]

    for artifact in virtualization_artifacts:
        if os.path.exists(artifact):
            return True
            print('This is running in a virtual machine')
    return False



def linktree(): # testuser = kevinrose https://linktr.ee/kevinrose
    print(f'\n\t<<<<< linktree users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)


    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - linktree')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://linktr.ee/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        if '404' not in str(pagestatus):
            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - linktree",'')

            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if fullname == 'linktree.com':
                ranking = '9 - linktree'
                fullname = ''
            fullname = fullname.replace(' | Linktree', '')


            if 1 == 1:
            # if ranking == '4 - linktree':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds

def lookups_io_email():     # https://lookups.io/email/kevinrose@gmail.com
    if len(emails) > 0:
        row_data = {}
        ranking ='7 - manual'
        url = ('https://lookups.io')
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)


def gab():  # kevinrose
    print(f'\n\t<<<<< gab users >>>>>')



    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:
        row_data = {}
        (query, ranking) = (user, '6 - gab')
        url = f"https://gab.com/{user}"
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ("")
        try:
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        if pagestatus == 200:


            fullname = str(titleurl)
            if ' (' in fullname:
                fullname = fullname.split(' (')[0]

            if ' ' in fullname:
                (fullname, firstname, lastname, middlename) = fullname_parse(fullname)

            print(f'{url}	   {fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl    
            
            data.append(row_data)

            
def ipinfo():    # testuser=    77.15.67.232
    from subprocess import call, Popen, PIPE
    print(f"\n\t<<<<< ipinfo.io IP's >>>>>")
    for ip in ips:    
        row_data = {}
        (query, note, dnsdomain) = (ip, '', '')
        (city, business, country, state, entity, zipcode) = ('', '', '', '', '', '')
        (Latitude, Longitude, Coordinate) = ('', '', '')

        (content, titleurl, pagestatus) = ('', '', '')
        url = (f'https://ipinfo.io/{ip}/json')        

        if sys.platform == 'win32' or sys.platform == 'win64':    
            response = requests.get(url, headers=HEADERS, timeout=10, verify=False)
            response.raise_for_status()
            dataip = response.json()            

            if isinstance(dataip, list) and dataip:
                dataip = dataip[0]

            if not isinstance(dataip, dict):
                 print(item(f"{Fore.RED}API returned unexpected dataip format (not a dictionary).{Fore.WHITE}"))
                 return


            dnsdomain = dataip.get('hostname', '')
            city = dataip.get('city', '')
            country = f"{dataip.get('country', '')}"
            state = f"{dataip.get('region', 'N/A')}"
            Coordinate = dataip.get('loc', 'N/A')
            note = dataip.get('org', 'N/A')            
            
            if ',' in Coordinate:
                parts = Coordinate.split(',')
            try:
                Latitude = float(parts[0].strip())
                Longitude = float(parts[1].strip())
            except: pass

        print(f'{ip}	{country}	{city}')

        if "Wrong ip" not in content:
            ranking = '6 - ipinfo'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["ip"] = ip
            row_data["note"] = note
            row_data["zipcode"] = zipcode
            row_data["url"] = url
            row_data["note"] = note
            row_data["query"] = query
            row_data["city"] = city
            row_data["country"] = country
            row_data["state"] = state
            row_data["Latitude"] = Latitude
            row_data["Longitude"] = Longitude
            row_data["Coordinate"] = Coordinate            
            row_data["dnsdomain"] = dnsdomain
            data.append(row_data)            
            
def keybase():    # testuser=    kevin
    print(f'\n\t<<<<< keybase.io users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - keybase')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (info, misc, associates) = ('', '', '')

        url = (f'https://keybase.io/{user}')
        
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        try:

            content = content.strip()
            titleurl = titleurl.strip()
            fullname = str(titleurl)
            if " (" in fullname:
                fullname = fullname.split(" (")[1].split(")")[0]
                if 'Keybase' in fullname:
                    fullname = ''
                
                if ' ' in fullname:
                    (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                    if 'Keybase' in fullname:
                        fullname = ''
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    fullname = eachline.split('\"')[1].split(' (')[0]
                    if 'Keybase' in fullname:
                        fullname = ''                    
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # Load the JSON data
                    datatemp = json.loads(eachline)

                    # Extract the description value and print it
                    note = datatemp['description']
                elif 'rel="me"' in eachline and 'a href' in eachline:
                    if '"' in eachline:
                        eachline = eachline.split('"')[1]
                    if note == '':
                        note = eachline
                    elif info == '':
                        info = eachline                        
                    elif misc == '':
                        misc = eachline   
                    elif associates == '':
                        associates = eachline 

        except:
            pass
            
        # time.sleep(1) # will sleep for 1 seconds
        # if 'what you are looking for...it does not exist' not in content:
        # if 'Your conversation will be end-to-end encrypted' in content:
        if 'Following' in content:
            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            row_data["info"] = info
            row_data["note"] = note
            row_data["misc"] = misc            
            row_data["associates"] = associates
            
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl             
            
            
            data.append(row_data)


def kick():    # kevin
    print(f'\n\t<<<<< kick users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:
        row_data = {}
        (query, ranking) = (user, '4 - kick')
        url = f"https://kick.com/{user}/about"
        (fullname, firstname, lastname, middlename, note) = ('', '', '', '', '')
        (note) = ("")

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        if 'About' in titleurl:
            print(f'{url}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["note"] = note
            # row_data["pagestatus"] = pagestatus   
            
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl 
            data.append(row_data)
            
    
def kik(): # testuser = kevinrose
    print(f'\n\t<<<<< kik users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - kik')
        (fullname, titleurl, pagestatus, content) = ('', '', '', '')
        (note, firstname, lastname, photo, misc, lastseen) = ('', '', '', '', '', '')
        (otherurl, info, misc) = ('', '', '')
        user = user.rstrip()
        url = (f'https://ws2.kik.com/user/{user}')
        
        misc = (f'https://kik.me/{user}')
               
        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}')
        for eachline in content.split(","):
            if "firstName" in eachline:
                firstname = eachline.strip().split(":")[1]
                firstname = firstname.strip('\"')
                firstname = firstname.title()
            elif "lastName" in eachline:
                lastname = eachline.strip().split(":")[1]
                lastname = lastname.strip('\"')
                lastname = lastname.upper()
                lastname = lastname.replace('"}', '')
            elif "displayPicLastModified" in eachline:
                note = eachline.strip().split(":")[1]
             
                
            elif "displayPic\"" in eachline:
                photo = eachline.strip().split(":\"")[1].split("\"")[0].replace("\\","")
                
            fullname = (f'{firstname} {lastname}')
            fullname = fullname.replace("\"}","")
            # if ' ' in fullname:
                # (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            # else:
                # fullname = ''

        if 1 == 1:
        # if pagestatus == 200:
        # if '404' not in str(pagestatus):
            print(f'{url}	{fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = misc
            row_data["note"] = note
            row_data["misc"] = url
            row_data["info"] = photo        
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            row_data["titleurl"] = titleurl 
            data.append(row_data)


def linkedin(): # testuser = kevinrose
    print(f'\n\t<<<<< linkedin users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - linkedin')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        try:
            user = user.strip()
        except:
            pass        
        url = (f'https://www.linkedin.com/in/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        fullname = str(titleurl).replace(' | LinkedIn','')

        for eachline in content.split("\n"):
            if "og:description" in eachline:
                note = eachline.strip()
                try:
                    note = note.split('"')[3]
                    print(f'note = {note}') # temp
                except:
                    pass

        if pagestatus == 200: 
            if ' - ' in fullname:
                fullname = fullname.split(' - ')[0]
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                fullname = ''

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user 
            row_data["note"] = note            
            data.append(row_data)   

def mastadon(): # testuser = kevinrose
    print(f'\n\t<<<<< mastadon users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - mastadon')
        (fullname, lastname, firstname) = ('','','')
        (note, info, content, pagestatus) = ('', '', '', '')
    
        user = user.rstrip()
        url = (f'https://mastodon.social/@{user}')
        note = (f'https://mastodon.social/api/v2/search?q={user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))            
        except TypeError as error:
            print(f'{error}')
            
        for eachline in content.split("\n"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1].split(' (')[0]

        if 'accounts\":[{' in content:        
            datatemp = json.loads(content)              # Convert JSON data to Python dictionary
            fullname = datatemp["accounts"][0]["display_name"]
            info = datatemp['accounts'][0]['avatar']

        if user == fullname:
            fullname = ''
               
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        pagestatus = str(pagestatus)
        if "uccess" in pagestatus and 'This resource could not be found' not in content:
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["note"] = note
            row_data["info"] = info            
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            # row_data["middlename"] = middlename
            row_data["lastname"] = lastname

            data.append(row_data)

def medium(): # testuser = kevinrose 
    print(f'\n\t<<<<< medium users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - medium')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://medium.com/@{user}/about')
        try:
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if 'About' in titleurl:
            ranking = '4 - medium'
            titleurl = titleurl.replace('About – ', '').replace(' – Medium', '')

            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if 1 == 1:
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        # time.sleep(5) #will sleep for 5 seconds


def myfitnesspal(): # testuser = kevinrose protected by cloudflare
    print(f'\n\t<<<<< myfitnesspal users >>>>>')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - myfitnesspal')
        
        user = user.rstrip()
        url = (f'https://www.myfitnesspal.com/profile/{user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if 1==1:
        # if "uccess" in pagestatus and 'This resource could not be found' not in content:
            print(f'{url}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            row_data["titleurl"] = titleurl 
            data.append(row_data)


def myshopify():    # testuser=    rothys
    print(f'\n\t<<<<< myshopify users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '5 - myshopify')
        url = (f'https://{user}.myshopify.com/')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (fullname, info, note) = ('', '', '')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    info = eachline.split('\"')[1].split(' (')[0]
                elif 'ProfilePage\",\"description' in eachline:
                    info = eachline
                    # # Load the JSON data
                    datatemp = json.loads(eachline)

                    # # Extract the description value and print it
                    note = datatemp['description']

        except:
            pass
            
        time.sleep(1) # will sleep for 1 seconds
        if 'Success' in pagestatus:

            response = requests.get(url)

            if response.history:
                otherurls = response.url
                note = (f'redirects to {otherurls}')

            print(f'{url}	{note}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["info"] = info
            row_data["note"] = note           
            row_data["fullname"] = fullname

            data.append(row_data)


def myspace_users():
    print(f'\n\t<<<<< myspace users >>>>>')


    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)
    else:
        print(f'Playwright is installed')

    for user in users:
        row_data = {}
        (query, ranking) = (user, '4 - myspace')
        url = f"https://myspace.com/{user}"
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ("")
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        # if 1==1:
        if pagestatus == 200:
        # if pagestatus == 200 and ('Your search did not return any results') not in content:
        # if 'Success' in pagestatus and ('Your search did not return any results') not in content:
            fullname = str(titleurl)
            if ' (' in fullname:
                fullname = fullname.split(' (')[0]

            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            print(f'{url}	   {fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl    
            
            data.append(row_data)

def main_email(): 

    for email in emails:
        row_data = {}
        (query, ranking) = (email, '1 - main')

        if '@' in email.lower():
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["email"] = email
            data.append(row_data)

def main_ip(): 

    for ip in ips:
        row_data = {}
        (query, ranking) = (ip, '1 - main')
        if 1==1:
        # if '.' in ip.lower():
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["ip"] = ip
            data.append(row_data)
            
def main_phone(): 

    for phone in phones:
        row_data = {}
        (query, ranking, state) = (phone, '1 - main', '')
        state = phone_state_check(phone, state)
    
        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["phone"] = phone
        row_data["state"] = state
        data.append(row_data)            


def main_user():
    for user in users:
        row_data = {}
        (query, ranking) = (user, '1 - main')

        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["user"] = user
        data.append(row_data)      


def main_website():
    for website in websites:
        row_data = {}
        (query, ranking) = (website, '1 - main')

        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["website"] = website
        data.append(row_data) 

def massageanywhere():    # testuser=   Misty0427
    print(f'\n\t<<<<< massageanywhere users >>>>>')

    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - massageanywhere')
        url = (f'https://www.massageanywhere.com/profile/{user}')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        (info, note, fulladdress) = ('', '', '')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        if 1==1:
        # if 'Profile for' in titleurl:
            if 1==1:
            # if 'MassageAnywhere.com Profile for ' in titleurl:  
                titleurl = titleurl.replace('MassageAnywhere.com Profile for ','')
                if ' of ' in titleurl:
                    # titleurl = titleurl.split(' of ')
                    fullname = str(titleurl).split(' of ')[0]
                    fulladdress = titleurl.split(' of ')[1]

                    if ' ' in fullname:
                        (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                    else:
                        (fullname, firstname, lastname, middlename) = ('', '', '', '')

            print(f'{url}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname

            
            # row_data["phone"] = phone
            # row_data["note"] = note
            
            # row_data["city"] = city
            # row_data["content"] = content
            row_data["fulladdress"] = fulladdress
            row_data["titleurl"] = titleurl            
            row_data["pagestatus"] = pagestatus            
                        
            data.append(row_data)

def message_square(message):
    horizontal_line = f"+{'-' * (len(message) + 2)}+"
    empty_line = f"| {' ' * (len(message))} |"

    print(horizontal_line)
    print(empty_line)
    print(f"| {message} |")
    print(empty_line)
    print(horizontal_line)
    print(f'')

def noInternetMsg():
    '''
    prints a pop-up that says "Connect to the Internet first"
    '''
    window = Tk()
    window.geometry("1x1")
      
    w = Label(window, text ='Translate-Inator', font = "100") 
    w.pack()
    messagebox.showwarning("Warning", "Connect to the Internet first") 

def osintIndustries_email(): 
    if len(emails) > 0:
        row_data = {}
        ranking = '9 - manual'
        url = ('https://app.osint.industries/')
        # app.osint.industries
        
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)

def osint_rocks():    # testEmail= kevinrose@gmail.com
    row_data = {}
    (query, note, url) = ('', '', 'https://osint.rocks/')
    ranking = ('7 - osint.rocks')
    # row_data["query"] = query
    row_data["ranking"] = ranking
    row_data["url"] = url
    data.append(row_data)



def patreon(): # testuser = kevinrose
    print(f'\n\t<<<<< patreon users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '5 - patreon')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        user = user.rstrip()
        url = (f'https://www.patreon.com/{user}/creators')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        if '404' not in str(pagestatus):
            print(f'{url}{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            # row_data["pagestatus"] = pagestatus    
            # row_data["content"] = content                
            # row_data["titleurl"] = titleurl 


            data.append(row_data)


def paypal(): # testuser = kevinrose
    print(f'\n\t<<<<< paypal users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - paypal')

        (city, country, fullname, titleurl, pagestatus, state) = ('', '', '', '', '', '')
        (fulladdress, lastname, firstname, photo) = ('', '', '', '')
        (email, phone, info, misc) = ('', '', '', '')
        user = user.rstrip()
        url = (f'https://www.paypal.com/paypalme/{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        (note) = ('')
        if '404' not in str(pagestatus):
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                else:
                    # define the regular expression pattern
                    pattern = r'{"userInfo":{(.*?)}}'

                    # match the pattern to the input string
                    match = re.search(pattern, content)

                    # extract the data variable from the match object
                    if match:
                        # datatemp = match
                        datatemp = match.group(1)
                        note = datatemp
                        note = note.replace("null",'\"\"')
        # else:
            # print(f'{user}') 
        if ':' in note:
            titleurl = titleurl.replace('PayPal.Me','').strip() # task
            # print(f'titleurl = {titleurl}   hello world')   # temp   
            # fullname = str(titleurl)       

            # Extract variables using regex
            try:
                fullname = re.search(r'"displayName":"(.*?)"', datatemp).group(1)

                if ' ' in fullname:
                    (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            except:pass                
            try:    
                email = re.search(r'"displayEmail":(null|".*?")', datatemp).group(1)
            except:pass                
            if "null" in email:
                email = ''
            try:    
                phone = re.search(r'"displayMobilePhone":(null|".*?")', datatemp).group(1)
            except:pass                
            if "null" in phone:
                phone = ''

            try:    
                photo = re.search(r'"coverPhotoUrl":(null|".*?")', datatemp).group(1)
                photo = photo.replace('"', '')
            except:pass                
            if "null" in photo:
                photo = ''
            
            try:    
                city = re.search(r'"displayAddress":"(.*?)"', datatemp).group(1)
                if ", " in city:
                   temp = city.split(", ")
                   city = temp[0]
                   state = temp[1]                
            except:pass                
            try:    
                info = re.search(r'"website":(null|".*?")', datatemp).group(1)
            except:pass
            if "null" in info:
                info = ''
                
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["phone"] = phone 
            row_data["email"] = email 
            row_data["city"] = city            
            row_data["state"] = state
            row_data["fulladdress"] = fulladdress            
            row_data["fullname"] = fullname
            # row_data["info"] = photo
            row_data["misc"] = misc           
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            row_data["note"] = note            
  
            data.append(row_data)


def ping_url(url):
    try:
        r = requests.get(url, timeout=5)
        return 200 <= r.status_code < 400
    except requests.exceptions.RequestException:
        return False


def phone_dashes(phone):
    # Remove any non-digit characters from the phone number
    cleaned_number = ''.join(filter(str.isdigit, phone))
    
    # Add dashes at appropriate positions
    formatted_number = '-'.join([cleaned_number[:3], cleaned_number[3:6], cleaned_number[6:]])
    
    return formatted_number


def phone_state_check(phone, state):
    """
    Checks if the phone number starts with any U.S. state area code and if the state is empty ('').
    If both conditions are met, sets the state accordingly.

    Args:
        phone (str): The phone number as a string.
        state (str): The state abbreviation.

    Returns:
        str: Updated state abbreviation (state based on area code if conditions are met, otherwise original state).
    """
    phone = phone.lstrip('1').strip()

    area_codes_by_state = {
        "AL?": ["205", "251", "256", "334", "938"],
        "AK?": ["907"],
        "AZ?": ["480", "520", "602", "623", "928"],
        "AR?": ["479", "501", "870"],
        "CA?": ["209", "213", "279", "310", "323", "341", "408", "415", "424", "442", "510", "530", "559", "562", "619", "626", "628", "650", "657", "661", "707", "714", "747", "760", "805", "818", "820", "831", "858", "909", "916", "925", "949", "951"],
        "CO?": ["303", "719", "720", "970"],
        "CT?": ["203", "475", "860", "959"],
        "DE?": ["302"],
        "FL?": ["239", "305", "321", "352", "386", "407", "561", "689", "727", "754", "772", "786", "813", "850", "863", "904", "941", "954"],
        "GA?": ["229", "404", "470", "478", "678", "706", "762", "770", "912"],
        "HI?": ["808"],
        "ID?": ["208", "986"],
        "IL?": ["217", "224", "309", "312", "331", "447", "464", "618", "630", "708", "730", "773", "779", "815", "847", "861", "872"], 
        "IN?": ["219", "260", "317", "463", "574", "765", "812", "930"],
        "IA?": ["319", "515", "563", "641", "712"],
        "KS?": ["316", "620", "785", "913"],
        "KY?": ["270", "364", "502", "606", "859"],
        "LA?": ["225", "318", "337", "504", "985"],
        "ME?": ["207"],
        "MD?": ["240", "301", "410", "443", "667"],
        "MA?": ["339", "351", "413", "508", "617", "774", "781", "857", "978"],
        "MI?": ["231", "248", "269", "313", "517", "586", "616", "734", "810", "906", "947", "989"],
        "MN?": ["218", "320", "507", "612", "651", "763", "952"],
        "MS?": ["228", "601", "662", "769"],
        "MO?": ["314", "417", "557", "573", "636", "660", "816", "975"],
        "MT?": ["406"],
        "NE?": ["308", "402", "531"],
        "NV?": ["702", "725", "775"],
        "NH?": ["603"],
        "NJ?": ["201", "551", "609", "640", "732", "848", "856", "862", "908", "973"],
        "NM?": ["505", "575"],
        "NY?": ["212", "315", "329", "332", "347", "363", "516", "518", "585", "607", "631", "646", "680", "716", "718", "838", "845", "914", "917", "929", "934"], 
        "NC?": ["252", "336", "704", "743", "828", "910", "919", "980", "984"],
        "ND?": ["701"],
        "OH?": ["216", "220", "234", "283", "326", "330", "380", "419", "440", "513", "567", "614", "740", "937"],
        "OK?": ["405", "539", "580", "918"],
        "OR?": ["458", "503", "541", "971"],
        "PA?": ["215", "223", "267", "272", "412", "445", "484", "570", "610", "717", "724", "814", "878"],
        "RI?": ["401"],
        "SC?": ["803", "843", "854", "864"],
        "SD?": ["605"],
        "TN?": ["423", "615", "629", "731", "865", "901", "931"],
        "TX?": ["210", "214", "254", "281", "325", "346", "361", "409", "430", "432", "469", "512", "682", "713", "737", "806", "817", "830", "832", "903", "915", "936", "940", "956", "972", "979"],
        "UT?": ["385", "435", "801"],
        "VT?": ["802"],
        "VA?": ["276", "434", "540", "571", "703", "757", "804"],
        "WA?": ["206", "253", "360", "425", "509", "564"],
        "WV?": ["304", "681"],
        "WI?": ["262", "414", "534", "608", "715", "920"],
        "WY?": ["307"]
    }

    if state == "":
        for state_code, area_codes in area_codes_by_state.items():
            if any(phone.startswith(code) for code in area_codes):
                return state_code
    
    return state


def pinterest():    # testuser=    kevinrose     # add city
    print(f'\n\t<<<<< pinterest users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - pinterest')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        (country, email, fullname, middlename, lastname, firstname) = ('', '', '','','', '')
        (success, note, photo, website, city, otherurls) = ('','','','','', '')

        url = (f'https://www.pinterest.com/{user}/')
        otherurls = (f'https://pinterest.com/search/users/?q={user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 


        try:
            parts = titleurl.split(' (', 1)

            if len(parts) > 1:
                titleurl = parts[0]
            fullname = str(titleurl)
            if ' ' in fullname:
                ranking = '4 - pinterest'
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        except:
            pass
            
        # fullname = fullname.replace('User AVATAR', '')  # test
        if pagestatus == 200:
            if ' ' in fullname:
            # if titleurl != '':
                print(f' {url}	   {fullname}	{note}')
                
                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["user"] = user
                row_data["url"] = url
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname
                row_data["middlename"] = middlename
                row_data["lastname"] = lastname
                row_data["note"] = note    
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl 
                

                data.append(row_data)                

def is_package_installed(package_name: str) -> bool:
    return importlib.util.find_spec(package_name) is not None

def playwright_installed() -> bool:
    """
    Checks if Playwright browser binaries are installed by verifying
    presence of versioned browser folders under Playwright's cache directory.
    """
    if platform.system() == "Windows":
        base_dir = Path(os.getenv("LOCALAPPDATA", "")) / "ms-playwright"
    else:
        base_dir = Path.home() / ".cache" / "ms-playwright"

    if not base_dir.exists():
        return False

    # Look for folders that start with browser names and a dash
    expected_prefixes = ["chromium-", "firefox-", "webkit-"]
    found = {prefix: False for prefix in expected_prefixes}

    for item in base_dir.iterdir():
        if item.is_dir():
            for prefix in expected_prefixes:
                if item.name.startswith(prefix):
                    found[prefix] = True

    return all(found.values())
    
def playwright_browsers_installed() -> bool:
    """
    Checks if Playwright browser binaries are installed.
    Looks for versioned folders like chromium-*, firefox-*, webkit-*.
    """
    if platform.system() == "Windows":
        base_dir = Path(os.getenv("LOCALAPPDATA", "")) / "ms-playwright"
    else:
        base_dir = Path.home() / ".cache" / "ms-playwright"

    if not base_dir.exists():
        return False

    expected_prefixes = ["chromium-", "firefox-", "webkit-"]
    found = {prefix: False for prefix in expected_prefixes}

    for item in base_dir.iterdir():
        if item.is_dir():
            for prefix in expected_prefixes:
                if item.name.startswith(prefix):
                    found[prefix] = True

    return all(found.values())

def playwright_ready() -> bool:
    """
    Returns True if both packages are installed and browser binaries exist.
    """
    return (
        is_package_installed("playwright") and
        is_package_installed("playwright_stealth") and
        playwright_browsers_installed()
    )



async def playwright_url(URL: str):
    async with Stealth().use_async(async_playwright()) as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()

        try:
            response = await page.goto(URL, wait_until="domcontentloaded", timeout=30000)
            await page.wait_for_timeout(1000)

            content = await page.content()
            referer = response.request.headers.get("referer", "") if response else ""


            osurl = page.url
            titleurl = await page.title()
            pagestatus = response.status if response else "No response"

            return (content, referer, osurl, titleurl, pagestatus)

        except Exception as e:
            return ("", "", "", "", f"Error: {e}")

        finally:
            await browser.close()



def poshmark():    # testuser=    kevinrose
    print(f'\n\t<<<<< poshmark users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - poshmark')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        url = (f'https://poshmark.com/closet/{user}')
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except:
            pass

        if 'Success' in pagestatus:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                elif "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1].replace('\'s Closet', '')
                    
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''                    
                    
                    # firstname = fullname

            if 1==1:

                print(f' {url}	   {fullname}')

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["user"] = user
                row_data["url"] = url
                row_data["firstname"] = firstname
                row_data["fullname"] = fullname
                # row_data["middlename"] = middlename
                row_data["lastname"] = lastname
                # row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                # row_data["titleurl"] = titleurl                 
                


                data.append(row_data)   

def print_logo():
    
    art = """
 ___    _            _   _ _         _   _             _   
|_ _|__| | ___ _ __ | |_(_) |_ _   _| | | |_   _ _ __ | |_ 
 | |/ _` |/ _ \ '_ \| __| | __| | | | |_| | | | | '_ \| __|
 | | (_| |  __/ | | | |_| | |_| |_| |  _  | |_| | | | | |_ 
|___\__,_|\___|_| |_|\__|_|\__|\__, |_| |_|\__,_|_| |_|\__|
                               |___/                       

  """
    print(f'{art}')

def public():    # testuser=    kevinrose
    print(f'\n\t<<<<< public users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)


    for user in users:    
        row_data = {}
        (query, ranking) = (user,'7 - public')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '','')
        url = (f'https://public.com/@{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        # if 1 == 1:
        if pagestatus == 200:
            for eachline in content.split("\n"):
                if eachline == "": pass                                             # skip blank lines
                elif "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[1]
                    fullname = fullname.split(" (")[0]
                    if ' ' in fullname:
                        (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                    else:
                        fullname = ''
            if 1==1:
                print(f' {url}	   {fullname}')

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["user"] = user
                row_data["url"] = url
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname
                # row_data["middlename"] = middlename
                row_data["lastname"] = lastname
                row_data["pagestatus"] = pagestatus
                row_data["titleurl"] = titleurl
                # row_data["content"] = content
                # row_data["titleurl"] = titleurl                
                data.append(row_data)   
 

def read_text(filename):
    """
        Reads the input file
        parses the data into lists
        exports it to the outputfile
    """

    # message = (f'Reading {filename} line 3460')
    # message_square(message)

    if not os.path.exists(filename):
        input(f"{filename} doesnt exist.")
        sys.exit()
    elif os.path.getsize(filename) == 0:
        input(f'{filename} is empty. Fill it with username, email, ip, phone and/or websites.')
        sys.exit()
    elif os.path.isfile(filename):
        inputfile = open(filename)
    else:
        input(f'See {filename} does not exist. Hit Enter to exit...')
        sys.exit()
        
    for eachline in inputfile:
        (query, ranking, fullname, url, email, user) = ('', '', '', '', '', '')
        (phone, business, fulladdress, city, state, country) = ('', '', '', '', '', '')
        (note, AKA, DOB, SEX, info, misc) = ('', '', '', '', '', '')
        (firstname, middlename, lastname, associates, case, sosfilenumber) = ('', '', '', '', '', '')
        (owner, president, sosagent, managers, Time, Latitude) = ('', '', '', '', '', '')
        (Longitude, Coordinate, original_file, Source, Source_file_information, Plate) = ('', '', '', '', '', '')
        (VIS, VIN, VYR, VMA, LIC, LIY) = ('', '', '', '', '', '')
        (DLN, DLS, content, referer, osurl, titleurl) = ('', '', '', '', '', '')
        (pagestatus, ip, dnsdomain) = ('', '', '')

        eachline = eachline + "\t" * 52
        eachline = eachline.split('\t')  # splits by tabs

        query = (eachline[0].strip())
        if ranking == '':
            ranking = '1 - main'
        fullname = (eachline[2].strip())
        url = (eachline[3].strip())
        email = (eachline[4].strip())
        user = (eachline[5].strip())
        phone = (eachline[6].strip())
        busines = (eachline[7].strip())
        fulladdress = (eachline[8].strip())
        city = (eachline[9].strip())
        state = (eachline[10].strip())
        country = (eachline[11].strip())
        note = (eachline[12].strip())
        AKA = (eachline[13].strip())
        DOB = (eachline[14].strip())
        SEX = (eachline[15].strip())
        info = (eachline[16].strip())
        misc = (eachline[17].strip())
        firstname = (eachline[18].strip())
        middlename = (eachline[19].strip())
        lastname = (eachline[20].strip())
        associates = (eachline[21].strip())
        case = (eachline[22].strip())
        sosfilenumber = (eachline[23].strip())
        owner = (eachline[24].strip())
        president = (eachline[25].strip())
        sosagent = (eachline[26].strip())
        managers = (eachline[27].strip())
        Time = (eachline[28].strip())
        Latitude = (eachline[29].strip())
        Longitude = (eachline[30].strip())
        Coordinate = (eachline[31].strip())
        original_file = (eachline[32].strip())
        Source = (eachline[33].strip())
        Source_file_information = (eachline[34].strip())
        Plate = (eachline[35].strip())
        VIS = (eachline[36].strip())
        VIN = (eachline[37].strip())
        VYR = (eachline[38].strip())
        VMA = (eachline[39].strip())
        LIC = (eachline[40].strip())
        LIY = (eachline[41].strip())
        DLN = (eachline[42].strip())
        DLS = (eachline[43].strip())
        content = (eachline[44].strip())
        referer = (eachline[45].strip())
        osurl = (eachline[46].strip())
        titleurl = (eachline[47].strip())
        pagestatus = (eachline[48].strip())
        ip = (eachline[49].strip())
        dnsdomain = (eachline[50].strip())

        # Regex data type
        # if re.search(regex_email, query):    # terrible. matches emails and phone numbers
        if bool(re.search(r"^[\w\.\+\-]+\@[\w]+\.[a-z]{2,3}$", query)):  # regex email    # works
            email = query
            user = email.split('@')[0]
            temp1 = [email]
            if query.lower() not in emails:            # don't add duplicates
                emails.append(email)
        # elif re.search(regex_url, query):   # test
            # url = query

        elif re.search(regex_host, query):  # regex_host (Finds url and dnsdomain) # False positives for emails    # todo
            url = query

            if url.lower().startswith('http'):
                if url.lower() not in websites:            # don't add duplicates
                    websites.append(url)            
            else:
                logsource = 'IOC-dnsdomain'
                dnsdomain = query
            url = url.rstrip('/')
            if "@" not in url and url.lower() not in websites:            # don't add duplicates
                websites.append(url)            
            dnsdomain = url.lower()
            dnsdomain = dnsdomain.replace("https://", "")
            dnsdomain = dnsdomain.replace("http://", "")
            dnsdomain = dnsdomain.split('/')[0]
            notes2 = dnsdomain.split('.')[-1]
            dnsdomain = dnsdomain.rstrip('/')
            if dnsdomain.lower() not in dnsdomains:            # don't add duplicates
                dnsdomains.append(dnsdomain)
            
        elif re.search(regex_ipv4, query):  # regex_ipv4
            (ip) = (query)
            if query.lower() not in ips:            # don't add duplicates
                ips.append(ip)

        elif re.search(regex_ipv6, query):  # regex_ipv6
            (ip) = (query)
            if query.lower() not in ips:            # don't add duplicates
                ips.append(ip)

        elif re.search(regex_phone, query) or re.search(regex_phone11, query) or re.search(regex_phone2, query):  # regex_phone
            (phone) = (query)
            phone = re.sub(r'[\+\- \(\)]', '', phone).strip()    # E165 standard
            # phone = phone.replace("-", '')  # E165 standard 
            # phone = phone.replace('(','').replace(')','').replace(' ','')   # E165 standard 
            # phone = phone.replace("+", "")          # E165 standard doesn't have a + (E.164 has a +)      
            # phone = phone.lstrip('1') # E.165 standard 16365551212

            if phone not in phones:            # don't add duplicates
                phones.append(phone)

            
        elif query.lower().startswith('http'):
            url = query
            if url.lower() not in websites:            # don't add duplicates
                websites.append(url)            
        elif query.strip() == '':
            print(f'blank input found')
        else:
            user = query.lower()
            if user not in users:
                users.append(user)

            # user = query
            # if query.lower() not in users:            # don't add duplicates
                # users.append(user)

    return emails,dnsdomains,ips,users,phones,websites


def read_xlsx(input_xlsx):

    """Read data from an xlsx file and return as a list of dictionaries.
    Read XLSX Function: The read_xlsx() function reads data from the input 
    Excel file using the openpyxl library. It extracts headers from the 
    first row and then iterates through the data rows, creating dictionaries 
    for each row with headers as keys and cell values as values.
    
    """
    print(f'Reading xlsx for hunting is still in beta') # temp
    message = (f'Reading {input_xlsx}')
    message_square(message)
 
    wb = openpyxl.load_workbook(input_xlsx, read_only=True, data_only=True, keep_links=False)
    ws = wb.active
    data = [] 

    # get header values from first row
    headers = [cell.value for cell in ws[1]]

    # get data rows
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row): continue
        row_data = {}   # test
        row_data = dict(zip(headers, row))

    # dnsdomains = []

        # url
        url = str(row_data.get("url") or '').strip()

        if url.lower() not in [u.lower() for u in websites]:
            websites.append(url)

        # dnsdomain
        dnsdomain = str(row_data.get("dnsdomain") or '').strip()

        # Only append non-empty and non-duplicate values
        if dnsdomain and dnsdomain.lower() not in [d.lower() for d in dnsdomains]:
            dnsdomains.append(dnsdomain)
            
        # user
        user = (row_data.get("user"))

        try:
            if user.lower() not in [u.lower() for u in users]:
            # if user and user.lower() not in [u.lower() for u in users]:
                users.append(user)
        except:pass
        # ip
        ip = str(row_data.get("ip") or '').replace('\n', '').strip()

        if ip and ip.lower() not in [i.lower() for i in ips]:
            ips.append(ip)
 
        # email
        email = str(row_data.get("email") or '').strip()

        if '@' in email and email.lower() not in [e.lower() for e in emails]:
            emails.append(email)
            
        # phone
        try:    
            phone = str(row_data.get("phone") or '').strip()
        except:
            phone = str(row_data.get("phone") or '')
            
        if phone:
            # Remove unwanted characters
            try:
                phone = re.sub(r'[^\d+]', '', phone)
            except:
                pass
                
            # Optional: validate E.164 format
            try:
                if re.match(r'^\+?[1-9]\d{1,14}$', phone) and phone not in phones:
                    phones.append(phone)
            except:pass

        # business
        business = str(
            row_data.get("business")
            or row_data.get("business/entity")
            or row_data.get("Business")
            or ''
        ).strip()

        # owner
        owner = str(row_data.get("owner") or '').strip()

        # AKA
        AKA = str(
            row_data.get("AKA") or
            row_data.get("aka") or
            row_data.get("alias") or
            ''
        ).strip()


        # city
        city = str(row_data.get("city") or '').strip().title()
    
        # state
        state = str(row_data.get("state") or '').strip()

        # DOB
        DOB = (
            row_data.get("DOB") or
            row_data.get("dob") or
            ''
        )   # .strip()

        # associates
        associates = str(
            row_data.get("associates") or
            row_data.get("friend") or
            ''
        ).strip()

        # SEX
        SEX = str(row_data.get("SEX") or row_data.get("gender") or '').strip()
        
        # firstname
        firstname = str(row_data.get("firstname") or '').strip().title()

        # lastname
        lastname = str(row_data.get("lastname") or '').strip().upper()


        # middlename
        middlename = row_data.get("middlename", "")

        # fullname
        fullname = str(row_data.get("fullname") or '').strip()

        if not fullname:
            if firstname and lastname and middlename:
                fullname = f'{firstname} {middlename} {lastname}'
            elif firstname and lastname:
                fullname = f'{firstname} {lastname}'


        # timestamp
        timestamp = (row_data.get("Time") or '')
        timestamp = timestamp or ""
        timestamp, time_orig, timezone = convert_timestamp(timestamp, '', '')




        row_data["user"] = user
        row_data["url"] = url
        row_data["email"] = email
        row_data["phone"] = phone
        row_data["fullname"] = fullname
        row_data["business"] = business
        row_data["city"] = city 
        row_data["state"] = state 
        row_data["DOB"] = DOB  
        row_data["SEX"] = SEX         
        row_data["AKA"] = AKA
        row_data["firstname"] = firstname
        row_data["middlename"] = middlename
        row_data["lastname"] = lastname
        row_data["Time"] = timestamp
        row_data["associates"] = associates
        row_data["owner"] = owner
        row_data["ip"] = ip
        row_data["dnsdomain"] = dnsdomain

        # data.append(row_data)
        try:
            data.append(row_data)
        except Exception as e:
            print(f"Error appending data: {str(e)}")

    return data

def read_xlsx_basic_old(input_xlsx):
    message = (f'Reading basic intel: {input_xlsx} line 3786')
    message_square(message)

    wb = openpyxl.load_workbook(input_xlsx)
    ws = wb.active

    for row in ws.iter_rows(min_row=2, values_only=True):  # Assuming headers are in the first row
        entry = dict(zip(headers_intel, row))
        reordered_entry = {key: entry[key] for key in headers_intel}  # Reorder the entry based on headers_intel
        data.append(reordered_entry)

    return data


def read_xlsx_basic(input_xlsx):
    message = (f'Reading basic intel: {input_xlsx} line 3801')
    message_square(message)

    # data = []
    wb = openpyxl.load_workbook(input_xlsx)
    ws = wb.active

    headers = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))  # Read headers from the first row

    for row in ws.iter_rows(min_row=2, values_only=True):  # Start iterating from the second row
        entry = dict(zip(headers, row))  # Use the headers read from the file
        filtered_entry = {key: entry[key] for key in headers_intel if key in headers}  # Filter out only the relevant headers
        data.append(filtered_entry)

    return data


def read_xlsx_basic_location(input_xlsx):
    message = (f'Reading basic intel: {input_xlsx} line 3819')
    message_square(message)

    # data = []
    wb = openpyxl.load_workbook(input_xlsx)
    ws = wb.active

    for row in ws.iter_rows(min_row=2, values_only=True):  # Assuming headers are in the first row
        entry = dict(zip(headers_intel, row))
        data.append(entry)

    return data
    
    
def reddit(): # testuser = kevinrose
    print(f'\n\t<<<<< reddit users >>>>>')

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - reddit')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.rstrip()
        url = (f'https://www.reddit.com/user/{user}/')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("  <"):
            if "nobody on Reddit goes by that name" in eachline or "This account has been suspended" in eachline or "This user has deleted their account" in eachline:
                ranking = '9 - reddit'
            elif "hasn't posted yet" in eachline :
                ranking ='7 - reddit'
                note = "hasn't posted yet"

        if '9' not in ranking:
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["user"] = user            
            row_data["note"] = note

            data.append(row_data)           


def request_url(url):
    
    fake_referer = 'https://www.google.com/'
    headers_url = {'Referer': fake_referer}

    
    (content, referer, osurl, titleurl, pagestatus)= ('blank', '', '', '', '')
    (response) = ('')
    
    if url.lower().startswith('http'):
        blah = ''
    else:
        url = ("https://" +url)

    try:
        response = requests.get(url, verify=False, headers=headers_url)        
        response.raise_for_status()
        pagestatus  = response.status_code
        content = response.content.decode()
        content = BeautifulSoup(content, 'html.parser')
        


    except requests.exceptions.RequestException as e:
        # print(f"Could not fetch URL {url}: {str(e)}")
        # raise requests.exceptions.RequestException(str(e))
        (pagestatus) = ('Fail')

# osurl
    try:
        osurl = response.headers['Server']
    except:
    # except KeyError:
        pass

# titleurl
    try:
        titleurl = titleurl_get(content)
        # titleurl = titleurl_og(content)
    except KeyError:
        titleurl = ''


#pagestatus
    
    if str(pagestatus).startswith('2') :    
        pagestatus = (f'Success - pagestatus')
    elif str(pagestatus).startswith('3') :    
        pagestatus = (f'Redirect - {pagestatus}')
    elif str(pagestatus).startswith('4') :    
        pagestatus = (f'Fail - {pagestatus}')
    elif str(pagestatus).startswith('5') :    
        pagestatus = (f'Fail - {pagestatus}')
    elif str(pagestatus).startswith('1') :    
        pagestatus = (f'Info - {pagestatus}')

    pagestatus = pagestatus.strip()    

    return (content, referer, osurl, titleurl, pagestatus)



def resolverRS():# testIP= 77.15.67.232
    print(f'\n\t<<<<< resolverRS ip >>>>>')
   
    for ip in ips:
        row_data = {}
        (query) = (ip)
        (country, city, zipcode, case, note, state) = ('', '', '', '', '', '')
        (misc, info) = ('', '')
        
        url = (f'https://resolve.rs/ip/geolocation.html?ip={ip}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "403 ERROR" in eachline:
                pagestatus = '403 Error'
                content = ''

            elif "\"code\": \"" in eachline :   # and zipcode != ''
                zipcode = eachline.split("\"")[3]
            elif "\"en\": \"" in eachline :   # city
                # print(f'')
                if city == '':
                    city = eachline.split("\"")[3]
                elif misc == '' and city != '': # continent
                    misc = eachline.split("\"")[3]

                elif misc != '' and country == '' and city != '': # country
                    country = eachline.split("\"")[3]
                elif info == '' and misc != '' and country != '' and city != '':    # registered country
                    info = eachline.split("\"")[3]
                elif state == '' and info != '' and misc != '' and country != '' and city != '':    # state
                    state = eachline.split("\"")[3]
            elif "COMCAST" in eachline :   # isp  >ASN</a>
                note = 'COMCAST'


        # pagestatus = ''                
        if url != '':
            print(f'{url}') 
            ranking = '6 - resolve.rs'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["ip"] = ip
            row_data["note"] = note
            row_data["city"] = city
            row_data["country"] = country
               
            data.append(row_data)

    # return data

def request(url):
    (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
    (string) = ('')
    fake_referer = 'https://www.google.com/'
    headers_url = {'Referer': fake_referer}
    url = url.replace("http://", "https://")     # test

    page  = requests.get(url, headers=headers_url)

    pagestatus  = page.status_code
    soup = BeautifulSoup(page.content, 'html.parser')
    content = soup.prettify()
    try:
        osurl = page.headers['Server']
    except:pass
    
    try:
        titleurl = soup.title.string
    except:pass
    

#pagestatus
    
    if str(pagestatus).startswith('2') :    
        pagestatus = (f'Success - {pagestatus}')
    elif str(pagestatus).startswith('3') :    
        pagestatus = (f'Redirect - {pagestatus}')
    elif str(pagestatus).startswith('4') :    
        pagestatus = (f'Fail - {pagestatus}')
    elif str(pagestatus).startswith('5') :    
        pagestatus = (f'Fail - {pagestatus}')
    elif str(pagestatus).startswith('1') :    
        pagestatus = (f'Info - {pagestatus}')
    try:
        pagestatus = str(pagestatus).strip()
    except Exception as e:
        print(f"Error striping pagestatus: {str(e)}")
# titleurl

    try:
        title_tag = content.find('title')
        if title_tag is not None:
            title = title_tag.text.strip()
            # The full name is often included in parentheses after the username
            # Example: "Tom Anderson (myspacetom)"
            # We'll split the title at the first occurrence of '(' to extract the full name
            parts = title.split(' (', 1)
            if len(parts) > 1:
                titleurl = parts[0]
    except Exception as e:
        # print(f"Error parsing title: {str(e)}")
        pass

    if titleurl !="":
        try:
            meta_tags = content.find_all('meta')
            for tag in meta_tags:
   
                if tag.get('property') == 'og:title':
                    titleurl = tag.get('content')
                    titleurl =  title.split(' (')[0]
        except Exception as e:
            # print(' ')
            # print(f'Error parsing metadata: {str(e)}')
            pass

    try:
        titleurl = str(titleurl)    #test
        titleurl = (titleurl.encode('utf8'))    # 'ascii' codec can't decode byte
        titleurl = (titleurl.decode('utf8'))    # get rid of bytes b''
    except TypeError as error:
        print(f'{error}')
        titleurl = ''

    try:
        titleurl = str(titleurl).strip()
        content = str(content).strip()
    except:pass
    return (content, referer, osurl, titleurl, pagestatus)   

def reversephonecheck():# testPhone= 
    print(f'\n\t<<<<< reversephonecheck phone numbers >>>>>')
    url = (f'https://www.reversephonecheck.com' )


    # If the site is down, skip the module
    if not ping_url(url):
        print(f'{url} is down, skipping reversephonecheck module.')
        return

    for phone in phones:
        row_data = {}
        (query) = (phone)
        ranking = '9 - reversephonecheck'
        (fulladdress, country, city, state, case, note) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        (areacode, prefix, line, count, match, match2) = ('', '', '', 1, '', '')
        (url) = ('')

        if len(phone.lstrip('1')) == 10:
            # print(f' {phone} has 10 digits')    # temp
            phone = phone.lstrip('1')
            phone = (phone[:3] + "-" + phone[3:6] + "-" + phone[6:])
            # print(f'phone  {phone}') # temp
            
        (line2) = ('')
        if "-" in phone:
            phone2 = phone.split("-")
            areacode = phone2[0]
            prefix = phone2[1]
            try:
                line = phone2[2]
                line2 = line
                line = line[:2]
                
            except:
                pass

        url = (f'https://www.reversephonecheck.com/1-{areacode}/{prefix}/{line}/#{phone[-2:]}' )

        (content, referer, osurl, titleurl, pagestatus) = request(url) 
        match = (f"{prefix} - {line2}")
        
        phone = phone.replace('-', '')
        for eachline in content.split("\n"):
            if match in eachline:
                pagestatus = 'research'
                ranking = '4 - reversephonecheck'
                count += 1

        if '404' in pagestatus:
            ranking = '99 - reversephonecheck'
        elif pagestatus == 'research' and count == 2:
            print(f'{url} {phone}')
            ranking = '5 - reversephonecheck'
        else:
            print(f'{url} {phone}')
            ranking = '9 - reversephonecheck'

        if state == '':
            state = phone_state_check(phone, state).replace('?', '')


        if '404' not in str(pagestatus):
        # if 1==1:    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone
            row_data["note"] = note
            
            row_data["city"] = city
            row_data["state"] = state
            
            
            row_data["country"] = country
            row_data["fulladdress"] = fulladdress
      
            data.append(row_data)

def roblox(): # testuser = kevinrose
    print(f'\n\t<<<<< roblox users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - roblox')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://www.roblox.com/user.aspx?username={user}')

        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("  <"):
            if "og:title" in eachline:
                fullname = eachline.strip().split("\"")[1]

        if '404' not in pagestatus and 'ail' not in pagestatus:
            if fullname.lower() == user.lower():
                fullname = ''
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                fullname = ''

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["fullname"] = fullname
            # row_data["firstname"] = firstname
            # row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user            
            # row_data["titleurl"] = titleurl
            # row_data["pagestatus"] = pagestatus
            # row_data["content"] = content
            
            data.append(row_data)  

def rumble(): # testuser = kevinrose
    print(f'\n\t<<<<< rumble users >>>>>')

    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - rumble')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        # url = (f'https://rumble.com/user/{user}/about')
        url = (f'https://rumble.com/c/{user}/about')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        fullname = str(titleurl)
        if 'umble' in fullname:
            fullname == ''

        if pagestatus == 200:
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                firstname = fullname

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user            
            # row_data["titleurl"] = titleurl
            # row_data["pagestatus"] = pagestatus
            # row_data["content"] = content
            
            data.append(row_data)  


def samples():
    print(f'''    
Alain_THIERRY
anh_usa
frizz925
gurkha_love
Hmei7
itdoesnthavetohappen
JLight29
kevinrose
KC3ELO
kevinwagner
kobegal
luckyjames844
maverick3819
merz356
Misty0427
MR-JEFF
N3tAtt4ck3r
Pattycakes98
rothys
ryanlwatkins
thekevinrose
williger
zazenergy
nullcrew
realDonaldTrump
Their1sn0freakingwaythisisreal12JT4321
        
77.15.67.232
92.20.236.78
255.255.255.255

2607:f0d0:1002:51::4
2607:f0d0:1002:0051:0000:0000:0000:0004

annemconnor@yahoo.com
kandyem@yahoo.com
kevinrose@gmail.com
craig@craigslist.org
ceo@zappos.com
lnd_whitaker@yahoo.com
gsmstocks@gmail.com
lydianorman1@hotmail.com
soniraj388@gmail.com
tanderson09@gmail.com
tin_max87@yahoo.com
Their1sn0freakingwaythisisreal12344321@fakedomain.com

385-347-1531
312-999-9999
5596833344
15596833344

'''
)    


def sherlock():    # testuser=    kevinrose
    print(f'\n\t<<<<< Manually check Sherlock users >>>>>')
    
    for user in users:    
        note = (f'sherlock {user} --csv')
        # note = (f'sherlock {user} -v --csv --print-all')

        info = (f'https://sherlockproject.xyz/installation')
        row_data = {}
        (query, ranking) = (user, '9 - manual')

        if 1==1:

            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["url"] = url
            row_data["note"] = note
            row_data["info"] = info
            
      
            data.append(row_data)


def slack(): # testuser = kevinrose
    print(f'\n\t<<<<< slack users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - slack')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://{user}.slack.com')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 


        # Regex to extract teamName
        fullname_match = re.search(r'"teamName":"(.*?)"', content)
        fullname = fullname_match.group(1) if fullname_match else ""

        # Regex to extract formattedEmailDomains
        dnsdomain_match = re.search(r'"formattedEmailDomains":"(.*?)"', content)
        dnsdomain = dnsdomain_match.group(1) if dnsdomain_match else ""
        dnsdomain = dnsdomain.lstrip('@')


        if pagestatus == 200:
        # if '404' not in pagestatus and 'ail' not in pagestatus:
            if fullname.lower() == user.lower():
                fullname = ''
            
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            if " " not in fullname:
                fullname = ''

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
            row_data["user"] = user            
            # row_data["titleurl"] = titleurl
            row_data["dnsdomain"] = dnsdomain
            # row_data["content"] = content
            # row_data["pagestatus"] = pagestatus
            data.append(row_data)   
            

def snapchat(): # testuser = kevinrose
    print(f'\n\t<<<<< snapchat users >>>>>')

    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - snapchat')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        user = user.rstrip()
        url = (f'https://www.snapchat.com/add/{user}?')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        # for eachline in content.split("\n"):

            # if "og:title" in eachline:
                # fullname = eachline.strip().split("\"")[1].replace(' on Snapchat','')
            # if ' ' in fullname:
                # (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                # ranking = ('5 - snapchat')
            # else:
                # fullname = ''


        fullname = str(titleurl)
        if '(' in fullname:
            fullname = fullname.split('(')[0].strip()
            ranking = ('5 - snapchat') 
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            ranking = ('3 - snapchat')           

        if pagestatus == 200:

            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["lastname"] = lastname
            row_data["user"] = user
            # row_data["pagestatus"] = pagestatus
            # row_data["titleurl"] = titleurl
            data.append(row_data)

def sportstracker():  # testuser = kevinrose
    print(f'\n\t<<<<< sports-tracker users >>>>>')

    for user in users:
        row_data = {}

        query = user.rstrip()
        ranking = '9 - sportstracker'
        fullname = firstname = middlename = lastname = ''
        business = city = state = country = ''
        email = ''
        SEX = ''
        info = ''
        note = ''
        misc = ''
        associates = AKA = owner = president = ''
        otherurls = ''
        pagestatus = ''
        titleurl = ''

        # API URL
        url = f'https://api.sports-tracker.com/apiserver/v1/user/name/{query}'

        # Skip non-alpha usernames
        if not any(char.isalpha() for char in query):
            continue

        # Request
        content, referer, osurl, titleurl, pagestatus = request(url)

        if 'ail' in pagestatus:
            continue

        # Parse JSON
        try:
            parsed_data = json.loads(content)
        except Exception as error:
            print(f'JSON error: {error}')
            continue

        # Sports‑Tracker payload
        try:
            payload = parsed_data['payload']
        except:
            continue

        # -----------------------------
        # USERNAME
        # -----------------------------
        try:
            user = payload.get('username', query)
        except:
            user = query

        # -----------------------------
        # FULL NAME
        # -----------------------------
        try:
            fullname = payload.get('realName', '')
        except:
            fullname = ''

        # -----------------------------
        # SEX / GENDER
        # -----------------------------
        try:
            gender = payload.get('gender', '')
            if gender.upper() == 'MALE':
                SEX = 'M'
            elif gender.upper() == 'FEMALE':
                SEX = 'F'
            else:
                SEX = ''
        except:
            SEX = ''

        # -----------------------------
        # NOTE (lastModified)
        # -----------------------------
        try:
            note = str(payload.get('lastModified', ''))
        except:
            note = ''

        # -----------------------------
        # INFO (uuid)
        # -----------------------------
        try:
            info = payload.get('uuid', '')
        except:
            info = ''

        # -----------------------------
        # NAME PARSING
        # -----------------------------
        if fullname and ' ' in fullname:
            fullname, firstname, middlename, lastname = fullname_parse(fullname)

        # -----------------------------
        # PROFILE URL
        # -----------------------------
        profile_url = f'https://www.sports-tracker.com/user/{user}'

        print(f'{profile_url}\t{fullname}')

        # Ranking bump
        if fullname or note:
            ranking = '3 - sportstracker'

            # -----------------------------
            # BUILD ROW DATA
            # -----------------------------
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["middlename"] = middlename
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = profile_url
            row_data["email"] = email
            row_data["business"] = business
            row_data["city"] = city
            row_data["state"] = state
            row_data["country"] = country
            row_data["user"] = user
            row_data["SEX"] = SEX
            row_data["info"] = info
            row_data["misc"] = misc
            row_data["note"] = note
            row_data["titleurl"] = titleurl
            row_data["associates"] = associates
            row_data["owner"] = owner
            row_data["president"] = president
            row_data["AKA"] = AKA
            row_data["pagestatus"] = pagestatus
            data.append(row_data)



def sportstrackerold(): # testuser = kevinrose https://www.sportstracker.net/kevinrose
    print(f'\n\t<<<<< sportstracker users >>>>>')
    (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)


    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - sportstracker')
        (firstname, lastname) = ('', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://api.sports-tracker.com/apiserver/v1/user/name/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request_url(url)
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
        if 1 == 1:
        # if '404' not in str(pagestatus):
            # grab display_name = fullname
            titleurl = titleurl.replace("'s favorite items - sportstracker",'')

            fullname = str(titleurl)
            
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                fullname = ''

            if fullname == 'sportstracker.com':
                ranking = '9 - sportstracker'
                fullname = ''

            if 1 == 1:
            # if ranking == '4 - sportstracker':
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["fullname"] = fullname
                row_data["firstname"] = firstname            
                row_data["lastname"] = lastname
                row_data["url"] = url
                row_data["user"] = user
                row_data["city"] = city            
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl         

                data.append(row_data)

        time.sleep(5) #will sleep for 5 seconds
        

def spotify(): # testuser = kevinrose
    print(f'\n\t<<<<< spotify users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - spotify')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        # (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://open.spotify.com/user/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}')
        pagestatus = str(pagestatus)
        
        if '404' not in str(pagestatus):
            titleurl = titleurl.replace(" on Spotify","").strip()
            fullname = str(titleurl)
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                ranking = ('5 - spotify')
            else:
                fullname = ''
            fullname = fullname.replace('Spotify – Web Player', '')
            
            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["user"] = user
            row_data["pagestatus"] = pagestatus
            row_data["titleurl"] = titleurl
            # row_data["content"] = content           
                        
            data.append(row_data)

def substack(): # testuser = kevinrose
    print(f'\n\t<<<<< substack users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking, content) = (user, '7 - substack', '')
        (fullname, firstname, lastname, middlename, note) = ('', '', '', '', '')
        # (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        # url = (f'https://{user}.substack.com')
        url = (f'https://substack.com/@{user}?utm_source=about-page')        
        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
        except TypeError as error:
            print(f'{error}')        

        
        for eachline in content.split("<"):
            if "og:title" in eachline:
                fullname = eachline
                # print(f'fullname = {fullname}') # temp
                fullname = eachline.strip().split("\"")[1]
                fullname = fullname.replace(" | Substack", "")
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]        
        
        
        
        if '404' not in str(pagestatus):
            titleurl = titleurl.replace(" | Substack","").strip()
            fullname = str(titleurl)
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                ranking = ('5 - substack')
            else:
                fullname = ''
            lastname = lastname.replace('SUBSTACK', '')
            if "’s Substack" in fullname:
                try:
                    fullname = fullname.split("’s Substack")[0]
                except:
                    pass
            if 'Substack' not in fullname:
            
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                row_data["url"] = url
                row_data["lastname"] = lastname
                row_data["firstname"] = firstname
                row_data["fullname"] = fullname
                row_data["user"] = user
                row_data["note"] = note            
                row_data["pagestatus"] = pagestatus
                data.append(row_data)
            
  
def thatsthememail():   # testEmail= smooth8101@yahoo.com 
    print(f'\n\t<<<<< thatsthem emails >>>>>')
     
    for email in emails:
        # print(f'{email}')
        row_data = {}
        (query, content, note) = (email, '', '')
        (country, city, zipcode, case, note) = ('', '', '', '', '')
        
        url = (f'https://thatsthem.com/email/{email}')
        # (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "Found 0 results for your query" in eachline and case == '':
                # print(f'not found')  # temp
                url = ('')

        pagestatus = ''                
        if url != '':
            # print(f'{url}	{email}') 
            ranking ='7 - thatsthem'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email
            # row_data["note"] = note
            
            # row_data["city"] = city
            # row_data["country"] = country
            # row_data["referer"] = city
            # row_data["titleurl"] = titleurl            
            # row_data["pagestatus"] = pagestatus            
                        
            data.append(row_data)
            # print(f'row_data = {row_data}') # temp


        else:
            print(f'{url}	{email}') 

def thatsthemip():# testIP= 8.8.8.8
    print(f'\n\t<<<<< thatsthem ip >>>>>')
       
    for ip in ips:
        row_data = {}
        (country, city, zipcode, case, note, state) = ('', '', '', '', '', '')
        
        url = (f'https://thatsthem.com/ip/{ip}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)

        for eachline in content.split("\n"):
            if "located in " in eachline:
                state = eachline
                note = eachline

            elif "403 ERROR" in eachline:
                pagestatus = '403 Error'
                content = ''
            elif "Found 0 results for your query" in eachline:
                print(f'Not found')  # temp
                url = ('')
        # pagestatus = ''                
        if url != '':
            ranking = '9 - thatsthem'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["ip"] = ip
            # row_data["note"] = note
            
            # row_data["city"] = city
            # row_data["country"] = country
            # row_data["referer"] = city
            # row_data["titleurl"] = titleurl            
            # row_data["pagestatus"] = pagestatus            
                        
            data.append(row_data)

def thatsthemphone():# testPhone=   
    # note: need to add dishes if there are none
    print(f'\n\t<<<<< thatsthem phone numbers >>>>>')
    for phone in phones:
        
        row_data = {}
        (query, ranking) = (phone, '9 - thatsthem')
        
        if '-' not in phone:
            phone = phone_dashes(phone.lstrip('1'))
        
        (country, city, zipcode, case, note, state) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        
        time.sleep(2) # will sleep for 10 seconds
        url = ('https://thatsthem.com/phone/%s' %(phone.lstrip('1')))    # https://thatsthem.com/reverse-phone-lookup


        if "Found 0 results for your query" in content or "The request could not be satisfied" in content:
            url = ('')
            note = ('captcha protected')

        for eachline in content.split("\n"):
            if "Found 0 results for your query" in eachline and case == '':
                print(f'Not found')  # temp
                url = ('')

        phone = phone.replace('-', '')
        
        if note == '':
            ranking = '6 - thatsthem'
            state = phone_state_check(phone, state).replace('?', '')
            
        else:   
            ranking = '9 - thatsthem'

        if 1==1:
        # if url != '':
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone
            row_data["state"] = state
            row_data["note"] = note
            data.append(row_data)

def telegram(): # testuser = kevinrose
    print(f'\n\t<<<<< telegram users >>>>>')
    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking, note) = (user, '9 - telegram', '')
        
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', 'research', '')
        user = user.rstrip()
        url = (f'https://t.me/{user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        try:

            for eachline in content.split("\n"):
                if "og:title" in eachline:
                    fullname = eachline.strip().split("\"")[3]

            if 1==1:
            # if 'nofollow' not in fullname:
                print(f'{url}	{fullname}') 

                row_data["query"] = query
                row_data["ranking"] = ranking
                # row_data["fullname"] = fullname                
                row_data["url"] = url
                row_data["user"] = user
                row_data["note"] = fullname
                row_data["pagestatus"] = pagestatus    
                # row_data["content"] = content                
                row_data["titleurl"] = titleurl                 
                data.append(row_data)

        except TypeError as error:
            print(f'{error}')


def threads():    # testuser=    kevinrose     # add info
    print(f'\n\t<<<<< threads users >>>>>')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - threads')
        (fullname, firstname, lastname, middlename, note)  = ('','','','', '')

        url = (f'https://www.threads.net/@{user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
            content = content.strip()
            titleurl = titleurl.strip()
            # for eachline in content.split("\n"):
                # if "@context" in eachline:
                    # content = eachline.strip()
                # elif 'og:title' in eachline and 'content=\"' in eachline:
                    # fullname = eachline.split('\"')[1].split(' (')[0]
                # elif "og:description" in eachline:
                    # note = eachline.strip()
                    # note = note.replace("\" property=\"og:description\"/>",'').replace("<meta content=\"",'')

        except:
            pass
        
        fullname = str(titleurl)
        if ' (' in fullname:
            fullname = fullname.split(' (')[0]
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        else:
            fullname = ''
        
        # if 1 == 1:
        if '@' in titleurl:
            print(f'{url} {fullname}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["middlename"] = middlename
            row_data["fullname"] = fullname
            row_data["note"] = note
            # row_data["pagestatus"] = pagestatus
            # row_data["titleurl"] = titleurl            
            row_data["user"] = user           
            data.append(row_data)


def tiktok(): # testuser = kevinrose
    print(f'\n\t<<<<< tiktok users>>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - tiktok')
        (fullname, firstname, lastname, middlename, note)  = ('','','','', '')


        # (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', 'research', '')
        user = user.rstrip()
        url = (f'https://tiktok.com/@{user}?')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        if 'uccess' in pagestatus:
            fullname = str(titleurl)
            fullname = fullname.split(' (')[0]
            if fullname == user:
                fullname = ''
            elif 'Make Your Day' in fullname:
                ranking ='7 - tiktok'
                fullname = ''
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
                ranking = '4 - tiktok'
            else:
                fullname = ''            
         
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["user"] = user
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["titleurl"] = titleurl
            row_data["pagestatus"] = pagestatus
            
                                   
            data.append(row_data)

def tinder():    # testuser=    john
    print(f'\n\t<<<<< tinder users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '7 - tinder')
        (fullname, firstname, lastname, middlename, note, DOB)  = ('','','','', '', '')
        (misc) = ('')

        url = (f'https://tinder.com/@{user}')

        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            
            
            
            content = content.strip()
            titleurl = titleurl.strip()
            for eachline in content.split("\n"):
                if "@context" in eachline:
                    content = eachline.strip()
                elif 'og:title' in eachline and 'content=\"' in eachline:
                    firstname = eachline.split('\"')[1].split(' (')[0]
                    
                elif 'schools\"' in eachline:
                # elif 'schools\":\[{\"name' in eachline:
                    datatemp = json.loads(eachline)
                    note = datatemp["schools"][0]["name"]
                    # Extract the description value and print it
                    # note = data['description']
                    misc = eachline
                    
                    DOB= datatemp["webProfile"]["user"]["birth_date"]
        except:
            pass
            
        # time.sleep(1) # will sleep for 1 seconds
        if 'alternate' in content:
            print(f'{url}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["note"] = note
            row_data["DOB"] = DOB
            row_data["misc"] = misc

            data.append(row_data)
 
def titleurl_get(content):

    titleurl = ''
    try:
        # soup = BeautifulSoup(html, 'html.parser')
        title_tag = content.find('title')
        if title_tag is not None:
            title = title_tag.text.strip()
            # The full name is often included in parentheses after the username
            # Example: "Tom Anderson (myspacetom)"
            # We'll split the title at the first occurrence of '(' to extract the full name
            parts = title.split(' (', 1)
            if len(parts) > 1:
                titleurl = parts[0]
    except Exception as e:
        # print(f'Error parsing title: {str(e)}')
        pass


    return titleurl

def titleurl_og(content):
    (titleurl) = ('')

    try:
        meta_tags = content.find_all('meta')
        for tag in meta_tags:
            if tag.get('property') == 'og:title':
                titleurl = tag.get('content')
                titleurl =  title.split(' (')[0]
    except Exception as e:
        pass
    return titleurl


def tripadvisor(): # testuser = kevinrose
    print(f'\n\t<<<<< tripadvisor users >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 
        
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - tripadvisor')
        (city, country, fullname, titleurl, pagestatus, content) = ('', '', '', '', '', '')
        (note) = ('')
        user = user.strip()
        url = (f'https://www.tripadvisor.com/Profile/{user}')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 


        if pagestatus == 200:


            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["user"] = user            
            row_data["pagestatus"] = pagestatus
            
            data.append(row_data)  
            
            
def truepeople_email(): 
    '''
    this is protected by javascript, cookies and cloud flair
    it would require using selenium and a hard coded webdriver
    '''
    
    print(f'\n\t<<<<< truepeoplesearch emails >>>>>')
    
    for email in emails:
        row_data = {}
        (query, content, note) = (email, '', '')
        url = f'https://www.truepeoplesearch.com/resultemail?email={email}'
        # (content, referer, osurl, titleurl, pagestatus) = request(url)
        if '@' in email.lower():

            if 'Enable JavaScript and cookies to continue' in content:
                print(f'blah')
            # if 'We could not find any records for that search criteria' in content:
                ranking = '99 - truepeoplesearch'    # needs work
            elif 'All retrieved results' in content:
                note = 'Events shown in time zone'
                ranking = '4 - truepeoplesearch'
                # print(f' {email}')                
            else:
                ranking = '9 - truepeoplesearch'
                # print(f' {email}  ')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["email"] = email
            row_data["note"] = note
            # row_data["content"] = content
            data.append(row_data)
            # print(f'row_data = {row_data}') # temp


def truthSocial(): # testuser = realdonaldtrump https://truthsocial.com/@realDonaldTrump
    print(f'\n\t<<<<< truthsocial users >>>>>')
    print(f'\n\t\t\tThis one one takes a while')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - truthsocial')

        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (note) = ('')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://truthsocial.com/@{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        (fullname,lastname,firstname, email, name, country) = ('','','', '', '', '')
        pagestatus = ''
        # time.sleep(3) #will sleep for 3 seconds
        for eachline in content.split("  <"):
            if 'This resource could not be found' in eachline:
                pagestatus = '404'
            elif "og:title" in eachline:
                titleurl = eachline.strip().split("\"")[1]
                fullname = str(titleurl).split(" (")[0]
                pagestatus = '200'
                ranking = '9 - truthsocial'
                if titleurl == 'Truth Social':
                    pagestatus = '404'
                else:
                    pagestatus = '200'
                    ranking = '3 - truthsocial'
            elif "og:description" in eachline:
                note = eachline.strip().split("\"")[1]

        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        else:
            fullname = ''

        if 1==1:
        # if '@' in titleurl: 
            print(f'{url}	{fullname}') 
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["user"] = user
            # row_data["lastname"] = lastname
            # row_data["firstname"] = firstname
            row_data["fullname"] = fullname
            row_data["note"] = note
            row_data["titleurl"] = titleurl
            row_data["pagestatus"] = pagestatus
            data.append(row_data)


def tumblr(): # testuser = kevinrose

    
    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)
    else:
        print(f'Playwright is installed')
        
    print(f'\n\t<<<<< Tumblr users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking, fullname, note) = (user, '4 - tumblr', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('','', '', '', '')
        (firstname, middlename, lastname) = ('','', '')
        user = user.rstrip()
        url = (f'https://www.tumblr.com/{user}')
        # url = (f'https://{user}.tumblr.com/')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            # (content, referer, osurl, titleurl, pagestatus) = playwright_url(url: str)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
            
        except TypeError as error:
            print(f'{error}')    

        for eachline in content.split("\n"):
            if 'og:title' in eachline:
                match = re.search(r'content="([^"]+)"', eachline)
                if match:
                    content_value = match.group(1)
                    parts = content_value.split('·')
                    fullname = parts[1].strip() if len(parts) > 1 else ""
                    if fullname == "Untitled":
                       fullname = ""
                # else:
                    # fullname = eachline # temp
        # print(f'fullname = {fullname}') # temp
        
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        # else:
            # fullname = ''


        if ' on Tumblr' not in titleurl:
            ranking = '7 - tumblr'

        if pagestatus == 200:

            print(f'{url}  {fullname}') 

            row_data["query"] = query
            
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["user"] = user
            row_data["url"] = url
            row_data["firstname"] = firstname           
            row_data["lastname"] = lastname  
            row_data["middlename"] = middlename  
            # row_data["titleurl"] = titleurl 
            # row_data["content"] = content
                                 
            
            data.append(row_data)
   



def twitch(): # testuser = kevinrose

    
    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)
        
    print(f'\n\t<<<<< twitch users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking, fullname, note) = (user, '4 - twitch.tv', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('','', '', '', '')
        (firstname, middlename, lastname) = ('','', '')
        user = user.rstrip()
        url = (f'https://www.twitch.tv/{user}')
        url2 = (f'https://twitchtracker.com/{user}')
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)

            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url2))
            
        except TypeError as error:
            print(f'{error}')    

        for eachline in content.split("\n"):
            if 'og:title' in eachline:
                match = re.search(r'content="([^"]+)"', eachline)
                if match:
                    content_value = match.group(1)
                    parts = content_value.split('·')
                    fullname = parts[1].strip() if len(parts) > 1 else ""
                    if fullname == "Untitled":
                       fullname = ""
                # else:
                    # fullname = eachline # temp
        # print(f'fullname = {fullname}') # temp
        
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
        # else:
            # fullname = ''


        # if ' on Tumblr' not in titleurl:
            # ranking = '7 - tumblr'

        # if 1==1:
        if pagestatus == 200:

            print(f'{url}  {fullname}') 

            row_data["query"] = query
            
            row_data["ranking"] = ranking
            row_data["fullname"] = fullname
            row_data["user"] = user
            row_data["url"] = url
            row_data["note"] = url2
            # row_data["firstname"] = firstname           
            # row_data["lastname"] = lastname  
            # row_data["middlename"] = middlename  
            # row_data["titleurl"] = titleurl 
            # row_data["content"] = content
            
            # row_data["pagestatus"] = pagestatus                                 
            
            data.append(row_data)
            
            
def twitter():    # testuser=    kevinrose     # add info
    print(f'\n\t<<<<< twitter users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - X')

        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        url = (f'https://x.com/{user}')
        (content, referer, osurl, titleurl, pagestatus) = ('','','', '', '')
        # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # print(titleurl, url, pagestatus)  # temp
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            # print(titleurl)  # temp
            titleurl = titleurl.replace(") on Twitter","")
            titleurl = titleurl.lower().replace(User.lower(),"")
            titleurl = titleurl.replace(" (","")
            fullname = str(titleurl)
            fullname = fullname.replace(" account suspended","")
            fullname = fullname.replace("twitter /","")
            titleurl = titleurl.lower().replace(fullname.lower(),"")

            print(f'{url}	   {fullname}	{titleurl}')

            ranking = '5 - X'
        except:
            # print(f'{url}	{fullname}') 

            print(f'{url}	   {fullname}	{titleurl}')
            ranking = '9 - X'


        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["user"] = user
        row_data["url"] = url
        row_data["lastname"] = lastname
        row_data["firstname"] = firstname
        row_data["fullname"] = fullname
        row_data["titleurl"] = titleurl
        row_data["pagestatus"] = pagestatus

        data.append(row_data)

        # time.sleep(10) #will sleep for 10 seconds


def breachbase(): 
    if len(emails) > 0:
        row_data = {}
        ranking ='7 - manual'
        url = ('https://breachbase.com/')
        row_data["ranking"] = ranking
        row_data["url"] = url
        data.append(row_data)
            

def vimeo():    # testuser=    kevinrose
    print(f'\n\t<<<<< vimeo users >>>>>')
    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 

    for user in users:    
        row_data = {}
        (query, ranking) = (user, '3 - vimeo')
        (fullname, firstname, lastname, middlename, note, DOB)  = ('','','','', '', '')
        (misc) = ('')

        url = (f'https://vimeo.com/{user}')
        (content, referer, osurl, titleurl, pagestatus) = ('','','','','')

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 
            
        # time.sleep(1) # will sleep for 1 seconds
        if pagestatus == 200:
            
            for eachline in content.split("  <"):
                if "og:description" in eachline:
                    note = eachline.strip().split("\"")[3]
            
            
            fullname = str(titleurl)
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)        
        
            print(f'{url} {fullname}')    

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["lastname"] = lastname
            row_data["firstname"] = firstname
            row_data["middlename"] = firstname
            row_data["fullname"] = fullname
            row_data["note"] = note
            row_data["user"] = user
            # row_data["misc"] = misc
            # row_data["content"] = content
            # row_data["pagestatus"] = pagestatus
            # row_data["titleurl"] = titleurl

            data.append(row_data)

            

def whatismyip():    # testuser= 77.15.67.232  
    print(f"\n\t<<<<< whatismyipaddress.com IP's >>>>>")
    for ip in ips:
        row_data = {}
        (query, city, state, zipcode, pagestatus, title) = (ip, '', '', '', '', '')
        url = (f'https://whatismyipaddress.com/ip/{ip}')

        ranking = '9 - whatismyipaddress'
        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["url"] = url
        row_data["ip"] = ip
        data.append(row_data)

def whatnot(): # testuser = kevinrose
    print(f'\n\t<<<<< whatnot users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '9 - whatnot')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (city, country, fullname, titleurl, pagestatus) = ('', '', '', '', '')
        user = user.rstrip()
        url = (f'https://www.whatnot.com/user/{user}')
        
        # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        
        
        if '404' not in str(pagestatus):
            titleurl = titleurl.replace("Just a moment...","").strip()

            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["titleurl"] = titleurl
            row_data["pagestatus"] = pagestatus
            data.append(row_data)


def whatsmyname():    # testuser=   kevinrose
    print(f'\n\t<<<<< Manually check whatsmyname users >>>>>')
    for user in users:    
        row_data = {}
        query = user
        url = ('https://whatsmyname.app/')
        
        note = (f'cd C:\Forensics\scripts\python\git-repo\WhatsMyName && python web_accounts_list_checker.py -u {user} -of C:\Forensics\scripts\python\output_{user}txt') 
        ranking = '7 - manual'
        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["url"] = url
        row_data["note"] = note
        row_data["user"] = user
        data.append(row_data)

def whitepagesphone():# testuser=    210-316-9435
    print(f'\n\t<<<<< whitepages phone numbers >>>>>')
    for phone in phones:    
        row_data = {}
        (query, ranking, state) = (phone, '9 - whitepages', '')
        url = ('https://www.whitepages.com/phone/1-%s' %(phone.lstrip('1')))
        state = phone_state_check(phone, state)
        # (content, referer, osurl, titleurl, pagestatus) = request(url)    # access denied cloudflare
       
        row_data["query"] = query
        row_data["ranking"] = ranking
        row_data["url"] = url
        row_data["phone"] = phone
        row_data["state"] = state
        data.append(row_data)    

def whocalld():# testPhone=  DROP THE LEADING 1
    print(f'\n\t<<<<< whocalld phone numbers >>>>>')
    url = ('https://whocalld.com' )

    # If the site is down, skip the module
    if not ping_url(url):
        print(f'{url} is down, skipping reversephonecheck module.')
        return

    for phone in phones:
        row_data = {}
        (query, note) = (phone, '')
        (country, city, state, zipcode, case, note) = ('', '', '', '', '', '')
        (fullname, content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '', '')
 
        url = ('https://whocalld.com/+1%s' %(phone.lstrip('1')))
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)    # protected by cloudflare

        for eachline in content.split("\n"):
            if "Not found" in eachline and case == '':

                url = ('')
            elif "This seems to be" in eachline:
                if ' in ' in eachline:
                    note = eachline.replace(". </p>",'').replace("<p>",'').strip().replace("This",phone)
                    city = eachline.split(" in ")[1].replace(". </p>",'').replace("<p>",'').strip()
                    if ", " in city:
                        state = city.split(", ")[1].replace(".",'')
                        city = city.split(", ")[0]
                    note = (f'According to {url} {note}')
            elif "The name of this caller seemed to be " in eachline:
                note = eachline
                fullname = eachline.replace("The name of this caller seemed to be ",'').split(",")[0].strip()
                if ' in ' in eachline:
                    # try:
                        # note = eachline.replace(". </p>",'').replace("<p>",'').strip().replace("This",phone)
                        # city = eachline.split(" in ")[2].replace(". </p>",'').replace("<p>",'').strip()
                    # except TypeError as error:
                        # print(f'{error}')

                    if ", " in city:
                        state = city.split(", ")[1].replace(".",'')
                        city = city.split(", ")[0]
        if state == '':
            state = phone_state_check(phone, state).replace('?', '')

        pagestatus = ''        
                
        if url != '':
            print(f'{url}	{fullname}  {city}  {state}') 
            
            ranking = '3 - spydialer'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = 'https://www.spydialer.com/'
            row_data["note"] = note
            row_data["fullname"] = fullname
            row_data["phone"] = phone
            row_data["note"] = note
            row_data["city"] = city
            row_data["country"] = country
            row_data["state"] = state
            row_data["zipcode"] = zipcode   
            data.append(row_data)

        else:
            ranking = '4 - spydialer'
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = 'https://www.spydialer.com'
            row_data["phone"] = phone
            row_data["state"] = state
            data.append(row_data)


        time.sleep(2) 

def whoisip():    # testuser=    77.15.67.232   only gets 403 Forbidden
    from subprocess import call, Popen, PIPE
    print(f"\n\t<<<<< whois IP's >>>>>")
    for ip in ips:    
        row_data = {}
        (query, ranking, note) = (ip, '9 - whois', '')
        (city, business, country, zipcode, state) = ('', '', '', '', '')
        (Latitude, Longitude) = ('', '')

        (content, titleurl, pagestatus) = ('', '', '')
        (email, phone, fullname, entity, fulladdress) = ('', '', '', '', '') 
        url = (f'https://www.ipaddress.com/ipv4/{ip}')      

        if regex_ipv6.match(ip):
            url = f'https://www.ipaddress.com/ipv6/{ip}'

        if sys.platform == 'win32' or sys.platform == 'win64':    
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            if '403 Forbidden' in content:
                pagestatus = '403 Forbidden'
                content = ''

            for eachline in content.split("\n"):
                 #
                if " is located in " in eachline:
                    (ranking) = ('4 - whois')
                    meta_match = re.search(r'content="([^"]+)"', eachline)
                    if meta_match:
                        note = meta_match.group(1)
                    note = note.replace(' | Public IP Address', '')
                if 'Reserved IP Address' in eachline:
                    note = (f'{ip} is a reserved IP address.')



                    
            time.sleep(3) #will sleep for 30 seconds
            if "Fail" in pagestatus:
                pagestatus = 'fail'
        else:
            WhoisArgs = (f'whois {ip}')
            response= Popen(WhoisArgs, shell=True, stdout=PIPE)
            for line in response.stdout:
                line = line.decode("utf-8")
                if ':' in line and "# " not in line and len(line) > 2:
                    line = line.strip()
                    content = (f'{content}\n{line}')
                if email == '':
                    if line.startswith('RAbuseEmail:'):
                        try:
                            email = (line.split(': ')[1].lstrip())
                        except:pass    
                    elif line.lower().startswith('abuse-mailbox:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabuseemail:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('Orgtechemail:'):email = (line.split(': ')[1].lstrip())
                
                if phone == '':
                    if line.lower().startswith('rabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('phone:'):phone = (line.split(': ')[1].lstrip())
                    phone = phone.replace("+", "")
                if line.lower().startswith('rtechname:'):fullname = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('person:'):fullname = (line.split(': ')[1].lstrip())                
                
                if line.lower().startswith('country:'):country = (line.split(': ')[1].lstrip())
                if line.lower().startswith('city:'):city = (line.split(': ')[1].lstrip())
                if line.lower().startswith('address:'):fulladdress = ('%s %s' %(fulladdress, line.split(': ')[1].lstrip()))
                if line.lower().startswith('stateprov:'):state = (line.split(': ')[1].lstrip())
                if line.lower().startswith('postalcode:'):zipcode = (line.split(': ')[1].lstrip())
                if line.lower().startswith('orgname:'):entity = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('org-name:'):entity = (line.split(': ')[1].lstrip())

        print(f'{ip}	{country}	{city}	{zipcode}')
        if note == '':
            note = entity


        if '.' in ip or ':' in ip:
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["ip"] = ip
            row_data["note"] = note
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["email"] = email
            row_data["phone"] = phone
            # row_data["note"] = entity
            row_data["fulladdress"] = fulladdress
            row_data["query"] = query
            
            row_data["city"] = city
            row_data["country"] = country
            row_data["state"] = state
            row_data["zipcode"] = zipcode            
            row_data["Latitude"] = Latitude
            row_data["Longitude"] = Longitude
            # row_data["content"] = content # temp
            # row_data["pagestatus"] = pagestatus # temp
           
           
            data.append(row_data)

def wordpress(): # testuser = kevinrose
    print(f'\n\t<<<<< wordpress users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '6 - wordpress')


        (Success, fullname, lastname, firstname, case, SEX) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city, note) = ('', '')
        user = user.rstrip()
        url = (f'https://wordpress.org/support/users/{user}/')
        note = (f'https://{user}.wordspress.com')        
        
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        
        try:
            (content, referer, osurl, titleurl, pagestatus) = request(url)
            # (content, referer, osurl, titleurl, pagestatus) = ('','','','','') # too many false positives
        except socket.error as ex:
            print(f'{ex}')

        if 'That page can' not in content:
        # if 'Do you want to register' not in content:
            titleUrl = titleurl.replace("'s Profile | WordPress.org","").strip()
            fullname = str(titleurl)
            fullname = fullname.split(" (")[0]
            print(f'{url}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["note"] = note
            row_data["user"] = user
            row_data["note"] = note
            row_data["pagestatus"] = pagestatus
            row_data["titleurl"] = titleurl
            # row_data["content"] = content            
            data.append(row_data)


def wordpress_profiles(): # testuser = kevinrose
    print(f'\n\t<<<<< wordpress users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user,'7 - wordpress')


        (Success, fullname, lastname, firstname, case, SEX) = ('','','','','','')
        (photo, country, website, email, language, username) = ('','','','','','')
        (city) = ('')
        user = user.rstrip()
        url = (f'https://profiles.wordpress.org/{user}/')
        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = ('','','','','')
        except:
            pass
        if '404' not in str(pagestatus):
            fullname = str(titleurl)
            fullname = fullname.split(" (")[0]
            if fullname.lower() in titleurl.lower():
                (fullname, titleurl) = ('', '')
            
            print(f'{url}	{fullname}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["user"] = user
            row_data["pagestatus"] = pagestatus
            row_data["titleurl"] = titleurl
            # row_data["content"] = content             
            data.append(row_data)


def write_blurb():
    '''
        read Intel.xlsx and write OSINT_.docx to describe what you found.
    '''
    fullname_column_index = ''
    docx_file = output_xlsx

    # message = (f'Writing blurb from {input_xlsx} line 5503')
    # message_square(message)

    if not os.path.exists(input_xlsx):
        input(f"{input_xlsx} doesnt exist.")
        sys.exit()
    elif os.path.getsize(input_xlsx) == 0:
        input(f'{input_xlsx} is empty. Fill it with intel you found.')
        sys.exit()
    elif os.path.isfile(input_xlsx):
        print(f' ')
    else:
        input(f'See {input_xlsx} does not exist. Hit Enter to exit...')
        sys.exit()


    # Open the Excel file
    wb = openpyxl.load_workbook(input_xlsx)
    sheet = wb.active
    
    # Find the column headers
    header_row = sheet[1]
    column_names = [cell.value for cell in header_row]    
    
    # Columns to skip
    columns_to_skip = [None, "None", "query", "ranking", "content", "referer", "osurl", "titleurl", "pagestatus", "city", "state", "country", "firstname", "middlename", "lastname", "Latitude", "Longitude", "Coordinate", "original_file", "Source file information", "Icon", "Type", "Tag"]

    
    for idx, cell in enumerate(header_row, start=1):
        if cell.value == "fullname":
            fullname_column_index = idx
            # print(f'Fullname: {fullname_column_index}')
            break

    if fullname_column_index is None:
        print("Fullname column not found in the Excel file.")
        return

    # Create a new Word document
    doc = Document()
 
    sentence = (f'An open-source search revealed the following details.\n\n')
    
    print(f'{sentence}')  
    doc.add_paragraph(sentence)    
    # Loop through rows in the Excel file and write to Word document
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not any(row): continue
        sentence = "\n".join(f"{column}: {value}" for column, value in zip(column_names, row) if column not in columns_to_skip and value is not None)
        if sentence:
            sentence = sentence.replace(' 00:00:00', '')    # test
            doc.add_paragraph(sentence)
            doc.add_paragraph("")  # Add an empty line between rows


    # Save the Word document
    doc.save(docx_file)

    message = (f'Data written to {docx_file}')
    message_square(message)


def write_intel(data):
    '''
    The write_locations() function receives the processed data as a list of 
    dictionaries and writes it to a new Excel file using openpyxl. 
    It defines the column headers, sets column widths, and then iterates 
    through each row of data, writing it into the Excel worksheet.
    '''
    message = (f'Writing {output_xlsx}')
    message_square(message)

    try:
        data = sorted(data, key=lambda x: (x.get("ranking", ""), x.get("fullname", ""), x.get("query", "")))
        print(f'sorted by ranking')
    except TypeError as error:

        print(f'{error}')

    global workbook
    workbook = Workbook()
    global worksheet
    worksheet = workbook.active

    worksheet.title = 'Intel'
    header_format = {'bold': True, 'border': True}
    worksheet.freeze_panes = 'B2'  # Freeze cells
    worksheet.selection = 'B2'

    log_headers = [
        "Date", "Subject", "Requesting Agency", "Requesting Agent", "Case"
        , "Summary of Findings", "Source", "Notes"
    ]


    # Write headers to the first row
    for col_index, header in enumerate(headers_intel):
        cell = worksheet.cell(row=1, column=col_index + 1)
        cell.value = header
        if col_index in [2, 3, 4, 5, 6]: 
            fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid") # orange?
            cell.fill = fill
        elif col_index in [7,8, 12, 13, 14, 15, 16, 17, 21, 22, 23, 24, 25, 26, 35, 36, 37, 38, 39, 40, 41, 42, 43, 49, 50]:  # yellow headers
            fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")  # Use yellow color
            cell.fill = fill
        # elif col_index == 27:  # Red for column 27
            # fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")  # Red color
            # cell.fill = fill




    ## Excel column width

    worksheet.column_dimensions['A'].width = 15 # query
    worksheet.column_dimensions['B'].width = 20 # ranking
    worksheet.column_dimensions['C'].width = 20 # fullname
    worksheet.column_dimensions['D'].width = 25 # url
    worksheet.column_dimensions['E'].width = 25 # email
    worksheet.column_dimensions['F'].width = 15 # user
    worksheet.column_dimensions['G'].width = 14 # phone
    worksheet.column_dimensions['H'].width = 16 # business
    worksheet.column_dimensions['I'].width = 24 # fulladdress
    worksheet.column_dimensions['J'].width = 12 # city
    worksheet.column_dimensions['K'].width = 10 # state
    worksheet.column_dimensions['L'].width = 8 # country
    worksheet.column_dimensions['M'].width = 20 # note
    worksheet.column_dimensions['N'].width = 14 # AKA
    worksheet.column_dimensions['O'].width = 11 # DOB
    worksheet.column_dimensions['P'].width = 5 # SEX
    worksheet.column_dimensions['Q'].width = 20 # info
    worksheet.column_dimensions['R'].width = 20 # misc
    worksheet.column_dimensions['S'].width = 10 # firstname
    worksheet.column_dimensions['T'].width = 11 # middlename
    worksheet.column_dimensions['U'].width = 10 # lastname
    worksheet.column_dimensions['V'].width = 10 # associates
    worksheet.column_dimensions['W'].width = 10 # case
    worksheet.column_dimensions['X'].width = 13 # sosfilenumber
    worksheet.column_dimensions['Y'].width = 10 # owner
    worksheet.column_dimensions['Z'].width = 10 # president
    worksheet.column_dimensions['AA'].width = 10 # sosagent
    worksheet.column_dimensions['AB'].width = 10 # managers
    worksheet.column_dimensions['AC'].width = 15 # Time
    worksheet.column_dimensions['AD'].width = 12 # Latitude
    worksheet.column_dimensions['AE'].width = 12 # Longitude
    worksheet.column_dimensions['AF'].width = 22 # Coordinate
    worksheet.column_dimensions['AG'].width = 12 # original_file
    worksheet.column_dimensions['AH'].width = 12 # Source
    worksheet.column_dimensions['AI'].width = 12 # Source file information
    worksheet.column_dimensions['AJ'].width = 10 # Plate
    worksheet.column_dimensions['AK'].width = 10 # VIS
    worksheet.column_dimensions['AL'].width = 10 # VIN
    worksheet.column_dimensions['AM'].width = 10 # VYR
    worksheet.column_dimensions['AN'].width = 10 # VMA
    worksheet.column_dimensions['AO'].width = 10 # LIC
    worksheet.column_dimensions['AP'].width = 10 # LIY
    worksheet.column_dimensions['AQ'].width = 10 # DLN
    worksheet.column_dimensions['AR'].width = 10 # DLS
    worksheet.column_dimensions['AS'].width = 10 # content
    worksheet.column_dimensions['AT'].width = 10 # referer
    worksheet.column_dimensions['AU'].width = 10 # osurl
    worksheet.column_dimensions['AV'].width = 10 # titleurl
    worksheet.column_dimensions['AW'].width = 12 # pagestatus
    worksheet.column_dimensions['AX'].width = 16 # ip
    worksheet.column_dimensions['AY'].width = 15 # dnsdomain

    for i in range(len(data)):
        if data[i] is None:
            data[i] = ''

    for row_index, row_data in enumerate(data):

        for col_index, col_name in enumerate(headers_intel):
            try:
                cell_data = row_data.get(col_name)
                worksheet.cell(row=row_index+2, column=col_index+1).value = cell_data
            except Exception as e:
                print(f"Error printing line: {str(e)}")

    # Create a new worksheet for color codes
    color_worksheet = workbook.create_sheet(title='ColorCode')
    color_worksheet.freeze_panes = 'B2'  # Freeze cells

    # Excel column width
    color_worksheet.column_dimensions['A'].width = 14# Color
    color_worksheet.column_dimensions['B'].width = 20# Description


    # Excel row height
    color_worksheet.row_dimensions[2].height = 22  # Adjust the height as needed
    color_worksheet.row_dimensions[3].height = 22
    color_worksheet.row_dimensions[4].height = 23
    color_worksheet.row_dimensions[5].height = 23
    color_worksheet.row_dimensions[6].height = 40   # truck

    color_worksheet.cell(row=1, column=1).value = 'Color'
    color_worksheet.cell(row=1, column=2).value = 'description'
    color_worksheet.cell(row=2, column=1).value = 'Red'
    color_worksheet.cell(row=3, column=1).value = 'Orange'
    color_worksheet.cell(row=4, column=1).value = 'Green'
    color_worksheet.cell(row=5, column=1).value = 'Yellow'

    color_worksheet.cell(row=7, column=1).value = 'ABBREVIATIONS'
    color_worksheet.cell(row=8, column=1).value = 'AKA'
    color_worksheet.cell(row=9, column=1).value = 'DOB'
    color_worksheet.cell(row=10, column=1).value = 'VIS'
    color_worksheet.cell(row=11, column=1).value = 'VIN'
    color_worksheet.cell(row=12, column=1).value = 'VYR'
    color_worksheet.cell(row=13, column=1).value = 'VMA'
    color_worksheet.cell(row=14, column=1).value = 'LIC'
    color_worksheet.cell(row=15, column=1).value = 'LIY'
    color_worksheet.cell(row=16, column=1).value = 'DLN'
    color_worksheet.cell(row=17, column=1).value = 'DLS'

       
    color_worksheet.cell(row=2, column=2).value = 'Bad Intel or dead link'
    color_worksheet.cell(row=3, column=2).value = 'Research'
    color_worksheet.cell(row=4, column=2).value = 'Good Intel'
    color_worksheet.cell(row=5, column=2).value = 'Highlighted'

    color_worksheet.cell(row=8, column=2).value = 'Also Known As (Alias)'
    color_worksheet.cell(row=9, column=2).value = 'Date of Birth'
    color_worksheet.cell(row=10, column=2).value = 'Vehicle State'
    color_worksheet.cell(row=11, column=2).value = 'Vehicle Identification Number'
    color_worksheet.cell(row=12, column=2).value = 'Vehicle Year'
    color_worksheet.cell(row=13, column=2).value = 'Vehicle Make'
    color_worksheet.cell(row=14, column=2).value = 'License'
    color_worksheet.cell(row=15, column=2).value = 'License Year'
    color_worksheet.cell(row=16, column=2).value = 'Drivers License Number'
    color_worksheet.cell(row=17, column=2).value = 'Drivers License State'


    # colored fills
    red_fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')
    orange_fill = PatternFill(start_color='FFA500', end_color='FFA500', fill_type='solid')
    green_fill = PatternFill(start_color='00FF00', end_color='00FF00', fill_type='solid')
    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')


    # Apply the orange fill to the cell in row 2, column 2
    color_worksheet.cell(row=2, column=2).fill = red_fill
    color_worksheet.cell(row=3, column=2).fill = orange_fill
    color_worksheet.cell(row=4, column=2).fill = green_fill
    color_worksheet.cell(row=5, column=2).fill = yellow_fill


    # Create a new worksheet for logs
    log_worksheet = workbook.create_sheet(title='Log')
    log_worksheet.freeze_panes = 'B2'  # Freeze cells

# Date, Subject, Requesting Agency, Requesting Agent, Case, Summary of Findings, Source, Notes, Requestor

    # Excel column width
    log_worksheet.column_dimensions['A'].width = 14# Date
    log_worksheet.column_dimensions['B'].width = 20# Subject
    log_worksheet.column_dimensions['C'].width = 24# Requesting Agency
    log_worksheet.column_dimensions['D'].width = 20# Requesting Agent
    log_worksheet.column_dimensions['E'].width = 14# Case
    log_worksheet.column_dimensions['F'].width = 20# Summary of Findings
    log_worksheet.column_dimensions['G'].width = 14# Source
    log_worksheet.column_dimensions['H'].width = 25# Notes

    log_worksheet.cell(row=1, column=1).value = 'Date'
    log_worksheet.cell(row=1, column=2).value = 'Subject'
    log_worksheet.cell(row=1, column=3).value = 'Requesting Agency'
    log_worksheet.cell(row=1, column=4).value = 'Requesting Agent'
    log_worksheet.cell(row=1, column=5).value = 'Case'
    log_worksheet.cell(row=1, column=6).value = 'Summary of Findings'
    log_worksheet.cell(row=1, column=7).value = 'Notes'



    workbook.save(output_xlsx)

# Save the workbook
# wb.save('output.xlsx')

def write_intel_basic(data, output_xlsx):
    message = (f'Writing intel to {output_xlsx}')
    message_square(message)

    wb = openpyxl.Workbook()
    ws = wb.active

    # ws.append(headers)  # Writing headers
    ws.append(headers_intel)  # Writing headers


    for row_data in data:
        row = [row_data.get(header, '') for header in headers_intel]
        ws.append(row)

    wb.save(output_xlsx)

def write_locations_basic(data, output_xlsx):
    message = (f'Writing locations to {output_xlsx}')
    message_square(message)

    wb = openpyxl.Workbook()
    ws = wb.active

    ws.append(headers_locations)  # Writing location headers

    for row_data in data:
        row = [row_data.get(header_l, '') for header_l in headers_locations]
        ws.append(row)

    wb.save(output_xlsx)

def write_locations(data):
    '''
    The write_locations() function receives the processed data as a list of 
    dictionaries and writes it to a new Excel file using openpyxl. 
    It defines the column headers, sets column widths, and then iterates 
    through each row of data, writing it into the Excel worksheet.
    '''
    message = (f'Writing locations to {output_xlsx}')
    message_square(message)

    global workbook
    workbook = Workbook()
    global worksheet
    worksheet = workbook.active

    worksheet.title = 'Locations'
    header_format = {'bold': True, 'border': True}
    worksheet.freeze_panes = 'B2'  # Freeze cells
    worksheet.selection = 'B2'

    headers_locations = [
        "#", "Time", "Latitude", "Longitude", "Address", "Group", "Subgroup"
        , "Description", "Type", "Source", "Deleted", "Tag", "Source file information"
        , "Service Identifier", "Carved", "Name", "business", "number", "street"
        , "city", "county", "state", "zipcode", "country", "fulladdress", "query"
        , "Sighting State", "Plate", "Capture Time", "Capture Network", "Highway Name"
        , "Coordinate", "Capture Location Latitude", "Capture Location Longitude"
        , "Container", "Sighting Location", "Direction", "Time Local", "End time"
        , "Category", "Manually decoded", "Account", "PlusCode", "Time Original", "Timezone"
        , "Icon", "original_file", "case", "Index"

    ]

    # Write headers to the first row
    for col_index, header in enumerate(headers_locations):
        cell = worksheet.cell(row=1, column=col_index + 1)
        cell.value = header
        if col_index in [2, 3, 4]: 
            fill = PatternFill(start_color="FFA500", end_color="FFA500", fill_type="solid") # orange?
            cell.fill = fill
        elif col_index in [1, 5, 6, 7, 8, 9, 15, 16, 24, 30, 31, 36, 38]:  # yellow headers
            fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")  # Use yellow color
            cell.fill = fill
        elif col_index == 27:  # Red for column 27
            fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")  # Red color
            cell.fill = fill
    try:
        ## Excel column width
        worksheet.column_dimensions['A'].width = 8# #
        worksheet.column_dimensions['B'].width = 19# Time
        worksheet.column_dimensions['C'].width = 18# Latitude
        worksheet.column_dimensions['D'].width = 18# Longitude
        worksheet.column_dimensions['E'].width = 45# Address
        worksheet.column_dimensions['F'].width = 14# Group
        worksheet.column_dimensions['G'].width = 13# Subgroup
        worksheet.column_dimensions['H'].width = 17# Description
        worksheet.column_dimensions['I'].width = 9# Type
        worksheet.column_dimensions['J'].width = 10# Source
        worksheet.column_dimensions['K'].width = 10# Deleted
        worksheet.column_dimensions['L'].width = 11# Tag
        worksheet.column_dimensions['M'].width = 20# Source file information
        worksheet.column_dimensions['N'].width = 15# Service Identifier
        worksheet.column_dimensions['O'].width = 7# Carved
        worksheet.column_dimensions['P'].width = 15# Name
        
        ## bonus
        worksheet.column_dimensions['Q'].width = 20# business 
        worksheet.column_dimensions['R'].width = 10# number
        worksheet.column_dimensions['S'].width = 20# street 
        worksheet.column_dimensions['T'].width = 15# city   
        worksheet.column_dimensions['Y'].width = 25# county    
        worksheet.column_dimensions['V'].width = 12# state   
        worksheet.column_dimensions['W'].width = 8# zipcode     
        worksheet.column_dimensions['X'].width = 6# country    
        worksheet.column_dimensions['Y'].width = 26# FullAddress   
        worksheet.column_dimensions['Z'].width = 26# query

        ##  Flock
        worksheet.column_dimensions['AA'].width = 11# Sighting State
        worksheet.column_dimensions['AB'].width = 11# Plate
        worksheet.column_dimensions['AC'].width = 22# Capture Time
        worksheet.column_dimensions['AD'].width = 15# Capture Network
        worksheet.column_dimensions['AE'].width = 21# Highway Name
        worksheet.column_dimensions['AF'].width = 30# Coordinate
        worksheet.column_dimensions['AG'].width = 20# Capture Location Latitude
        worksheet.column_dimensions['AH'].width = 20# Capture Location Longitude

        ##
        worksheet.column_dimensions['AI'].width = 10# Container
        worksheet.column_dimensions['AJ'].width = 14# Sighting Location
        worksheet.column_dimensions['AK'].width = 10# Direction
        worksheet.column_dimensions['AL'].width = 11# Time Local
        worksheet.column_dimensions['AM'].width = 25# End time
        worksheet.column_dimensions['AN'].width = 10# Category
        worksheet.column_dimensions['AO'].width = 18# Manually decoded
        worksheet.column_dimensions['AP'].width = 10# Account
        worksheet.column_dimensions['AQ'].width = 25 # PlusCode
        worksheet.column_dimensions['AR'].width = 21 # Time Original
        worksheet.column_dimensions['AS'].width = 9 # Timezone
        worksheet.column_dimensions['AT'].width = 10 # Icon   
        worksheet.column_dimensions['AU'].width = 20 # original_file
        worksheet.column_dimensions['AV'].width = 10 # case
        worksheet.column_dimensions['AW'].width = 6 # Index
    except:pass
    
    for row_index, row_data in enumerate(data):

        for col_index, col_name in enumerate(headers_locations):
            cell_data = row_data.get(col_name)
            try:
                worksheet.cell(row=row_index+2, column=col_index+1).value = cell_data
            except Exception as e:
                print(f"Error printing line: {str(e)}")


    # Create a new worksheet for color codes
    color_worksheet = workbook.create_sheet(title='Icons')
    color_worksheet.freeze_panes = 'B2'  # Freeze cells

    # Excel column width
    color_worksheet.column_dimensions['A'].width = 8# Icon sample
    color_worksheet.column_dimensions['B'].width = 9# Name
    color_worksheet.column_dimensions['C'].width = 29# Description

    # Excel row height
    color_worksheet.row_dimensions[2].height = 22  # Adjust the height as needed
    color_worksheet.row_dimensions[3].height = 22
    color_worksheet.row_dimensions[4].height = 23
    color_worksheet.row_dimensions[5].height = 23
    color_worksheet.row_dimensions[6].height = 40   # truck
    color_worksheet.row_dimensions[7].height = 6
    color_worksheet.row_dimensions[8].height = 24
    color_worksheet.row_dimensions[9].height = 22
    color_worksheet.row_dimensions[10].height = 22
    color_worksheet.row_dimensions[11].height = 22
    color_worksheet.row_dimensions[12].height = 23
    color_worksheet.row_dimensions[13].height = 23
    color_worksheet.row_dimensions[14].height = 25
    color_worksheet.row_dimensions[15].height = 25
    color_worksheet.row_dimensions[16].height = 23
    color_worksheet.row_dimensions[17].height = 6
    color_worksheet.row_dimensions[18].height = 38
    color_worksheet.row_dimensions[19].height = 38
    color_worksheet.row_dimensions[20].height = 38
    color_worksheet.row_dimensions[21].height = 38
    color_worksheet.row_dimensions[22].height = 38
    color_worksheet.row_dimensions[23].height = 6
    color_worksheet.row_dimensions[24].height = 15
    color_worksheet.row_dimensions[25].height = 6
    color_worksheet.row_dimensions[26].height = 15


    
    # Define color codes
    color_worksheet['A1'] = ' '
    color_worksheet['B1'] = 'Icon'
    color_worksheet['C1'] = 'Icon Description'

    icon_data = [

        ('', 'Car', 'Lpr red car (License Plate Reader)'),
        ('', 'Car2', 'Lpr yellow car'),
        ('', 'Car3', 'Lpr greeen car with circle'),
        ('', 'Car4', 'Lpr red car with circle'),
        ('', 'Truck', 'Lpr truck'),         
        ('', '', ''),
        ('', 'Calendar', 'Calendar'), 
        ('', 'Home', 'Home'),                
        ('', 'Images', 'Photo'),
        ('', 'Intel', 'I'),  
        ('', 'Locations', 'Reticle'),  
        ('', 'default', 'Yellow flag'),  
        ('', 'Office', 'Office'),         
        ('', 'Searched', 'Searched Item'),          
        ('', 'Videos', 'Video clip'),        
        ('', '', ''),
        ('', 'Toll', 'Blue square'), 
        ('', 'N', 'Northbound blue arrow'),
        ('', 'E', 'Eastbound blue arrow'),
        ('', 'S', 'Southbound blue arrow'),
        ('', 'W', 'Westbound blue arrow'),
        ('', '', ''),
        ('', 'Yellow font', 'Tagged'),
        ('', 'Chats', 'Chats'),   # 


        ('', '', ''),
        ('', 'NOTE', 'visit https://earth.google.com/ <file><Import KML> select gps.kml <open>'),
    ]

    for row_index, (icon, tag, description) in enumerate(icon_data):
        color_worksheet.cell(row=row_index + 2, column=1).value = icon
        color_worksheet.cell(row=row_index + 2, column=2).value = tag
        color_worksheet.cell(row=row_index + 2, column=3).value = description

    car_icon = 'https://maps.google.com/mapfiles/kml/pal4/icon15.png'   # red car
    car2_icon = 'https://maps.google.com/mapfiles/kml/pal2/icon47.png'  # yellow car
    car3_icon = 'https://maps.google.com/mapfiles/kml/pal4/icon54.png'  # green car with circle
    car4_icon = 'https://maps.google.com/mapfiles/kml/pal4/icon7.png'  # red car with circle
    truck_icon = 'https://maps.google.com/mapfiles/kml/shapes/truck.png'    # blue truck
    default_icon = 'https://maps.google.com/mapfiles/kml/pal2/icon13.png'   # yellow flag
    calendar_icon = 'https://maps.google.com/mapfiles/kml/pal2/icon23.png' # paper
    chat_icon = 'https://maps.google.com/mapfiles/kml/shapes/post_office.png' # email
    locations_icon = 'https://maps.google.com/mapfiles/kml/pal3/icon28.png'    # yellow paddle
    home_icon = 'https://maps.google.com/mapfiles/kml/pal3/icon56.png'
    images_icon = 'https://maps.google.com/mapfiles/kml/pal4/icon46.png'
    intel_icon = 'https://maps.google.com/mapfiles/kml/pal3/icon44.png'
    office_icon = 'https://maps.google.com/mapfiles/kml/pal3/icon21.png'
    searched_icon = 'https://maps.google.com/mapfiles/kml/pal4/icon0.png'  #  
    toll_icon = 'https://earth.google.com/images/kml-icons/track-directional/track-none.png'
    videos_icon = 'https://maps.google.com/mapfiles/kml/pal2/icon30.png'
    n_icon = 'https://earth.google.com/images/kml-icons/track-directional/track-0.png'
    e_icon = 'https://earth.google.com/images/kml-icons/track-directional/track-4.png'
    s_icon = 'https://earth.google.com/images/kml-icons/track-directional/track-8.png'
    w_icon = 'https://earth.google.com/images/kml-icons/track-directional/track-12.png'

    try:
        # Insert graphic from URL into cell of color_worksheet

        response = requests.get(car_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A2')

        response = requests.get(car2_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A3')

        response = requests.get(car3_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A4')

        response = requests.get(car4_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A5')

        response = requests.get(truck_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A6')

        response = requests.get(calendar_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A8')
        
        response = requests.get(home_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A9')

        response = requests.get(images_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A10')

        response = requests.get(intel_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A11')

        response = requests.get(locations_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A12')

        response = requests.get(default_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A13')

        response = requests.get(office_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A14')        

        response = requests.get(searched_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A15')

        response = requests.get(videos_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A16')
        
        response = requests.get(toll_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A18')

        response = requests.get(n_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A19')

        response = requests.get(e_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A20')

        response = requests.get(s_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A21')

        response = requests.get(w_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A22')

        response = requests.get(chat_icon)
        img = Image(io.BytesIO(response.content))
        color_worksheet.add_image(img, 'A24')     
        
    except:
        pass

    
    workbook.save(output_xlsx)




def youtube(): # testuser = kevinrose
    print(f'\n\t<<<<< youtube users >>>>>')
    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - youtube')

        (fullname) = ('')
        user = user.rstrip()
        url = (f'https://www.youtube.com/{user}')
        (content, referer, osurl, titleurl, pagestatus) = request(url)
        titleurl = titleurl.replace(' - YouTube','')
        if '404' not in str(pagestatus):
            fullname = str(titleurl)
            
            if fullname.lower() == user.lower():
                fullname = ''
            if ' ' in fullname:
                (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            else:
                (fullname, firstname, lastname, middlename) = ('', '', '', '')


            print(f'{url}	{fullname}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["url"] = url
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname

            data.append(row_data)


def robtex():
    print(f'\n\t<<<<<robtex dns lookup >>>>>')    

    for website in websites:    
        row_data = {}
        (query, ranking) = (website, '9 - robtexDNS-lookup')

        (final_url, dnsdomain, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '', '')
        (otherurl, ip) = ('', '')
        url = website
        url = url.replace("http://", "https://")
        if "http" not in url.lower():
            url = (f'https://{url}')
        
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]

        ip = ip_address(dnsdomain)

        note = url  

        url = (f'https://www.robtex.com/dns-lookup/{dnsdomain}#quick')
        
        if 1==1:
        # if dnsdomain not in final_url:
            print(f'{website}	{ip}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["ip"] = ip            
          
            row_data["note"] = note            
            row_data["dnsdomain"] = dnsdomain    
            row_data["referer"] = referer   
            row_data["osurl"] = osurl              
            row_data["titleurl"] = titleurl              
            row_data["pagestatus"] = pagestatus              

            data.append(row_data)    


def titles():    # testsite= google.com
    from subprocess import call, Popen, PIPE
    print(f'\n\t<<<<< Titles grab Website\'s >>>>>')

    if playwright_ready() is False:
        print(f'pip install playwright playwright_stealth')
        return 
        
    for website in websites:    
        row_data = {}
        (query, ranking) = (website, '7 - website')

        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        (ip, note) = ('', '')
        
        url = website

        fake_referer = 'https://www.google.com/'
        headers_url = {'Referer': fake_referer}
        url = url.replace("http://", "https://")     # test
        if "http" not in url.lower():
            url = (f'https://{url}')

        url = url.replace("https://", "http://")

        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}') 

        # dnsdomain
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]
        
        # ip
        ip = ip_address(dnsdomain)

        if pagestatus == 200:
            print(f'{website}	   {pagestatus}	{titleurl}')
            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["ip"] = ip            
          
            row_data["note"] = note            
            row_data["dnsdomain"] = dnsdomain    
            row_data["referer"] = referer   
            row_data["osurl"] = osurl              
            row_data["titleurl"] = titleurl              
            row_data["pagestatus"] = pagestatus              
          
            data.append(row_data) 
        else:
            print(f'{website}	   {pagestatus}	{titleurl}')            

def validnumber():# testPhone= 
    print(f'\n\t<<<<< validnumber phone numbers >>>>>')

    # https://validnumber.com/phone-number/3124377966/
    for phone in phones:
        row_data = {}
        (query, ranking) = (phone, '9 - validnumber')
        (country, city, state, zipcode, case, note) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        (query) = (phone)
 
        url = ('https://validnumber.com/phone-number/%s/' %(phone.lstrip("1")))
        
        (content, referer, osurl, titleurl, pagestatus) = request(url)    # protected by cloudflare
        pagestatus = ''
        for eachline in content.split("\n"):
            if "No name associated with this number" in eachline and case == '':
                print(f'not found')  # temp
                
                # url = ('')
            elif "Find out who owns" in eachline:
                if 'This device is registered in ' in eachline:
                    ranking = '5 - validnumber'
                    note = eachline.split('\"')[1]
                    note = note.split('Free owner details for')[0]
                    city = eachline.split("This device is registered in ")[1].split("Free owner details")[0]
                    state = city.split(',')[1].strip()
                    city = city.split(',')[0]
        if city != '':        
            city = city.title()
            if city == "Directory Assistance":
                ranking ='7 - validnumber'
        state = state.replace('..','').replace('Illinois','IL')
        if state == '':
            state = phone_state_check(phone, state).replace('?', '')

        if url != '':        

            print(f'{url}') 

            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone            
            row_data["note"] = note            
            row_data["state"] = state    
            row_data["city"] = city   

            data.append(row_data) 


def venmo(): # testuser = kevinrose
    print(f'\n\t<<<<< venmo users >>>>>')
    if playwright_ready() is False:
        # print(f'pip install playwright playwright_stealth')
        return (content, referer, osurl, titleurl, pagestatus)
    # else:
        # print(f'Playwright is installed')


    for user in users:    
        row_data = {}
        (query, ranking) = (user, '4 - venmo')
        (fullname, firstname, lastname, middlename) = ('', '', '', '')
        (fullname) = ('')
        user = user.rstrip()
        url = (f'https://account.venmo.com/u/{user}')
        
        try:
            # (content, referer, osurl, titleurl, pagestatus) = request(url)
            (content, referer, osurl, titleurl, pagestatus) = asyncio.run(playwright_url(url))
        except TypeError as error:
            print(f'{error}')         
        
        
        titleurl = titleurl.replace('Venmo | ','')
        fullname = str(titleurl).strip()
        try:
            fullname = fullname.split(' | ')[0]
        except:pass
        if ' ' in fullname:
            (fullname, firstname, middlename, lastname) = fullname_parse(fullname)
            
        # if '404' not in str(pagestatus):
        if pagestatus == 200:
            print(f'{url}   {fullname}')
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["user"] = user
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["middlename"] = middlename
            row_data["url"] = url
            # row_data["pagestatus"] = pagestatus
            # row_data["content"] = content
            # row_data["titleurl"] = titleurl
            data.append(row_data)
            

def viewdnsdomain():
    print(f'\n\t<<<<<viewdns lookup >>>>>')    

    for website in websites:    
        row_data = {}
        (query, ranking) = (website, '9 - viewdns')

        (final_url, dnsdomain, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus) = ('', '', '', '', '')
        (otherurl, ip) = ('', '')
        url = website
        url = url.replace("http://", "https://")
        if "http" not in url.lower():
            url = (f'https://{url}')
        
        dnsdomain = url.lower()
        dnsdomain = dnsdomain.replace("https://", "")
        dnsdomain = dnsdomain.replace("http://", "")
        dnsdomain = dnsdomain.split('/')[0]

        ip = ip_address(dnsdomain)

        note = url  

        url = (f'https://viewdns.info/whois/?domain={dnsdomain}')
        # (content, referer, osurl, titleurl, pagestatus) = request(url)
        # time.sleep(10) #will sleep for 10 seconds
        if 1==1:
        # if dnsdomain not in final_url:
            print(f'{website}	{ip}')

            row_data["query"] = query
            row_data["ranking"] = ranking
            # row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["ip"] = ip            
          
            row_data["note"] = note            
            row_data["dnsdomain"] = dnsdomain    
            row_data["referer"] = referer   
            row_data["osurl"] = osurl              
            row_data["titleurl"] = titleurl              
            row_data["pagestatus"] = pagestatus              

            data.append(row_data)    


def whoiswebsite():    # testsite= google.com
    from subprocess import call, Popen, PIPE
    print(f'\n\t<<<<< whois Website\'s >>>>>')

    for dnsdomain in dnsdomains:    
        row_data = {}
        (query, ranking) = (dnsdomain, '7 - whois')

        # query = website
        # website = website.replace('http://','')
        # website = website.replace('https://','')
        # website = website.replace('www.','')
        
        url = (f'https://www.ip-adress.com/website/{dnsdomain}')
        url2 = ('https://whois.domaintools.com/%s' %(dnsdomain.replace('www.','')))
        (email,phone,fullname,country,city,state) = ('','','','','','')
        (city, country, zipcode, state, ip) = ('', '', '', '', '')
        (content, titleurl, pagestatus) = ('', '', '')
        (email, phone, fullname, entity, fulladdress) = ('', '', '', '', '') 

        if sys.platform == 'win32' or sys.platform == 'win64':    
            print(f'skipping whois query from windows')  # temp

            row_data["query"] = query
            row_data["ranking"] = '9 - whois.domaintools.com'
            row_data["fullname"] = fullname
            row_data["url"] = url
            row_data["city"] = city            
            row_data["country"] = country            
            row_data["state"] = state            
            row_data["dnsdomain"] = dnsdomain            

            data.append(row_data) 

        elif dnsdomain.endswith('.com') or dnsdomain.endswith('.edu') or dnsdomain.endswith('.net'):
            WhoisArgs = (f'whois {dnsdomain}')
            response= Popen(WhoisArgs, shell=True, stdout=PIPE)
            for line in response.stdout:
                line = line.decode("utf-8")
                if ':' in line and "# " not in line and len(line) > 2:
                    line = line.strip()
                    content = (f'{content}\n{line}')
                if email == '':
                    if line.startswith('RAbuseEmail:'):
                        try:
                            email = (line.split(': ')[1].lstrip())
                        except:pass    
                    elif line.lower().startswith('abuse-mailbox:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabuseemail:'):email = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('Orgtechemail:'):email = (line.split(': ')[1].lstrip())
                
                if phone == '':
                    if line.lower().startswith('rabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('orgabusephone:'):phone = (line.split(': ')[1].lstrip())
                    elif line.lower().startswith('phone:'):phone = (line.split(': ')[1].lstrip())
                    phone = phone.replace("+", "")
                if line.lower().startswith('rtechname:'):fullname = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('person:'):fullname = (line.split(': ')[1].lstrip())                
                
                if line.lower().startswith('country:'):country = (line.split(': ')[1].lstrip())
                if line.lower().startswith('city:'):city = (line.split(': ')[1].lstrip())
                if line.lower().startswith('address:'):fulladdress = ('%s %s' %(fulladdress, line.split(': ')[1].lstrip()))
                if line.lower().startswith('stateprov:'):state = (line.split(': ')[1].lstrip())
                if line.lower().startswith('postalcode:'):zipcode = (line.split(': ')[1].lstrip())
                if line.lower().startswith('orgname:'):entity = (line.split(': ')[1].lstrip())
                elif line.lower().startswith('org-name:'):entity = (line.split(': ')[1].lstrip())

            print(f'"whois {dnsdomain}	   {email}	{phone}')

            row_data["query"] = query
            row_data["ranking"] = '7 - whois'
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
         
            row_data["titleurl"] = titleurl            
            row_data["city"] = city            
            row_data["country"] = country            
            row_data["note"] = note            
            row_data["state"] = state            
            row_data["SEX"] = SEX            
            row_data["zipcode"] = zipcode            
            row_data["dnsdomain"] = dnsdomain            
            row_data["titleurl"] = titleurl            
            row_data["pagestatus"] = pagestatus            


            data.append(row_data) 


        else:
            print(f'{dnsdomain} not an edu net or edu site?')

            row_data["query"] = query
            row_data["ranking"] = '7 - whois'
            row_data["fullname"] = fullname
            row_data["firstname"] = firstname
            row_data["lastname"] = lastname
            row_data["url"] = url
         
            row_data["titleurl"] = titleurl            
            row_data["city"] = city            
            row_data["country"] = country            
            row_data["note"] = note            
            row_data["state"] = state            
            row_data["SEX"] = SEX            
            row_data["zipcode"] = zipcode            
            row_data["dnsdomain"] = dnsdomain            
            row_data["titleurl"] = titleurl            
            row_data["pagestatus"] = pagestatus            


            data.append(row_data) 

        time.sleep(7)

def zabasearch():# testPhone= 
    print(f'\n\t<<<<< zabasearch phone numbers >>>>>')

    # https://www.zabasearch.com/phone/9175776069/
    for phone in phones:
        row_data = {}
        (query, ranking) = (phone, '9 - zabasearch')
        (country, city, state, zipcode, case, note) = ('', '', '', '', '', '')
        (content, referer, osurl, titleurl, pagestatus)  = ('', '', '', '', '')
        (query) = (phone)
 
        url = (f'https://www.zabasearch.com/phone/{phone}/')
        
        # (content, referer, osurl, titleurl, pagestatus) = request(url) 

        if 1==1:
            row_data["query"] = query
            row_data["ranking"] = ranking
            row_data["url"] = url
            row_data["phone"] = phone            
            row_data["note"] = note            
            row_data["state"] = state    
            row_data["city"] = city  
            # row_data["content"] = content
            row_data["pagestatus"] = pagestatus            
            row_data["referer"] = referer   
            row_data["osurl"] = osurl   
            row_data["titleurl"] = titleurl   
            data.append(row_data)

if __name__ == '__main__':
    main()

# <<<<<<<<<<<<<<<<<<<<<<<<<<Revision History >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
3.0.3 - playwright scraper
3.0.2 - slack, roblox, ham_radio
3.0.1 - Convert area code to state if it doesn't exist, instantusername
3.0.0 - switched to openpyxl, added log sheet
2.8.7 - fixed regex_phone, skip internet check if it's a virtual machine
2.8.6 - made the .py and .exe version dummy proof. Just double click and it runs
2.7.6 - internet checker, removed -I and -O requirement
2.8.0 - kik
"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<Future Wishlist  >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
geospy.ai $$
https://socialeye.net/  $$
hackcheck.io
deflock.me
intelx.io
leakix.net
fofa.so

https://magabook.com/Mike
https://gamejolt.com/site-api/web/profile/@menutkart
https://www.digitalfootprintcheck.com/free-checker.html?q=kevinrose@gmail.com
https://www.kickstarter.com/profile/kevinrose


https://www.deviantart.com/kevinrose/gallery

implement convert_timestamp()


https://start.me/p/0Pqbdg/osint-500-tools   # reviewed

https://start.me/p/1kJKR9/commandergirl-s-suggestions



# phone : whatsapp, group me, weibo, chime, crickwick,
walkie talkie, apple,okru, 


https://www.textnow.com/
https://www.talkatone.com/


add timestamp to log sheet
currently only reads input.txt. add input.xlsx input.
python identityhunt.py  -E -i -p -U -I Intel_test.xlsx  # works
fix -c convert

populate log sheet with todays date


NAM = last name(comma) first name (space) Middle initial


add ID and photo , phone2 column

tkinter purely gui interface
.replace isn't working in several modules
instagramtwo()
create a new identity_hunt with xlsx and requests instead of urllib2


https://opengovus.com/search?q=kevinrose%2C+LLC
https://api.allen-live.in/api/v1/user/identities/{email}
https://public-api.wordpress.com/rest/v1.1/users/{email}/auth-options
https://seguro.marca.com/ueregistro/v2/usuarios/comprobacion/{email}/2
https://{user}.blogspot.com/
https://{user}.livejournal.com
https://500px.com/{user}
https://7dach.ru/profile/{user}
https://admireme.vip/{user}/
https://albicla.com/{user}
https://albicla.com/{user}/post/1
https://ameblo.jp/{user}
https://anonup.com/@{user}
https://apclips.com/{user}
https://apex.tracker.gg/apex/profile/origin/{user}/overview
https://api.500px.com/graphql?query=query%28%24username%3AString%21%29%7BuserByUsername%28username%3A%24username%29%7Bid%20legacyId%20username%20displayName%20firstName%20lastName%20registeredAt%20userProfile%7Bfirstname%20lastname%20about%20country%20city%20state%7DsocialMedia%7Bwebsite%20twitter%20facebook%20instagram%7D%7D%7D&variables=%7B%22username%22%3A%22{user}%22%7D
https://api.boosty.to/v1/blog/{user}
https://api.cropty.io/v1/auth/{user}
https://api.dailymotion.com/user/{user}?fields=id,username,screenname,description,avatar_720_url,cover_250_url,followers_total,following_total,videos_total,country,created_time,verified,url
https://api.destream.net/siteapi/v2/live/details/{user}
https://api.discogs.com/users/{user}
https://api.fotka.com/v2/user/dataStatic?login={user}
https://api.mojang.com/minecraft/profile/lookup/name/{user} # kevinrose
https://api.niftygateway.com/user/profile-and-offchain-nifties-by-url/?profile_url={user} # kevinrose
https://api.omg.lol/address/{user}/info
https://api.picsart.com/users/show/{user}.json # kevinrose
https://api.scratch.mit.edu/users/{user}
https://api.sports-tracker.com/apiserver/v1/user/name/{user}
https://api.stats.fm/api/v1/users/{user}
https://api.tracker.gg/api/v2/apex/standard/profile/origin/{user}
https://api.warframe.market/v2/user/{user}
https://api.zhihu.com/books/people/{user}/publications?offset=0&limit=5 # kevinrose
https://apiv2.fansly.com/api/v1/account?usernames={user}
https://asciinema.org/~{user}
https://atcoder.jp/api/users/exists/?userScreenName={user}
https://atcoder.jp/users/{user}
https://audiojungle.net/user/{user} # kevinrose
https://bandcamp.com/{user} # kevinrose
https://bentbox.co/{user}
https://bit.ly/{user}
https://bitbucket.org/{user}/
https://boosty.to/{user}
https://boot.dev/u/{user}
https://bsky.app/profile/{user}.bsky.social
https://calendly.com/api/booking/profiles/{user} # kevinrose
https://client.warpcast.com/v2/user-by-username?username={user} # kevinrose
https://codeberg.org/{user}
https://codeberg.org/api/v1/users/{user}
https://codeforces.com/api/user.info?handles={user}
https://codeforces.com/profile/{user}
https://coderlegion.com/user/{user}
https://coderwall.com/{user}
https://coderwall.com/{user}.json
https://crates.io/api/v1/users/{user}
https://crates.io/users/{user}
https://cyber.harvard.edu/people/{user}
https://daily.dev/{user} # kevinrose
https://destream.net/live/{user}
https://dev.to/{user}
https://dev.to/api/users/by_username?url={user}
https://discourse.jupyter.org/u/{user}
https://discourse.jupyter.org/u/{user}.json
https://discourse.mozilla.org/u/{user}
https://discourse.mozilla.org/u/{user}.json
https://discuss.hashicorp.com/u/{user}
https://discuss.hashicorp.com/u/{user}.json
https://discuss.kotlinlang.org/u/{user}
https://discuss.kotlinlang.org/u/{user}.json
https://discuss.python.org/u/{user}
https://discuss.python.org/u/{user}.json
https://donatello.to/{user}
https://dribbble.com/{user} # kevinrose
https://dribbble.com/{user}/about
https://en.liberapay.com/{user}
https://en.wikipedia.org/wiki/User:{user}   # kevinrose
https://fansly.com/{user} # kevinrose
https://forum.arduino.cc/u/{user}
https://forum.arduino.cc/u/{user}.json
https://forum.elixirforum.com/u/{user}
https://forum.elixirforum.com/u/{user}.json
https://forum.f-droid.org/u/{user}
https://forum.f-droid.org/u/{user}.json
https://forum.ghost.org/u/{user}
https://forum.ghost.org/u/{user}.json
https://forums.opera.com/api/user/{user}
https://forums.opera.com/user/{user}
https://fotka.com/profil/{user}
https://foursquare.com/{user}
https://freesound.org/people/{user}/
https://gitea.com/{user}
https://gitea.com/api/v1/users/{user}
https://gitee.com/{user} # kevinrose
https://gitee.com/api/v5/users/{user} # kevinrose
https://gitlab.com/{user} # kevinrose
https://gitlab.com/api/v4/users?username={user} # kevinrose
https://gpodder.net/user/{user}/
https://habr.com/ru/users/{user}/
https://hamaha.net/{user}
https://hamaha.net/{user}/tab:info
https://hashnode.com/@{user}
https://hub.docker.com/u/{user}
https://hub.docker.com/v2/users/{user}/
https://ifttt.com/p/{user} # kevinrose
https://independent.academia.edu/{user}
https://issuu.com/{user}
https://issuu.com/query?format=json&_=3210224608766&profileUsername={user}&action=issuu.user.get_anonymous
https://itch.io/profile/{user}
https://launchpad.net/~{user} # kevinrose
https://learn.microsoft.com/api/profiles/{user}
https://learn.microsoft.com/en-us/users/{user}/
https://leetcode.com/u/{user}/
https://lemmy.world/api/v3/user?username={user}
https://lemmy.world/u/{user}
https://lichess.org/@/{user} # kevinrose
https://lichess.org/api/user/{user} # kevinrose
https://linktr.ee/{user}
https://meta.discourse.org/u/{user}
https://meta.discourse.org/u/{user}.json
https://mix.com/{user}
https://music.yandex.ru/handlers/library.jsx?owner={user}
https://music.yandex.ru/users/{user}
https://namemc.com/profile/{user}
https://naturalnews.com/author/{user}/
https://news.ycombinator.com/user?id={user}
https://niftygateway.com/profile/{user}
https://odysee.com/@{user}
https://ok.ru/{user}
https://omg.lol/{user}
https://osu.ppy.sh/users/{user}
https://packagist.org/users/{user}/
https://paragraph.com/@{user}
https://paragraph.com/api/blogs/@{user}
https://pastebin.com/u/{user}
https://play.google.com/store/apps/developer?id={user}
https://pr0gramm.com/api/profile/info?name={user}
https://pypi.org/user/{user}
https://soundcloud.com/{user} # kevinrose
https://sourceforge.net/u/{user}/
https://stats.fm/{user}
https://t.me/{user}
https://themeforest.net/user/{user}
https://trello.com/{user} # kevinrose
https://trello.com/1/Members/{user} # kevinrose
https://ubuntu-mate.community/u/{user}
https://ubuntu-mate.community/u/{user}.json
https://uk.advfn.com/forum/profile/{user}
https://users.rust-lang.org/u/{user}
https://users.rust-lang.org/u/{user}.json
https://vimeo.com/api/v2/{user}/info.json
https://virgool.io/@{user}
https://vk.com/{user} # kevinrose
https://warframe.market/profile/{user}
https://warpcast.com/{user}
https://wiki.archlinux.org/title/User:{user}
https://www.7cups.com/@{user}
https://www.adultism.com/profile/{user}
https://www.adultism.com/profile/{user}/friends
https://www.airliners.net/user/{user}/profile
https://www.allthelyrics.com/forum/member.php?username={user}
https://www.allthelyrics.com/forum/members/{user}.html
https://www.americanthinker.com/author/{user}/
https://www.babepedia.com/user/{user}
https://www.bandlab.com/{user} # kevinrose
https://www.bandlab.com/api/v1.3/users/{user} # kevinrose
https://www.bdsmsingles.com/members/{user}/
https://www.beatstars.com/{user} # kevinrose
https://www.behance.net/{user}/appreciated
https://www.bitchute.com/channel/{user}/
https://www.boot.dev/u/{user}
https://www.defensivecarry.com/members/?username={user}
https://www.discogs.com/user/{user} # kevinrose
https://www.donationalerts.com/api/v1/user/{user}/donationpagesettings
https://www.donationalerts.com/r/{user}
https://www.etoro.com/api/logininfo/v1.1/users/{user} # kevinrose
https://www.etoro.com/people/{user} # kevinrose
https://www.figma.com/@{user}
https://www.freepik.com/author/{user}
https://www.instructables.com/member/{user}/
https://www.kaggle.com/{user} # kevinrose
https://www.last.fm/user/{user}
https://www.minds.com/{user} # kevinrose
https://www.minds.com/api/v3/register/validate?username={user} # kevinrose
https://www.newamerica.org/our-people/{user}/
https://www.npmjs.com/~{user}
https://www.openstreetmap.org/user/{user}
https://www.producthunt.com/@{user}
https://www.speedrun.com/api/v1/users/{user}
https://www.speedrun.com/users/{user}
https://www.sports-tracker.com/view_profile/{user}
https://www.thefirearmsforum.com/members/?username={user}
https://www.threads.net/api/v1/users/web_profile_info/?username={user} # kevinrose
https://www.vinted.pt/member/general/search?search_text={user} # kevinrose
https://www.weforum.org/people/{user}
https://www.zhihu.com/people/{user}
https://www.zomato.com/{user}/reviews
https://zmarsa.com/uzytkownik/{user}


https://github.com/kaifcodec/user-scanner

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<     notes            >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
if 1- main, then write a report of your findings.
change the order of columns by modifying headers in the write module

Protected by cookies: dailymotion, trello
"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<     The End        >>>>>>>>>>>>>>>>>>>>>>>>>>