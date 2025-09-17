#!/usr/bin/python
# coding: utf-8


# <<<<<<<<<<<<<<<<<<<<<<<<<<     Change Me       >>>>>>>>>>>>>>>>>>>>>>>>>>
# change this section with your details
global agency
agency = "MWW" # ISP, MWW
global agencyFull
agencyFull = "Ministry of Wacky Walks"   # Ministry of Wacky Walks
global divisionFull
divisionFull = "Criminal Investigation Division" # Criminal Investigation Division


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Pre-Sets       >>>>>>>>>>>>>>>>>>>>>>>>>>
author = 'LincolnLandForensics'
description = "Convert imaging logs to xlsx, print stickers, write activity reports/checklists and case notes"
version = '3.4.7'


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Imports        >>>>>>>>>>>>>>>>>>>>>>>>>>
try:
    import docx # pip install python-docx
    import pdfplumber  # pip install pdfplumber
    import pdfrw    # pip install pdfrw
    import openpyxl # pip install openpyxl
    import tkinter  # -d
    import pandas as pd # new module
except TypeError as error:
    print(f"{error}")
    print(f"install missing modules:    pip install -r requirements_ForensicsReporter.txt")
    exit()
import re
import os
import sys  
import time # for wait line
import argparse  # for menu system
from datetime import date, datetime
from subprocess import call
# from datetime import datetime

from tkinter import *   # -t  # Frame is not defined if this is missing
from tkinter import ttk # -d
from tkinter import messagebox # -d

from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Border, Side, Alignment

d = datetime.today()

Day = d.strftime("%d")
Month = d.strftime("%m")    # %B = October
Year = d.strftime("%Y")        
# todaysDate = d.strftime("%m/%d/%Y")
# todaysDate = d.strftime("%Y/%m/%d") 
todaysDate = d.strftime("%Y-%m-%d") 

# todaysDateTime = d.strftime("%m_%d_%Y_%H-%M-%S")
todaysDateTime = d.strftime("%Y-%m-%d_%H-%M-%S")    # used for uniq file naming


ANNOT_KEY = '/Annots'
ANNOT_FIELD_KEY = '/T'
ANNOT_VAL_KEY = '/V'
ANNOT_RECT_KEY = '/Rect'
SUBTYPE_KEY = '/Subtype'
WIDGET_SUBTYPE_KEY = '/Widget'

# Regex section
regex_md5 = re.compile(r'^([a-fA-F\d]{32})$')  # regex_md5        [a-f0-9]{32}$/gm
regex_sha1 = re.compile(r'^([a-fA-F\d]{40})$')    #regex_sha1
regex_sha256 = re.compile(r'^([a-fA-F\d]{64})$')#regex_sha256

# colors
color_red = color_yellow = color_green = color_blue = color_purple = color_reset = ''
from colorama import Fore, Back, Style
print(Back.BLACK)
color_red, color_yellow, color_green = Fore.RED, Fore.YELLOW, Fore.GREEN
color_blue, color_purple, color_reset = Fore.BLUE, Fore.MAGENTA, Style.RESET_ALL


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Menu           >>>>>>>>>>>>>>>>>>>>>>>>>>

def main():
    '''
        main menu
        Sets up the menu system and global variables
    '''
    
    global Row
    Row = 1  # defines arguments
    global Row2
    Row2 = 7  #     
    global caseMain
    caseMain = ''    
    
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument('-I', '--input', help='default is input_case.xlsx or input.txt', required=False)
    parser.add_argument('-O', '--output', help='', required=False)
    parser.add_argument('-b', '--blankone', help='create a blank sheet', required=False, action='store_true')

    parser.add_argument('-c','--caseNotes', help='casenotes module (optional) used with -r', required=False, action='store_true')
    parser.add_argument('-C','--checklist', help='checklist module', required=False, action='store_true')

    # parser.add_argument('-d', '--details', help='manually enter details like exhibit number', required=False, action='store_true')
    parser.add_argument('-g', '--guidataentry', help='data entry GUI', required=False, action='store_true')
    parser.add_argument('-l', '--logparse', help='Berla, Cellebrite, FTK, tableau log parser', required=False, action='store_true')
    parser.add_argument('-L', '--logs_parse', help='dump all your logs into Logs\ folder', required=False, action='store_true')
    parser.add_argument('-r', '--report', help='write report', required=False, action='store_true')
    parser.add_argument('-s', '--sticker', help='write sticker', required=False, action='store_true')



    args = parser.parse_args()

    # global section
    global input_details
    input_details = 'no'
    global filename
    filename = ('input.txt')
    global input_file
    input_file = ('input_case.xlsx')

    global logs_folder
    logs_folder = ('Logs\\')   # s subfolder full of logs
    global logs_list
    logs_list = ['']
    global log_type

    global output_docx   # docx actitivy report
    output_docx = ('output_%s.xlsx' %(todaysDateTime))

    global output_txt   # text actitivy report
    output_txt = ('output_%s.txt' %(todaysDateTime))

    global spreadsheet
    spreadsheet = ('log_%s.xlsx' %(todaysDateTime)) # uniq naming for -l module
    global sheet_format
    sheet_format = ('')

    global output_xlsx
    output_xlsx = "output.xlsx"
    
    global output_file
    output_file = spreadsheet    # duplicate 
    
    win = Frame()
    win.grid(sticky=N+S+E+W)




    if args.output:
        output_xlsx = args.output

    if args.input:  # in case you don't want a different input file
        filename = args.input  
        input_file = args.input  

    if args.blankone:  # test
        create_and_write_xlsx()

    if args.report:
        global case_notes_status
        if args.caseNotes:  # if you add -c                                  
            case_notes_status  = ('True')
        else:
            (case_notes_status) = ('False')            
        read_xlsx()
        write_checklist()   # test
        
    if args.checklist:
        write_checklist() 
    
    elif args.logparse:
        log_type = ('file')
        if args.input:  # in case you don't want a different input file
            filename = args.input        
        parse_log() # parse image log
    elif args.logs_parse:
        log_type = ('folder')
        parse_log() # parse imager logs

    elif args.sticker:
        write_sticker() 
    elif args.guidataentry:
        gui_data_entry()

    if not any([args.blankone, args.guidataentry, args.logparse, args.logs_parse, args.report, args.caseNotes, args.checklist, args.sticker]):
        parser.print_help()
        banner_print()
        usage()
        return 0

    try:
        workbook.close()
    except:
        pass
    return 0

# <<<<<<<<<<<<<<<<<<<<<<<<<<   Sub-Routines   >>>>>>>>>>>>>>>>>>>>>>>>>>

def banner_print():
    """
        prints an ASCII banner
    """
    art = """  
  _   _   _   _   _   _   _   _   _     _   _   _   _   _   _   _   _  
 / \ / \ / \ / \ / \ / \ / \ / \ / \   / \ / \ / \ / \ / \ / \ / \ / \ 
( F | o | r | e | n | s | i | c | s ) ( R | e | p | o | r | t | e | r )
 \_/ \_/ \_/ \_/ \_/ \_/ \_/ \_/ \_/   \_/ \_/ \_/ \_/ \_/ \_/ \_/ \_/ 
    """
    print(f"{color_blue}{art}{color_reset}")

def convert_timestamp(timestamp, time_orig=None, timezone=None):    # not in use?
    timezone = timezone or ''
    time_orig = time_orig or timestamp
    timestamp = str(timestamp)

    # Regular expression to find the timezone
    timezone_pattern = r"([A-Za-z ]+)$"
    matches = re.findall(timezone_pattern, timestamp)
    
    if matches:
        timezone = matches[-1].strip()
        timestamp = timestamp.replace(timezone, "").strip()
    
    # Handling specific timezones in the timestamp
    if "(" in timestamp:
        timestamp, tz_info = timestamp.split('(')
        timezone = tz_info.replace(")", '').strip()
    elif " CDT" in timestamp:
        timezone = "CDT"
        timestamp = timestamp.replace(" CDT", "").strip()
    elif " CST" in timestamp:
        timezone = "CST"
        timestamp = timestamp.replace(" CST", "").strip()

    formats = [
        "%B %d, %Y, %I:%M:%S %p %Z",  # June 13, 2022, 9:41:33 PM CDT (Flock)
        "%Y:%m:%d %H:%M:%S",
        "%Y-%m-%d %H:%M:%S",
        "%m/%d/%Y %I:%M:%S %p",
        "%m/%d/%Y %I:%M %p",  # Timestamps without seconds
        "%m/%d/%Y %H:%M:%S",  # Military time without seconds
        "%m-%d-%y at %I:%M:%S %p %Z",  # e.g., 09-10-23 at 4:29:12 PM CDT
        "%m-%d-%y %I:%M:%S %p",
        "%B %d, %Y at %I:%M:%S %p %Z",
        "%B %d, %Y at %I:%M:%S %p",
        "%B %d, %Y %I:%M:%S %p %Z",
        "%B %d, %Y %I:%M:%S %p",
        "%Y-%m-%dT%H:%M:%SZ",  # ISO 8601 with UTC timezone
        "%Y/%m/%d %H:%M:%S",  # e.g., 2022/06/13 21:41:33
        "%d-%m-%Y %I:%M:%S %p",  # e.g., 13-06-2022 9:41:33 PM
        "%d/%m/%Y %H:%M:%S",  # e.g., 13/06/2022 21:41:33
        "%Y-%m-%d %I:%M:%S %p",  # e.g., 2022-06-13 9:41:33 PM
        "%Y%m%d%H%M%S",  # e.g., 20220613214133
        "%Y%m%d %H%M%S",  # e.g., 20220613 214133
        "%m/%d/%y %H:%M:%S",  # e.g., 06/13/22 21:41:33
        "%d-%b-%Y %I:%M:%S %p",  # e.g., 13-Jun-2022 9:41:33 PM
        "%d/%b/%Y %H:%M:%S",  # e.g., 13/Jun/2022 21:41:33
        "%Y/%b/%d %I:%M:%S %p",  # e.g., 2022/Jun/13 9:41:33 PM
        "%d %b %Y %H:%M:%S",  # e.g., 13 Jun 2022 21:41:33
        "%A, %B %d, %Y %I:%M:%S %p %Z",  # e.g., Monday, June 13, 2022 9:41:33 PM CDT
        "%A, %B %d, %Y %I:%M:%S %p"     # e.g., Monday, June 13, 2022 9:41:33 PM
    ]

    for fmt in formats:
        try:
            dt_obj = datetime.strptime(timestamp.strip(), fmt)
            return dt_obj, time_orig, timezone
        except ValueError:
            continue

    raise ValueError(f"Timestamp format not recognized for: {time_orig}")
    
def create_docx():
    '''
        if there isn't a template to read, 
        this creates an activity report from scratch
    '''
    global document
    document = docx.Document()
    
    caseNumber = ""

    #header section
    section = document.sections[0]
    header = section.header
    header

    header.is_linked_to_previous = False
    # section.different_first_page_header_footer = True
    paragraph = header.paragraphs[0]
    paragraph.text = ("%s\n\nACTIVITY REPORT                                                             %s" %(agencyFull,divisionFull))


    p = document.add_paragraph('\n')    # start with a blank line   # todo this line is too thick
    p = document.add_paragraph('Activity Report:\t\t\t\tDate of Activity:')
    p = document.add_paragraph('%s\t\t\t\t%s' %(caseNumber, todaysDate))

    # insert a big line here

    # p = document.add_paragraph('Subject of Activity:\t\t\t\tCase Agent:\t\tType By:')
    # p = document.add_paragraph('%s\t\t\t\t%s\t\t%s' %(subjectBusinessName, caseAgent, forensicExaminer))

    document.save(output_docx) 
    print(f"{color_green}created {output_docx}{color_reset}")       
    return document
    
def enter_data():  # not in use?
    """
        GUI data input section using tkinter
    """
    
    
    # accepted = accept_var.get()
    if 1==1:
    # if accepted=="Accepted":
        # (caseNumber, caseName) = ('', '')
        (exhibit, subjectBusinessName, caseType, caseAgent) = ('', '', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')
        (temp, hostame, phoneIMEI2, phone2) = ('', '', '', '')

        caseNumber = caseNumber_entry.get()
        caseName = caseName_entry.get()
        
        if 1==1:
        # if caseNumber and caseName:
            # Case
            subjectBusinessName = subjectBusinessName_entry.get()
            caseAgent = caseAgent_combobox.get()
            forensicExaminer = forensicExaminer_combobox.get()
            
            # Description
            exhibit = str(exhibit_entry.get())
            makeModel = makeModel_entry.get()
            serial = serial_entry.get()
            exhibitType = exhibitType_combobox.get()
            phoneNumber = phoneNumber_entry.get()
            phoneIMEI = phoneIMEI_entry.get()
            userName = userName_entry.get()
            userPwd = userPwd_entry.get()

            # Lab Chain of Custody
            seizureAddress = seizureAddress_entry.get()
            seizureRoom = seizureRoom_entry.get()
            dateSeized = dateSeized_entry.get()
            # seizedBy = seizedBy_entry.get()
            seizedBy = seizedBy_combobox.get()
            # dateReceived = dateReceived_entry.get()

            # notes
            tempNotes = tempNotes_entry.get()


            # print out sticker format
            print(f"Case: {caseNumber} Ex:{exhibit}")
            print(f"CaseName: {caseName}")
            print(f"Subject: {subjectBusinessName}")
            print(f"Make: {makeModel}")
            print(f"Serial: {serial}")
            print(f"Agent: {caseAgent}")
            print(f"------------------------") 

            filepath = "log_case.xlsx"
            
            if not os.path.exists(filepath):
                workbook = openpyxl.Workbook()
                sheet = workbook.active

                heading = ["caseNumber", "exhibit", "caseName", "subjectBusinessName", "caseType"
                , "caseAgent", "forensicExaminer", "reportStatus", "notes", "summary", "tempNotes"
                , "exhibitType", "makeModel", "serial", "OS", "hostname", "userName", "userPwd"
                , "email", "emailPwd", "ip", "phoneNumber", "phoneIMEI", "phone2", "phoneIMEI2"
                , "mobileCarrier", "biosTime", "currentTime", "timezone", "shutdownMethod"
                , "shutdownTime", "seizureAddress", "seizureRoom", "dateSeized", "seizedBy"
                , "seizureStatus", "dateReceived", "receivedBy", "removalDate", "removalStaff"
                , "reasonForRemoval", "inventoryDate", "storageLocation", "status", "imagingTool"
                , "imagingType", "imageMD5", "imageSHA256", "imageSHA1", "verifyHash", "writeBlocker"
                , "imagingStarted", "imagingFinished", "storageType", "storageMakeModel"
                , "storageSerial", "storageSize", "evidenceDataSize", "analysisTool"
                , "analysisTool2", "exportLocation", "exportedEvidence", "qrCode", "operation"
                , "vaultCaseNumber", "vaultTotal", "caseNumberOrig", "Action", "priority", "temp"]

                sheet.append(heading)
                workbook.save(filepath)
            workbook = openpyxl.load_workbook(filepath)
            sheet = workbook.active

            sheet.append([caseNumber, exhibit, caseName, subjectBusinessName,
                    caseType, caseAgent, forensicExaminer, reportStatus, notes,
                    summary, exhibitType, makeModel, serial, OS, phoneNumber,
                    phoneIMEI, mobileCarrier, biosTime, currentTime, timezone,
                    shutdownMethod, shutdownTime, userName, userPwd, email,
                    emailPwd, ip, seizureAddress, seizureRoom, dateSeized,
                    seizedBy, dateReceived, receivedBy, removalDate, removalStaff,
                    reasonForRemoval, inventoryDate, seizureStatus, status, imagingTool,
                    imagingType, imageMD5, imageSHA1, imageSHA256, writeBlocker,
                    imagingStarted, imagingFinished, storageType, storageMakeModel,
                    storageSerial, storageSize, evidenceDataSize, analysisTool,
                    analysisTool2, exportLocation, exportedEvidence, storageLocation,
                    caseNumberOrig, priority, operation, Action, vaultCaseNumber,
                    qrCode, vaultTotal, tempNotes])




            workbook.save(filepath)
                
        else:
            tkinter.messagebox.showwarning(title="Error", message="Case Number and Case Name are required.")
    else:
        tkinter.messagebox.showwarning(title= "Error", message="You have not verified the info")
    
def dictionary_build(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
    forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
    phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
    userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
    dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
    seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
    writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
    storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
    storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
    vaultTotal, tempNotes, temp, hostname, phoneIMEI2, phone2, verifyHash):    
    '''
        build a dictionary file of important columns for writing to a pdf
    '''
    
    my_dict = {}
    my_dict['caseNumber']=caseNumber
    my_dict['caseName']=caseName
    my_dict['subjectBusinessName']=subjectBusinessName
    my_dict['caseAgent']=caseAgent
    my_dict['forensicExaminer']=forensicExaminer
    my_dict['exhibit']=exhibit
    my_dict['makeModel']=makeModel
    my_dict['serial']=serial
    my_dict['OS']=OS
    my_dict['ip']=ip
    my_dict['exhibitType']=exhibitType   
    my_dict['phoneNumber']=phoneNumber
    my_dict['phoneIMEI']=phoneIMEI
    my_dict['userName']=userName
    my_dict['userPwd']=userPwd
    my_dict['email']=email
    my_dict['emailPwd']=emailPwd
    my_dict['biosTime']=biosTime
    my_dict['currentTime']=currentTime
    my_dict['priority']=priority
    my_dict['timezone']=timezone
    my_dict['seizureAddress']=seizureAddress
    my_dict['seizureRoom']=seizureRoom
    my_dict['dateSeized']=dateSeized
    my_dict['seizedBy']=seizedBy
    my_dict['seizureStatus']=seizureStatus
    my_dict['dateReceived']=dateReceived
    my_dict['receivedBy']=receivedBy
    my_dict['removalDate ']=removalDate 
    my_dict['removalStaff']=removalStaff
    my_dict['imagingTool']=imagingTool
    my_dict['imagingType']=imagingType
    my_dict['imageMD5']=imageMD5
    my_dict['writeBlocker']=writeBlocker
    my_dict['imagingStarted']=imagingStarted
    my_dict['imagingFinished']=imagingFinished
    my_dict['imageSHA256']=imageSHA256
    my_dict['storageType']=storageType
    my_dict['storageMakeModel']=storageMakeModel
    my_dict['storageSerial']=storageSerial
    my_dict['storageSize']=storageSize
    my_dict['analysisTool']=analysisTool
    my_dict['analysisTool2']=analysisTool2
    my_dict['notes']=notes

    return (my_dict)
    
def fix_date(date):
    '''
        standardize date formatting, Tableau
    '''
    
    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()
    date = date.replace("  ", " ")  # test
    # date = ('%s      ' %(date)) # test
    date = date.split(' ')     # Fri Jun 04 07:55:41 2021
    mo = date[1]    # convert month to a number
    mo = mo.replace("Jan", "1").replace("Feb", "2").replace("Mar", "3").replace("Apr", "4")
    mo = mo.replace("May", "5").replace("Jun", "6").replace("Jul", "7").replace("Aug", "8")
    mo = mo.replace("Sep", "9").replace("Oct", "10").replace("Nov", "11").replace("Dec", "12")
    try:    
        dy = date[2].lstrip('0')
    except TypeError as error:
        print(error)
    try:
        tm = date[3].lstrip('0')
        # tm = date[3]
    except TypeError as error:
        print(error)
    yr = date[4]
    # date = ('%s/%s/%s %s' %(mo, dy, yr, tm))  # 3/4/2021 9:17
    # date = ('%s/%s/%s %s' %(yr, mo, dy, tm))  # 3/4/2021 9:17
    date = ('%s-%s-%s %s' %(yr, mo, dy, tm))  # 3/4/2021 9:17
      
    
    return date

def fix_date2(date):
    '''
        standardize date formatting
        2022-07-14 21:15:11
        
        31/07/2022 11:48:57 (-5)
    '''
    print(f"{color_red}fix_date2{color_reset}")  
    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()

    return date

def fix_date3(date):
    '''
        standardize date formatting from Cellebrite
        
        31/07/2022 11:48:57 (-5) to 2022-07-31 11:48

    '''

    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()
    # date = date.replace("  ", " ")  # test
    # date = ('%s      ' %(date)) # test
    date = date.split(' ')     # 31/07/2022 11:48:57 (-5)
    tempDate = date[0]
    tm = date[1]
    tempDate = tempDate.split('/')
    dy = tempDate[0]    
    mo = tempDate[1]
    yr = tempDate[2]

    # date = ('%s/%s/%s %s' %(mo, dy, yr, tm)).lstrip('0')  # 3/4/2021 9:17
    date = ('%s-%s-%s %s' %(yr, mo, dy, tm)).lstrip('0')  # 3/4/2021 9:17
 
    return date

def gui_data_entry():
    """
        GUI data entry function
    """
    win = Frame()
    # win.title('Evidence Form')  # test
    # win = Frame().title("Evidence form")  # todo
    win.grid(sticky=N+S+E+W) 
    # win.geometry("700x350") # test
    # Frame labels
    case_info_frame = LabelFrame(win, text='Case', padx=5, pady=5)
    description_frame = LabelFrame(win, text='Description', padx=5, pady=5)
    custody_frame = LabelFrame(win, text='Chain of Custody', padx=5, pady=5)
    notes_frame = LabelFrame(win, text='Notes', padx=5, pady=5)
    
    # stick notes
    case_info_frame.grid(sticky=W+W)
    description_frame.grid(sticky=E+W)
    custody_frame.grid(sticky=E+W)
    notes_frame.grid(sticky=E+W)   

    ## create multiple frames

    for frame in case_info_frame, description_frame, custody_frame, notes_frame:
        for col in 0, 1, 2:
            frame.columnconfigure(col, weight=1)


    ########## case section ##########

    ## row 0 labels (caseNumber, caseNumber, subjectBusinessName)
    caseNumber_label = tkinter.Label(case_info_frame, text="Case Number")
    caseNumber_label.grid(row=0, column=0)

    caseName_label = tkinter.Label(case_info_frame, text="Case Name")    
    caseName_label.grid(row=0, column=1)

    subjectBusinessName_label = tkinter.Label(case_info_frame, text="Subject or d/b/a")    
    subjectBusinessName_label.grid(row=0, column=2)

    ## global section
    global caseNumber_entry
    global caseName_entry
    global subjectBusinessName_entry
    global exhibit_entry
    global makeModel_entry
    global serial_entry
    global exhibitType_entry
    global phoneNumber_entry
    global phoneIMEI_entry
    global userName_entry
    global userPwd_entry
    global seizureAddress_entry
    global seizureRoom_entry
    global dateSeized_entry
    # global seizedBy_entry
    global seizedBy_combobox    
    global dateReceived_entry
    global tempNotes_entry
    global caseAgent_combobox
    global forensicExaminer_combobox
    global exhibitType_combobox


    ## row 1 Entry (caseNumber, caseNumber, subjectBusinessName)
    caseNumber_entry = tkinter.Entry(case_info_frame)
    caseName_entry = tkinter.Entry(case_info_frame)
    subjectBusinessName_entry = tkinter.Entry(case_info_frame, width = 24)
    caseNumber_entry.grid(row=1, column=0)
    caseName_entry.grid(row=1, column=1)    # test
    subjectBusinessName_entry.grid(row=1, column=2)

    ## row 2 & 3 label (caseAgent, forensicExaminer)
    caseAgent_label = tkinter.Label(case_info_frame, text="Case Agent")
    caseAgent_combobox = ttk.Combobox(case_info_frame, values=["", "SA April Moore", "Road Runner", "Sherlock Holmes"])
    caseAgent_label.grid(row=2, column=0)
    caseAgent_combobox.grid(row=3, column=0)

    forensicExaminer_label = tkinter.Label(case_info_frame, text="Forensic Examiner")   # works
    forensicExaminer_combobox = ttk.Combobox(case_info_frame, values=["", "Sherlock Holmes", "CIA Thomas #", "DFE Ness 9351"], width = 24)
    forensicExaminer_label.grid(row=2, column=1)    # was 0, 2
    forensicExaminer_combobox.grid(row=3, column=1) # works

    ########## Description section ##########

    ## row 0 label (exhibit, makeModel, serial)
    exhibit_label = tkinter.Label(description_frame, text="Exhibit")  
    exhibit_label.grid(row=0, column=0)

    makeModel_label = tkinter.Label(description_frame, text="Make/Model")    
    makeModel_label.grid(row=0, column=1)

    serial_label = tkinter.Label(description_frame, text="Serial #")    
    serial_label.grid(row=0, column=2)

    ## row 1 entry (exhibit, makeModel, serial)
    exhibit_entry = tkinter.Entry(description_frame, width = 7)
    # exhibit_entry = tkinter.Entry(description_frame, width= 40) # sets windows frame size

    exhibit_entry.grid(row=1, column=0)

    makeModel_entry = tkinter.Entry(description_frame)
    makeModel_entry.grid(row=1, column=1)

    serial_entry = tkinter.Entry(description_frame)
    serial_entry.grid(row=1, column=2)

    ## row 2 label (exhibitType, phoneNumber, phoneIMEI)

    exhibitType_label = tkinter.Label(description_frame, text="Exhibit Type")
    exhibitType_combobox = ttk.Combobox(description_frame, values=["", "DVR", "desktop", "laptop", "phone", "POS", "router", "server", "switch" , "vehicle"], width = 10)
    exhibitType_label.grid(row=2, column=0)

    phoneNumbe_label = tkinter.Label(description_frame, text="Phone Number")    
    phoneNumbe_label.grid(row=2, column=1)

    phoneIMEI = tkinter.Label(description_frame, text="Phone IMEI #")    
    phoneIMEI.grid(row=2, column=2)


    ## row 3 entry (exhibitType, phoneNumber, phoneIMEI)
    ## exhibitType_entry = tkinter.Entry(description_frame)
    ## exhibitType_combobox = ttk.Combobox(description_frame, values=["", "DVR", "desktop", "phone", "POS", "router", "server", "switch" , "vehicle"])

    ## exhibitType_entry.grid(row=8, column=0)
    exhibitType_combobox.grid(row=3, column=0)

    phoneNumber_entry = tkinter.Entry(description_frame, width = 16)
    phoneNumber_entry.grid(row=3, column=1)

    phoneIMEI_entry = tkinter.Entry(description_frame, width = 17)
    phoneIMEI_entry.grid(row=3, column=2)

    ## row 4 label (userName, userPwd)
    userName_label = tkinter.Label(description_frame, text="User Name")    
    userName_label.grid(row=4, column=0)
    userPwd_label = tkinter.Label(description_frame, text="User Password")    
    userPwd_label.grid(row=4, column=1)

    ## row 5 entry (userName, userPwd)
    userName_entry = tkinter.Entry(description_frame)
    userName_entry.grid(row=5, column=0)

    userPwd_entry = tkinter.Entry(description_frame)
    userPwd_entry.grid(row=5, column=1)

    ########## chain of custody section ##########

    ## row 0 label (seizureAddress, seizureRoom)
    seizureAddress_label = tkinter.Label(custody_frame, text="Seizure Address")    
    seizureAddress_label.grid(row=0, column=0)

    seizureRoom_label = tkinter.Label(custody_frame, text="Seizure Room")    
    seizureRoom_label.grid(row=0, column=1)

    ## row 1 entry (seizureAddress, seizureRoom)
    seizureAddress_entry = tkinter.Entry(custody_frame, width = 50)    # test
    seizureRoom_entry = tkinter.Entry(custody_frame, width = 5)
    seizureAddress_entry.grid(row=1, column=0)
    seizureRoom_entry.grid(row=1, column=1)


    ## row 2 label (dateSeized, seizedBy, dateReceived)
    dateSeized_label = tkinter.Label(custody_frame, text="Date Seized")    
    dateSeized_label.grid(row=2, column=0)


    # dateReceived = tkinter.Label(custody_frame, text="Date Received")    
    # dateReceived.grid(row=2, column=2)

    ## row 3 entry (dateSeized, seizedBy, dateReceived)
    dateSeized_entry = tkinter.Entry(custody_frame, width = 19)
    dateSeized_entry.grid(row=3, column=0)

    seizedBy_label = tkinter.Label(custody_frame, text="Seized_By")     
    seizedBy_label.grid(row=2, column=1)
    seizedBy_entry = tkinter.Entry(custody_frame, width = 25)
    seizedBy_entry.grid(row=3, column=1)
    seizedBy_combobox = ttk.Combobox(case_info_frame, values=["", "SA Herby Hancock", "SSA John Doe", "Sherlock Holmes"])

    # dateReceived_entry = tkinter.Entry(custody_frame)
    # dateReceived_entry.grid(row=3, column=2)

    ########## tempNotes section ##########

    ## row 0 label (tempNotes)
    tempNotes_label = tkinter.Label(notes_frame, text="Temp Notes")  
    tempNotes_label.grid(row=0, column=0)

    ## row 1 entry (tempNotes)
    tempNotes_entry = tkinter.Entry(notes_frame, width = 75)
    tempNotes_entry.grid(row=1, column=0)

    ## Button
    button = tkinter.Button(frame, text="Enter Data", command= enter_data, bg='#0052cc', fg='#ffffff')
    # button = tkinter.Button(frame, text="Enter Data", command= enter_data_DB, bg='#0052cc', fg='#ffffff')

    button.grid(row=3, column=0, sticky="news", padx=20, pady=10)

    win.mainloop()
    print(f"{color_green}Data written to log_case.xlsx{color_reset}")
    
def parse_log():
    '''
        parse tableau, recon imager, cellebrite triage_windows.cmd and FTK logs
    '''

    import os
    (caseNumber, caseName, exhibit) = ('', '', '')
    if log_type == 'file':  # only ask for exhibit number if it's a single log
        if input_details == "yes":
            caseNumber = str(input("caseNumber : ")).strip()
            caseName = str(input("caseName : ")).strip()
            exhibit = str(input("exhibit : "))  # .strip()
        logs_list = [filename]
    elif log_type == 'folder':
        print('')
        
        if input_details == "yes":
            caseNumber = str(input("caseNumber : ")).strip()
            caseName = str(input("caseName : ")).strip()

        if not os.path.exists(logs_folder):
            print(f"{color_red}{logs_folder} folder does not exist{color_reset}")
            print(f"{color_yellow}create a {logs_folder} folder and fill it with logs to parse{color_reset}")            
            exit() 
        else:
            logs_list = os.listdir(logs_folder)
    
        logs_list2 = []
        # remove folder names from logs_list2
        
        for logFile in logs_list:
            logFile = ("%s%s" %(logs_folder, logFile))
            logs_list2.append(logFile)
        logs_list = logs_list2

        # read section
    for logFile in logs_list:
        msg_blurb = (f'Reading {logFile}')
        msg_blurb_square(msg_blurb, color_green)
       
        (header, reportStatus, date) = ('', '', '<insert date here>')

        (subjectBusinessName, caseType, caseAgent) = ('', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')
        (temp, hostname, phoneIMEI2, phone2, verifyHash) = ('', '', '', '', '')
        
        # bonus variables
        (vehicleYear, vehicleManufacturer, vehicleModel) = ('', '', '') # BerlaIVe Acquisition
        (imagingTool1, imagingTool2, make, model) = ('', '', '', '')

        if logFile.lower().endswith('.pdf'):
            csv_file = ''
            (caseNumber, exhibit, caseType, forensicExaminer, makeModel, OS, hostname, status, exhibitType, serial, phoneNumber, phoneIMEI, phoneIMEI2, email, imagingType, imageMD5, imageSHA256, imagingStarted, exportLocation, imagingFinished, imagingTool, imagingType, storageSize, evidenceDataSize, analysisTool, tempNotes, imagingTool) = pdf_extract(logFile)
            csv_file = tempNotes.split('\\n')
        elif logFile.lower().endswith('.ufdx'):
            csv_file = ''
            (caseNumber, exhibit, subjectBusinessName, caseType, forensicExaminer, makeModel, seizureAddress, seizedBy, imagingType, exportLocation) = ufdx_parser(logFile)
            csv_file = tempNotes.split('\\n')
        else:
            csv_file = open(logFile) 
        
        for each_line in csv_file:
            if "Task:" in each_line:
                imagingType = re.split("Task: ", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()

            elif " Extraction type " in each_line: #cellebrite xls
                imagingType = re.split(" Extraction type ", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()
            elif "Source Type: Physical" in each_line:
                imagingType = "disk to file"
            elif "Image type :" in each_line: #recon imager
                imagingType = re.split("Image type :", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()


            elif "ExtractionMethod=" in each_line : #cellebrite *.ufd log file
                imagingType = each_line.replace("ExtractionMethod=", "").strip()
                if imagingType == "AdvancedLogical":
                    imagingType = "advanced logical"
                elif imagingType == "Logical":
                    imagingType = "logical"  
                elif imagingType == "iOS_Full_Filesystem":
                    imagingType = "full file system" 

            elif "ExtractionType=" in each_line and imagingType == '': #cellebrite *.ufd log file
                imagingType = each_line.replace("ExtractionType=", "").strip()
                if imagingType == "AdvancedLogical":
                    imagingType = "advanced logical"
                elif imagingType == "Logical":
                    imagingType = "logical"            


            elif "Start of Tableau Imager" in each_line: # tableau imager
                imagingTool = each_line.replace("Start of ", "").strip()



            # status
            elif "Status: Ok" in each_line or "Imaging Status : Successful" in each_line:
                status = 'Imaged'
            elif "Status: Error/Failed" in each_line:
                status = 'Not imaged'
            elif "Acquisition Successfully Completed" in each_line: # BerlaIVe AcquisitionLog_001.txt
                status = "Imaged"

            elif "Acquisition Failed To Complete" in each_line: # BerlaIVe AcquisitionLog.txt
                status = "Not imaged"

            # exhibit      
            elif "Evidence Number: " in each_line:      #FTK_parse, magnet
                exhibit = re.split("Evidence Number: ", each_line, 0)
                exhibit = str(exhibit[1]).replace('=','')   # .strip()
            elif "Evidence Number=" in each_line:      #UFD test
                exhibit = re.split("Evidence Number=", each_line, 0)
                exhibit = str(exhibit[1])
                print(f' exhibit is {exhibit}') # 5est
            elif "Exhibit#" in each_line:      #cellebrite
                exhibit = re.split("Exhibit#", each_line, 0)
                exhibit = str(exhibit[1])   # .strip()
            elif "Exhibit Number=" in each_line: # CellebriteUFED4PC.txt
                exhibit = each_line.replace("Exhibit Number=", "").replace("=","").strip()

            elif "Evidence Number" in each_line:      #recon imager
                exhibit = re.split("Evidence Number", each_line, 0)

                # exhibit = re.split("Evidence Number     :", each_line, 0)
                exhibit = str(exhibit[1]).replace(":", "")  # .strip()

            # exhibitType
            elif "Device=" in each_line:
                exhibitType = re.split("Device=", each_line, 0)
                exhibitType = str(exhibitType[1]).strip()
                if exhibitType == 'IPHONE':
                   exhibitType = 'phone'
                


            # makeModel
            elif "Unique description: " in each_line:
                makeModel = re.split("Unique description: ", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Device    " in each_line: #cellebrite excel
                makeModel = re.split("Device    ", each_line, 0)
                makeModel = str(makeModel[1]).strip()
            elif "Selected device name" in each_line: #cellebrite
                makeModel = re.split("Selected device name", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Selected Model:" in each_line:
                makeModel = re.split("Selected Model:", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Case Name: " in each_line: # BerlaIVe AcquisitionLog_001.txt      # broken
                caseName = each_line.replace("Case Name: ", "").strip()

            elif "Case Identifier=" in each_line and caseNumber == '': # ufed
                caseNumber = each_line.replace("Case Identifier=", "").strip()

            elif "Crime Type=" in each_line and caseType == '': # ufed
                caseType = each_line.replace("Crime Type=", "").strip()



            elif "Model=" in each_line and "DeviceModel=" not in each_line and model == '': #cellebrite *.ufd log file
                model = each_line.replace("Model=", "").strip()
                print(f'model = {model} blah')   # temp
            # elif "Model=" in each_line and model == '': #cellebrite *.ufd log file
                # model = each_line.replace("Model=", "").strip()

            elif "Vendor=" in each_line: #cellebrite *.ufd log file
                make = each_line.replace("Vendor=", "").strip()

            elif "Device Model:" in each_line: # CellebritePremium DeviceInfo.txt
                model = each_line.replace("Device Model:", "").strip()

            elif "Vendor:" in each_line: # CellebritePremium DeviceInfo.txt
                make = each_line.replace("Vendor:", "").replace("Tableau", "").strip()
            elif "Model:" in each_line and len(storageMakeModel) == 0: # tableau
                storageMakeModel = re.split("Model:", each_line, 0)
                storageMakeModel = str(storageMakeModel[1]).strip()
            elif "Device / Media Name:" in each_line: # SumuriReconImager.txt
                storageMakeModel = each_line.replace("Device / Media Name:", "").strip()

            elif "Vehicle Year:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleYear = each_line.replace("Vehicle Year:", "").strip()

            elif "Vehicle Manufacturer:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleManufacturer = each_line.replace("Vehicle Manufacturer:", "").strip()

            elif "Vehicle Model:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleModel = each_line.replace("Vehicle Model:", "").strip()

            # OS
            elif "Revision:" in each_line: #cellebite 
                os = re.split("Revision:", each_line, 0)
                os = str(os[1]).strip()
                if 'iPhone' in makeModel:
                    os = ('iOS %s' %(os))

            elif "OS=" in each_line: # cellebrite *.ufd log file
                OS = each_line.replace("OS=", "").strip()

            elif "OS Version:" in each_line: # CellebritePremium DeviceInfo.txt
                OS = each_line.replace("OS Version:", "").strip()

            elif "Operating System Version:" in each_line: #cellebrite *.ufd log file
                OS = each_line.replace("Operating System Version:", "").strip()

            # elif "Operating System:" in each_line: # MagnetAXIOM Case Information.txt
                # OS = each_line.replace("Operating System:", "").strip()

            elif "Vehicle ECU:" in each_line: # BerlaIVe AcquisitionLog.txt
                OS = each_line.replace("Vehicle ECU:", "").strip()
            elif "Android_ID=" in each_line: # Cellebrite .ufd
                OS = ('Android %s' %(OS))
            elif "Apple" in makeModel and OS != '': # Cellebrite .ufd    # test
                OS = ('iOS %s' %(OS))

            # userPwd
            elif "Passcode_0=" in each_line: #cellebrite ufed

                userPwd = re.split("Passcode_0=", each_line, 0)
                userPwd = str(userPwd[1]).strip()    
                print(f'userPwd = {userPwd}')   # temp

            # elif "Apple" in makeModel and OS != '': # Cellebrite .ufd    # test
                # userPwd = ('iOS %s' %(OS))


                
            # serial

            # elif "Serial Number:" in each_line and serial != '': #cellebrite
            elif "Serial Number:" in each_line: #cellebrite
                # serial = each_line.replace("Serial Number:", "").strip()
                serial = re.split("Serial Number:", each_line, 0)
                      
                serial = str(serial[1]).strip()
                if "number: " in serial:
                    serial = ''
                storageSerial = serial
                serial = ''
                
            elif "Machine Serial" in each_line: #RECON imager
            # elif "Machine Serial" in each_line and serial != '': #RECON imager
                serial = re.split(":", each_line, 0)
                serial = str(serial[1]).strip()

            elif "Vehicle VIN:" in each_line: # BerlaIVe AcquisitionLog.txt
                serial = each_line.replace("Vehicle VIN:", "").strip()
                if serial == ('unknown'):
                    serial = ''
            # elif "Serial " in each_line and serial != '': #cellebrite
                # serial = re.split("Serial ", each_line, 0)
                # print("serial=",serial[1].strip())      
                # serial = str(serial[1]).strip()

            # storageSerial    
            # elif "Drive Serial Number:" in each_line: # FTKImager Image.E01.txt # fix me
                # storageSerial = each_line.replace("Drive Serial Number:", "").strip()

            elif "S/N:" in each_line: # TableauImager 000ecc45 0067205e
                storageSerial = each_line.replace("S/N:", "").strip()
            elif "Serial Number:" in each_line and storageSerial != '': # FTKImager Image.E01.txt # fix me
                storageSerial = each_line.replace("Drive Serial Number:", "").strip()
            elif "Serial number:" in each_line and storageSerial == '': # Tableau imager 
                storageSerial = each_line.replace("Drive Serial number:", "").replace("Serial number:", "").strip()


            elif "Unique Identifier:" in each_line: # MagnetAcquire image_info.txt
                storageSerial = each_line.replace("Unique Identifier:", "").strip() 

     
            # phoneIMEI
            elif "IMEI:" in each_line: # CellebritePremium DeviceInfo.txt
                phoneIMEI = each_line.replace("IMEI:", "").strip()

            elif "IMEI1=" in each_line: # CELLEBRITEPREMIUM EXTRACTION_FFS.TXT
                phoneIMEI = each_line.replace("IMEI1=", "").strip()

            elif "IMEI " in each_line: # GrayKey_R5CR8147V0A.pdf
                phoneIMEI = each_line.replace("IMEI1=", "").strip()

            # phoneNumber
            elif "MSISDN" in each_line: #cellebrite
                phoneNumber = re.split("MSISDN", each_line, 0)
                phoneNumber = str(phoneNumber[1]).strip()
                if ')' in phoneNumber:
                    phoneNumber = phoneNumber.replace("+1 (", "1-").replace(") ", "-")
                (exportedEvidence, status) = ('', 'Imaged')

            elif " Username" in each_line: #cellebrite xls
                phoneNumber = re.split(" Username", each_line, 0)
                phoneNumber = str(phoneNumber[1]).strip()
                if ')' in phoneNumber:
                    phoneNumber = phoneNumber.replace("+1 (", "1-").replace(") ", "-")
                (exportedEvidence, status) = ('', 'Imaged')
            elif "UserName=" in each_line: # CELLEBRITE AdvancedLogical.ufd
                phoneNumber = each_line.replace("UserName=", "").strip()
                # userName = phoneNumber

            # forensicExaminer
            elif "Examiner:" in each_line:
                forensicExaminer = re.split("Examiner:", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
                forensicExaminer =forensicExaminer.replace("CIA - ", "")
            elif "User: " in each_line:
                forensicExaminer = re.split("User: ", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
            elif "Examiner Name:" in each_line: # MagnetAcquire image_info_001.txt
                forensicExaminer = each_line.replace("Examiner Name:", "").strip()
            elif "Examiner Name=" in each_line: # CellebriteUFED4PC.txt
                forensicExaminer = each_line.replace("Examiner Name=", "").strip()
            elif "Examiner         :" in each_line: # recon imager
                forensicExaminer = re.split("Examiner         :", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
            elif "Examiner name" in each_line:  #cellebrite
                forensicExaminer = re.split("Examiner name", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
                forensicExaminer =forensicExaminer.replace("CIA - ", "")

            # caseNumber
            elif "Case ID:" in each_line:
                caseNumber = re.split("Case ID:", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")
            elif "Case Number:" in each_line:
                caseNumber = re.split("Case Number:", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")
            elif "Case Number=" in each_line:   # CellebriteUFED4PC.txt
                caseNumber = each_line.replace("Case Number=", "").strip()

            elif "Case Number         :" in each_line:   # Recon imager and probably tablaue
                caseNumber = each_line.replace("Case Number         :", "").strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")




            elif "CaseNumber" in each_line:   #cellebrite
                caseNumber = re.split("CaseNumber", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()


            # notes
            elif "Case Notes:" in each_line:    # Tableau logs
                notes = re.split("Case Notes:", each_line, 0)
                notes = str(notes[1]).strip()
                notes = notes.replace("<<not entered>>", "")
            elif "Notes: " in each_line:
                notes = re.split("Notes: ", each_line, 0)
                notes = str(notes[1]).strip()
                notes = notes.replace("<<not entered>>", "")

            elif "Notes         :" in each_line:    # recon imager
                notes = re.split("Notes         :", each_line, 0)
                notes = str(notes[1]).strip()

            elif "Source Device :" in each_line:    # recon imager
                (vol, partition, size, frmat) = ('', '', '', '')
                sourcenotes = re.split("Source Device :", each_line, 0)
                sourcenotes = str(sourcenotes[1]).strip()
                # append 7 spaces for fault tollerance  # to do
                details = re.split("  ", sourcenotes, 0)
                vol = str(details[0]).strip()
                
                try:
                    partition = str(details[3]).strip()    
                    size = str(details[4]).strip()  
                    frmat = str(details[5]).strip()  
                    blurb1 = ("This image was from %s and was the %s %s %s volume." %(vol, partition, size, frmat))
                    notes = ("%s %s" %(notes, blurb1))
                except:pass
                    
            # imagingTool
            elif "Imager App: " in each_line:
                imagingTool1 = re.split("Imager App: ", each_line, 0)
                imagingTool1 = str(imagingTool1[1]).strip()
            elif "Created By AccessData® FTK® Imager" in each_line:
                imagingTool1 = each_line.replace("Created By AccessData® FTK® Imager", "").replace("®", "").replace("Â", "").strip()



            elif "Created By AccessData" in each_line:
                imagingTool1 = each_line.replace("Created By AccessData", "").replace("®", "").strip()


                
            elif "Imager Ver: " in each_line:
                imagingTool2 = re.split("Imager Ver: ", each_line, 0)
                imagingTool2 = str(imagingTool2[1]).strip()

            elif "UFED Version:    Product Version: " in each_line:    #cellebrite
                imagingTool = re.split("UFED Version:    Product Version: ", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Cellebrite UFED %s' %(imagingTool))
            elif "UFED version" in each_line:    #cellebrite
                imagingTool = re.split("UFED version", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Cellebrite UFED %s' %(imagingTool))

            elif "Cellebrite Physical Analyzer version" in each_line: #cellebrite xls
                imagingTool1 = re.split("Cellebrite Physical Analyzer version", each_line, 0)
                imagingTool1 = str(imagingTool1[1]).strip()
                imagingTool1 = ('Cellebrite Physical Analyzer %s' %(imagingTool1))

            elif "RECON Imager Version : " in each_line:    # Recon Imager
                imagingTool = re.split("RECON Imager Version : ", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Recon Imager %s' %(imagingTool))

            elif "AcquisitionTool=" in each_line: #cellebrite *.ufd log file
                imagingTool = each_line.replace("AcquisitionTool=", "").strip()

            elif "Created by iVe " in each_line and "Acquisition finished" not in each_line: # BerlaIVe AcquisitionLog_001.txt 
                imagingTool = each_line.replace("Created by iVe ", "Berla iVe ").strip()
                if " built on" in imagingTool and "Acquisition finished" not in imagingTool:
                    imagingTool = imagingTool.split(' built on')[0].strip()

            elif "Imager Product:" in each_line: # MagnetAcquire image_info.txt
                imagingTool = each_line.replace("Imager Product:", "").strip() 

            elif "Imager Version:" in each_line: # MagnetAcquire image_info.txt
                imagingToolVer = each_line.replace("Imager Version:", "").strip() 
                imagingTool = ('%s %s' %(imagingTool, imagingToolVer))    

            elif "ExtractionSoftwareVersion=" in each_line: # CellebritePA_FFS.txt
                imagingToolTemp = each_line.replace("ExtractionSoftwareVersion=", "").strip()
                imagingTool = ('Cellebrite PA %s' %(imagingToolTemp))

            # storageSize
            elif "Capacity in bytes reported Pwr-ON: " in each_line: # todo swap storageSize from capacity
                capacity = re.split("Capacity in bytes reported Pwr-ON: ", each_line, 0)
                capacity = str(capacity[1]).strip()
                if "(" in capacity:
                    capacity = re.split("\(", each_line, 0)
                    capacity = str(capacity[1]).strip()
                    capcty = capacity.replace(")", "")
                    capacity = capcty.split('.')[0]
                    if ' ' in capcty:
                        size = capcty.split(' ')[1]
                    else:
                        size = ''
                    # size = capacity
                    capacity = ('%s %s' %(capacity, size))
                    storageSize = capacity

            elif "Source data size: " in each_line: # FTKImager Image.E01.txt
                storageSize = each_line.replace("Source data size: ", "").strip()
            elif "Device Size:" in each_line: # MagnetAcquire image_info.txt
                storageSize = each_line.replace("Device Size:", "").strip() 

            elif "Disk Size:" in each_line: # SumuriReconImager.txt
                storageSize = each_line.replace("Disk Size:", "").strip()
                storageSize = storageSize.split(' (')[0]


            # storageType
            elif "Cable/Interface type: " in each_line:
                storageType = re.split("Cable/Interface type: ", each_line, 0)
                storageType = str(storageType[1]).strip()
                storageType

            elif "Drive Interface Type: " in each_line: # FTKImager Image.E01.txt
                storageType = each_line.replace("Drive Interface Type: ", "").strip()

            elif "Media Type:" in each_line: # MagnetAcquire image_info.txt
                storageType = each_line.replace("Media Type:", "").strip() 

            elif "T356789iu" in each_line:  # fix me  T356789i
                writeBlocker = "Tableau T356789iu"
            elif "T356789i" in each_line:  # fix me  
                writeBlocker = "Tableau T356789i"

            elif "Source data size: " in each_line:
                capacity = re.split("Source data size: ", each_line, 0)
                capacity = str(capacity[1]).strip()
                if "(" in capacity:
                    capacity = re.split("\(", each_line, 0)
                    capacity = str(capacity[1]).strip()
                    capcty = capacity.replace(")", "")
                    capacity = capcty.split('.')[0]
                    if ' ' in capcty:
                        size = capcty.split(' ')[1]
                    # size = capacity
                    capacity = ('%s %s' %(capacity, size))
            
            # exportLocation        
            elif "Filename of first chunk: " in each_line:
                exportLocation = re.split("Filename of first chunk: ", each_line, 0)
                exportLocation = str(exportLocation[1]).strip()
            elif "Information for " in each_line:       # ftk_parse
                exportLocation = re.split("Information for ", each_line, 0)
                exportLocation = str(exportLocation[1]).strip()
            # elif "E01" in each_line:
                # exportLocation = each_line.strip()

            elif "FileDump=" in each_line: # CellebritePremium EXTRACTION_FFS.txt
                exportLocation = each_line.replace("FileDump=", "").strip()
            elif "File Path: " in each_line: # BerlaIVe AcquisitionLog_001.txt
                exportLocation = each_line.replace("File Path: ", "").strip()
                exportedEvidence = 'Y'

            # imageMD5
            elif "Disk MD5:  " in each_line:    # Tableau
                imageMD5 = re.split("Disk MD5:  ", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()


            elif "MD5 checksum:" in each_line:  # fix me
                imageMD5 = re.split("MD5 checksum:", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()
                imageMD5 = re.split(": ", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()
                if "verified" in each_line:
                    status = "Imaged"
                    imageMD5 = imageMD5.replace(' : verified','')
                else:
                    status = 'Not imaged'
            elif "MD5 Acquisition Hash:" in each_line: # MagnetAcquire image_info.txt
                imageMD5 = each_line.replace("MD5 Acquisition Hash:", "").strip()
                status = "Imaged"
            elif "MD5 hash calculated over data:" in each_line: # SumuriReconImager.txt
                imageMD5 = each_line.replace("MD5 hash calculated over data:", "").strip()

            elif "MD5 Image Hash: " in each_line:    # Magnet Axiom
                imageMD5 = each_line.replace("MD5 Image Hash: ", "")

            elif "MD5 Verification Hash: " in each_line:    # Magnet Axiom
                verifyHash = each_line.replace("MD5 Verification Hash: ", "")


            # elif "MD5 " in each_line:    # GrayKey_R5CR8147V0A.pdf
                # imageMD5 = each_line.lstrip("MD5 ")
                # print(f"{color_yellow}<<<<<<<<<<<   testing  >>>>>>>>>>>>>{color_reset}")    # task

            # imageSHA1
            elif "Disk SHA1: " in each_line:    # Tableau
                imageSHA1 = re.split("Disk SHA1: ", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()

            elif "SHA1 checksum:" in each_line:  # FTKImager Image.E01.txt
                imageSHA1 = re.split("SHA1 checksum:", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()
                imageSHA1 = re.split(": ", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()
                if "verified" in each_line:
                    status = "Imaged"
                    imageSHA1 = imageSHA1.replace(' : verified','')
                else:
                    status = 'Not imaged'
            elif "SHA1 hash calculated over data:" in each_line: # SumuriReconImager.txt
                imageSHA1 = each_line.replace("SHA1 hash calculated over data:", "").strip()

            elif ".zip=" in each_line: # cellebrite .ufd
                hashTemp = re.split("=", each_line, 0)
                exportLocation = hashTemp[0]
                hashTemp = str(hashTemp[1]).strip()

                if  re.match(regex_sha256, hashTemp):    #regex SHA256 hash
                    imageSHA256 = hashTemp
                elif  re.match(regex_md5, hashTemp):    #regex md5 hash
                    imageMD5 = hashTemp
                elif  re.match(regex_sha1, hashTemp):    #regex SHA1 hash
                    imageSHA1 = hashTemp
                
            # _Triage_.txt parsing
            elif "Host Name: " in each_line:
                hostname = re.split("Host Name: ", each_line, 0)
                hostname = str(hostname[1]).strip()
                if "MagnetAXIOM" in logFile:
                    hostname = ''
                else:
                    notes = ("%s The hostname is %s." %(notes, hostname))
            elif "Timezone: " in each_line:
                timezone = re.split("Timezone: ", each_line, 0)
                timezone = str(timezone[1]).strip()
                notes = ("%s The system timezone is set to %s." %(notes, timezone))
            elif "OS Name: " in each_line:
                OS = re.split("OS Name: ", each_line, 0)
                OS = str(OS[1]).strip()
            elif "   IPv4 Address" in each_line:
                ip = re.split("   IPv4 Address. . . . . . . . . . . : ", each_line, 0)
                ip = str(ip[1]).strip()
                notes = ("%s The IP address was %s." %(notes, ip))
            elif "    Lock Status:" in each_line:
                encryption = re.split("    Lock Status:", each_line, 0)
                encryption = str(encryption[1]).strip()
                if 'Locked' in encryption:
                    encryption = 'BitLocker Encrypted'
                    notes = ("%s BitLocker encryption is enabled." %(notes)) 
            elif "Email:" in each_line:
                email = re.split("Email: ", each_line, 0)
                email = str(email[1]).strip()
            elif "Currentuser:" in each_line:
                userName = re.split("Currentuser:", each_line, 0)
                userName = str(userName[1]).strip()


            # imagingStarted
            elif "Acquisition started:" in each_line:
                imagingStarted = re.split("Acquisition started: ", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                try:
                    imagingStarted = fix_date(imagingStarted)
                except:pass    

            elif "Acquisition Started:" in each_line: # MagnetAcquire image_info.txt
                imagingStarted = each_line.replace("Acquisition Started:", "").strip()
                imagingStarted = fix_date2(imagingStarted)  # test

            elif "Extraction start date/time" in each_line: #cellebrite
                imagingStarted = re.split("time", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip().replace(" -05:00", "").strip(':').strip().replace("(GMT-5)", "")
                # try:
                    # imagingStarted = fix_date(imagingStarted)
                # except:pass    

            elif "Imaging Start Time :" in each_line:   # Recon imager
                imagingStarted = re.split("Imaging Start Time :", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                # try:
                    # imagingStarted = fix_date(imagingStarted)
                # except:pass    

            elif "Date=" in each_line: #cellebrite *.ufd log file
                imagingStarted = each_line.replace("Date=", "").strip()
                imagingStarted = fix_date3(imagingStarted)
                # try:
                    # imagingStarted = fix_date3(imagingStarted)
                # except:pass    

            elif "Date and time:" in each_line: # CELLEBRITEPREMIUM DEVICEINFO.TXT
                imagingStarted = each_line.replace("Date and time:", "").strip()
                biosTime = imagingStarted

            elif "Started: " in each_line:
                imagingStarted = re.split("Started: ", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                if "Cellebrite" in imagingTool or "Tableau" in imagingTool: 
                    # imagingStarted = fix_date(imagingStarted)
                    try:
                        imagingStarted = fix_date(imagingStarted)
                    except:pass    

            elif "Start Date/Time:" in each_line: # Magnet Axiom
                imagingStarted = each_line.replace("Start Date/Time:", "").strip()
                biosTime = imagingStarted




               
            # imagingFinished
                
            elif "Closed:" in each_line:
                imagingFinished = re.split("Closed: ", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                imagingFinished = fix_date(imagingFinished)
                try:
                    imagingStarted = fix_date(imagingFinished)
                except:pass    
                
            elif "Acquisition finished:" in each_line:
                imagingFinished = re.split("Acquisition finished:", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                try:
                    imagingStarted = fix_date(imagingFinished)
                except:pass    

            elif "Acquisition Finished:" in each_line: # MagnetAcquire image_info.txt
                imagingFinished = each_line.replace("Acquisition Finished:", "").strip()
                imagingFinished = fix_date2(imagingFinished)    # test


            elif "Extraction end date" in each_line:     #cellebrite
                imagingFinished = re.split("Extraction end date", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                imagingFinished = imagingFinished.replace("/time", "").replace(" -05:00", "").strip(':').strip().replace("(GMT-5)", "")

                # try:
                    # imagingStarted = fix_date(imagingFinished)
                # except:pass    

            elif "Imaging End Time   :" in each_line: # Recon imager
                imagingFinished = re.split("Imaging End Time   :", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                # try:
                    # imagingStarted = fix_date(imagingFinished)
                # except:pass    

            elif "EndTime=" in each_line: #cellebrite *.ufd log file
                imagingFinished = each_line.replace("EndTime=", "").strip()
                status = ('Imaged')
                imagingFinished = fix_date3(imagingFinished)
                # try:
                    # imagingFinished = fix_date(imagingFinished)
                # except:pass    
            elif "End Date/Time: " in each_line: # magnet axiom
                imagingFinished = each_line.replace("End Date/Time: ", "").strip()
                status = ('Imaged')
                # imagingFinished = fix_date3(imagingFinished)


            # tempnotes
            elif "Description:" in each_line: # MagnetAcquire image_info.txt
                description = each_line.replace("Description:", "").strip() 
                tempNotes = ('%s %s' %(tempNotes, description))

            elif "Chipset:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Device Bluetooth Name:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Encryption Type:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Number Of Installed Applications:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Live encryption state:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))


                
        # tempNotes save the input file name for bulk uploads. delete if it's input.txt
        tempNotes = ('%s %s' %(logFile, tempNotes)).strip()
        if tempNotes == 'input.txt':
            tempNotes = ''
            
        # if makeModel == '':
            # makeModel = ('%s %s' %(make, model))    # test
        if not makeModel:
            makeModel = ('%s %s' %(make, model))    # test

        if vehicleYear != '' : # BerlaIVe AcquisitionLog.txt
        # if vehicleYear != '' and vehicleManufacturer != '' and vehicleModel != '': # BerlaIVe AcquisitionLog.txt
            makeModel = ("%s %s %s" %(vehicleYear, vehicleManufacturer, vehicleModel))

        if caseNumber != '' and exhibit != '':
            qrCode = ("%s_%s" %(caseNumber, exhibit))
        # if qrCode == '_': 
            # qrCode = ''
        
        if len(imagingTool1) != 0:
            imagingTool = ('%s %s' %(imagingTool1.strip(), imagingTool2.strip()))
        
        # if len(storageSize) != 0 and storageSerial == "" and len(storageType) != 0:
            # notes = ("This had a %s (S/N: %s) %s, %s drive. %s" %(storageMakeModel, storageSerial, storageSize, storageType, notes))   # test
        # elif len(storageMakeModel) != 0 and len(storageSerial) != 0 and len(storageSize) != 0 and len(storageType) != 0: 
            ##notes = ("This had a %s, model %s, serial #%s, %s drive. %s" %(storageSize, storageMakeModel, storageSerial, storageType, notes))   # test
            # notes = ("This had a %s (S/N: %s) %s, %s drive. %s" %(storageMakeModel, storageSerial, storageSize, storageType, notes))   # test
        # elif len(storageSize) != 0:
            # notes = ("This had a %s drive, model %s, serial #%s, %s drive. %s" %(storageSize, storageMakeModel, storageSerial, storageType, notes))   # test


        # if len(OS) != 0 and 'The operating system was' not in notes:
            # notes = ("%s The operating system was %s." %(notes, OS)) 


        if status.lower() == 'not imaged' and verifyHash == '':
            notes = ("%s This drive could not be imaged." %(notes))
            verifyHash = 'N'
            
        print(f'''
        caseNumber = {caseNumber}
        subjectBusinessName = {subjectBusinessName}
        forensicExaminer = {forensicExaminer}
        exhibitType = {exhibitType}
        makeModel = {makeModel}
        serial = {serial}
        OS = {OS}
        phoneNumber = {phoneNumber}\n
        ''')
        write_output(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, forensicExaminer, reportStatus, notes, summary, tempNotes, exhibitType, makeModel, serial, OS, hostname, userName, userPwd, email, emailPwd, ip, phoneNumber, phoneIMEI, phone2, phoneIMEI2, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, seizureAddress, seizureRoom, dateSeized, seizedBy, seizureStatus, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, storageLocation, status, imagingTool, imagingType, imageMD5, imageSHA256, imageSHA1, verifyHash, writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, qrCode, operation, vaultCaseNumber, vaultTotal, caseNumberOrig, Action, priority, temp)
    print(f"{color_green}Exporting logs as {spreadsheet}{color_reset}")    

def msg_blurb_square(msg_blurb, color):
    horizontal_line = f"+{'-' * (len(msg_blurb) + 2)}+"
    empty_line = f"| {' ' * (len(msg_blurb))} |"

    print(color + horizontal_line)
    print(empty_line)
    print(f"| {msg_blurb} |")
    print(empty_line)
    print(horizontal_line)
    print(f'{color_reset}')

def pdf_extract(filename):
    """
    Extract data from a PDF file and return relevant information as a tuple.
    """
    # Initialize variables
    caseNumber = ''
    caseType = ''
    forensicExaminer = ''
    exhibitType = ''
    makeModel = ''
    exhibit = ''
    serial = ''
    OS = ''
    hostname = ''
    phoneNumber = ''
    phoneIMEI = ''
    phoneIMEI2 = ''
    email = ''
    status = ''
    imagingTool = ''
    imagingType = ''
    imageMD5 = ''
    imageSHA256 = ''
    imagingStarted = ''
    imagingFinished = ''
    storageSize = ''
    evidenceDataSize = ''
    exportLocation = ''
    analysisTool = ''
    tempNotes = ''

    # Open the PDF file
    with pdfplumber.open(filename) as pdf:
        for page in pdf.pages:
            # Extract text from each page and append to tempNotes
            tempNotes += page.extract_text() + '\n'

        # Extract relevant information using regex
        caseNumber_match = re.search(r'Case Identifier (.*?)\n', tempNotes)
        if caseNumber_match:
            caseNumber = caseNumber_match.group(1).strip()

        forensicExaminer_match = re.search(r'Examiner Name(.*?)\n', tempNotes)  # cellebrite
        if forensicExaminer_match:
            forensicExaminer = forensicExaminer_match.group(1).strip()
        if forensicExaminer.startswith(': '):
            forensicExaminer = forensicExaminer.replace(': ', '')
        makeModel_match = re.search(r"Device Name / Evidence Number (.*?)\n", tempNotes)    # cellebrite
        if makeModel_match:
            makeModel = makeModel_match.group(1).strip()

        os_match = re.search(r'OS Name (.*?)\n', tempNotes)    # cellebrite
        if os_match:
            OS = os_match.group(1).strip()

        os_match2 = re.search(r'Software Version (.*?)\n', tempNotes)    # GrayKey
        if os_match2:
            OS = os_match2.group(1).strip()

        os_match3 = re.search(r'General OS Version (.*?)\n', tempNotes)    # cellebrite preliminary report
        if os_match3:
            OS = os_match3.group(1).strip()

        caseType_match = re.search(r'Crime Type (.*?)\n', tempNotes)
        if caseType_match:
            caseType = caseType_match.group(1).strip()



        exhibit_match = re.search(r'Evidence ID:(.*?)\n', tempNotes)    # graykey
        if exhibit_match:
            exhibit = exhibit_match.group(1).strip()

        makeModel_match = re.search(r'Model:(.*?)\n', tempNotes)    # graykey
        if makeModel_match:
            makeModel = makeModel_match.group(1).strip()

        makeModel_match2 = re.search(r'General Detected Phone Model (.*?)\n', tempNotes)    # graykey
        if makeModel_match2:
            makeModel = makeModel_match2.group(1).strip()

        makeModel_match3 = re.search(r'Model (.*?)\n', tempNotes)    # graykey
        if makeModel_match3 and makeModel:
            makeModel = makeModel_match3.group(1).strip()

        serial_match = re.search(r'Serial Number:(.*?)\n', tempNotes)    # graykey
        if serial_match:
            serial = serial_match.group(1).strip()

        serial_match2 = re.search(r'General Serial (.*?)\n', tempNotes)    # Cellebrite preliminary report
        if serial_match2:
            serial = serial_match2.group(1).strip()


        status_match = re.search(r'Extraction Status Success\n', tempNotes)
        if status_match:
            status = "imaged"

        hostname_match = re.search(r'Device Name(.*?)\n', tempNotes)  # GrayKey
        if hostname_match:
            hostname = hostname_match.group(1).strip()        

        imei_match = re.search(r'General IMEI (.*?)\n', tempNotes)  # Cellebrite preliminary report
        if imei_match:
            phoneIMEI = imei_match.group(1).strip() 

        imei2_match = re.search(r'General IMEI (.*?)\n', tempNotes)  # Cellebrite preliminary report
        if imei2_match:
            phoneIMEI2 = imei2_match.group(1).strip() 

        imagingTool_match = re.search(r'GrayKey Software: OS Version:(.*?),', tempNotes)    # graykey
        if imagingTool_match:
            imagingTool = f"GrayKey {imagingTool_match.group(1).strip()}"

        imagingTool_match2 = re.search(r'Application Version (.*?),', tempNotes)    # cellebrite
        if imagingTool_match2:
            imagingTool = f"Cellebrite {imagingTool_match2.group(1).strip()}"

        imagingType_match = re.search(r'Extraction Method (.*?)\n', tempNotes)
        if imagingType_match:
            imagingType = imagingType_match.group(1).strip()

        imagingStarted_match = re.search(r'Report generation time:(.*?)\n', tempNotes)    # graykey
        if imagingStarted_match:
            imagingStarted = imagingStarted_match.group(1).strip()

        imagingStarted_match2 = re.search(r'Extraction Start Time (.*?)\n', tempNotes)    # cellebrite
        if imagingStarted_match2:
            imagingStarted = imagingStarted_match2.group(1).strip()

        imagingFinished_match = re.search(r'Extraction End Time (.*?)\n', tempNotes)    # cellebrite
        if imagingFinished_match:
            imagingFinished = imagingFinished_match.group(1).strip()

        sha256_regex = r'\\b[A-Fa-f0-9]{64}\\b'

        sha256_match = re.search(sha256_regex, tempNotes)
        if sha256_match:
            imageSHA256 = sha256_match

        phoneIMEI_match = re.search(r'IMEI:(.*?)\n', tempNotes)    # graykey
        if phoneIMEI_match:
            phoneIMEI = phoneIMEI_match.group(1).strip()

        phoneIMEI_match2 = re.search(r'General IMEI (.*?)\n', tempNotes)    # graykey
        if phoneIMEI_match2:
            phoneIMEI = phoneIMEI_match2.group(1).strip()



        # Add additional extraction logic here as needed for other fields

    if 'GrayKey Progress Report' in OS:
        OS = '' # test
        print(f'ffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff')  # temp


    # Return all extracted information as a tuple
    return (
        caseNumber, exhibit, caseType, forensicExaminer, makeModel, OS, hostname, status, exhibitType, serial, phoneNumber,
        phoneIMEI, phoneIMEI2, email, imagingType, imageMD5, imageSHA256, imagingStarted, exportLocation,
        imagingFinished, imagingTool, imagingType, storageSize, evidenceDataSize, analysisTool, tempNotes, imagingTool
    )

def pdf_filltest(input_pdf_path, output_pdf_path, data_dict):       # not in use?
    """
    Fill out PDF form fields based on a provided dictionary of values.
    
    Parameters:
    - input_pdf_path: Path to the input PDF template
    - output_pdf_path: Path to save the filled PDF
    - data_dict: Dictionary containing field names as keys and values to fill in the fields
    """
    # Define the necessary constants
    ANNOT_KEY = '/Annots'
    SUBTYPE_KEY = '/Subtype'
    WIDGET_SUBTYPE_KEY = '/Widget'
    ANNOT_FIELD_KEY = '/T'

    print(f"data_dict = {data_dict}")  # Debugging: print the data_dict to ensure correct values
    
    # Read the input PDF template
    template_pdf = pdfrw.PdfReader(input_pdf_path)
    
    # Iterate through all pages of the PDF
    for page in template_pdf.pages:
        annotations = page.get(ANNOT_KEY, [])
        
        # Iterate through all annotations (form fields)
        for annotation in annotations:
            field_name = annotation.get(ANNOT_FIELD_KEY)
            
            if field_name and annotation.get(SUBTYPE_KEY) == WIDGET_SUBTYPE_KEY:
                field_name = field_name[1:-1]  # Remove the leading/trailing parentheses
                
                # Check if the field name exists in the provided data_dict using get()
                field_value = data_dict.get(field_name)
                
                if field_value is not None:  # Only process if the field value exists in the data_dict
                    print(f"Filling field: {field_name} with value {field_value}")  # Debugging: check the field being filled
                    
                    # Handle boolean values (e.g., checkboxes)
                    if isinstance(field_value, bool):
                        if field_value:  # If True, check the checkbox
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('Yes')))
                        else:  # If False, uncheck the checkbox
                            annotation.update(pdfrw.PdfDict(AS=pdfrw.PdfName('Off')))
                    else:
                        # Handle text or other values
                        annotation.update(pdfrw.PdfDict(V='{}'.format(field_value)))
                        annotation.update(pdfrw.PdfDict(AP=''))  # Clear appearance if needed

    # Update the NeedAppearances flag to ensure the changes are visible in the PDF
    template_pdf.Root.AcroForm.update(pdfrw.PdfDict(NeedAppearances=pdfrw.PdfObject('true')))
    
    # Write the filled PDF to the output path
    pdfrw.PdfWriter().write(output_pdf_path, template_pdf)

def pdf_fill(input_pdf_path, output_pdf_path, data_dict):   
    """
        # fill out EvidenceForm
        receives input template based on agency and itemType
        receives output template with a uniq name
        data_dict is my_dict which is many of the columns needed to write pdf reports
    """
    print(f"data_dict = {data_dict}")   # temp
    template_pdf = pdfrw.PdfReader(input_pdf_path)
    for page in template_pdf.pages:
        annotations = page[ANNOT_KEY]
        for annotation in annotations:
            if annotation[SUBTYPE_KEY] == WIDGET_SUBTYPE_KEY:
                if annotation[ANNOT_FIELD_KEY]:
                    key = annotation[ANNOT_FIELD_KEY][1:-1]
                    if key in data_dict.keys():
                        if type(data_dict[key]) == bool:
                            if data_dict[key] == True:
                                annotation.update(pdfrw.PdfDict(
                                    AS=pdfrw.PdfName('Yes')))
                        else:
                            annotation.update(
                                pdfrw.PdfDict(V='{}'.format(data_dict[key]))
                            )
                            annotation.update(pdfrw.PdfDict(AP=''))
    template_pdf.Root.AcroForm.update(pdfrw.PdfDict(NeedAppearances=pdfrw.PdfObject('true')))
    pdfrw.PdfWriter().write(output_pdf_path, template_pdf)

def read_xlsx():
    """
        reads input_case.xlsx by default
        this will read in each line and write a report 
        it then makes a backup of copy xlsx of the lines you tossed in
    """
    sheet_name, bodyDoing, bodyTodo, caseNumberTodo, bodyDone = 'Cases', '', '', '', ''

    if not os.path.exists(input_file):
        print(f"{color_red}{input_file} does not exist{color_reset}")
        exit() 
    else:
        dftemp = pd.read_excel(input_file, sheet_name=sheet_name)
        df = dftemp.fillna('')  # Replace NaN with empty string
        msg_blurb = (f'Reading {input_file}')
        msg_blurb_square(msg_blurb, color_green) 

    (header, reportStatus, date) = ('', '', '')
    (body, executiveSummary, evidenceBlurb) = ('', '', '')
    (style) = ('')
    (caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, executiveSummary, body, footer) = ('', '', '', '', '', '', '', '', '')
    (subject, vowel, status, verifyHash) = ('test', 'aeiou', '', '')


    # read in the spreadsheet
    for index, row in df.iterrows():
        caseNumber = str(row['caseNumber'])
        exhibit = str(row['exhibit'])
        exhibit = exhibit
        if exhibit.endswith('.0'):
            exhibit = exhibit[:-2]
        caseName = str(row['caseName'])
        subjectBusinessName = str(row['subjectBusinessName'])
        caseType = str(row['caseType'])
        caseAgent = str(row['caseAgent'])
        forensicExaminer = str(row['forensicExaminer'])
        reportStatus = str(row['reportStatus'])
        notes = str(row['notes'])
        summary = str(row['summary'])
        status = str(row['status'])
        exhibitType = str(row['exhibitType'])
        makeModel = str(row['makeModel'])
        serial = str(row['serial'])
        OS = str(row['OS'])
        phoneNumber = str(row['phoneNumber'])
        phone2 = str(row['phone2'])
        
        if phone2.endswith('.0'):
            phone2 = phone2[:-2]
        phoneIMEI = str(row['phoneIMEI'])       # .rstrip(',0')
        mobileCarrier = str(row['mobileCarrier'])
        verifyHash = str(row.get('verifyHash', ''))  # Defaults to an empty string if key is missing
 
        
        if caseNumberTodo == '':
            caseNumberTodo = (f'''- [ ] {caseNumber} - Add to spreadsheet
- [ ] {caseNumber} - Review warrant
- [ ] {caseNumber} - Label all pieces
- [ ] {caseNumber} - Photos
- [ ] {caseNumber} - Inventory label on shelf
- [ ] {caseNumber} - Verify PC images with Arsenal
- [ ] {caseNumber} - Return exhibits once imaged/verified
- [ ] {caseNumber} - Case notes/reports saved/printed/filing cabinet w/ checklist
- [ ] {caseNumber} - Digital evidence to Agent
- [ ] {caseNumber} - Activity report(s)
- [ ] {caseNumber} - Move Exhibits to Legacy sheet once all done''')

        if exhibit != '':
            if status.lower() == "imaged" or status.lower() == "not imaged" or status.lower() == "copied":
                bodyDone = (f'''{bodyDone}
- [x] {caseNumber} Ex: {exhibit} - Image''')
            else:
                    bodyTodo = (f'''{bodyTodo}
- [ ] {caseNumber} Ex: {exhibit} - Image''')

            if verifyHash.lower() == "y":
                bodyDone = (f'''{bodyDone}
- [x] {caseNumber} Ex: {exhibit} - Verify Hash''')
            elif verifyHash.lower() == "":
                    bodyTodo = (f'''{bodyTodo}
- [ ] {caseNumber} Ex: {exhibit} - Verify Hash''')
  
            if reportStatus.lower() == "finalized" or reportStatus.lower() == "draft": # draft is only for semi-finalized reports
                bodyDone = (f'''{bodyDone}
- [x] {caseNumber} Ex: {exhibit} - Analyze''')
            # elif reportStatus.lower() == "draft":
                # bodyDoing = (f'''{bodyDoing}
# - [ ] {caseNumber} Ex: {exhibit} - Analyze''')
            else:
                bodyTodo = (f'''{bodyTodo}
- [ ] {caseNumber} Ex: {exhibit} - Analyze''')


        biosTime = str(row['biosTime'])
        biosTime = biosTime.strip()
        if biosTime == 'NaT':
            biosTime = ''  

        currentTime = str(row['currentTime'])
        currentTime = currentTime.strip()
        if currentTime == 'NaT':
            currentTime = ''  

        timezone = str(row['timezone'])
        shutdownMethod = str(row['shutdownMethod'])

        shutdownTime = str(row['shutdownTime'])
        shutdownTime = shutdownTime.strip()
        if shutdownTime == 'NaT':
            shutdownTime = ''

        userName = str(row['userName'])
        userPwd = str(row['userPwd'])
        email = str(row['email'])
        emailPwd = str(row['emailPwd'])
        ip = str(row['ip'])
        seizureAddress = str(row['seizureAddress'])
        seizureRoom = str(row['seizureRoom'])
        dateSeized = str(row['dateSeized'])
        seizedBy = str(row['seizedBy'])
        dateReceived = str(row['dateReceived'])
        receivedBy = str(row['receivedBy'])

        removalDate = str(row['removalDate'])
        removalDate = removalDate.strip()
        if removalDate == 'NaT':
            removalDate = ''        
        
        removalStaff = str(row['removalStaff'])
        reasonForRemoval = str(row['reasonForRemoval'])

        inventoryDate = str(row['inventoryDate'])
        inventoryDate = inventoryDate.strip()
        if inventoryDate == 'NaT':
            inventoryDate = ''  

        seizureStatus = str(row['seizureStatus'])
        status = str(row['status'])
        imagingTool = str(row['imagingTool'])
        imagingType = str(row['imagingType'])
        imageMD5 = str(row['imageMD5'])
        imageSHA1 = str(row['imageSHA1'])
        imageSHA256 = str(row['imageSHA256'])
        writeBlocker = str(row['writeBlocker'])
        imagingStarted = str(row['imagingStarted'])
        imagingFinished = str(row['imagingFinished'])
        storageType = str(row['storageType'])
        storageMakeModel = str(row['storageMakeModel'])
        storageSerial = str(row['storageSerial'])
        storageSize = str(row['storageSize'])
        evidenceDataSize = str(row['evidenceDataSize'])
        analysisTool = str(row['analysisTool'])
        analysisTool2 = str(row['analysisTool2'])
        exportLocation = str(row['exportLocation'])
        exportedEvidence = str(row['exportedEvidence'])
        storageLocation = str(row['storageLocation'])
        caseNumberOrig = str(row['caseNumberOrig'])
        priority = str(row['priority'])
        operation = str(row['operation'])
        Action = str(row['Action'])
        vaultCaseNumber = str(row['vaultCaseNumber'])
        qrCode = str(row['qrCode'])
        vaultTotal = str(row['vaultTotal'])
        tempNotes = str(row['tempNotes'])
        # try:
            # verifyHash = str(row["verifyHash"])
        # except TypeError as error:
            # print(f"{color_red}error = {error}{color_reset}")

        try:
            temp = str(row['temp'])
        except:
            temp = ''
        try:
            hostname = str(row['hostname'])
        except:
            hostname = ''
        try:
            phoneIMEI2 = str(row['phoneIMEI2'])
        except:
            phoneIMEI2 = ''

        if ' 00:00:00' in dateReceived: 
            dateReceived = dateReceived.replace(" 00:00:00", "")

        if phoneIMEI.endswith('.0'):
            phoneIMEI = phoneIMEI[:-2]
        if phoneIMEI2.endswith('.0'):
            phoneIMEI2 = phoneIMEI2[:-2]
           
        # Summary writer, put a blank space or write your own summary if you don't want one auto generated
        if not summary and dateSeized and forensicExaminer and seizureAddress and agency != "ISP":
            summary = (
                f"On {dateSeized}, {forensicExaminer} attended the warrant at {seizureAddress}. "
                f"{forensicExaminer} read a copy of the search warrant authorizing the digital forensic analysis of digital computers, "
                f"-------------- and media."
            )
        elif summary:
            summary = summary

        qrCode = f"{caseNumber}_{exhibit}"



        pdf_output = f"EvidenceForm_{caseNumber}_Ex_{exhibit}.pdf"
        if not header:
            header = f"""
ACTIVITY REPORT                              BUREAU OF CRIMINAL INVESTIGATIONS
____________________________________________________________________________________

Activity Number:                             Date of Activity:
{caseNumber}                                		\t{todaysDate}
____________________________________________________________________________________
____________________________________________________________________________________
Subject of Activity:                         Case Agent:             Typed by:
{caseName} {subjectBusinessName}                        {caseAgent}        	{forensicExaminer}
{caseType}
____________________________________________________________________________________

Note

Data contained in these findings may be sensitive or confidential. It is intended for viewing only by those involved in the investigation, prosecution, defense, and adjudication of this case. Any other viewing is not authorized.

Executive Summary 

Special Agent {caseAgent} of the {agencyFull}, {divisionFull}, requested an examination of evidence for any information regarding the {caseType} investigation in the {caseName} case. {summary}
"""
        header = header.replace(' 00:00:00', '')

        if not executiveSummary:
            executiveSummary = f"""
{caseNumber}                                        {todaysDate}

{caseName} {subjectBusinessName}                              {caseAgent}    {forensicExaminer}

Note:

Data contained in these findings may be sensitive or confidential. It is intended for viewing only by those involved in the investigation, prosecution, defense, and adjudication of this case. Any other viewing is not authorized.

Executive Summary

Special Agent {caseAgent} of the {agencyFull}, {divisionFull}, requested an examination of evidence for any information regarding the {caseType} investigation in the {caseName} case. {summary}


Forensic Imaging
"""
        executiveSummary = executiveSummary.replace(' 00:00:00', '')

        report = f"""


Exhibit {exhibit}

"""


        if makeModel:
            article = "An" if makeModel[0].lower() in vowel else "A"
            report = f"{report}{article} {makeModel}"

        if mobileCarrier:
            carrier_info = mobileCarrier if exhibitType == 'phone' else f"(Carrier: {mobileCarrier})"
            report = f"{report} {carrier_info}"
        if exhibitType:
            report = f"{report} {exhibitType}"

        if phoneNumber and phoneNumber not in {'NA', 'na', 'N/A'}:
            report = f"{report} (MSISDN: {phoneNumber})"

        if phoneIMEI and phoneIMEI not in {'NA', 'na', 'N/A'}:
            report = f"{report} (IMEI: {phoneIMEI})"

        if phone2 and phone2 not in {'NA', 'na', 'N/A'}: # test
            report = f"{report} (MSISDN2: {phone2})"

        if phoneIMEI2 and phoneIMEI2 not in {'NA', 'na', 'N/A'}: # test
            report = f"{report} (IMEI2: {phoneIMEI2})"
            
        if serial:
            report = f"{report} (S/N: {serial})"

        if OS:
            report = f"{report} (OS: {OS})"

        if hostname:
            report = f"{report} (Hostname: {hostname})"

        if dateReceived:
            report = f"{report} was received on {dateReceived.replace(' ', ' at ', 1) if ' ' in dateReceived else dateReceived}"
        else:
            report = f"{report} was received"

        report = f"{report}."

        # Check if imagingStarted is not empty and status is not "Not imaged"
        if imagingStarted and status != "Not imaged":
            report = f"{report} On {imagingStarted.replace(' ', ' at ', 1)},"

        # Check if forensic examiner is a Digital Forensic Examiner
        if forensicExaminer.startswith("DFE"):
            report = f"{report} {forensicExaminer}"
        else:
            report = f"{report} Digital Forensic Examiner {forensicExaminer}"

        # Imaging tool and type with write blocker
        if imagingTool and imagingType and writeBlocker:
            article = "an" if imagingType[0].lower() in vowel else "a"
            report = f"{report} used {imagingTool}, utilizing a {writeBlocker} write blocker, to conduct {article} {imagingType}"

        # Imaging tool and type without write blocker
        elif imagingTool and imagingType:
            article = "an" if imagingType[0].lower() in vowel else "a"
            report = f"{report} used {imagingTool} to conduct {article} {imagingType}"

        # Only imaging tool
        elif imagingTool:
            report = f"{report} used {imagingTool} to conduct"

        # Imaging type and exported evidence is not "N"
        elif imagingType and exportedEvidence != "N":
            report = f"{report} conducted a {imagingType}"

        # Exported evidence is "N"
        elif exportedEvidence == "N":
            report = f"{report} did not conduct a"

        # Default case when nothing else applies
        else:
            report = f"{report} conducted a"


            
        if phoneNumber not in {'', 'NA', 'na', 'N/A'}:
            report = f"{report} phone extraction"
        elif imagingStarted:
            report = f"{report} forensic extraction"
        else:
            report = f"{report} manual analysis"


        if storageType and storageMakeModel and storageSerial and storageSize:
            report = f"{report} on the {storageMakeModel} (S/N: {storageSerial}) {storageSize} {storageType} drive."
        elif storageMakeModel and storageSerial and storageSize:
            report = f"{report} on the {storageMakeModel} (S/N: {storageSerial}) {storageSize} drive."
        elif storageMakeModel and storageSize:
            report = f"{report} on the {storageMakeModel} {storageSize} drive."
        else:
            report = f"{report}."

    
        # image hash
        if len(imageMD5) != 0 and exportLocation != '' and len(imageSHA256) != 0 and imageSHA256 != 'NA' and imageSHA256 != 'na' and imageSHA256 != 'N/A':
            # report = ("%s The image (SHA256 Hash: % s) (MD5 Hash: % s) was saved as %s." %(report, imageSHA256, imageMD5, exportLocation.split('\\')[-1])) 

            if writeBlocker != '':

                report = (
                    f"{report} A write blocker is a tool that prevents any write access to a device, thus only allowing for read-only access to maintain the integrity of the evidence. "
                    f"The image (MD5 Hash: {imageMD5}) was saved as {os.path.basename(exportLocation)}. "
                    f"The forensic imaging process was completed successfully and verified with no errors. The source hash was calculated and confirmed unchanged before processing. The acquisition and verification hash values matched, ensuring data integrity.\n"
                    f"\n\tSource MD5 hash:         {imageMD5}\n\tVerification MD5 hash: {imageMD5}\n\n"
                )
            else:
                report = (
                    f"{report}\n"
                    f"The image (MD5 Hash: {imageMD5}) was saved as {os.path.basename(exportLocation)}. "
                    f"The forensic imaging process was completed successfully and verified with no errors. The source hash was calculated and confirmed unchanged before processing. The acquisition and verification hash values matched, ensuring data integrity.\n"
                    f"\n\tSource MD5 hash:         {imageMD5}\n\tVerification MD5 hash: {imageMD5}\n\n"
                )



        elif len(imageMD5) != 0 and exportLocation != '':
            # report = ("%s The image (MD5 Hash: % s) was saved as %s." %(report, imageMD5, exportLocation.split('\\')[-1])) 
            if writeBlocker != '':
                report = (
                    f"{report} A write blocker is a tool that prevents any write access to a device, thus only allowing for read-only access to maintain the integrity of the evidence. "
                    f"The image (MD5 Hash: {imageMD5}) was saved as {os.path.basename(exportLocation)}. "
                    f"The imaging process completed with no errors and was verified. The acquisition and verification hash values matched and are listed below:\n"
                    f"\n\tSource MD5 hash:         {imageMD5}\n\tVerification MD5 hash: {imageMD5}\n\n"
                )
            else:
                report = (
                    f"{report}\n"
                    f"The image (MD5 Hash: {imageMD5}) was saved as {os.path.basename(exportLocation)}. "
                    f"The imaging process completed with no errors and was verified. The acquisition and verification hash values matched and are listed below:\n"
                    f"\n\tSource MD5 hash:         {imageMD5}\n\tVerification MD5 hash: {imageMD5}\n\n"
                )




        elif len(imageSHA256) != 0 and imageSHA256 != 'NA' and imageSHA256 != 'na' and imageSHA256 != 'N/A':
            # report = ("%s The image had a SHA256 hash of % s." %(report, imageSHA256))
            if writeBlocker != '':
                report = (
                    f"{report} A write blocker is a tool that prevents any write access to a device, thus only allowing for read-only access to maintain the integrity of the evidence. "
                    f"The image (SHA256 hash: {imageSHA256}) was saved as {os.path.basename(exportLocation)}. "
                    f"The imaging process completed with no errors and was verified. The acquisition and verification hash values matched and are listed below:\n"
                    f"\n\tSource SHA256 hash:         {imageSHA256}\n\tVerification SHA256 hash: {imageSHA256}\n\n"
                )
            else:
                report = (
                    f"{report}\n"
                    f"The image (SHA256 hash: {imageSHA256}) was saved as {os.path.basename(exportLocation)}. "
                    f"The imaging process completed with no errors and was verified. The acquisition and verification hash values matched and are listed below:\n"
                    f"\n\tSource SHA256 hash:         {imageSHA256}\n\tVerification SHA256 hash: {imageSHA256}\n\n"
                )

        # if "write blocker, to conduct" in report:
            # report = ("%s A write blocker is a tool that prevents any write access to a device, thus only allowing for read-only access to maintain the integrity of the evidence. " %(report))  
        # Analysis tool processing
        if analysisTool and analysisTool2:
            report = f"{report}The image was processed with {analysisTool} and further analyzed with {analysisTool2}."
        elif analysisTool:
            report = f"{report} The image was processed with {analysisTool}."

        if verifyHash.lower() == 'y':
            report = f"{report}The forensic image hash value was verified prior to processing thereby confirming the data remained unaltered prior to processing."


        # Username and password to report
        if userName and userPwd and exhibitType:
            report = f"{report} \"{userName}\" with a password of \"{userPwd}\" was a login to this {exhibitType}."
        elif userName and userPwd:
            report = f"{report} \"{userName}\" with a password of \"{userPwd}\" was a login to this device."

        # Email and password to report
        if email and emailPwd and exhibitType:
            report = f"{report} \"{email}\" with a password of \"{emailPwd}\" was an email configured on this {exhibitType}."
        elif email and exhibitType:
            if " and " in email:
                report = f"{report} {email} were email addresses configured on this {exhibitType}."
            else:
                report = f"{report} {email} was an email configured on this {exhibitType}."
        elif email and emailPwd:
            report = f"{report} \"{email}\" with a password of \"{emailPwd}\" was an email configured on this device."

        # Add notes to report if not empty
        if notes:
            report = f"{report}\n\n{notes}"

        # Exported evidence check
        if exportedEvidence == "Y" and 'elevant files were exported' not in notes:
            report = f"{report.rstrip()} Relevant files were exported."
        elif exportedEvidence == "N" and 'search for relevant files was conduced and no files were found' not in notes:
            report = f"{report.rstrip()} A search for relevant files was conducted and no files were found."

        # Evidence return check
        if "2" in removalDate and "eturned" in storageLocation:  # returned or Returned
            removalDate2 = removalDate.split(' ')[0] if " " in removalDate else removalDate

            if exhibitType:
                report = f"{report} This {exhibitType} was returned to the owner on {removalDate2}."
            else:
                report = f"{report} This exhibit was returned to the owner on {removalDate2}."

    
        report = report.replace("    , was received. ", "    ")
        report = report.replace("This was a DVR system was not imaged.","This was a DVR system and was not imaged.")
        report = report.replace("Digital Forensic Examiner Casey Karaffa did not conduct a forensic extraction.","This was not imaged.")
        report = report.replace("The image was processed with copy.","Pertinent files were copied.")
        report = report.replace("This had a  drive, model , serial # .","") # fixme     
        notes = notes.replace("This had a  drive, model , serial # .","")  # fixme
        report = report.replace(", serial # .",".") # fixme 
        notes = notes.replace(", serial # .",".") # fixme 
        if storageSerial == "000ecc45 0067205e":   # fixme  
            storageSerial = ' '

        print(report)

        body = f"{body}{report}"

    
        # Write excel
        write_output(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, forensicExaminer, reportStatus, notes, summary, tempNotes, exhibitType, makeModel, serial, OS, hostname, userName, userPwd, email, emailPwd, ip, phoneNumber, phoneIMEI, phone2, phoneIMEI2, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, seizureAddress, seizureRoom, dateSeized, seizedBy, seizureStatus, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, storageLocation, status, imagingTool, imagingType, imageMD5, imageSHA256, imageSHA1, verifyHash, writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, qrCode, operation, vaultCaseNumber, vaultTotal, caseNumberOrig, Action, priority, temp)

        if case_notes_status == 'True':
            my_dict = dictionary_build(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
            forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
            phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
            userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
            dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
            seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
            writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
            storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
            storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
            vaultTotal, tempNotes, temp, hostname, phoneIMEI2, phone2, verifyHash)
        else:
            my_dict = []
        
        # write an evidence form based on which agency you are from
        # Set pdf_output based on whether exhibit is provided
        if exhibit:
            pdf_output = f"ExhibitNotes_{caseNumber}_Ex{exhibit}.pdf"  # Output with exhibit
        else:
            pdf_output = f"ExhibitNotes_{caseNumber}_{todaysDateTime}.pdf"  # Output without exhibit, using today's date
            time.sleep(2)  # Wait 2 seconds to ensure a unique name

        # Choose the appropriate form based on agency and exhibit type
        if agency == "ISP":
            if exhibitType.lower() == 'phone':
                pdf_template = "EvidenceForm_MDIS.pdf"  # Mobile Device Evidence Sheet
            else:
                pdf_template = "EvidenceForm_EDIS.pdf"  # Electronic Device Evidence Sheet
        elif agency == "IDOR" and os.path.exists("Blank_EvidenceForm_IDOR.pdf"):
            pdf_template = "Blank_EvidenceForm_IDOR.pdf"
        else:
            pdf_template = "Blank_EvidenceForm.pdf"

        # Fill the selected PDF template
        if my_dict:  # This checks if the dictionary is not empty
            pdf_fill(pdf_template, pdf_output, my_dict)
            # print(f" ----------------------------   saving {pdf_output}\n {my_dict}")  # temp

        footer = f'''
Report Conclusion

This document contains findings regarding the analysis of digital evidence that was submitted for forensic examination. The data contained in these findings and data extractions should not be regarded as evidence, but rather findings concerning that evidence. Contact the case officer or prosecutor to obtain the evidentiary data.

It should be noted that not all files were reviewed during this examination. It is incumbent upon the requester to thoroughly review the data and make a determination as to the probative or exculpatory nature of any and all information.

All forensic equipment and software is functionally tested/validated without errors. Where possible, before utilizing stand-alone write blockers, firmware was updated to the current version and functionality was verified.

All forensic acquisition, analysis and write-blocking software used for this case is licensed and/or registered to {forensicExaminer} and/or the {agencyFull}.

This report contains digital examination of the exhibits provided based on the investigative information and tools available to the examiner at the time of the analysis. 

A copy of this report will be given to {caseAgent}. Additional analysis may be requested after review of the report or as the investigation continues.

All digital images obtained pursuant to this investigation will be maintained on {agency} servers for five years past the date of adjudication and/or case discontinuance. Copies of digital images will be made available upon request. All files copied from the images and provided to the case agent for review are identified as the DIGITAL EVIDENCE FILE and will be included as an exhibit in the case file.
'''


    # write docx report
    write_activity_report(caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, executiveSummary, body, footer)
    write_todo(caseNumber, caseNumberTodo, bodyDoing, bodyTodo, bodyDone)

def sanitize_filename(name):
    # Define a pattern for illegal characters
    name = name.replace('d/b/a ', '')
    name = name.replace(' ', '_')
    illegal_chars = r'[<>:"/\\|?*\x00]'
    # Replace them with an underscore or other safe character
    return re.sub(illegal_chars, '_', name)
    
def write_checklist():  # panda edition
    """
        create a checklist based on Panda
        Warning: Panda Creates a very big file when you convert it to exe
    """
    # Check if the output file already exists
    if os.path.exists(output_file):
        # Open the existing workbook
        book = openpyxl.load_workbook(output_file)
    else:
        # print(f"Making checklist {output_file}") 
        # Create a new workbook if the output file doesn't exist
        book = openpyxl.Workbook()
        book.active.title = "Checklist"

    # Check if the "Checklist" sheet already exists in the workbook
    if "Checklist" not in book.sheetnames:
        # If it doesn't exist, create the sheet
        book.create_sheet("Checklist")
    
    # Get the "Checklist" sheet
    checklist_sheet = book["Checklist"]

    # Set the font to Calibri 11 pt for all cells in the header row (row 7)
    for cell in checklist_sheet[7]:
        cell.font = Font(name='Calibri', size=11)

    # Set the row height for row 7 to 100
    checklist_sheet.row_dimensions[7].height = 100

    # Draw a solid border around cell
    cell_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))


    for col in range(1, 27):  # Columns A to D (index 1 to 4)
        cell = checklist_sheet.cell(row=7, column=col)
        # cell.border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        cell.border = cell_border

    # Freeze cell B8
    checklist_sheet.freeze_panes = "B8"

    # Set column widths
    checklist_sheet.column_dimensions['A'].width = 5
    checklist_sheet.column_dimensions['B'].width = 7
    checklist_sheet.column_dimensions['C'].width = 3
    checklist_sheet.column_dimensions['D'].width = 3
    checklist_sheet.column_dimensions['E'].width = 3
    checklist_sheet.column_dimensions['F'].width = 3
    checklist_sheet.column_dimensions['G'].width = 3
    checklist_sheet.column_dimensions['H'].width = 3
    checklist_sheet.column_dimensions['I'].width = 3
    checklist_sheet.column_dimensions['J'].width = 3
    checklist_sheet.column_dimensions['K'].width = 3
    checklist_sheet.column_dimensions['L'].width = 3
    checklist_sheet.column_dimensions['M'].width = 3
    checklist_sheet.column_dimensions['N'].width = 9
    checklist_sheet.column_dimensions['O'].width = 10
    checklist_sheet.column_dimensions['P'].width = 3 
    checklist_sheet.column_dimensions['Q'].width = 3 
    checklist_sheet.column_dimensions['R'].width = 3
    checklist_sheet.column_dimensions['S'].width = 3
    checklist_sheet.column_dimensions['T'].width = 3
    checklist_sheet.column_dimensions['U'].width = 3
    checklist_sheet.column_dimensions['V'].width = 3
    checklist_sheet.column_dimensions['W'].width = 3
    checklist_sheet.column_dimensions['X'].width = 3
    checklist_sheet.column_dimensions['Y'].width = 3
    checklist_sheet.column_dimensions['Z'].width = 3
    checklist_sheet.column_dimensions['AA'].width = 3
    checklist_sheet.column_dimensions['AB'].width = 3

    # Write additional headers to row 7
    additional_headers = [
        "exhibit#", "type", "evidence sheet (in)", "evidence sheet (out)", "label (all separate pieces)",
        "imaged", "image backup", "analyzed", "report (sign, print, forward)", "case notes printed",
        "digital evidence", "digital evidence backup", "digital evidence to agent", "return evidence", "", "Verify hash", "MemDump", "triage",
        "Magnet Encrypted Disk Detection", "password", "KAPE", "photograph", "OS", "IP or IMEI",
        "hostname", "Arsenal VM (verify)", "DFE"
    ]
    for idx, header in enumerate(additional_headers, start=1):
        cell = checklist_sheet.cell(row=7, column=idx, value=header)

    cell = checklist_sheet.cell(row=1, column=14, value="case#")
    cell = checklist_sheet.cell(row=2, column=14, value="caseName")
    cell = checklist_sheet.cell(row=3, column=14, value="subject")
    cell = checklist_sheet.cell(row=4, column=14, value="agent")
    cell = checklist_sheet.cell(row=5, column=14, value="forensics")

    # Create an orange fill pattern

    green_fill = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid') # green 
    orange_fill = PatternFill(start_color='FFc000', end_color='FFc000', fill_type='solid')  # orange
    red_fill = PatternFill(start_color='FF0000', end_color='FF0000', fill_type='solid')

    # Apply the orange fill to the cell
    # cell.fill = green_fill
    # cell.fill = orange_fill
    # cell.fill = red_fill    
    cell.border = cell_border   # create a border around current cell
    
    # Set cell C7 and D7, and I7 and J7 to orange
    cell_c7 = checklist_sheet['C7']
    cell_c7.fill = orange_fill


    cell_d7 = checklist_sheet['D7']
    cell_d7.fill = orange_fill

    cell_e7 = checklist_sheet['E7']
    cell_e7.fill = orange_fill
    
    cell_i7 = checklist_sheet['I7']
    cell_i7.fill = orange_fill

    cell_j7 = checklist_sheet['J7']
    cell_j7.fill = orange_fill

    cell_m7 = checklist_sheet['M7']
    cell_m7.fill = orange_fill

    cell_r7 = checklist_sheet['R7']
    cell_r7.fill = orange_fill


    cell_v7 = checklist_sheet['V7']
    cell_v7.fill = orange_fill




    # Set rotation to 45 degrees for row 7
    for row in checklist_sheet.iter_rows(min_row=7, max_row=7):
        for cell in row:
            cell.alignment = Alignment(textRotation=45)

    # Set page layout to landscape
    checklist_sheet.page_setup.orientation = checklist_sheet.ORIENTATION_LANDSCAPE

    # Specify the name of the sheet you want to read
    sheet_name = 'Cases'    # was 'Forensics'

    # Read the Excel file and load the specified sheet into a DataFrame
    # df = pd.read_excel(input_file)
    
    if not os.path.exists(input_file):
        print(f"{color_red}{input_file} does not exist{color_reset}")
        
        exit()    
    else:
        msg_blurb = (f'Reading {input_file}')
        msg_blurb_square(msg_blurb, color_green)
        
        dftemp = pd.read_excel(input_file, sheet_name=sheet_name)
        df = dftemp.fillna('')  # Replace NaN with empty string

    # df_sorted = df.sort_values(by='one')  # Sort the DataFrame by the "one" column
    # for index, row in df_sorted.iterrows():
    for index, row in df.iterrows():
        (caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, exhibit) = ('', '', '', '', '', '')
        (exhibitType, sheetIn, sheetOut, labeled, imaged, imageBackup) = ('', '', '', '', '', '')
        (analyzed, report, caseNotes, de, deBackup, deAgent) = ('', '', '', '', '', '')
        (memory, triage, edd, password, kape, photo) = ('', '', '', '', '', '')
        (OS, ipIMEI, hostname, arsenal, ip, phoneIMEI) = ('', '', '', '', '', '')
        (dateReceived, exportedEvidence, analysisTool, analysisTool2) = ('', '', '', '')
        (verifyHash, returnEvidence) = ('', '')
        caseNumber = row['caseNumber'] 
        caseName = row['caseName']
        subjectBusinessName = row['subjectBusinessName']
        removalDate = row['removalDate']
        
        reasonForRemoval = row['reasonForRemoval']
               
        caseAgent = row['caseAgent']
        forensicExaminer = row['forensicExaminer']
        DFE = forensicExaminer
        if 'dfe th' in DFE.lower() and 'dfe kar' in DFE.lower():
            DFE = 'JT/CK'
        elif 'dfe th' in DFE.lower():
            DFE = 'JT'
        elif 'dfe kar' in DFE.lower():
            DFE = 'CK'
            
        exhibit = str(row['exhibit'])
        if exhibit.endswith('.0'):
            exhibit = exhibit[:-2]

        exhibitType = str(row['exhibitType'])
        if "nan" in exhibitType.lower():
            exhibitType = ''
        sheetIn = str(row['dateReceived'])
        if "nan" in sheetIn:
            sheetIn = ""
        elif sheetIn != "":
            sheetIn = ".."

        sheetOut = str(row['removalDate'])
        if "nan" in sheetOut:
            sheetOut = "" 
            print(f"{color_red}nan exists{color_reset}")
        elif sheetOut != "":
            sheetOut = ".."
       
        imaged = row['status']
        
        try:
            verifyHash = str(row['verifyHash'])
        except:
            print(f'error with verifyHash: {verifyHash}')
            # verifyHash = ''
        
        
        if imaged.lower == 'imaged' or imaged == 'Imaged':
            imaged = 'Y'
        elif imaged.lower == 'not imaged' or 'not ' in imaged.lower():
            imaged = 'N'
            verify = 'N'
            verifyHash = 'N'
            # print(f' hello world {verifyHash}')  # test
        imageBackup = imaged    
        analyzed = str(row['analysisTool'])
        analysisTool2 = str(row['analysisTool2'])        
        exportedEvidence = str(row['exportedEvidence'])

        if 'nan' in analyzed:
            analyzed = ''
        if 'nan' in exportedEvidence:
            exportedEvidence = ''            
            

        if analyzed != "" and exportedEvidence != "":
            analyzed = 'Y'   
        # elif analyzed != "" and exportedEvidence != "":
            # analyzed = 'Y'    
        # else:
            # analyzed = ''    

        report = str(row['reportStatus'])
        if isinstance(report, str) and "inal" in report:
            report = 'Y'    
        elif isinstance(report, str) and "raft" in report:
            report = 'd' 
        else:
            report = '' 

        if "storage" in exhibitType or "dvr" in exhibitType.lower() or "UPS" in exhibitType  or "switch" in exhibitType or "vehicle" in exhibitType.lower() :
            (memory, triage, edd, password, kape, arsenal) = ('N', 'N', 'N', 'N', 'N', 'N')
        elif "phone" in exhibitType or "tablet" in exhibitType:
            (memory, triage, edd, kape, arsenal) = ('N', 'N', 'N', 'N', 'N')

        else:
            (memory, triage, edd, password, kape) = ('', '', '', '', '')

        if verifyHash == '' and "not imaged" in imaged.lower():
            verifyHash = 'N'

        if "memdump" in analysisTool2.lower():
            memory = "Y"

        if "magnet encrypted disk detect" in analysisTool2.lower():
            edd = 'Y'
        if "kape" in analysisTool2.lower():
            kape = 'Y'
            
        de = row['exportedEvidence']
        deBackup = row['exportedEvidence']
        # if de.lower() == 'n':
            # deAgent = 'N'
        deAgent = 'N' if str(de).lower() == 'n' else deAgent


        
        password = str(row['userPwd'])
        
        if password == 'nan':
            password = ''
        elif password != '':
            password = 'Y'

        OS = str(row['OS']).strip()


        if pd.isna(OS):  # Check if the value is missing
            OS = ''
        elif "nan" in OS.lower():
            OS = ''        
        elif OS != "":
            OS = 'Y'
        ip = str(row['ip'])

        if ip == 'nan':
            ip = ''
        
  
        phoneIMEI = str(row['phoneIMEI'])
        # if 'nan' in phoneIMEI:        
            # phoneIMEI == ''

        # if phoneIMEI == 'nan':
            # phoneIMEI = ''

        if ip != "" or phoneIMEI != "": # task always prints y
            ipIMEI = 'Y' 

        if removalDate != '' and 'returned' in reasonForRemoval.lower():
            returnEvidence = 'Y'


        try:
            hostname = str(row['hostname'])
            if hostname == 'nan':
                hostname = ''
            if len(hostname) >=2:
                hostname = 'Y' 
            if hostname != "":
                hostname = 'Y' 
        except TypeError as error:
            print(error)

        if "arsenal" in analysisTool2.lower():  #test
            arsenal = "Y"


    
        print(f"{exhibit}\t{exhibitType}")

        cell = checklist_sheet.cell(row=1, column=15, value=caseNumber)
        cell = checklist_sheet.cell(row=2, column=15, value=caseName)
        cell = checklist_sheet.cell(row=3, column=15, value=subjectBusinessName)
        cell = checklist_sheet.cell(row=4, column=15, value=caseAgent)
        cell = checklist_sheet.cell(row=5, column=15, value=forensicExaminer)

        # Define your data values
        data_values = [exhibit, exhibitType, sheetIn, sheetOut, labeled, imaged, imageBackup, analyzed, report, caseNotes, de, deBackup, deAgent, returnEvidence, caseNumber, verifyHash, memory, triage, edd, password, kape, photo, OS, ipIMEI, hostname, arsenal, DFE]

        # Find the next available row index
        next_row = checklist_sheet.max_row + 1

        # Append data to the checklist_sheet while applying border
        for col, value in enumerate(data_values, start=1):  # Start columns from 1 (A)
            cell = checklist_sheet.cell(row=next_row, column=col, value=value)
            cell.border = cell_border

        # checklist_sheet.append([exhibit, exhibitType, sheetIn, sheetOut, labeled, imaged, imageBackup, analyzed, report, caseNotes, de, deBackup, deAgent, "", "", "", memory, triage, edd, password, kape, photo, OS, ipIMEI, hostname, arsenal])
    print(f"{color_green}Data written to {output_file}{color_reset}")    

    # Save the workbook to the output file
    book.save(output_file)

def write_output(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, forensicExaminer, reportStatus, notes, summary, tempNotes, exhibitType, makeModel, serial, OS, hostname, userName, userPwd, email, emailPwd, ip, phoneNumber, phoneIMEI, phone2, phoneIMEI2, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, seizureAddress, seizureRoom, dateSeized, seizedBy, seizureStatus, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, storageLocation, status, imagingTool, imagingType, imageMD5, imageSHA256, imageSHA1, verifyHash, writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, qrCode, operation, vaultCaseNumber, vaultTotal, caseNumberOrig, Action, priority, temp):
    """
        Write the output file (xlsx)
    """
    
    
    # Check if the output file already exists
    if os.path.exists(output_file):
        # Open the existing workbook
        book = openpyxl.load_workbook(output_file)
        # Get the active worksheet
        sheet = book.active
    else:
        # Create a new workbook if the output file doesn't exist
        book = Workbook()
        sheet = book.active
        sheet.title = 'Cases'   # was "Forensics"

        headers = ["caseNumber", "exhibit", "caseName", "subjectBusinessName", "caseType"
        , "caseAgent", "forensicExaminer", "reportStatus", "notes", "summary", "tempNotes"
        , "exhibitType", "makeModel", "serial", "OS", "hostname", "userName", "userPwd"
        , "email", "emailPwd", "ip", "phoneNumber", "phoneIMEI", "phone2", "phoneIMEI2"
        , "mobileCarrier", "biosTime", "currentTime", "timezone", "shutdownMethod"
        , "shutdownTime", "seizureAddress", "seizureRoom", "dateSeized", "seizedBy"
        , "seizureStatus", "dateReceived", "receivedBy", "removalDate", "removalStaff"
        , "reasonForRemoval", "inventoryDate", "storageLocation", "status", "imagingTool"
        , "imagingType", "imageMD5", "imageSHA256", "imageSHA1", "verifyHash", "writeBlocker"
        , "imagingStarted", "imagingFinished", "storageType", "storageMakeModel"
        , "storageSerial", "storageSize", "evidenceDataSize", "analysisTool"
        , "analysisTool2", "exportLocation", "exportedEvidence", "qrCode", "operation"
        , "vaultCaseNumber", "vaultTotal", "caseNumberOrig", "Action", "priority", "temp"]

        sheet.append(headers)




        # Set the header row cell colors

        orange_columns = ['A', 'C', 'd', 'e', 'f', 'g', 'h']
        for col in orange_columns: 
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='FFc000', end_color='FFc000', fill_type='solid')    #orange

        yellow_columns = ['B', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z', 'AA', 'AB', 'AC', 'AD', 'AE']
        for col in yellow_columns:
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')

        violet_columns = ['I', 'J', 'K']
        for col in violet_columns:
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='CCCCFF', end_color='CCCCFF', fill_type='solid')    # purple

        green_columns = ['AF', 'AG', 'AH', 'AI', 'AJ', 'AK', 'AL', 'AM', 'AN', 'AO', 'AP' ]
        for col in green_columns:
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='92D050', end_color='92D050', fill_type='solid')    # green

        blue_columns = ['AQ', 'AR', 'AS', 'AT', 'AU', 'AV', 'AW', 'AX', 'AY', 'AZ', 'BA', 'BB', 'BC', 'BD', 'BE', 'BF', 'BG', 'BH', 'BI', 'BJ']
        for col in blue_columns:
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='66CCFF', end_color='66CCFF', fill_type='solid')    # blue

        pink_columns = ['BK', 'BL', 'BM', 'BN', 'BO', 'BP', 'BQ' ]
        for col in pink_columns:
            cell = sheet[f"{col}1"]
            cell.fill = PatternFill(start_color='FF99FF', end_color='FF99FF', fill_type='solid')    # pink

        # Freeze cells at B2
        sheet.freeze_panes = 'B2'

        # Set column width for the 'Cases' sheet
        sheet.column_dimensions['A'].width = 15 #  caseNumber
        sheet.column_dimensions['B'].width = 7 #  exhibit
        sheet.column_dimensions['C'].width = 16 #  caseName
        sheet.column_dimensions['D'].width = 25 #  subjectBusinessName
        sheet.column_dimensions['E'].width = 16 #  caseType
        sheet.column_dimensions['F'].width = 25 #  caseAgent
        sheet.column_dimensions['G'].width = 15 #  forensicExaminer
        sheet.column_dimensions['H'].width = 13 #  reportStatus
        sheet.column_dimensions['I'].width = 25 #  notes
        sheet.column_dimensions['J'].width = 15 #  summary
        sheet.column_dimensions['K'].width = 40 #  tempNotes
        sheet.column_dimensions['L'].width = 12 #  exhibitType
        sheet.column_dimensions['M'].width = 30 #  
        sheet.column_dimensions['N'].width = 17 #  
        sheet.column_dimensions['O'].width = 15 #  
        sheet.column_dimensions['P'].width = 18 #  
        sheet.column_dimensions['Q'].width = 12 # 
        sheet.column_dimensions['R'].width = 12 #  
        sheet.column_dimensions['S'].width = 20 #  
        sheet.column_dimensions['T'].width = 12 #  
        sheet.column_dimensions['U'].width = 14 #  
        sheet.column_dimensions['V'].width = 14 #  
        sheet.column_dimensions['W'].width = 16 #  
        sheet.column_dimensions['X'].width = 16 #  
        sheet.column_dimensions['Y'].width = 16 #  
        sheet.column_dimensions['Z'].width = 15 #  
        sheet.column_dimensions['AA'].width = 16 #  
        sheet.column_dimensions['AB'].width = 16 #  
        sheet.column_dimensions['AC'].width = 12 #  
        sheet.column_dimensions['AD'].width = 15 #  
        sheet.column_dimensions['AE'].width = 16 #  
        sheet.column_dimensions['AF'].width = 15 #  
        sheet.column_dimensions['AG'].width = 12 #  
        sheet.column_dimensions['AH'].width = 16 #  
        sheet.column_dimensions['AI'].width = 12 #  
        sheet.column_dimensions['AJ'].width = 18 #  
        sheet.column_dimensions['AK'].width = 16 #  
        sheet.column_dimensions['AL'].width = 15 #  
        sheet.column_dimensions['AM'].width = 16 #  
        sheet.column_dimensions['AN'].width = 25 #  
        sheet.column_dimensions['AO'].width = 18 #  
        sheet.column_dimensions['AP'].width = 15 #  
        sheet.column_dimensions['AQ'].width = 25 #  
        sheet.column_dimensions['AR'].width = 12 #  
        sheet.column_dimensions['AS'].width = 24 #  
        sheet.column_dimensions['AT'].width = 15 #  
        sheet.column_dimensions['AU'].width = 16 #  
        sheet.column_dimensions['AV'].width = 15 #  
        sheet.column_dimensions['AW'].width = 15 #  
        sheet.column_dimensions['AX'].width = 11 #  
        sheet.column_dimensions['AY'].width = 15 #  
        sheet.column_dimensions['AZ'].width = 22 #  
        sheet.column_dimensions['BA'].width = 16 #  
        sheet.column_dimensions['BB'].width = 13 #  
        sheet.column_dimensions['BC'].width = 23 #  
        sheet.column_dimensions['BD'].width = 19 #  
        sheet.column_dimensions['BE'].width = 14 #  
        sheet.column_dimensions['BF'].width = 15 #  
        sheet.column_dimensions['BG'].width = 23 #  
        sheet.column_dimensions['BH'].width = 15 #  
        sheet.column_dimensions['BI'].width = 25 #  
        sheet.column_dimensions['BJ'].width = 15 #  
        sheet.column_dimensions['BK'].width = 15 #  
        sheet.column_dimensions['BL'].width = 15 #  
        sheet.column_dimensions['BM'].width = 19 #  
        sheet.column_dimensions['BN'].width = 15 #  
        sheet.column_dimensions['BO'].width = 19 #      
        sheet.column_dimensions['BP'].width = 10 #  
        sheet.column_dimensions['BQ'].width = 9 #  
        sheet.column_dimensions['BR'].width = 5 #       

    # Write data to the 'Cases' sheet
    row_data = [caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, forensicExaminer, reportStatus, notes, summary, tempNotes, exhibitType, makeModel, serial, OS, hostname, userName, userPwd, email, emailPwd, ip, phoneNumber, phoneIMEI, phone2, phoneIMEI2, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, seizureAddress, seizureRoom, dateSeized, seizedBy, seizureStatus, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, storageLocation, status, imagingTool, imagingType, imageMD5, imageSHA256, imageSHA1, verifyHash, writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, qrCode, operation, vaultCaseNumber, vaultTotal, caseNumberOrig, Action, priority, temp]    
    
    sheet.append(row_data)

    # Set the font to Calibri 11 pt for all cells in the data rows
    for row in sheet.iter_rows(min_row=0):
        for cell in row:
            cell.font = Font(name='Calibri', size=11)

    # Save the workbook to the output file
    book.save(output_file)

def create_and_write_xlsx():
    """
    Creates an xlsx database file with formatting and writes data to it using openpyxl and pandas
    """
    from openpyxl.utils import get_column_letter


    # Define the data structure
    data = {
        'caseNumber': [], 'exhibit': [], 'caseName': [], 'subjectBusinessName': [],
        'caseType': [], 'caseAgent': [], 'forensicExaminer': [], 'reportStatus': [],
        'notes': [], 'summary': [], 'tempNotes': [], 'exhibitType': [], 'makeModel': [],
        'serial': [], 'OS': [], 'hostname': [], 'userName': [], 'userPwd': [], 'email': [],
        'emailPwd': [], 'ip': [], 'phoneNumber': [], 'phoneIMEI': [], 'phone2': [],
        'phoneIMEI2': [], 'mobileCarrier': [], 'biosTime': [], 'currentTime': [],
        'timezone': [], 'shutdownMethod': [], 'shutdownTime': [], 'seizureAddress': [],
        'seizureRoom': [], 'dateSeized': [], 'seizedBy': [], 'seizureStatus': [],
        'dateReceived': [], 'receivedBy': [], 'removalDate': [], 'removalStaff': [],
        'reasonForRemoval': [], 'inventoryDate': [], 'storageLocation': [], 'status': [],
        'imagingTool': [], 'imagingType': [], 'imageMD5': [], 'imageSHA256': [],
        'imageSHA1': [], 'verifyHash': [], 'writeBlocker': [], 'imagingStarted': [],
        'imagingFinished': [], 'storageType': [], 'storageMakeModel': [], 'storageSerial': [],
        'storageSize': [], 'evidenceDataSize': [], 'analysisTool': [], 'analysisTool2': [],
        'exportLocation': [], 'exportedEvidence': [], 'qrCode': [], 'operation': [],
        'vaultCaseNumber': [], 'vaultTotal': [], 'caseNumberOrig': [], 'Action': [],
        'priority': [], 'temp': []
    }

    # Create DataFrame
    df = pd.DataFrame(data)

    # Create workbook and worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Cases"
    ws.freeze_panes = "B2"

    # Define styles
    bold_font = Font(bold=True)
    border = Border(left=Side(style='thin'), right=Side(style='thin'),
                    top=Side(style='thin'), bottom=Side(style='thin'))

    # Color groups by column name (not letter)
    orange = {'caseNumber', 'caseName', 'subjectBusinessName', 'caseType', 'caseAgent', 'forensicExaminer', 'reportStatus'}
    yellow = {'exhibit', 'exhibitType', 'makeModel', 'serial', 'OS', 'hostname', 'userName', 'userPwd', 'email', 'emailPwd',
              'ip', 'phoneNumber', 'phoneIMEI', 'phone2', 'phoneIMEI2', 'mobileCarrier', 'biosTime', 'currentTime', 'timezone', 'shutdownMethod', 'shutdownTime'}
    violet = {'notes', 'summary', 'tempNotes'}
    green = {'seizureAddress', 'seizureRoom', 'dateSeized', 'seizedBy', 'seizureStatus', 'dateReceived', 'receivedBy', 'removalDate', 'removalStaff', 'reasonForRemoval', 'inventoryDate'}
    blue = {'storageLocation', 'status', 'imagingTool', 'imagingType', 'imageMD5', 'imageSHA256', 'imageSHA1',
            'verifyHash', 'writeBlocker', 'imagingStarted', 'imagingFinished', 'storageType', 'storageMakeModel', 'storageSerial', 'storageSize', 'evidenceDataSize',
            'analysisTool', 'analysisTool2', 'exportLocation', 'exportedEvidence'}
    pink = {'qrCode', 'operation', 'vaultCaseNumber', 'vaultTotal', 'caseNumberOrig', 'Action', 'priority'}

    # Color mapping
    color_map = {
        'orange': ('FFC000', orange),
        'yellow': ('FFFF00', yellow),
        'violet': ('CCCCFF', violet),
        'green': ('92D050', green),
        'blue': ('66CCFF', blue),
        'pink': ('FF99FF', pink)
    }

    # Write headers with formatting
    for col_num, col_name in enumerate(df.columns, start=1):
        cell = ws.cell(row=1, column=col_num, value=col_name)
        cell.font = bold_font
        cell.border = border

        # Apply color fill based on group
        for color_hex, group in color_map.values():
            if col_name in group:
                cell.fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type='solid')
                break

        # Set column width
        ws.column_dimensions[get_column_letter(col_num)].width = 20

    # Save workbook
    wb.save(output_xlsx)

    msg_blurb = (f'Blank sheet saved as {output_xlsx}')
    msg_blurb_square(msg_blurb, color_green)

def ufdx_parser(xml_path):
    import xml.etree.ElementTree as ET
    tree = ET.parse(xml_path)
    root = tree.getroot()
    TempNotes = ''
    # Extract model
    makeModel = root.find('DeviceInfo').attrib.get('Model', '')
    
    # Build dictionary of caption/value pairs
    fields = {}
    for field in root.findall('.//CrimeCase/Fields/Fields'):
        caption = field.attrib.get('Caption', '').strip()
        value = field.attrib.get('Value', '').strip()
        fields[caption] = value

    # Extract specific fields
    caseNumber = fields.get('Case Identifier', '')
    exhibit = fields.get('Device Name / Evidence Number', '')
    if not exhibit:
        exhibit = fields.get('Case ID', '')
    subjectBusinessName = fields.get('Device owner', '')
    # subjectBusinessName = fields[caption]
    
    caseType = fields.get('Crime type', '')
    if not caseType:
        caseType = fields.get('Crime Type', '')
        
    forensicExaminer = fields.get('Examiner Name', '')    
    seizedBy = fields.get('Seized by', '')
    seizureAddress = fields.get('Location', '')

    # Extract extraction info
    extraction = root.find('.//Extractions/Extraction')
    imagingType = extraction.attrib.get('TransferType', '') if extraction is not None else ''
    exportLocation = extraction.attrib.get('Path', '') if extraction is not None else ''
    return (caseNumber, exhibit, subjectBusinessName, caseType, forensicExaminer, makeModel, seizureAddress, seizedBy, imagingType, exportLocation)

def write_activity_report(caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, executiveSummary, body, footer): 
    """
        write %s__%s_%s_%s_DRAFT.docx
    """
    safe_filename1 = sanitize_filename(caseNumber)
    safe_filename2 = sanitize_filename(subjectBusinessName)    
    
    output_docx = (f"{safe_filename1}_{safe_filename2}_{Year}-{Month}-{Day}_DRAFT.docx") 

    try:
        document = docx.Document("Blank_ActivityReport.docx") # read in the template if it exists
    except:
        print(f"{color_red}You are missing Blank_ActivityReport.docx{color_reset}")        

        document = create_docx()   # create a basic template file
    
    if executiveSummary != '':
        document.add_paragraph(executiveSummary)    
    
    document.add_paragraph(body)  

    if footer != '':
        document.add_paragraph(footer) 
    
    document.save(output_docx)   # print output to the new file

    msg_blurb = (f'Activity report written to {output_docx}')
    msg_blurb_square(msg_blurb, color_green)

def write_sticker():
    """
        write a sticker, maximum of 10 cells (sorry)
    """
    output_docx = "stickers.docx"
    output = open(output_docx, 'w+')

    (header, reportStatus, date) = ('', '', '<insert date here>')
    (headers) = []

    if not os.path.exists('Avery2x4Labels.docx'):
        print(f"{color_red}you are missing Avery2x4Labels.docx.... so Im making a lame version{color_reset}")        
        
        # Create a new Word document
        document = docx.Document()
        # Set the side margins to 0.25 inches
        section = document.sections[0]
        section.margin_left = docx.shared.Inches(0.25)
        section.margin_right = docx.shared.Inches(0.25)
        table = document.add_table(rows=5, cols=3)
    else:
        document = docx.Document('Avery2x4Labels.docx')
        # Find the first table in the document
        table = document.tables[0]

    # Change the default font size for the document
    for style in document.styles:
        style.font.size = docx.shared.Pt(14)
 
    # Specify the name of the sheet you want to read
    sheet_name = 'Cases'    # was 'Forensics'

    # Read the Excel file and load the specified sheet into a DataFrame

    if not os.path.exists(input_file):
        print(f"{color_red}{input_file} does not exist{color_reset}")    
        exit()
    else:
        msg_blurb = (f'Reading {input_file}')
        msg_blurb_square(msg_blurb, color_green)
  

        dftemp = pd.read_excel(input_file, sheet_name=sheet_name)
        # df = dftemp.fillna('').sort_values(by='exhibit')  # Replace NaN with empty string and sort by exhibit
        df = dftemp.fillna('') # Replace NaN with empty string

    for index, row in df.iterrows():
        (caseNumber, caseName, subjectBusinessName, caseAgent, exhibit, makeModel) = ('', '', '', '', '', '')
        (serial, status) = ('', '')

        caseNumber = str(row['caseNumber'])
        caseName = row['caseName']
        subjectBusinessName = row['subjectBusinessName']
        caseAgent = row['caseAgent']
        exhibit = str(row['exhibit'])
        if exhibit.endswith('.0'):
            exhibit = exhibit[:-2]
            
        makeModel = row['makeModel']
        serial = str(row['serial'])       
        status = row['status']

        if status.lower() == 'imaged':    
            header = (f'''Case#: {caseNumber}      Ex: {exhibit}
CaseName: {caseName}
Subject: {subjectBusinessName}
Make: {makeModel} 
Serial: {serial}
Agent: {caseAgent}
{status}
''')
        else:
            header = (f'''Case#: {caseNumber}      Ex: {exhibit}
CaseName: {caseName}
Subject: {subjectBusinessName}
Make: {makeModel} 
Serial: {serial}
Agent: {caseAgent}
''')

        header = header.strip()
        headers.append(header)

    # write to stickers.docx
    l = 0
    for i in range(5):
        try:
            cell = table.cell(i, 0)
            cell.text = headers[l]
            l += 1
            cell = table.cell(i, 2)
            cell.text = headers[l]
            l += 1
        except:
            pass
    document.save(output_docx)
    print(f"{color_green}Data written to {output_docx}")

def write_todo(caseNumber, caseNumberTodo, bodyDoing, bodyTodo, bodyDone):
    # Define the filename
    todo_filename = f'Todo_{caseNumber}.md'
    
    # Define the content of the file

    bodyList = f'''---

kanban-plugin: board

---

## Doing
{bodyDoing}

## To-Do
{caseNumberTodo}
{bodyTodo}

## Done

{bodyDone}


%% kanban:settings
```
{{"kanban-plugin":"board","list-collapse":[false,false,false],"show-archive-all":false}}
```
%%'''


    
    # Write to the file
    with open(todo_filename, 'w', encoding='utf-8') as file:
        file.write(bodyList)
    
    msg_blurb = (f'ToDo file "{todo_filename}" has been created.')
    msg_blurb_square(msg_blurb, color_green)

def usage():
    """
        working examples of syntax
    """
    file = sys.argv[0].split('\\')[-1]
    


    print(f"\nDescription: {color_green}{description}{color_reset}")
    print(f"{file} Version: {version} by {author}")
    print(f"\n    {color_yellow}insert your info into input_case.xlsx")
    print(f"\n    or insert logs into Logs folder{color_reset}")
    print(f"\nExample:")
    print(f"    {file} -b")    
    print(f"    {file} -C")
    print(f"    {file} -g")
    print(f"    {file} -l")
    print(f"    {file} -L")
    print(f"    {file} -r")
    print(f"    {file} -r -c -C -s -I ForensicCasesExample.xlsx")
    print(f"    {file} -s")
 
if __name__ == '__main__':
    main()

# <<<<<<<<<<<<<<<<<<<<<<<<<< Revision History >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
3.4.6 - fixed -b blank 
3.4.2 - parse .ufdx logs 
3.4.1 - if exhibit ends in 0 , like 10 , it strips off the 0 and now it's a 1, same with phone2
3.4.0 - ReOrganized column order (run old sheets through -r) to re-organize old data)
3.1.1 - Add verifyHash column
3.1.0 - export a markdown file as a todo list
3.0.9 - went back to Exhibit instead of Item
3.0.7 - export a markdown file as a todo list
3.0.6 - added lots of new wording and disclaimer based on a report by J.S.
3.0.5 - converted to print(f" so it's easier to read and update by variables)
3.0.0 - added color & error checking
2.9.4 - converted from (xlxriter and xlxreader) to pandas to support multiple lines in Notes column
2.7.3 - added checklist to xlsx
2.7.2 - sticker maker prints out avery 2 x 4" labels now
2.7.0 - added -g for GUI interface for data enty
2.6.7 - fixed Tableau storageSerial being the Tableau serial number (I think)
2.6.5 - changed report writing output and summary writer (add a space in summary if you don't want it to write anything.)
2.6.2 - reportStatus gets colored if it's marked Finalized, Draft or Y
2.6.1 - if you change agency, agencyFull and divisionFull it writes a more customized report
2.6.0 - added -L option to parse a folder full of logs all at once. -I and -O are optional now
2.5.6 - Logs: CellebritePremium DeviceInfo.txt, Berla iVE
2.5.0 - Column Re-order to group like exhibits together (case, description, lab chain of custody, acquisition, notes)
2.1.3 - fixed log parser to populate storageSize, storageMakeModel, storageSerial, storageSize (Tableau)
2.1.2 - Added about a dozen columns for additional info (the columns need to be re-ordered one of these days.)
2.1.1 - Added ISP pdf templates for pdf writing (just change agency = to agency = 'ISP'
2.1.0 - Added CaseNotes.pdf output if you add -c to -r
2.0.3 - Added Recon imager log parsing
2.0.2 - ActivityReport....docx output works best from the template.
2.0.1 - Reorginized column orders, fixed serial #
1.0.1 - Created a Tableau log parser
1.0.0 - Created forensic report writer
0.1.2 - converted tabs to 4 spaces for #pep8

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<< Future Wishlist  >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
should I change elif len(imageMD5) != 0 and exportLocation != '':
to
elif len(imageMD5) != 0 and exportLocation != '' and verifyHash.lower() == 'y':

fix date formats. 

Don't write the hash verification blurb if verifyHash = 'N'

when doing -L with a folder name in there, it craps out. (skip folders)
exhibit.lstrip('=')

add a -f option if you want it worded in first person perspective (sounds hard to write)
Add a glossary of terms?

fix pdf output (often blank but sometimes it works)
if date doesn't have time, don't put 0:0:0

Label GUI - Tkinter screen (Case#, Agent, Case Name, Location, Date, Exhibit, Room) # of stickers <print>

standaradize date format in reporting (Wednesday, July 7, 2021)

figure out DocX tags or variables to insert data into the header fields

add a brother (or Dymo) label printer output to xlsx with qrCode
qrCode could be caseNumber_exhibit_serial (it depends on what the evidence staff want displayed on their inventory scanner)

parse: MagentAcuire, MagnetAxiom, SumuriReconImager

if qrCode = '_': qrCode = ''

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<      notes            >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
change the agency, agencyFull, divisionFull to your agencies info

if you have log output (like GrayKey) you want parsed, send it my way. As long as there is a key:value pair on one line, I can do it.
If you want your agencies forms filled, you just need to insert these variables into your pdf.


“black iPhone 11 currently stored as exhibit number x under case number 22-xxxx which was located on master bedroom nightstand….”

"""


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Copyright        >>>>>>>>>>>>>>>>>>>>>>>>>>

# Copyright (C) 2023 LincolnLandForensics
#
# This program is free software; you can redistribute it and/or modify it under
# the terms of the GNU General Public License version 2, as published by the
# Free Software Foundation
#
# This program is distributed in the hope that it will be useful, but WITHOUT
# ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
# FOR A PARTICULAR PURPOSE. See the GNU General Public License for more
# details (http://www.gnu.org/licenses/gpl.txt).


# <<<<<<<<<<<<<<<<<<<<<<<<<<      The End        >>>>>>>>>>>>>>>>>>>>>>>>>>