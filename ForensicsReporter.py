#!/usr/bin/python
# coding: utf-8


# <<<<<<<<<<<<<<<<<<<<<<<<<<     Change Me       >>>>>>>>>>>>>>>>>>>>>>>>>>
# change this section with your details
global agency
agency = "MWW" # ISP, MWW
global agencyFull
agencyFull = "Ministry of Wacky Walks"   # Ministry of Wacky Walks
global divisionFull
divisionFull = "Criminal Investigation Division" # Criminal Investigation Division


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Imports        >>>>>>>>>>>>>>>>>>>>>>>>>>

try:
    import docx # pip install python-docx
    import xlrd # read xlsx
    import pdfrw    # pip install pdfrw
    import hashlib # pip install hashlib
    import openpyxl # pip install openpyxl
    import xlsxwriter   
    import pdfplumber   # pip install pdfplumber

except:
    print('install missing modules:    pip install -r requirements_ForensicsReporter.txt')
    exit()
import re
import os
import sys  
import time # for wait line
import argparse  # for menu system
from datetime import date
from subprocess import call
from datetime import datetime

from tkinter import *   # -t  # Frame is not defined if this is missing
import tkinter  # -d
from tkinter import ttk # -d
from tkinter import messagebox # -d

d = datetime.today()

Day    = d.strftime("%d")
Month = d.strftime("%m")    # %B = October
Year  = d.strftime("%Y")        
todaysDate = d.strftime("%m/%d/%Y")
todaysDateTime = d.strftime("%m_%d_%Y_%H-%M-%S")

ANNOT_KEY = '/Annots'
ANNOT_FIELD_KEY = '/T'
ANNOT_VAL_KEY = '/V'
ANNOT_RECT_KEY = '/Rect'
SUBTYPE_KEY = '/Subtype'
WIDGET_SUBTYPE_KEY = '/Widget'

# Regex section
regex_md5 = re.compile(r'^([a-fA-F\d]{32})$')  # regex_md5        [a-f0-9]{32}$/gm
regex_sha1 = re.compile(r'^([a-fA-F\d]{40})$')    #regex_sha1
regex_sha256 = re.compile(r'^([a-fA-F\d]{64})$')#regex_sha256

# <<<<<<<<<<<<<<<<<<<<<<<<<<      Pre-Sets       >>>>>>>>>>>>>>>>>>>>>>>>>>

author = 'LincolnLandForensics'
description = "convert imaging logs to xlsx, print stickers and write activity reports/ case notes"
version = '2.7.3'

# <<<<<<<<<<<<<<<<<<<<<<<<<<      Menu           >>>>>>>>>>>>>>>>>>>>>>>>>>

def main():
    '''
    main menu
    '''
    
    global Row
    Row = 1  # defines arguments
    global Row2
    Row2 = 7  #     
    global caseMain
    caseMain = ''    
    
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument('-I', '--input', help='default is input.txt', required=False)
    parser.add_argument('-O', '--output', help='', required=False)
    parser.add_argument('-d', '--details', help='manually enter details like exhibit number', required=False, action='store_true')
    parser.add_argument('-g', '--guidataentry', help='data entry GUI', required=False, action='store_true')
    parser.add_argument('-l', '--logparse', help='Berla, Cellebrite, FTK, tableau log parser', required=False, action='store_true')
    parser.add_argument('-L', '--logs_parse', help='dump all your logs into Logs\ folder', required=False, action='store_true')
    parser.add_argument('-r', '--report', help='write report', required=False, action='store_true')
    parser.add_argument('-c','--caseNotes', help='casenotes module (optional) used with -r', required=False, action='store_true')
    parser.add_argument('-s', '--sticker', help='write sticker', required=False, action='store_true')

    args = parser.parse_args()

    # global section
    global inputDetails
    inputDetails = 'no'
    global filename
    filename = ('input.txt')
    global logsFolder
    logsFolder = ('Logs\\')   # s subfolder full of logs
    global logsList
    logsList = ['']
    global log_type
    global outputFileXlsx   # docx actitivy report
    outputFileXlsx = ('output_.docx')
    global spreadsheet
    spreadsheet = ('log_%s.xlsx' %(todaysDateTime)) # uniq naming for -l module
    global sheet_format
    sheet_format = ('')

    # global outputFileXlsx
    outputFileXlsx = "output.xlsx"
    win = Frame()
    win.grid(sticky=N+S+E+W)


    if args.output:
        outputFileXlsx = args.input
        
    
    # input section
    if args.details:
        inputDetails = ("yes")
    if args.report or args.logparse or args.logs_parse or args.sticker:
        create_xlsx()

    if args.report:
        global caseNotesStatus
        if args.caseNotes:  # if you add -c                                  
            caseNotesStatus  = ('True')
        else:
            (caseNotesStatus) = ('False')            
        read_text()
    elif args.logparse:
        log_type = ('file')
        if args.input:  # in case you don't want a different input file
            filename = args.input        
        parse_log() # parse image log
    elif args.logs_parse:
        log_type = ('folder')
        parse_log() # parse imager logs

    elif args.sticker:
        write_sticker() 
    elif args.guidataentry:
        guiDataEntry()
    else:
        if not args.input:  
            parser.print_help() 
            usage()
            return 0
    try:
        workbook.close()
    except:
        pass
    return 0


# <<<<<<<<<<<<<<<<<<<<<<<<<<   Sub-Routines   >>>>>>>>>>>>>>>>>>>>>>>>>>

def create_docx():
    '''
    if there isn't a template to read, 
    this creates an activity report from scratch
    '''
    global document
    document = docx.Document()
    
    caseNumber = "2022-0159"


    #header section
    section = document.sections[0]
    header = section.header
    header

    header.is_linked_to_previous = False
    # section.different_first_page_header_footer = True
    paragraph = header.paragraphs[0]
    paragraph.text = ("%s\n\nACTIVITY REPORT                                                             %s" %(agencyFull,divisionFull))


    p = document.add_paragraph('\n')    # start with a blank line   # todo this line is too thick
    p = document.add_paragraph('Activity Report:\t\t\t\tDate of Activity:')
    p = document.add_paragraph('%s\t\t\t\t%s' %(caseNumber, todaysDate))

    # insert a big line here

    # p = document.add_paragraph('Subject of Activity:\t\t\t\tCase Agent:\t\tType By:')
    # p = document.add_paragraph('%s\t\t\t\t%s\t\t%s' %(subjectBusinessName, caseAgent, forensicExaminer))

    document.save(outputFileXlsx)
    return document
    
def create_xlsx():  # BCI output (Default)
    '''
    Creates an xlsx database file with formatting
    '''
    
    global workbook
    workbook = xlsxwriter.Workbook(spreadsheet)
    global Sheet1
    Sheet1 = workbook.add_worksheet('Forensics')
    global Sheet2   
    Sheet2 = workbook.add_worksheet('Checklist')    # .set_landscape()
    
    header_format = workbook.add_format({'bold': True, 'border': 1})
    header_formatCase = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#FFc000'})   # orange Case items
    header_formatDescription = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'yellow'})   # yellow Description items
    header_formatNotes = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#CCCCFF'})   # purple Notes items
    header_formatCustody = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#92D050'})   # green Custody items
    header_formatAcquisition = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#66CCFF'})   # blue Acquisition items
    header_formatExtra = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#FF99FF'})   # pink Extra items

    header_checklist = workbook.add_format({'bold': True, 'border': 1})
    header_checklist.set_rotation(45)

    header_checklist2 = workbook.add_format({'bold': True, 'border': 1, 'bg_color':'#FFc000'})
    header_checklist2.set_rotation(45)


    global cell_format
    cell_format = workbook.add_format({'bold': False, 'border': 1})


    Sheet1.freeze_panes(1, 2)  # Freeze cells
    Sheet1.set_selection('B2')
    
    Sheet2.freeze_panes(7, 1)  # Freeze cells
    Sheet2.set_selection('D2')    

    # Excel column width
    Sheet1.set_column(0, 0, 15) # caseNumber
    Sheet1.set_column(1, 1, 7) # exhibit
    Sheet1.set_column(2, 2, 16) # caseName
    Sheet1.set_column(3, 3, 20) # subjectBusinessName
    Sheet1.set_column(4, 4, 16) # caseType
    Sheet1.set_column(5, 5, 25) # caseAgent
    Sheet1.set_column(6, 6, 15) # forensicExaminer
    Sheet1.set_column(7, 7, 13) # reportStatus
    Sheet1.set_column(8, 8, 25) # notes
    Sheet1.set_column(9, 9, 15) # summary
    Sheet1.set_column(10, 10, 12) # exhibitType
    Sheet1.set_column(11, 11, 30) # makeModel
    Sheet1.set_column(12, 12, 17) # serial
    Sheet1.set_column(13, 13, 15) # OS
    Sheet1.set_column(14, 14, 14) # phoneNumber
    Sheet1.set_column(15, 15, 16) # phoneIMEI
    Sheet1.set_column(16, 16, 15) # mobileCarrier
    Sheet1.set_column(17, 17, 16) # biosTime
    Sheet1.set_column(18, 18, 16) # currentTime
    Sheet1.set_column(19, 19, 12) # timezone
    Sheet1.set_column(20, 20, 15) # shutdownMethod
    Sheet1.set_column(21, 21, 16) # shutdownTime
    Sheet1.set_column(22, 22, 12) # userName
    Sheet1.set_column(23, 23, 12) # userPwd
    Sheet1.set_column(24, 24, 20) # email
    Sheet1.set_column(25, 25, 12) # emailPwd
    Sheet1.set_column(26, 26, 14) # ip
    Sheet1.set_column(27, 27, 15) # seizureAddress
    Sheet1.set_column(28, 28, 12) # seizureRoom
    Sheet1.set_column(29, 29, 16) # dateSeized
    Sheet1.set_column(30, 30, 12) # seizedBy
    Sheet1.set_column(31, 31, 16) # dateReceived
    Sheet1.set_column(32, 32, 15) # receivedBy
    Sheet1.set_column(33, 33, 16) # removalDate
    Sheet1.set_column(34, 34, 25) # removalStaff
    Sheet1.set_column(35, 35, 18) # reasonForRemoval
    Sheet1.set_column(36, 36, 15) # inventoryDate
    Sheet1.set_column(37, 37, 18) # seizureStatus
    Sheet1.set_column(38, 38, 12) # status
    Sheet1.set_column(39, 39, 24) # imagingTool
    Sheet1.set_column(40, 40, 15) # imagingType
    Sheet1.set_column(41, 41, 16) # imageMD5
    Sheet1.set_column(42, 42, 15) # imageSHA1
    Sheet1.set_column(43, 43, 15) # imageSHA256  #25
    Sheet1.set_column(44, 44, 15) # writeBlocker
    Sheet1.set_column(45, 45, 22) # imagingStarted
    Sheet1.set_column(46, 46, 16) # imagingFinished
    Sheet1.set_column(47, 47, 13) # storageType
    Sheet1.set_column(48, 48, 23) # storageMakeModel
    Sheet1.set_column(49, 49, 19) # storageSerial
    Sheet1.set_column(50, 50, 14) # storageSize
    Sheet1.set_column(51, 51, 15) # evidenceDataSize
    Sheet1.set_column(52, 52, 23) # analysisTool
    Sheet1.set_column(53, 53, 15) # analysisTool2
    Sheet1.set_column(54, 54, 25) # exportLocation
    Sheet1.set_column(55, 55, 15) # exportedEvidence
    Sheet1.set_column(56, 56, 20) # storageLocation
    Sheet1.set_column(57, 57, 19) # caseNumberOrig
    Sheet1.set_column(58, 58, 9) # priority
    Sheet1.set_column(59, 59, 15) # operation
    Sheet1.set_column(60, 60, 10) # Action
    Sheet1.set_column(61, 61, 19) # vaultCaseNumber
    Sheet1.set_column(62, 62, 15) # qrCode
    Sheet1.set_column(63, 63, 15) # vaultTotal
    Sheet1.set_column(64, 64, 40) # tempNotes

    Sheet2.set_column(0, 0, 5) # exhibit
    Sheet2.set_column(1, 1, 7) # type
    Sheet2.set_column(2, 2, 3) # 
    Sheet2.set_column(3, 3, 3) # 
    Sheet2.set_column(4, 4, 3) # 
    Sheet2.set_column(5, 5, 3) # 
    Sheet2.set_column(6, 6, 3) # 
    Sheet2.set_column(7, 7, 3) # 
    Sheet2.set_column(8, 8, 3) # 
    Sheet2.set_column(9, 9, 3) # 
    Sheet2.set_column(10, 10, 3) # 
    Sheet2.set_column(11, 11, 3) # 
    Sheet2.set_column(12, 12, 3) 
    Sheet2.set_column(13, 13, 9) 
    Sheet2.set_column(14, 14, 24) # case details
    
    Sheet2.set_row(6, 100)  # Set the height of Row 1 to 100.

    # hidden columns
    Sheet1.set_column(57, 57, None, None, {'hidden': 1}) # caseNumberOrig
    Sheet1.set_column(58, 58, None, None, {'hidden': 1}) # priority
    Sheet1.set_column(59, 59, None, None, {'hidden': 1}) # operation
    Sheet1.set_column(60, 60, None, None, {'hidden': 1}) # Action
    Sheet1.set_column(61, 61, None, None, {'hidden': 1}) # vaultCaseNumber
    Sheet1.set_column(62, 62, None, None, {'hidden': 1}) # qrCode
    Sheet1.set_column(63, 63, None, None, {'hidden': 1}) # vaultTotal
    
    # Write column headers

    Sheet1.write(0, 0, 'caseNumber', header_formatCase)
    Sheet1.write(0, 1, 'exhibit', header_formatDescription)
    Sheet1.write(0, 2, 'caseName', header_formatCase)
    Sheet1.write(0, 3, 'subjectBusinessName', header_formatCase)
    Sheet1.write(0, 4, 'caseType', header_formatCase)
    Sheet1.write(0, 5, 'caseAgent', header_formatCase)
    Sheet1.write(0, 6, 'forensicExaminer', header_formatCase)
    Sheet1.write(0, 7, 'reportStatus', header_formatCase)
    Sheet1.write(0, 8, 'notes', header_formatNotes)
    Sheet1.write(0, 9, 'summary', header_formatNotes)
    Sheet1.write(0, 10, 'exhibitType', header_formatDescription)
    Sheet1.write(0, 11, 'makeModel', header_formatDescription)
    Sheet1.write(0, 12, 'serial', header_formatDescription)
    Sheet1.write(0, 13, 'OS', header_formatDescription)
    Sheet1.write(0, 14, 'phoneNumber', header_formatDescription)
    Sheet1.write(0, 15, 'phoneIMEI', header_formatDescription)
    Sheet1.write(0, 16, 'mobileCarrier', header_formatDescription)
    Sheet1.write(0, 17, 'biosTime', header_formatDescription)
    Sheet1.write(0, 18, 'currentTime', header_formatDescription)
    Sheet1.write(0, 19, 'timezone', header_formatDescription)
    Sheet1.write(0, 20, 'shutdownMethod', header_formatDescription)
    Sheet1.write(0, 21, 'shutdownTime', header_formatDescription)
    Sheet1.write(0, 22, 'userName', header_formatDescription)
    Sheet1.write(0, 23, 'userPwd', header_formatDescription)
    Sheet1.write(0, 24, 'email', header_formatDescription)
    Sheet1.write(0, 25, 'emailPwd', header_formatDescription)
    Sheet1.write(0, 26, 'ip', header_formatDescription)
    Sheet1.write(0, 27, 'seizureAddress', header_formatCustody)
    Sheet1.write(0, 28, 'seizureRoom', header_formatCustody)
    Sheet1.write(0, 29, 'dateSeized', header_formatCustody)
    Sheet1.write(0, 30, 'seizedBy', header_formatCustody)
    Sheet1.write(0, 31, 'dateReceived', header_formatCustody)
    Sheet1.write(0, 32, 'receivedBy', header_formatCustody)
    Sheet1.write(0, 33, 'removalDate', header_formatCustody)
    Sheet1.write(0, 34, 'removalStaff', header_formatCustody)
    Sheet1.write(0, 35, 'reasonForRemoval', header_formatCustody)
    Sheet1.write(0, 36, 'inventoryDate', header_formatCustody)
    Sheet1.write(0, 37, 'seizureStatus', header_formatCustody)
    Sheet1.write(0, 38, 'status', header_formatAcquisition)
    Sheet1.write(0, 39, 'imagingTool', header_formatAcquisition)
    Sheet1.write(0, 40, 'imagingType', header_formatAcquisition)
    Sheet1.write(0, 41, 'imageMD5', header_formatAcquisition)
    Sheet1.write(0, 42, 'imageSHA1', header_formatAcquisition)
    Sheet1.write(0, 43, 'imageSHA256', header_formatAcquisition)
    Sheet1.write(0, 44, 'writeBlocker', header_formatAcquisition)
    Sheet1.write(0, 45, 'imagingStarted', header_formatAcquisition)
    Sheet1.write(0, 46, 'imagingFinished', header_formatAcquisition)
    Sheet1.write(0, 47, 'storageType', header_formatAcquisition)
    Sheet1.write(0, 48, 'storageMakeModel', header_formatAcquisition)
    Sheet1.write(0, 49, 'storageSerial', header_formatAcquisition)
    Sheet1.write(0, 50, 'storageSize', header_formatAcquisition)
    Sheet1.write(0, 51, 'evidenceDataSize', header_formatAcquisition)
    Sheet1.write(0, 52, 'analysisTool', header_formatAcquisition)
    Sheet1.write(0, 53, 'analysisTool2', header_formatAcquisition)
    Sheet1.write(0, 54, 'exportLocation', header_formatAcquisition)
    Sheet1.write(0, 55, 'exportedEvidence', header_formatAcquisition)
    Sheet1.write(0, 56, 'storageLocation', header_formatAcquisition)
    Sheet1.write(0, 57, 'caseNumberOrig', header_formatExtra)
    Sheet1.write(0, 58, 'priority', header_formatExtra)
    Sheet1.write(0, 59, 'operation', header_formatExtra)
    Sheet1.write(0, 60, 'Action', header_formatExtra)
    Sheet1.write(0, 61, 'vaultCaseNumber', header_formatExtra)
    Sheet1.write(0, 62, 'qrCode', header_formatExtra)
    Sheet1.write(0, 63, 'vaultTotal', header_formatExtra) # redundant with exhibit
    Sheet1.write(0, 64, 'tempNotes', header_format)

    Sheet2.write(6, 0, 'exhibit#', header_checklist)
    Sheet2.write(6, 1, 'type', header_checklist)
    Sheet2.write(6, 2, 'evidence sheet (in)', header_checklist2)
    Sheet2.write(6, 3, 'evidence sheet (out)', header_checklist2)
    Sheet2.write(6, 4, 'label (all separate pieces)', header_checklist)
    Sheet2.write(6, 5, 'imaged', header_checklist)
    Sheet2.write(6, 6, 'image backup', header_checklist)
    Sheet2.write(6, 7, 'analyzed', header_checklist)
    Sheet2.write(6, 8, 'report (sign, print, forward)', header_checklist2)
    Sheet2.write(6, 9, 'case notes printed', header_checklist2)
    Sheet2.write(6, 10, 'digital evidence', header_checklist)
    Sheet2.write(6, 11, 'digital evidence backup', header_checklist)
    Sheet2.write(6, 12, 'digital evidence to agent', header_checklist)

    Sheet2.write(1, 13, 'case#', header_format)
    Sheet2.write(2, 13, 'caseName', header_format)
    Sheet2.write(3, 13, 'subject', header_format)
    Sheet2.write(4, 13, 'agent', header_format)
    Sheet2.write(5, 13, 'forensics', header_format)


def enter_data():
    # accepted = accept_var.get()
    if 1==1:
    # if accepted=="Accepted":
        # (caseNumber, caseName) = ('', '')
        (exhibit, subjectBusinessName, caseType, caseAgent) = ('', '', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')


        caseNumber = caseNumber_entry.get()
        caseName = caseName_entry.get()
        
        if 1==1:
        # if caseNumber and caseName:
            # Case
            subjectBusinessName = subjectBusinessName_entry.get()
            caseAgent = caseAgent_combobox.get()
            forensicExaminer = forensicExaminer_combobox.get()
            
            # Description
            exhibit = exhibit_entry.get()
            makeModel = makeModel_entry.get()
            serial = serial_entry.get()
            exhibitType = exhibitType_combobox.get()
            phoneNumber = phoneNumber_entry.get()
            phoneIMEI = phoneIMEI_entry.get()
            userName = userName_entry.get()
            userPwd = userPwd_entry.get()

            # Lab Chain of Custody
            seizureAddress = seizureAddress_entry.get()
            seizureRoom = seizureRoom_entry.get()
            dateSeized = dateSeized_entry.get()
            # seizedBy = seizedBy_entry.get()
            seizedBy = seizedBy_combobox.get()
            # dateReceived = dateReceived_entry.get()

            # notes
            tempNotes = tempNotes_entry.get()


            # print out sticker format
            print("Case:", caseNumber, "Ex:", exhibit)
            print("CaseName:", caseName)
            print("Subject:", subjectBusinessName)
            print("Make:", makeModel)
            print("Serial:", serial)
            print("Agent:", caseAgent)
            # if status != '':
                # print(status)
            print("------------------------")

            filepath = "ForensicCasesTemp2.xlsx"
            
            if not os.path.exists(filepath):
                workbook = openpyxl.Workbook()
                sheet = workbook.active

                heading = ["caseNumber", "exhibit", "caseName", "subjectBusinessName",
                    "caseType", "caseAgent", "forensicExaminer", "reportStatus", "notes",
                    "summary", "exhibitType", "makeModel", "serial", "OS", "phoneNumber",
                    "phoneIMEI", "mobileCarrier", "biosTime", "currentTime", "timezone",
                    "shutdownMethod", "shutdownTime", "userName", "userPwd", "email",
                    "emailPwd", "ip", "seizureAddress", "seizureRoom", "dateSeized",
                    "seizedBy", "dateReceived", "receivedBy", "removalDate", "removalStaff",
                    "reasonForRemoval", "inventoryDate", "seizureStatus", "status", "imagingTool",
                    "imagingType", "imageMD5", "imageSHA1", "imageSHA256", "writeBlocker",
                    "imagingStarted", "imagingFinished", "storageType", "storageMakeModel",
                    "storageSerial", "storageSize", "evidenceDataSize", "analysisTool",
                    "analysisTool2", "exportLocation", "exportedEvidence", "storageLocation",
                    "caseNumberOrig", "priority", "operation", "Action", "vaultCaseNumber",
                    "qrCode", "vaultTotal", "tempNotes"]

                sheet.append(heading)
                workbook.save(filepath)
            workbook = openpyxl.load_workbook(filepath)
            sheet = workbook.active

            sheet.append([caseNumber, exhibit, caseName, subjectBusinessName,
                    caseType, caseAgent, forensicExaminer, reportStatus, notes,
                    summary, exhibitType, makeModel, serial, OS, phoneNumber,
                    phoneIMEI, mobileCarrier, biosTime, currentTime, timezone,
                    shutdownMethod, shutdownTime, userName, userPwd, email,
                    emailPwd, ip, seizureAddress, seizureRoom, dateSeized,
                    seizedBy, dateReceived, receivedBy, removalDate, removalStaff,
                    reasonForRemoval, inventoryDate, seizureStatus, status, imagingTool,
                    imagingType, imageMD5, imageSHA1, imageSHA256, writeBlocker,
                    imagingStarted, imagingFinished, storageType, storageMakeModel,
                    storageSerial, storageSize, evidenceDataSize, analysisTool,
                    analysisTool2, exportLocation, exportedEvidence, storageLocation,
                    caseNumberOrig, priority, operation, Action, vaultCaseNumber,
                    qrCode, vaultTotal, tempNotes])




            workbook.save(filepath)
                
        else:
            tkinter.messagebox.showwarning(title="Error", message="Case Number and Case Name are required.")
    else:
        tkinter.messagebox.showwarning(title= "Error", message="You have not verified the info")


def FormatFunction(bg_color = 'white'):
	global Format
	Format=workbook.add_format({
	'bg_color' : bg_color
	}) 

    
def dictionaryBuild(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
    forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
    phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
    userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
    dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
    seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
    writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
    storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
    storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
    vaultTotal, tempNotes):    
    '''
    build a dictionary file of important columns for writing to a pdf
    '''
    
    my_dict = {}
    my_dict['caseNumber']=caseNumber
    my_dict['caseName']=caseName
    my_dict['subjectBusinessName']=subjectBusinessName
    my_dict['caseAgent']=caseAgent
    my_dict['forensicExaminer']=forensicExaminer
    my_dict['exhibit']=exhibit
    my_dict['makeModel']=makeModel
    my_dict['serial']=serial
    my_dict['OS']=OS
    my_dict['ip']=ip
    my_dict['exhibitType']=exhibitType   
    my_dict['phoneNumber']=phoneNumber
    my_dict['phoneIMEI']=phoneIMEI
    my_dict['userName']=userName
    my_dict['userPwd']=userPwd
    my_dict['email']=email
    my_dict['emailPwd']=emailPwd
    my_dict['biosTime']=biosTime
    my_dict['currentTime']=currentTime
    my_dict['priority']=priority
    my_dict['timezone']=timezone
    my_dict['seizureAddress']=seizureAddress
    my_dict['seizureRoom']=seizureRoom
    my_dict['dateSeized']=dateSeized
    my_dict['seizedBy']=seizedBy
    my_dict['seizureStatus']=seizureStatus
    my_dict['dateReceived']=dateReceived
    my_dict['receivedBy']=receivedBy
    my_dict['removalDate ']=removalDate 
    my_dict['removalStaff']=removalStaff
    my_dict['imagingTool']=imagingTool
    my_dict['imagingType']=imagingType
    my_dict['imageMD5']=imageMD5
    my_dict['writeBlocker']=writeBlocker
    my_dict['imagingStarted']=imagingStarted
    my_dict['imagingFinished']=imagingFinished
    my_dict['imageSHA256']=imageSHA256
    my_dict['storageType']=storageType
    my_dict['storageMakeModel']=storageMakeModel
    my_dict['storageSerial']=storageSerial
    my_dict['storageSize']=storageSize
    my_dict['analysisTool']=analysisTool
    my_dict['analysisTool2']=analysisTool2
    my_dict['notes']=notes

    return (my_dict)
    
def format_function(bg_color='white'):
    '''
    xlsx color formatting
    currently just does white
    '''
    global format
    format = workbook.add_format({
        'bg_color': bg_color
    })

def fix_date(date):
    '''
    standardize date formatting, Tableau
    '''
    
    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()
    date = date.replace("  ", " ")  # test
    # date = ('%s      ' %(date)) # test
    date = date.split(' ')     # Fri Jun 04 07:55:41 2021
    mo = date[1]    # convert month to a number
    mo = mo.replace("Jan", "1").replace("Feb", "2").replace("Mar", "3").replace("Apr", "4")
    mo = mo.replace("May", "5").replace("Jun", "6").replace("Jul", "7").replace("Aug", "8")
    mo = mo.replace("Sep", "9").replace("Oct", "10").replace("Nov", "11").replace("Dec", "12")
    try:    
        dy = date[2].lstrip('0')
    except TypeError as error:
        print(error)
    try:
        tm = date[3].lstrip('0')
        # tm = date[3]
    except TypeError as error:
        print(error)
    yr = date[4]
    date = ('%s/%s/%s %s' %(mo, dy, yr, tm))  # 3/4/2021 9:17
    return date

def fix_date2(date):
    '''
    standardize date formatting
    2022-07-14 21:15:11
    
    31/07/2022 11:48:57 (-5)
 
    '''
    print('fix_date2')  # temp
    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()

    return date

def fix_date3(date):
    '''
    standardize date formatting from Cellebrite
    
    31/07/2022 11:48:57 (-5) to 7/31/2022 11:48

    '''

    (mo, dy, yr, tm) = ('', '', '', '')
    date = date.strip()
    # date = date.replace("  ", " ")  # test
    # date = ('%s      ' %(date)) # test
    date = date.split(' ')     # 31/07/2022 11:48:57 (-5)
    tempDate = date[0]
    tm = date[1]
    tempDate = tempDate.split('/')
    dy = tempDate[0]    
    mo = tempDate[1]
    yr = tempDate[2]

    date = ('%s/%s/%s %s' %(mo, dy, yr, tm)).lstrip('0')  # 3/4/2021 9:17
    print('fix_date_, %s %s' %(date, tempDate))  # temp

    return date

def guiDataEntry():
    win = Frame()
    # win.title('Evidence Form')  # test
    # win = Frame().title("Evidence form")  # todo
    win.grid(sticky=N+S+E+W) 
    # win.geometry("700x350") # test
    # Frame labels
    case_info_frame = LabelFrame(win, text='Case', padx=5, pady=5)
    description_frame = LabelFrame(win, text='Description', padx=5, pady=5)
    custody_frame = LabelFrame(win, text='Chain of Custody', padx=5, pady=5)
    notes_frame = LabelFrame(win, text='Notes', padx=5, pady=5)
    
    # stick notes
    case_info_frame.grid(sticky=W+W)
    description_frame.grid(sticky=E+W)
    custody_frame.grid(sticky=E+W)
    notes_frame.grid(sticky=E+W)   

    ## create multiple frames

    for frame in case_info_frame, description_frame, custody_frame, notes_frame:
        for col in 0, 1, 2:
            frame.columnconfigure(col, weight=1)


    ########## case section ##########

    ## row 0 labels (caseNumber, caseNumber, subjectBusinessName)
    caseNumber_label = tkinter.Label(case_info_frame, text="Case Number")
    caseNumber_label.grid(row=0, column=0)

    caseName_label = tkinter.Label(case_info_frame, text="Case Name")    
    caseName_label.grid(row=0, column=1)

    subjectBusinessName_label = tkinter.Label(case_info_frame, text="Subject or d/b/a")    
    subjectBusinessName_label.grid(row=0, column=2)

    ## global section
    global caseNumber_entry
    global caseName_entry
    global subjectBusinessName_entry
    global exhibit_entry
    global makeModel_entry
    global serial_entry
    global exhibitType_entry
    global phoneNumber_entry
    global phoneIMEI_entry
    global userName_entry
    global userPwd_entry
    global seizureAddress_entry
    global seizureRoom_entry
    global dateSeized_entry
    # global seizedBy_entry
    global seizedBy_combobox    
    global dateReceived_entry
    global tempNotes_entry
    global caseAgent_combobox
    global forensicExaminer_combobox
    global exhibitType_combobox


    ## row 1 Entry (caseNumber, caseNumber, subjectBusinessName)
    caseNumber_entry = tkinter.Entry(case_info_frame)
    caseName_entry = tkinter.Entry(case_info_frame)
    subjectBusinessName_entry = tkinter.Entry(case_info_frame, width = 24)
    caseNumber_entry.grid(row=1, column=0)
    caseName_entry.grid(row=1, column=1)    # test
    subjectBusinessName_entry.grid(row=1, column=2)

    ## row 2 & 3 label (caseAgent, forensicExaminer)
    caseAgent_label = tkinter.Label(case_info_frame, text="Case Agent")
    caseAgent_combobox = ttk.Combobox(case_info_frame, values=["", "SA April Moore", "Road Runner", "Sherlock Holmes"])
    caseAgent_label.grid(row=2, column=0)
    caseAgent_combobox.grid(row=3, column=0)

    forensicExaminer_label = tkinter.Label(case_info_frame, text="Forensic Examiner")   # works
    forensicExaminer_combobox = ttk.Combobox(case_info_frame, values=["", "Sherlock Holmes", "Elliott Ness"], width = 24)
    forensicExaminer_label.grid(row=2, column=1)    # was 0, 2
    forensicExaminer_combobox.grid(row=3, column=1) # works

    ########## Description section ##########

    ## row 0 label (exhibit, makeModel, serial)
    exhibit_label = tkinter.Label(description_frame, text="Exhibit")  
    exhibit_label.grid(row=0, column=0)

    makeModel_label = tkinter.Label(description_frame, text="Make/Model")    
    makeModel_label.grid(row=0, column=1)

    serial_label = tkinter.Label(description_frame, text="Serial #")    
    serial_label.grid(row=0, column=2)

    ## row 1 entry (exhibit, makeModel, serial)
    exhibit_entry = tkinter.Entry(description_frame, width = 7)
    # exhibit_entry = tkinter.Entry(description_frame, width= 40) # sets windows frame size

    exhibit_entry.grid(row=1, column=0)

    makeModel_entry = tkinter.Entry(description_frame)
    makeModel_entry.grid(row=1, column=1)

    serial_entry = tkinter.Entry(description_frame)
    serial_entry.grid(row=1, column=2)

    ## row 2 label (exhibitType, phoneNumber, phoneIMEI)

    exhibitType_label = tkinter.Label(description_frame, text="Exhibit Type")
    exhibitType_combobox = ttk.Combobox(description_frame, values=["", "DVR", "desktop", "laptop", "phone", "POS", "router", "server", "switch" , "vehicle"], width = 10)
    exhibitType_label.grid(row=2, column=0)

    phoneNumbe_label = tkinter.Label(description_frame, text="Phone Number")    
    phoneNumbe_label.grid(row=2, column=1)

    phoneIMEI = tkinter.Label(description_frame, text="Phone IMEI #")    
    phoneIMEI.grid(row=2, column=2)


    ## row 3 entry (exhibitType, phoneNumber, phoneIMEI)
    ## exhibitType_entry = tkinter.Entry(description_frame)
    ## exhibitType_combobox = ttk.Combobox(description_frame, values=["", "DVR", "desktop", "phone", "POS", "router", "server", "switch" , "vehicle"])

    ## exhibitType_entry.grid(row=8, column=0)
    exhibitType_combobox.grid(row=3, column=0)

    phoneNumber_entry = tkinter.Entry(description_frame, width = 16)
    phoneNumber_entry.grid(row=3, column=1)

    phoneIMEI_entry = tkinter.Entry(description_frame, width = 17)
    phoneIMEI_entry.grid(row=3, column=2)

    ## row 4 label (userName, userPwd)
    userName_label = tkinter.Label(description_frame, text="User Name")    
    userName_label.grid(row=4, column=0)
    userPwd_label = tkinter.Label(description_frame, text="User Password")    
    userPwd_label.grid(row=4, column=1)

    ## row 5 entry (userName, userPwd)
    userName_entry = tkinter.Entry(description_frame)
    userName_entry.grid(row=5, column=0)

    userPwd_entry = tkinter.Entry(description_frame)
    userPwd_entry.grid(row=5, column=1)

    ########## chain of custody section ##########

    ## row 0 label (seizureAddress, seizureRoom)
    seizureAddress_label = tkinter.Label(custody_frame, text="Seizure Address")    
    seizureAddress_label.grid(row=0, column=0)

    seizureRoom_label = tkinter.Label(custody_frame, text="Seizure Room")    
    seizureRoom_label.grid(row=0, column=1)

    ## row 1 entry (seizureAddress, seizureRoom)
    seizureAddress_entry = tkinter.Entry(custody_frame, width = 50)    # test
    seizureRoom_entry = tkinter.Entry(custody_frame, width = 5)
    seizureAddress_entry.grid(row=1, column=0)
    seizureRoom_entry.grid(row=1, column=1)


    ## row 2 label (dateSeized, seizedBy, dateReceived)
    dateSeized_label = tkinter.Label(custody_frame, text="Date Seized")    
    dateSeized_label.grid(row=2, column=0)


    # dateReceived = tkinter.Label(custody_frame, text="Date Received")    
    # dateReceived.grid(row=2, column=2)

    ## row 3 entry (dateSeized, seizedBy, dateReceived)
    dateSeized_entry = tkinter.Entry(custody_frame, width = 19)
    dateSeized_entry.grid(row=3, column=0)

    seizedBy_label = tkinter.Label(custody_frame, text="Seized_By")     
    seizedBy_label.grid(row=2, column=1)
    seizedBy_entry = tkinter.Entry(custody_frame, width = 25)
    seizedBy_entry.grid(row=3, column=1)
    seizedBy_combobox = ttk.Combobox(case_info_frame, values=["", "SA Herby Hancock", "SSA John Doe", "Sherlock Holmes"])


    # dateReceived_entry = tkinter.Entry(custody_frame)
    # dateReceived_entry.grid(row=3, column=2)

    ########## tempNotes section ##########

    
    ## row 0 label (tempNotes)
    tempNotes_label = tkinter.Label(notes_frame, text="Temp Notes")  
    tempNotes_label.grid(row=0, column=0)

    ## row 1 entry (tempNotes)
    tempNotes_entry = tkinter.Entry(notes_frame, width = 75)
    tempNotes_entry.grid(row=1, column=0)

    ## Button
    button = tkinter.Button(frame, text="Enter Data", command= enter_data, bg='#0052cc', fg='#ffffff')
    # button = tkinter.Button(frame, text="Enter Data", command= enter_data_DB, bg='#0052cc', fg='#ffffff')

    button.grid(row=3, column=0, sticky="news", padx=20, pady=10)

    win.mainloop()
    
def parse_log():
    '''
    parse tableau, recon imager, cellebrite triage_windows.cmd and FTK logs
    '''

    import os
    (caseNumber, caseName, exhibit) = ('', '', '')
    if log_type == 'file':  # only ask for exhibit number if it's a single log
        if inputDetails == "yes":
            caseNumber = str(input("caseNumber : ")).strip()
            caseName = str(input("caseName : ")).strip()
            exhibit = str(input("exhibit : ")).strip()
        logsList = [filename]
    elif log_type == 'folder':
        print('Reading logs from %s folder' %(logsFolder))
        if inputDetails == "yes":
            caseNumber = str(input("caseNumber : ")).strip()
            caseName = str(input("caseName : ")).strip()
        logsList = os.listdir(logsFolder)
        logsList2 = []
        for logFile in logsList:
            logFile = ("%s%s" %(logsFolder, logFile))
            logsList2.append(logFile)
            # print('logsFolder = %s  logFile = %s   logsList = %s    logsList2 = %s' %(logsFolder, logFile, logsList, logsList2)) # temp
        logsList = logsList2

        # read section
    for logFile in logsList:
        print('<<<<<< %s >>>>>>' %(logFile))
        style = workbook.add_format()
        (header, reportStatus, date) = ('', '', '<insert date here>')

        # (caseNumber, exhibit, caseName) = ('', '', '')
        (subjectBusinessName, caseType, caseAgent) = ('', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')
        
        # bonus variables
        (vehicleYear, vehicleManufacturer, vehicleModel) = ('', '', '') # BerlaIVe Acquisition
        (imagingTool1, imagingTool2, make, model) = ('', '', '', '')

        if logFile.lower().endswith('.pdf'):
            # print('Can\'t process .pdf files at this time, sorry: %s' %(logFile))
            csv_file = ''
            # tempNotes = ("failed %s" %(logFile))
            (forensicExaminer, exhibit, exhibitType, makeModel, serial, OS, phoneNumber, phoneIMEI, email, status, imagingType, imageMD5, imageSHA256, imagingStarted, imagingFinished, imagingTool, storageSize, evidenceDataSize, analysisTool, tempNotes) = pdfExtract(logFile)
            csv_file = tempNotes.split('\\n')
        else:
            csv_file = open(logFile)
            # csv_file = open(logFile, encoding='utf8')
            
            # with open(logFile, 'rb') as f:
              # csv_file = f.read()
        
        for each_line in csv_file:
            print(each_line)    # temp
            if "Task:" in each_line:
                imagingType = re.split("Task: ", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()

            elif " Extraction type " in each_line: #cellebrite xls
                imagingType = re.split(" Extraction type ", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()
            elif "Source Type: Physical" in each_line:
                imagingType = "disk to file"
            elif "Image type :" in each_line: #recon imager
                imagingType = re.split("Image type :", each_line, 0)
                imagingType = str(imagingType[1]).strip().lower()

            elif "ExtractionType=" in each_line: #cellebrite *.ufd log file
                imagingType = each_line.replace("ExtractionType=", "").strip()
                if imagingType == "AdvancedLogical":
                    imagingType = "advanced logical"
                elif imagingType == "Logical":
                    imagingType = "logical"            


            elif "Start of Tableau Imager" in each_line: # tableau imager
                imagingTool = each_line.replace("Start of ", "").strip()



            # status
            elif "Status: Ok" in each_line or "Imaging Status : Successful" in each_line:
                status = 'Imaged'
            elif "Status: Error/Failed" in each_line:
                status = 'Not imaged'
            elif "Acquisition Successfully Completed" in each_line: # BerlaIVe AcquisitionLog_001.txt
                status = "Imaged"

            elif "Acquisition Failed To Complete" in each_line: # BerlaIVe AcquisitionLog.txt
                status = "Not imaged"

            # exhibit      
            elif "Evidence Number: " in each_line:      #FTK_parse, magnet
                exhibit = re.split("Evidence Number: ", each_line, 0)
                exhibit = str(exhibit[1]).strip()
            elif "Exhibit#" in each_line:      #cellebrite
                exhibit = re.split("Exhibit#", each_line, 0)
                exhibit = str(exhibit[1]).strip()
            elif "Exhibit Number=" in each_line: # CellebriteUFED4PC.txt
                makeModel = each_line.replace("Exhibit Number=", "").strip()

            elif "Evidence Number" in each_line:      #recon imager
                exhibit = re.split("Evidence Number", each_line, 0)

                # exhibit = re.split("Evidence Number     :", each_line, 0)
                exhibit = str(exhibit[1]).replace(":", "").strip()

            # makeModel
            elif "Unique description: " in each_line:
                makeModel = re.split("Unique description: ", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Device    " in each_line: #cellebrite excel
                makeModel = re.split("Device    ", each_line, 0)
                makeModel = str(makeModel[1]).strip()
            elif "Selected device name" in each_line: #cellebrite
                makeModel = re.split("Selected device name", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Selected Model:" in each_line:
                makeModel = re.split("Selected Model:", each_line, 0)
                makeModel = str(makeModel[1]).strip()

            elif "Case Name: " in each_line: # BerlaIVe AcquisitionLog_001.txt      # broken
                makeModel = each_line.replace("Case Name: ", "").strip()

            elif "Model=" in each_line: #cellebrite *.ufd log file
                model = each_line.replace("Model=", "").strip()

            elif "Vendor=" in each_line: #cellebrite *.ufd log file
                make = each_line.replace("Vendor=", "").strip()

            elif "Device Model:" in each_line: # CellebritePremium DeviceInfo.txt
                model = each_line.replace("Device Model:", "").strip()

            elif "Vendor:" in each_line: # CellebritePremium DeviceInfo.txt
                make = each_line.replace("Vendor:", "").replace("Tableau", "").strip()
            elif "Model:" in each_line and len(storageMakeModel) == 0: # tableau
                storageMakeModel = re.split("Model:", each_line, 0)
                storageMakeModel = str(storageMakeModel[1]).strip()
            elif "Device / Media Name:" in each_line: # SumuriReconImager.txt
                storageMakeModel = each_line.replace("Device / Media Name:", "").strip()

            elif "Vehicle Year:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleYear = each_line.replace("Vehicle Year:", "").strip()

            elif "Vehicle Manufacturer:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleManufacturer = each_line.replace("Vehicle Manufacturer:", "").strip()

            elif "Vehicle Model:" in each_line: # BerlaIVe AcquisitionLog.txt
                vehicleModel = each_line.replace("Vehicle Model:", "").strip()

            # OS
            elif "Revision:" in each_line: #cellebite 
                os = re.split("Revision:", each_line, 0)
                os = str(os[1]).strip()
                if 'iPhone' in makeModel:
                    os = ('iOS %s' %(os))

            elif "OS=" in each_line: # cellebrite *.ufd log file
                OS = each_line.replace("OS=", "").strip()

            elif "OS Version:" in each_line: # CellebritePremium DeviceInfo.txt
                OS = each_line.replace("OS Version:", "").strip()

            elif "Operating System Version:" in each_line: #cellebrite *.ufd log file
                OS = each_line.replace("Operating System Version:", "").strip()

            elif "Operating System:" in each_line: # MagnetAXIOM Case Information.txt
                OS = each_line.replace("Operating System:", "").strip()

            elif "Vehicle ECU:" in each_line: # BerlaIVe AcquisitionLog.txt
                OS = each_line.replace("Vehicle ECU:", "").strip()
            elif "Android_ID=" in each_line: # Cellebrite .ufd
                OS = ('Android %s' %(OS))
            elif "Apple" in makeModel and OS != '': # Cellebrite .ufd    # test
                OS = ('iOS %s' %(OS))


                
            # serial

            # elif "Serial Number:" in each_line and serial != '': #cellebrite
            elif "Serial Number:" in each_line: #cellebrite
                # serial = each_line.replace("Serial Number:", "").strip()
                serial = re.split("Serial Number:", each_line, 0)
                      
                serial = str(serial[1]).strip()
                if "number: " in serial:
                    serial = ''
                storageSerial = serial
                serial = ''
                
            elif "Machine Serial" in each_line: #RECON imager
            # elif "Machine Serial" in each_line and serial != '': #RECON imager
                serial = re.split(":", each_line, 0)
                serial = str(serial[1]).strip()

            elif "Vehicle VIN:" in each_line: # BerlaIVe AcquisitionLog.txt
                serial = each_line.replace("Vehicle VIN:", "").strip()
                if serial == ('unknown'):
                    serial = ''
            # elif "Serial " in each_line and serial != '': #cellebrite
                # serial = re.split("Serial ", each_line, 0)
                # print("serial=",serial[1].strip())      
                # serial = str(serial[1]).strip()

            # storageSerial    
            # elif "Drive Serial Number:" in each_line: # FTKImager Image.E01.txt # fix me
                # storageSerial = each_line.replace("Drive Serial Number:", "").strip()

            elif "S/N:" in each_line: # TableauImager 000ecc45 0067205e
                storageSerial = each_line.replace("S/N:", "").strip()
            elif "Serial Number:" in each_line and storageSerial != '': # FTKImager Image.E01.txt # fix me
                storageSerial = each_line.replace("Drive Serial Number:", "").strip()
            elif "Serial number:" in each_line and storageSerial == '': # Tableau imager 
                storageSerial = each_line.replace("Drive Serial number:", "").replace("Serial number:", "").strip()


            elif "Unique Identifier:" in each_line: # MagnetAcquire image_info.txt
                storageSerial = each_line.replace("Unique Identifier:", "").strip() 

     
            # phoneIMEI
            elif "IMEI:" in each_line: # CellebritePremium DeviceInfo.txt
                phoneIMEI = each_line.replace("IMEI:", "").strip()

            elif "IMEI1=" in each_line: # CELLEBRITEPREMIUM EXTRACTION_FFS.TXT
                phoneIMEI = each_line.replace("IMEI1=", "").strip()

            # phoneNumber
            elif "MSISDN" in each_line: #cellebrite
                phoneNumber = re.split("MSISDN", each_line, 0)
                phoneNumber = str(phoneNumber[1]).strip()
                if ')' in phoneNumber:
                    phoneNumber = phoneNumber.replace("+1 (", "1-").replace(") ", "-")
                (exportedEvidence, status) = ('', 'Imaged')

            elif " Username" in each_line: #cellebrite xls
                phoneNumber = re.split(" Username", each_line, 0)
                phoneNumber = str(phoneNumber[1]).strip()
                if ')' in phoneNumber:
                    phoneNumber = phoneNumber.replace("+1 (", "1-").replace(") ", "-")
                (exportedEvidence, status) = ('', 'Imaged')
            elif "UserName=" in each_line: # CELLEBRITE AdvancedLogical.ufd
                phoneNumber = each_line.replace("UserName=", "").strip()
                # userName = phoneNumber

            # forensicExaminer
            elif "User: " in each_line:
                forensicExaminer = re.split("User: ", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
            elif "Examiner:" in each_line:
                forensicExaminer = re.split("Examiner:", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
                forensicExaminer =forensicExaminer.replace("CIA - ", "")
            elif "Examiner Name:" in each_line: # MagnetAcquire image_info_001.txt
                forensicExaminer = each_line.replace("Examiner Name:", "").strip()
            elif "Examiner Name=" in each_line: # CellebriteUFED4PC.txt
                forensicExaminer = each_line.replace("Examiner Name=", "").strip()

            elif "Examiner         :" in each_line: # recon imager
                forensicExaminer = re.split("Examiner         :", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
            elif "Examiner name" in each_line:  #cellebrite
                forensicExaminer = re.split("Examiner name", each_line, 0)
                forensicExaminer = str(forensicExaminer[1]).strip()
                forensicExaminer =forensicExaminer.replace("CIA - ", "")

            # caseNumber
            elif "Case ID:" in each_line:
                caseNumber = re.split("Case ID:", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")
            elif "Case Number:" in each_line:
                caseNumber = re.split("Case Number:", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")
            elif "Case Number=" in each_line:   # CellebriteUFED4PC.txt
                caseNumber = each_line.replace("Case Number=", "").strip()

            elif "Case Number         :" in each_line:   # Recon imager and probably tablaue
                caseNumber = each_line.replace("Case Number         :", "").strip()
                caseNumber = caseNumber.replace("<<not entered>>", "")




            elif "CaseNumber" in each_line:   #cellebrite
                caseNumber = re.split("CaseNumber", each_line, 0)
                caseNumber = str(caseNumber[1]).strip()


            # notes
            elif "Case Notes:" in each_line:    # Tableau logs
                notes = re.split("Case Notes:", each_line, 0)
                notes = str(notes[1]).strip()
                notes = notes.replace("<<not entered>>", "")
            elif "Notes: " in each_line:
                notes = re.split("Notes: ", each_line, 0)
                notes = str(notes[1]).strip()
                notes = notes.replace("<<not entered>>", "")

            elif "Notes         :" in each_line:    # recon imager
                notes = re.split("Notes         :", each_line, 0)
                notes = str(notes[1]).strip()

            elif "Source Device :" in each_line:    # recon imager
                (vol, partition, size, frmat) = ('', '', '', '')
                sourcenotes = re.split("Source Device :", each_line, 0)
                sourcenotes = str(sourcenotes[1]).strip()
                # append 7 spaces for fault tollerance  # to do
                details = re.split("  ", sourcenotes, 0)
                vol = str(details[0]).strip()
                
                try:
                    partition = str(details[3]).strip()    
                    size = str(details[4]).strip()  
                    frmat = str(details[5]).strip()  
                    blurb1 = ("This image was from %s and was the %s %s %s volume." %(vol, partition, size, frmat))
                    notes = ("%s %s" %(notes, blurb1))
                except:pass
                    
            # imagingTool
            elif "Imager App: " in each_line:
                imagingTool1 = re.split("Imager App: ", each_line, 0)
                imagingTool1 = str(imagingTool1[1]).strip()
            elif "Created By AccessData® FTK® Imager" in each_line:
                imagingTool1 = each_line.replace("Created By AccessData® FTK® Imager", "").replace("®", "").replace("Â", "").strip()



            elif "Created By AccessData" in each_line:
                imagingTool1 = each_line.replace("Created By AccessData", "").replace("®", "").strip()


                
            elif "Imager Ver: " in each_line:
                imagingTool2 = re.split("Imager Ver: ", each_line, 0)
                imagingTool2 = str(imagingTool2[1]).strip()

            elif "UFED Version:    Product Version: " in each_line:    #cellebrite
                imagingTool = re.split("UFED Version:    Product Version: ", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Cellebrite UFED %s' %(imagingTool))
            elif "UFED version" in each_line:    #cellebrite
                imagingTool = re.split("UFED version", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Cellebrite UFED %s' %(imagingTool))

            elif "Cellebrite Physical Analyzer version" in each_line: #cellebrite xls
                imagingTool1 = re.split("Cellebrite Physical Analyzer version", each_line, 0)
                imagingTool1 = str(imagingTool1[1]).strip()
                imagingTool1 = ('Cellebrite Physical Analyzer %s' %(imagingTool1))

            elif "RECON Imager Version : " in each_line:    # Recon Imager
                imagingTool = re.split("RECON Imager Version : ", each_line, 0)
                imagingTool = str(imagingTool[1]).strip()
                imagingTool = re.split(" ", imagingTool, 0)
                imagingTool = str(imagingTool[0]).strip()
                imagingTool = ('Recon Imager %s' %(imagingTool))

            elif "AcquisitionTool=" in each_line: #cellebrite *.ufd log file
                imagingTool = each_line.replace("AcquisitionTool=", "").strip()

            elif "Created by iVe " in each_line and "Acquisition finished" not in each_line: # BerlaIVe AcquisitionLog_001.txt 
                imagingTool = each_line.replace("Created by iVe ", "Berla iVe ").strip()
                if " built on" in imagingTool and "Acquisition finished" not in imagingTool:
                    imagingTool = imagingTool.split(' built on')[0].strip()

            elif "Imager Product:" in each_line: # MagnetAcquire image_info.txt
                imagingTool = each_line.replace("Imager Product:", "").strip() 

            elif "Imager Version:" in each_line: # MagnetAcquire image_info.txt
                imagingToolVer = each_line.replace("Imager Version:", "").strip() 
                imagingTool = ('%s %s' %(imagingTool, imagingToolVer))    

            elif "ExtractionSoftwareVersion=" in each_line: # CellebritePA_FFS.txt
                imagingToolTemp = each_line.replace("ExtractionSoftwareVersion=", "").strip()
                imagingTool = ('Cellebrite PA %s' %(imagingToolTemp))

            # storageSize
            elif "Capacity in bytes reported Pwr-ON: " in each_line: # todo swap storageSize from capacity
                capacity = re.split("Capacity in bytes reported Pwr-ON: ", each_line, 0)
                capacity = str(capacity[1]).strip()
                if "(" in capacity:
                    capacity = re.split("\(", each_line, 0)
                    capacity = str(capacity[1]).strip()
                    capcty = capacity.replace(")", "")
                    capacity = capcty.split('.')[0]
                    if ' ' in capcty:
                        size = capcty.split(' ')[1]
                    else:
                        size = ''
                    # size = capacity
                    capacity = ('%s %s' %(capacity, size))
                    storageSize = capacity

            elif "Source data size: " in each_line: # FTKImager Image.E01.txt
                storageSize = each_line.replace("Source data size: ", "").strip()
            elif "Device Size:" in each_line: # MagnetAcquire image_info.txt
                storageSize = each_line.replace("Device Size:", "").strip() 

            elif "Disk Size:" in each_line: # SumuriReconImager.txt
                storageSize = each_line.replace("Disk Size:", "").strip()
                storageSize = storageSize.split(' (')[0]


            # storageType
            elif "Cable/Interface type: " in each_line:
                storageType = re.split("Cable/Interface type: ", each_line, 0)
                storageType = str(storageType[1]).strip()
                storageType

            elif "Drive Interface Type: " in each_line: # FTKImager Image.E01.txt
                storageType = each_line.replace("Drive Interface Type: ", "").strip()

            elif "Media Type:" in each_line: # MagnetAcquire image_info.txt
                storageType = each_line.replace("Media Type:", "").strip() 

            elif "T356789iu" in each_line:  # fix me  T356789i
                writeBlocker = "Tableau T356789iu"
            elif "T356789i" in each_line:  # fix me  
                writeBlocker = "Tableau T356789i"

            elif "Source data size: " in each_line:
                capacity = re.split("Source data size: ", each_line, 0)
                capacity = str(capacity[1]).strip()
                if "(" in capacity:
                    capacity = re.split("\(", each_line, 0)
                    capacity = str(capacity[1]).strip()
                    capcty = capacity.replace(")", "")
                    capacity = capcty.split('.')[0]
                    if ' ' in capcty:
                        size = capcty.split(' ')[1]
                    # size = capacity
                    capacity = ('%s %s' %(capacity, size))
            
            # exportLocation        
            elif "Filename of first chunk: " in each_line:
                exportLocation = re.split("Filename of first chunk: ", each_line, 0)
                exportLocation = str(exportLocation[1]).strip()
            elif "Information for " in each_line:       # ftk_parse
                exportLocation = re.split("Information for ", each_line, 0)
                exportLocation = str(exportLocation[1]).strip()
            # elif "E01" in each_line:
                # exportLocation = each_line.strip()

            elif "FileDump=" in each_line: # CellebritePremium EXTRACTION_FFS.txt
                exportLocation = each_line.replace("FileDump=", "").strip()
            elif "File Path: " in each_line: # BerlaIVe AcquisitionLog_001.txt
                exportLocation = each_line.replace("File Path: ", "").strip()
                exportedEvidence = 'Y'

            # imageMD5
            elif "Disk MD5:  " in each_line:    # Tableau
                imageMD5 = re.split("Disk MD5:  ", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()


            elif "MD5 checksum:" in each_line:  # fix me
                imageMD5 = re.split("MD5 checksum:", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()
                imageMD5 = re.split(": ", each_line, 0)
                imageMD5 = str(imageMD5[1]).strip()
                if "verified" in each_line:
                    status = "Imaged"
                    imageMD5 = imageMD5.replace(' : verified','')
                else:
                    status = 'Not imaged'
            elif "MD5 Acquisition Hash:" in each_line: # MagnetAcquire image_info.txt
                imageMD5 = each_line.replace("MD5 Acquisition Hash:", "").strip()
                status = "Imaged"
            elif "MD5 hash calculated over data:" in each_line: # SumuriReconImager.txt
                imageMD5 = each_line.replace("MD5 hash calculated over data:", "").strip()
                
            # imageSHA1
            elif "Disk SHA1: " in each_line:    # Tableau
                imageSHA1 = re.split("Disk SHA1: ", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()

            elif "SHA1 checksum:" in each_line:  # FTKImager Image.E01.txt
                imageSHA1 = re.split("SHA1 checksum:", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()
                imageSHA1 = re.split(": ", each_line, 0)
                imageSHA1 = str(imageSHA1[1]).strip()
                if "verified" in each_line:
                    status = "Imaged"
                    imageSHA1 = imageSHA1.replace(' : verified','')
                else:
                    status = 'Not imaged'
            elif "SHA1 hash calculated over data:" in each_line: # SumuriReconImager.txt
                imageSHA1 = each_line.replace("SHA1 hash calculated over data:", "").strip()

            elif ".zip=" in each_line: # cellebrite .ufd
                hashTemp = re.split("=", each_line, 0)
                exportLocation = hashTemp[0]
                hashTemp = str(hashTemp[1]).strip()

                if  re.match(regex_sha256, hashTemp):    #regex SHA256 hash
                    imageSHA256 = hashTemp
                elif  re.match(regex_md5, hashTemp):    #regex md5 hash
                    imageMD5 = hashTemp
                elif  re.match(regex_sha1, hashTemp):    #regex SHA1 hash
                    imageSHA1 = hashTemp
                
            # _Triage_.txt parsing
            elif "Host Name: " in each_line:
                hostname = re.split("Host Name: ", each_line, 0)
                hostname = str(hostname[1]).strip()
                notes = ("%s The hostname is %s." %(notes, hostname))
            elif "Timezone: " in each_line:
                timezone = re.split("Timezone: ", each_line, 0)
                timezone = str(timezone[1]).strip()
                notes = ("%s The system timezone is set to %s." %(notes, timezone))
            elif "OS Name: " in each_line:
                OS = re.split("OS Name: ", each_line, 0)
                OS = str(OS[1]).strip()
            elif "   IPv4 Address" in each_line:
                ip = re.split("   IPv4 Address. . . . . . . . . . . : ", each_line, 0)
                ip = str(ip[1]).strip()
                notes = ("%s The IP address was %s." %(notes, ip))
            elif "    Lock Status:" in each_line:
                encryption = re.split("    Lock Status:", each_line, 0)
                encryption = str(encryption[1]).strip()
                if 'Locked' in encryption:
                    encryption = 'BitLocker Encrypted'
                    notes = ("%s BitLocker encryption is enabled." %(notes)) 
            elif "Email:" in each_line:
                email = re.split("Email: ", each_line, 0)
                email = str(email[1]).strip()
            elif "Currentuser:" in each_line:
                userName = re.split("Currentuser:", each_line, 0)
                userName = str(userName[1]).strip()


            # imagingStarted
            elif "Acquisition started:" in each_line:
                imagingStarted = re.split("Acquisition started: ", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                try:
                    imagingStarted = fix_date(imagingStarted)
                except:pass    

            elif "Acquisition Started:" in each_line: # MagnetAcquire image_info.txt
                imagingStarted = each_line.replace("Acquisition Started:", "").strip()
                imagingStarted = fix_date2(imagingStarted)  # test

            elif "Extraction start date/time" in each_line: #cellebrite
                imagingStarted = re.split("time", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip().replace(" -05:00", "").strip(':').strip().replace("(GMT-5)", "")
                # try:
                    # imagingStarted = fix_date(imagingStarted)
                # except:pass    

            elif "Imaging Start Time :" in each_line:   # Recon imager
                imagingStarted = re.split("Imaging Start Time :", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                # try:
                    # imagingStarted = fix_date(imagingStarted)
                # except:pass    

            elif "Date=" in each_line: #cellebrite *.ufd log file
                imagingStarted = each_line.replace("Date=", "").strip()
                imagingStarted = fix_date3(imagingStarted)
                # try:
                    # imagingStarted = fix_date3(imagingStarted)
                # except:pass    

            elif "Date and time:" in each_line: # CELLEBRITEPREMIUM DEVICEINFO.TXT
                imagingStarted = each_line.replace("Date and time:", "").strip()
                biosTime = imagingStarted

            elif "Started: " in each_line:
                imagingStarted = re.split("Started: ", each_line, 0)
                imagingStarted = str(imagingStarted[1]).strip()
                if "Cellebrite" in imagingTool or "Tableau" in imagingTool: 
                    # imagingStarted = fix_date(imagingStarted)
                    try:
                        imagingStarted = fix_date(imagingStarted)
                    except:pass    

               
            # imagingFinished
                
            elif "Closed:" in each_line:
                imagingFinished = re.split("Closed: ", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                imagingFinished = fix_date(imagingFinished)
                try:
                    imagingStarted = fix_date(imagingFinished)
                except:pass    
                
            elif "Acquisition finished:" in each_line:
                imagingFinished = re.split("Acquisition finished:", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                try:
                    imagingStarted = fix_date(imagingFinished)
                except:pass    

            elif "Acquisition Finished:" in each_line: # MagnetAcquire image_info.txt
                imagingFinished = each_line.replace("Acquisition Finished:", "").strip()
                imagingFinished = fix_date2(imagingFinished)    # test


            elif "Extraction end date" in each_line:     #cellebrite
                imagingFinished = re.split("Extraction end date", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                imagingFinished = imagingFinished.replace("/time", "").replace(" -05:00", "").strip(':').strip().replace("(GMT-5)", "")

                # try:
                    # imagingStarted = fix_date(imagingFinished)
                # except:pass    

            elif "Imaging End Time   :" in each_line: # Recon imager
                imagingFinished = re.split("Imaging End Time   :", each_line, 0)
                imagingFinished = str(imagingFinished[1]).strip()
                # try:
                    # imagingStarted = fix_date(imagingFinished)
                # except:pass    

            elif "EndTime=" in each_line: #cellebrite *.ufd log file
                imagingFinished = each_line.replace("EndTime=", "").strip()
                status = ('Imaged')
                imagingFinished = fix_date3(imagingFinished)
                print('imagingStarted = %s' %(imagingFinished))   # temp
                # try:
                    # imagingFinished = fix_date(imagingFinished)
                # except:pass    

            # tempnotes
            elif "Description:" in each_line: # MagnetAcquire image_info.txt
                description = each_line.replace("Description:", "").strip() 
                tempNotes = ('%s %s' %(tempNotes, description))

            elif "Chipset:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Device Bluetooth Name:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Encryption Type:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Number Of Installed Applications:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))

            elif "Live encryption state:" in each_line: # CellebritePA_FFS.txt  or ExtractionType=
                tempNotes = ('%s %s' %(tempNotes, each_line.strip() ))


                
        # tempNotes save the input file name for bulk uploads. delete if it's input.txt
        tempNotes = ('%s %s' %(logFile, tempNotes)).strip()
        if tempNotes == 'input.txt':
            tempNotes = ''
            
        # if makeModel == '':
            # makeModel = ('%s %s' %(make, model))    # test
        makeModel = ('%s %s' %(make, model))    # test

        if vehicleYear != '' : # BerlaIVe AcquisitionLog.txt
        # if vehicleYear != '' and vehicleManufacturer != '' and vehicleModel != '': # BerlaIVe AcquisitionLog.txt
            makeModel = ("%s %s %s" %(vehicleYear, vehicleManufacturer, vehicleModel))

        if caseNumber != '' and exhibit != '':
            qrCode = ("%s_%s" %(caseNumber, exhibit))
        # if qrCode == '_': 
            # qrCode = ''
        
        if len(imagingTool1) != 0:
            imagingTool = ('%s %s' %(imagingTool1.strip(), imagingTool2.strip()))
        
        # if len(storageSize) != 0 and storageSerial == "" and len(storageType) != 0:
            # notes = ("This had a %s (S/N: %s) %s, %s drive. %s" %(storageMakeModel, storageSerial, storageSize, storageType, notes))   # test
        # elif len(storageMakeModel) != 0 and len(storageSerial) != 0 and len(storageSize) != 0 and len(storageType) != 0: 
            ##notes = ("This had a %s, model %s, serial #%s, %s drive. %s" %(storageSize, storageMakeModel, storageSerial, storageType, notes))   # test
            # notes = ("This had a %s (S/N: %s) %s, %s drive. %s" %(storageMakeModel, storageSerial, storageSize, storageType, notes))   # test
        # elif len(storageSize) != 0:
            # notes = ("This had a %s drive, model %s, serial #%s, %s drive. %s" %(storageSize, storageMakeModel, storageSerial, storageType, notes))   # test


        # if len(OS) != 0 and 'The operating system was' not in notes:
            # notes = ("%s The operating system was %s." %(notes, OS)) 


        if status == 'Not imaged':
            notes = ("%s This drive could not be imaged." %(notes))

        print('%s\t%s\t%s\t\t\t%s\t\t\t%s\t%s\t%s\t\t%s\t%s\t%s' %(caseNumber, exhibit, caseName, subjectBusinessName, forensicExaminer, exhibitType, makeModel, serial, OS, phoneNumber))
        write_report(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
            forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
            phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
            userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
            dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
            seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
            writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
            storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
            storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
            vaultTotal, tempNotes)

def pdfExtract(filename):
    (forensicExaminer, exhibitType, makeModel, serial, OS, phoneNumber) = ('', '', '', '', '', '')
    (phoneIMEI, email, status, imagingType, imageMD5, imageSHA256) = ('', '', '', '', '', '')
    (imagingStarted, imagingFinished, imagingTool, storageSize, evidenceDataSize, analysisTool) = ('', '', '', '', '', '')
    (tempNotes, exhibit) = ('', '')

    with pdfplumber.open(filename) as pdf:
        totalpages = len(pdf.pages)
        for i in range(0 ,totalpages):
            page = pdf.pages[i]
            tempNotes = ('''%s %s''') %(tempNotes, page.extract_text())
            phoneIMEI = 'Duncan Donuts'
            forensicExaminer = re.search(r'Examiner Name:(.*?)\n', tempNotes)
            forensicExaminer = str(forensicExaminer[1]).strip()

            exhibit = re.search(r'Evidence ID:(.*?)\n', tempNotes)
            exhibit = str(exhibit[1]).strip()

            makeModel = re.search(r'Model(.*?)\n', tempNotes)   # todo
            makeModel = str(makeModel[1]).strip()

            serial = re.search(r'Serial Number (.*?)\n', tempNotes)
            serial = str(serial[1]).strip()

            imagingTool = re.search(r'GrayKey Software: OS Version:(.*?),', tempNotes)
            imagingTool = str(imagingTool[1]).strip()
            imagingTool = ('GrayKey %s' %(imagingTool))

            imagingStarted = re.search(r'Report generation time:(.*?)\n', tempNotes)
            imagingStarted = str(imagingStarted[1]).strip()

            # OS = re.search(r'Software Version (.*?)\n', tempNotes)  # this gets overwritten
            # OS = str(OS[1]).strip()

            # phoneNumber = re.search(r'Phone Number \+(.*?)\n', tempNotes)   # todo don't grab +
            # phoneNumber = str(phoneNumber[1]).strip()
            
            try:
                phoneIMEI = re.search(r'IMEI(.*?)\n', tempNotes)
                phoneIMEI = str(phoneIMEI[1]).strip()
            except:
                pass

            # try:
                # email = re.search(r'Accounts (.*?)\n', tempNotes)
                # email = str(email[1]).strip()
            # except:
                # pass
            
            # try:
                # evidenceDataSize = re.search(r'Extraction size(.*?)\n', tempNotes)
                # evidenceDataSize = evidenceDataSize
            # except:
                # pass
            # imageSHA256 = re.search(r'SHA256 (.*?)\n', tempNotes)
            # imageSHA256 = str(imageSHA256[1]).strip()
            
            # imageMD5 = re.search(r'MD5 (.*?)\n', tempNotes)
            # imageMD5 = str(imageMD5[1]).strip()
            
                    
        
    return(forensicExaminer, exhibit, exhibitType, makeModel, serial, OS, phoneNumber, phoneIMEI, email, status, imagingType, imageMD5, imageSHA256, imagingStarted, imagingFinished, imagingTool, storageSize, evidenceDataSize, analysisTool, tempNotes)


def pdf_fill(input_pdf_path, output_pdf_path, data_dict):   
    '''
    # fill out EvidenceForm
    receives input template based on agency and itemType
    receives output template with a uniq name
    data_dict is my_dict which is many of the columns needed to write pdf reports
    '''
    
    template_pdf = pdfrw.PdfReader(input_pdf_path)
    for page in template_pdf.pages:
        annotations = page[ANNOT_KEY]
        for annotation in annotations:
            if annotation[SUBTYPE_KEY] == WIDGET_SUBTYPE_KEY:
                if annotation[ANNOT_FIELD_KEY]:
                    key = annotation[ANNOT_FIELD_KEY][1:-1]
                    if key in data_dict.keys():
                        if type(data_dict[key]) == bool:
                            if data_dict[key] == True:
                                annotation.update(pdfrw.PdfDict(
                                    AS=pdfrw.PdfName('Yes')))
                        else:
                            annotation.update(
                                pdfrw.PdfDict(V='{}'.format(data_dict[key]))
                            )
                            annotation.update(pdfrw.PdfDict(AP=''))
    template_pdf.Root.AcroForm.update(pdfrw.PdfDict(NeedAppearances=pdfrw.PdfObject('true')))
    pdfrw.PdfWriter().write(output_pdf_path, template_pdf)


def read_text():
    '''
    dump your exhibit rows from xlsx to input.txt
    this will read in each line and write a report 
    it then makes a backup of copy xlsx of the lines you tossed in
    '''
    (header, reportStatus, date) = ('', '', '<insert date here>')
    (body, executiveSummary, evidenceBlurb) = ('', '', '')
    (style) = ('')
    csv_file = open(filename, encoding='utf8') 
    
    outputFileXlsx = "report.txt"
    output = open(outputFileXlsx, 'w+')
    (subject, vowel) = ('test', 'aeiou')

    footer = ('''
Evidence:
    All digital images obtained pursuant to this investigation will be maintained on %s servers for five years past the date of adjudication and/or case discontinuance. Copies of digital images will be made available upon request. All files copied from the images and provided to the case agent for review are identified as the DIGITAL EVIDENCE FILE and will be included as an exhibit in the case file. 
    ''') %(agency)
    
    for each_line in csv_file:
        (caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent) = ('', '', '', '', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')

        if phoneNumber != '':
            print('_____________%s is from a phone extract') %(phoneNumber) #temp

        (color) = ('white')
        # style.set_bg_color('white')  # test
        each_line = each_line +  "\t" * 60
        each_line = each_line.split('\t')  # splits by tabs

        value = each_line
        note = value

        if each_line[1]:                        ### checks to see if there is an each_line[1] before preceeding
            Value = note

            caseNumber = each_line[0]
            exhibit = each_line[1]
            caseName = each_line[2]
            subjectBusinessName = each_line[3]
            caseType = each_line[4]         
            # caseType = each_line[4].lower
            caseAgent = each_line[5]
            forensicExaminer = each_line[6]
            reportStatus = each_line[7]
            notes = each_line[8]
            summary = each_line[9]
            exhibitType = each_line[10]
            makeModel = each_line[11].strip()
            serial = each_line[12].strip()
            OS = each_line[13]
            phoneNumber = each_line[14]
            phoneIMEI = each_line[15]
            mobileCarrier = each_line[16]
            biosTime = each_line[17]
            currentTime = each_line[18]
            timezone = each_line[19]
            shutdownMethod = each_line[20]
            shutdownTime = each_line[21]
            userName = each_line[22]
            userPwd = each_line[23]
            email = each_line[24]
            emailPwd = each_line[25]
            ip = each_line[26]
            seizureAddress = each_line[27]
            seizureRoom = each_line[28]
            dateSeized = each_line[29]
            seizedBy = each_line[30]
            dateReceived = each_line[31]
            receivedBy = each_line[32]
            removalDate = each_line[33]
            removalStaff = each_line[34]
            reasonForRemoval = each_line[35]
            inventoryDate = each_line[36]
            seizureStatus = each_line[37]
            status = each_line[38]
            imagingTool = each_line[39]
            # imagingType = each_line[40].lower
            imagingType = each_line[40]
            imageMD5 = each_line[41]
            imageSHA1 = each_line[42]
            imageSHA256 = each_line[43]
            writeBlocker = each_line[44]
            imagingStarted = each_line[45]
            imagingFinished = each_line[46]
            storageType = each_line[47]
            storageMakeModel = each_line[48]
            storageSerial = each_line[49]
            storageSize = each_line[50]
            evidenceDataSize = each_line[51]
            analysisTool = each_line[52]
            analysisTool2 = each_line[53]
            exportLocation = each_line[54]
            exportedEvidence = each_line[55]
            storageLocation = each_line[56]
            caseNumberOrig = each_line[57]
            priority = each_line[58]
            operation = each_line[59]
            Action = each_line[60]
            vaultCaseNumber = each_line[61]
            qrCode = each_line[62]
            vaultTotal = each_line[63]
            tempNotes = each_line[64]
            
            # Summary writer, put a blank space or write your own summary if you don't want one auto generated
            if summary == '' and dateSeized != '' and forensicExaminer != '' and seizureAddress != '' and agency != "ISP":
                summary = ('On %s %s attended the warrant at %s.' %(dateSeized, forensicExaminer, seizureAddress))
            elif summary != '':
                summary == summary
            # else:
                # summary = ' ' 

            qrCode = ("%s_%s" %(caseNumber, exhibit))

        pdf_output = ("EvidenceForm_%s_Ex_%s.pdf" %(caseNumber, exhibit))
        if header == '':
            header = ('''
ACTIVITY REPORT                              BUREAU OF CRIMINAL INVESTIGATIONS
____________________________________________________________________________________

Activity Number:                             Date of Activity:
%s                               %s
____________________________________________________________________________________
____________________________________________________________________________________
Subject of Activity:                         Case Agent:             Typed by:
%s %s                               %s        %s
%s
____________________________________________________________________________________

Executive Summary 
    Special Agent %s of the %s, %s, requested an examination of evidence for any information regarding the %s investigation in the %s case. %s
''') %(caseNumber, todaysDate, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, caseAgent, agencyFull, divisionFull, caseType, caseName, summary)

            output.write(header+'\n')
        
        if executiveSummary == '':
            executiveSummary = ('''
%s                                    %s

%s %s                           %s    %s 

Executive Summary 
    Special Agent %s of the %s, %s, requested an examination of evidence for any information regarding the %s investigation in the %s case. %s
''') %(caseNumber, todaysDate, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseAgent, agencyFull, divisionFull, caseType, caseName, summary)


        report = ('''
        
Exhibit %s
    ''') %(exhibit)

        if makeModel != '':
            if makeModel[0].lower() in vowel:
                report = ('''%sAn %s''') %(report, makeModel)
            else:
                report = ('''%sA %s''') %(report, makeModel)
        if len(mobileCarrier) != 0:
            if exhibitType == 'phone':
                report = ("%s %s" %(report, mobileCarrier))
            else:
                report = ("%s (Carrier: %s)" %(report, mobileCarrier))



        if len(exhibitType) != 0:
            report = ("%s %s" %(report, exhibitType))

        if phoneNumber != '' and phoneNumber != 'NA' and phoneNumber != 'na' and phoneNumber != 'N/A':
            report = ("%s (MSISDN: %s)" %(report, phoneNumber))

        if phoneIMEI != '' and phoneIMEI != 'NA' and phoneIMEI != 'na' and phoneIMEI != 'N/A':
            report = ("%s (IMEI: %s)" %(report, phoneIMEI))


        if len(serial) != 0:
            report = ("%s (S/N: %s)" %(report, serial))
            # report = ("%s, serial # %s" %(report, serial))

        if len(OS) != 0:
            report = ("%s (OS: %s)" %(report, OS))
            # if OS[0].lower() in vowel:
                # report = ("%s, with an %s OS" %(report, OS))
            # else:
                # report = ("%s, with a %s OS" %(report, OS))          

        if len(dateReceived) != 0:
            report = ("%s was received on %s" %(report, dateReceived.replace(" ", " at ", 1)))
        else:
            report = ("%s was received" %(report))
        report = ("%s." %(report))
        
        # if len(imagingStarted) != 0:
        if len(imagingStarted) != 0 and status != "Not imaged":
            report = ("%s On %s," %(report, imagingStarted.replace(" ", " at ", 1)))
        report = ("%s Digital Forensic Examiner %s" %(report, forensicExaminer))

        if len(imagingTool) != 0 and imagingType != '' and writeBlocker != '': 
            if imagingType[0].lower() in vowel:
                report = ("%s used %s, utilizing a %s write blocker, to conduct an %s" %(report, imagingTool, writeBlocker, imagingType))  
            elif imagingType[0].lower() not in vowel:
                report = ("%s used %s, utilizing a %s write blocker, to conduct a %s" %(report, imagingTool, writeBlocker, imagingType))  

        elif len(imagingTool) != 0 and imagingType != '':
            if imagingType[0].lower() in vowel:
                report = ("%s used %s to conduct an %s" %(report, imagingTool, imagingType))  
            elif imagingType[0].lower() not in vowel:
                report = ("%s used %s to conduct a %s" %(report, imagingTool, imagingType))  

        elif imagingTool != '':
            report = ("%s used %s to conduct " %(report, imagingTool))  

        elif imagingType != '' and exportedEvidence != "N":
            report = ("%s conducted a %s" %(report, imagingType))  
        elif exportedEvidence == "N":
            report = ("%s did not conduct a" %(report))  
        else:
            report = ("%s conducted a" %(report))  

            
        if phoneNumber != '' and phoneNumber != 'NA' and phoneNumber != 'na' and phoneNumber != 'N/A':
            report = ("%s phone extraction." %(report))
        elif imagingStarted != '':        
            report = ("%s forensic extraction" %(report))

        else:        
            report = ("%s manual analysis" %(report))

        if len(storageType) != 0 and storageMakeModel != '' and storageSerial != '' and storageSize != '': 
            report = ("%s on the %s (S/N: %s) %s %s drive." %(report, storageMakeModel, storageSerial, storageSize, storageType))  
        elif storageMakeModel != '' and storageSerial != '' and storageSize != '': 
            report = ("%s on the %s (S/N: %s) %s drive." %(report, storageMakeModel, storageSerial, storageSize))  
        elif storageMakeModel != '' and storageSize != '': 
            report = ("%s on the %s %s drive." %(report, storageMakeModel, storageSize))  


        else: 
            report = ("%s." %(report))  
    
        # image hash
        if len(imageMD5) != 0 and exportLocation != '' and len(imageSHA256) != 0 and imageSHA256 != 'NA' and imageSHA256 != 'na' and imageSHA256 != 'N/A':
            report = ("%s The image (SHA256 Hash: % s) (MD5 Hash: % s) was saved as %s." %(report, imageSHA256, imageMD5, exportLocation.split('\\')[-1])) 
        elif len(imageMD5) != 0 and exportLocation != '':
            report = ("%s The image (MD5 Hash: % s) was saved as %s." %(report, imageMD5, exportLocation.split('\\')[-1])) 
        elif len(imageSHA256) != 0 and imageSHA256 != 'NA' and imageSHA256 != 'na' and imageSHA256 != 'N/A':
            report = ("%s The image had a SHA256 hash of % s." %(report, imageSHA256))

        # analysisTool
        if analysisTool != '' and analysisTool2 != '':      # analysisTool2
            report = ("%s The image was processed with %s and further analyzed with %s." %(report, analysisTool, analysisTool2))

        elif analysisTool != '':
            report = ("%s The image was processed with %s." %(report, analysisTool))

        # add username and password to report
        if len(userName) != 0 and userPwd != '' and exhibitType != '': 
            report = ("%s \"%s\" with a password of \"%s\" was a login to this %s." %(report, userName, userPwd, exhibitType)) 
        elif len(userName) != 0 and userPwd != '': 
            report = ("%s \"%s\" with a password of \"%s\" was a login to this device." %(report, userName, userPwd)) 

        # add email / password to report
        if len(email) != 0 and emailPwd != '' and exhibitType != '':  
            report = ("%s \"%s\" with a password of \"%s\" was an email configured on this %s." %(report, email, emailPwd, exhibitType)) 
        elif len(email) != 0 and exhibitType != '':  
            if " and " in email:
                report = ("%s %s were email addresses configured on this %s." %(report, email, exhibitType)) 
            else:
                report = ("%s %s was an email configured on this %s." %(report, email, exhibitType)) 

        elif len(email) != 0 and emailPwd != '':  
            report = ("%s \"%s\" with a password of \"%s\" was an email configured on this device." %(report, email, userPwd)) 
 
        if notes != '':
            report = ("%s %s" %(report, notes))
          
        # exportedEvidence
        if exportedEvidence == "Y" and 'elevant files were exported' not in notes:
            # report = ("%s Relevant files were exported." %(report.strip()))
            report = ("%s Relevant files were exported." %(report.rstrip()))
        elif exportedEvidence == "N" and 'search for relevant files was made and no files were found' not in notes:
            report = ("%s A search for relevant files was made and no files were found." %(report.rstrip()))

        # evidence return
        if "2" in removalDate and "eturned" in storageLocation: # returned or Returned
            
            if " " in removalDate:
                removalDate2 = removalDate.split(' ')[0]
            else:
                removalDate2 = removalDate

            if exhibitType != '':
                report = ("%s This %s was returned to the owner on %s." %(report, exhibitType, removalDate2)) # test
            else:
                report = ("%s This item was returned to the owner on %s." %(report, removalDate2))  
    
        report = report.replace("    , was received. ", "    ")
        report = report.replace("This was a DVR system was not imaged.","This was a DVR system and was not imaged.")
        report = report.replace("Digital Forensic Examiner Casey Karaffa did not conduct a forensic extraction.","This was not imaged.")
        report = report.replace("The image was processed with copy.","Pertinent files were copied.")
        report = report.replace("This had a  drive, model , serial # .","") # fixme     
        notes = notes.replace("This had a  drive, model , serial # .","")  # fixme
        report = report.replace(", serial # .",".") # fixme 
        notes = notes.replace(", serial # .",".") # fixme 
        if storageSerial == "000ecc45 0067205e":   # fixme  
            storageSerial = ' '

        print(report)
        output.write(report)
        # body = ("%s\n%s" %(body, report))
        body = ("%s%s" %(body, report))


        #Colors
        if 'Final' in  reportStatus or 'final' in  reportStatus or reportStatus == 'Y' or reportStatus == 'y' :        
            # FormatFunction(bg_color = 'green')
            FormatFunction(bg_color = '#92D050')  # green #92D050
            
        elif 'Draft' in  reportStatus or 'draft' in  reportStatus:        
            # FormatFunction(bg_color = 'orange')
            FormatFunction(bg_color = '#FFc000')         # orange       
        else:
            FormatFunction(bg_color = 'white')

        
        # Write excel
        write_report(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent,
            forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber,
            phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime,
            userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy,
            dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate,
            seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256,
            writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial,
            storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence,
            storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode,
            vaultTotal, tempNotes)
        # write_pdf(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent,
            # forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
            # phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
            # userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
            # dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
            # seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
            # writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
            # storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
            # storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
            # vaultTotal, tempNotes)
        
        if caseNotesStatus == 'True':
            my_dict = dictionaryBuild(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
            forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
            phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
            userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
            dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
            seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
            writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
            storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
            storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
            vaultTotal, tempNotes)
        # write an evidence form based on which agency you are from
            
            if exhibit != '':
                pdf_output = ("ExhibitNotes_%s_Ex%s.pdf" %(caseNumber, exhibit))    # output
            else:
                pdf_output = ("ExhibitNotes_%s_%s.pdf" %(caseNumber, todaysDateTime))    # output in case exhibit is empty
                time.sleep(2)  # wait 2 seconds so the name is uniq
                
            # choose which form you fill out based on agency acronym
            if agency == "ISP":            
                if exhibitType == 'phone':  # lower(exhibitType)
                    pdf_template = "EvidenceForm_MDIS.pdf"  # Mobile Device Evidence Sheet
                else:
                    pdf_template = "EvidenceForm_EDIS.pdf"  # Electronic Device Evidence Sheet
            else:   
                pdf_template = "Blank_EvidenceForm.pdf"
            pdf_fill(pdf_template, pdf_output, my_dict)
            pdf_fill(pdf_template, pdf_output, my_dict)

    # write docx report
    write_activity_report(caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, executiveSummary, body, footer)

    output.write(footer+'\n')

def write_report(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
        forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
        phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
        userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
        dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
        seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
        writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
        storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
        storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
        vaultTotal, tempNotes):
    '''
    write out_log_.xlsx
    '''
    global Row
    global Row2
    global caseMain
    
    # if caseMain == '':
    if Row == 1:
    
        caseMain = caseNumber
        print('caseMain = %s' %(caseMain))
    elif caseMain != caseNumber:
        caseMain = caseNumber
        print('hello world')    # temp
        Row2 += 2
        print('caseMain 2 = %s' %(caseMain))
        
    Sheet1.write_string(Row, 0, caseNumber)
    Sheet1.write_string(Row, 1, exhibit)
    Sheet1.write_string(Row, 2, caseName)
    Sheet1.write_string(Row, 3, subjectBusinessName)
    try:
        Sheet1.write_string(Row, 4, caseType)       # with .lower()   <built-in method lower of str object at 0x0000020B6D23DA70>
    except TypeError as error:
        print(error)
    Sheet1.write_string(Row, 5, caseAgent)
    Sheet1.write_string(Row, 6, forensicExaminer)
    Sheet1.write_string(Row, 7, reportStatus)
    Sheet1.write_string(Row, 7, reportStatus)   
    # Sheet1.write_string(Row, 7, reportStatus, {'validate': 'list', 'source': ['Finalized', 'Draft', '']})  

    Sheet1.write_string(Row, 8, notes)
    Sheet1.write_string(Row, 9, summary)
    Sheet1.write_string(Row, 10, exhibitType)
    Sheet1.write_string(Row, 11, makeModel)
    Sheet1.write_string(Row, 12, serial)
    Sheet1.write_string(Row, 13, OS)
    Sheet1.write_string(Row, 14, phoneNumber)
    Sheet1.write_string(Row, 15, phoneIMEI)
    Sheet1.write_string(Row, 16, mobileCarrier)
    Sheet1.write_string(Row, 17, biosTime)
    Sheet1.write_string(Row, 18, currentTime)
    Sheet1.write_string(Row, 19, timezone)
    Sheet1.write_string(Row, 20, shutdownMethod)
    Sheet1.write_string(Row, 21, shutdownTime)
    Sheet1.write_string(Row, 22, userName)
    Sheet1.write_string(Row, 23, userPwd)
    Sheet1.write_string(Row, 24, email)
    Sheet1.write_string(Row, 25, emailPwd)
    Sheet1.write_string(Row, 26, ip)
    Sheet1.write_string(Row, 27, seizureAddress)
    Sheet1.write_string(Row, 28, seizureRoom)
    Sheet1.write_string(Row, 29, dateSeized)
    Sheet1.write_string(Row, 30, seizedBy)
    Sheet1.write_string(Row, 31, dateReceived)
    Sheet1.write_string(Row, 32, receivedBy)
    Sheet1.write_string(Row, 33, removalDate)
    Sheet1.write_string(Row, 34, removalStaff)
    Sheet1.write_string(Row, 35, reasonForRemoval)
    Sheet1.write_string(Row, 36, inventoryDate)
    Sheet1.write_string(Row, 37, seizureStatus)
    Sheet1.write_string(Row, 38, status)
    Sheet1.write_string(Row, 39, imagingTool)
    
    try:
        Sheet1.write_string(Row, 40, imagingType)   #errors with .lower() <built-in method lower of str object at 0x0000020B6D21EDF0>
    except TypeError as error:
        print(error)
    
    Sheet1.write_string(Row, 41, imageMD5)
    Sheet1.write_string(Row, 42, imageSHA1)
    Sheet1.write_string(Row, 43, imageSHA256)
    Sheet1.write_string(Row, 44, writeBlocker)
    Sheet1.write_string(Row, 45, imagingStarted)
    Sheet1.write_string(Row, 46, imagingFinished)
    Sheet1.write_string(Row, 47, storageType)
    Sheet1.write_string(Row, 48, storageMakeModel)
    Sheet1.write_string(Row, 49, storageSerial)
    Sheet1.write_string(Row, 50, storageSize)
    Sheet1.write_string(Row, 51, evidenceDataSize)
    Sheet1.write_string(Row, 52, analysisTool)
    Sheet1.write_string(Row, 53, analysisTool2)
    Sheet1.write_string(Row, 54, exportLocation)
    Sheet1.write_string(Row, 55, exportedEvidence)
    Sheet1.write_string(Row, 56, storageLocation)
    Sheet1.write_string(Row, 57, caseNumberOrig)
    Sheet1.write_string(Row, 58, priority)
    Sheet1.write_string(Row, 59, operation)
    Sheet1.write_string(Row, 60, Action)
    Sheet1.write_string(Row, 61, vaultCaseNumber)
    Sheet1.write_string(Row, 62, qrCode)
    Sheet1.write_string(Row, 63, vaultTotal)
    Sheet1.write_string(Row, 64, tempNotes)

    Sheet2.write_string(1, 14, caseNumber, cell_format)
    Sheet2.write_string(2, 14, caseName, cell_format)
    Sheet2.write_string(3, 14, subjectBusinessName, cell_format)
    Sheet2.write_string(4, 14, caseAgent, cell_format)
    Sheet2.write_string(5, 14, forensicExaminer, cell_format)

    Sheet2.write_string(Row2, 0, exhibit, cell_format)
    Sheet2.write_string(Row2, 1, exhibitType, cell_format)
    if dateReceived != "":
        Sheet2.write_string(Row2, 2, '..', cell_format)    
    else:
        Sheet2.write_string(Row2, 2, '', cell_format)    

    if removalDate != "":
        Sheet2.write_string(Row2, 3, '..', cell_format)    
    else:
        Sheet2.write_string(Row2, 3, '', cell_format)    

    
    Sheet2.write_string(Row2, 4, '', cell_format)    

    if status == "Imaged":
        Sheet2.write_string(Row2, 5, 'Y', cell_format)
    elif status == "Not Imaged":         
        Sheet2.write_string(Row2, 5, 'N', cell_format)
    else:
        Sheet2.write_string(Row2, 5, '', cell_format)
    Sheet2.write_string(Row2, 6, '', cell_format)    

    if analysisTool != "" and exportedEvidence != "":
        Sheet2.write_string(Row2, 7, 'Y', cell_format)
    else:
        Sheet2.write_string(Row2, 7, '', cell_format)        
    if reportStatus == "Finalized":
        Sheet2.write_string(Row2, 8, 'Y', cell_format)
    else:
        Sheet2.write_string(Row2, 8, '', cell_format)
    Sheet2.write_string(Row2, 9, '', cell_format)    
    Sheet2.write_string(Row2, 10, exportedEvidence, cell_format)    
    Sheet2.write_string(Row2, 11, '', cell_format)    
    Sheet2.write_string(Row2, 12, '', cell_format)    
    Sheet2.write_string(Row2, 14, caseNumber, cell_format) 


    Row += 1
    Row2 += 1
    
def write_activity_report(caseNumber, caseName, subjectBusinessName, caseAgent, forensicExaminer, caseType, executiveSummary, body, footer): 
    '''
    write ActivityReport_%s__%s_%s_%s_DRAFT.docx
    '''
    
    outputDocx = ('ActivityReport_%s__%s_%s_%s_DRAFT.docx' %(caseNumber, Month, Day, Year)) 

    try:
        document = docx.Document("Blank_ActivityReport.docx") # read in the template if it exists
    except:
        print("you are missing Blank_ActivityReport.docx")
        document = create_docx()   # create a basic template file
    
    if executiveSummary != '':
        document.add_paragraph(executiveSummary)    
    
    document.add_paragraph(body)  

    if footer != '':
        document.add_paragraph(footer) 
    
    document.save(outputDocx)   # print output to the new file
    
    # print('your activity report is saved as %s' %(outputDocx))   # temp

def write_sticker():
    '''
    print all the details for a sticker/label. Make 4 copies and attach to all removable pieces of the PC.
    some day it will print to an avery label style PDF
    '''

    style = workbook.add_format()
    (header, reportStatus, date) = ('', '', '<insert date here>')
    (headers) = []
    csv_file = open(filename, encoding='utf8')
    outputFileXlsx = "sticker.txt"
    output = open(outputFileXlsx, 'w+')

    if not os.path.exists('Avery2x4Labels.docx'):
        print('you are missing Avery2x4Labels.docx.... so Im making a lame version')
        # Create a new Word document
        document = docx.Document()
        # Set the side margins to 0.25 inches
        section = document.sections[0]
        section.margin_left = docx.shared.Inches(0.25)
        section.margin_right = docx.shared.Inches(0.25)
        table = document.add_table(rows=5, cols=3)
    else:
        document = docx.Document('Avery2x4Labels.docx')
        # Find the first table in the document
        table = document.tables[0]

    # Change the default font size for the document
    for style in document.styles:
        style.font.size = docx.shared.Pt(14)
    

    # footer = '''  

# The images of all the devices will be retained. The case agent may request additional analysis or files to be exported if new evidence of probative value is determined, at a future date.
    
# Evidence:
    # Reports and supporting files were exported and given to the case agent.
    # '''

    footer = '''  

All digital images obtained pursuant to this investigation will be maintained on %s servers for five years past the date of adjudication and/or case discontinuance. Copies of digital images will be made available upon request. All files copied from the images and provided to the case agent for review are identified as the DIGITAL EVIDENCE FILE and will be included as an exhibit in the case file. 
    ''' %(agency)
    
    for each_line in csv_file:
        (caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent) = ('', '', '', '', '', '')
        (forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel) = ('', '', '', '', '', '')
        (serial, OS, phoneNumber, phoneIMEI, mobileCarrier, biosTime) = ('', '', '', '', '', '')
        (currentTime, timezone, shutdownMethod, shutdownTime, userName, userPwd) = ('', '', '', '', '', '')
        (email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized) = ('', '', '', '', '', '')
        (seizedBy, dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval) = ('', '', '', '', '', '')
        (inventoryDate, seizureStatus, status, imagingTool, imagingType, imageMD5) = ('', '', '', '', '', '')
        (imageSHA1, imageSHA256, writeBlocker, imagingStarted, imagingFinished, storageType) = ('', '', '', '', '', '')
        (storageMakeModel, storageSerial, storageSize, evidenceDataSize, analysisTool, analysisTool2) = ('', '', '', '', '', '')
        (exportLocation, exportedEvidence, storageLocation, caseNumberOrig, priority, operation) = ('', '', '', '', '', '')
        (Action, vaultCaseNumber, qrCode, vaultTotal, tempNotes) = ('', '', '', '', '')

        if phoneNumber != '':
            print('_____________%s is from a phone extract') %(phoneNumber) #temp

        (color) = ('white')
        # style.set_bg_color('white')  # test
        each_line = each_line +  "\t" * 27
        each_line = each_line.split('\t')  # splits by tabs

        value = each_line
        note = value

        if each_line[1]:                        ### checks to see if there is an each_line[1] before preceeding
            Value = note
            caseNumber = each_line[0]
            exhibit = each_line[1]
            caseName = each_line[2]
            subjectBusinessName = each_line[3]
            caseType = each_line[4]
            caseAgent = each_line[5]
            forensicExaminer = each_line[6]
            reportStatus = each_line[7]
            notes = each_line[8]
            summary = each_line[9]
            exhibitType = each_line[10]
            makeModel = each_line[11].strip()
            serial = each_line[12].strip()
            OS = each_line[13]
            phoneNumber = each_line[14]
            phoneIMEI = each_line[15]
            mobileCarrier = each_line[16]
            biosTime = each_line[17]
            currentTime = each_line[18]
            timezone = each_line[19]
            shutdownMethod = each_line[20]
            shutdownTime = each_line[21]
            userName = each_line[22]
            userPwd = each_line[23]
            email = each_line[24]
            emailPwd = each_line[25]
            ip = each_line[26]
            seizureAddress = each_line[27]
            seizureRoom = each_line[28]
            dateSeized = each_line[29]
            seizedBy = each_line[30]
            dateReceived = each_line[31]
            receivedBy = each_line[32]
            removalDate = each_line[33]
            removalStaff = each_line[34]
            reasonForRemoval = each_line[35]
            inventoryDate = each_line[36]
            seizureStatus = each_line[37]
            status = each_line[38]
            imagingTool = each_line[39]
            imagingType = each_line[40]
            imageMD5 = each_line[41]
            imageSHA1 = each_line[42]
            imageSHA256 = each_line[43]
            writeBlocker = each_line[44]
            imagingStarted = each_line[45]
            imagingFinished = each_line[46]
            storageType = each_line[47]
            storageMakeModel = each_line[48]
            storageSerial = each_line[49]
            storageSize = each_line[50]
            evidenceDataSize = each_line[51]
            analysisTool = each_line[52]
            analysisTool2 = each_line[53]
            exportLocation = each_line[54]
            exportedEvidence = each_line[55]
            storageLocation = each_line[56]
            caseNumberOrig = each_line[57]
            priority = each_line[58]
            operation = each_line[59]
            Action = each_line[60]
            vaultCaseNumber = each_line[61]
            qrCode = each_line[62]
            vaultTotal = each_line[63]
            tempNotes = each_line[64]

        if status == 'Imaged':    
            header = ('''   Case#: %s      Ex: %s
   CaseName: %s
   Subject: %s
   Make: %s 
   Serial: %s
   Agent: %s
   %s
''') %(caseNumber, exhibit, caseName, subjectBusinessName, makeModel, serial, caseAgent, status)
        else:
            header = ('''  Case#: %s      Ex: %s
  CaseName: %s
  Subject: %s
  Make: %s 
  Serial: %s
  Agent: %s
''') %(caseNumber, exhibit, caseName, subjectBusinessName, makeModel, serial, caseAgent)

        header = header.strip()
        headers.append(header)
        
# write it one line at at time. If phone isn't blank, include it

        # Write excel
        write_report(caseNumber, exhibit, caseName, subjectBusinessName, caseType, caseAgent, 
            forensicExaminer, reportStatus, notes, summary, exhibitType, makeModel, serial, OS, phoneNumber, 
            phoneIMEI, mobileCarrier, biosTime, currentTime, timezone, shutdownMethod, shutdownTime, 
            userName, userPwd, email, emailPwd, ip, seizureAddress, seizureRoom, dateSeized, seizedBy, 
            dateReceived, receivedBy, removalDate, removalStaff, reasonForRemoval, inventoryDate, 
            seizureStatus, status, imagingTool, imagingType, imageMD5, imageSHA1, imageSHA256, 
            writeBlocker, imagingStarted, imagingFinished, storageType, storageMakeModel, storageSerial, 
            storageSize, evidenceDataSize, analysisTool, analysisTool2, exportLocation, exportedEvidence, 
            storageLocation, caseNumberOrig, priority, operation, Action, vaultCaseNumber, qrCode, 
            vaultTotal, tempNotes)
        
        # write to stickers.txt
        output.write(header+'\n\n')
        
    # write to stickers.docx
    l = 0
    for i in range(5):
        try:
            cell = table.cell(i, 0)
            cell.text = headers[l]
            l += 1
            cell = table.cell(i, 2)
            cell.text = headers[l]
            l += 1
        except:
            pass
    document.save('stickers.docx')


def usage():
    '''
    working examples of syntax
    '''
    file = sys.argv[0].split('\\')[-1]
    print("\nDescription: " + description)
    print(file + " Version: %s by %s" % (version, author))
    print("\n\tinsert your info into input.txt")
    print("\n\t or insert logs into Logs folder")
    print("\nExample:")
    print("\t" + file + " -g     \t\t")
    print("\t" + file + " -l     \t\t")
    print("\t" + file + " -l -d     \t\t")
    print("\t" + file + " -L     \t\t")
    print("\t" + file + " -r     \t\t")
    print("\t" + file + " -r -c  \t\t")
    print("\t" + file + " -s     \t\t")


    
if __name__ == '__main__':
    main()

# <<<<<<<<<<<<<<<<<<<<<<<<<< Revision History >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
2.7.3 - added checklist to xlsx
2.7.2 - sticker maker prints out avery 2 x 4" labels now
2.7.0 - added -g for GUI interface for data enty
2.6.7 - fixed Tableau storageSerial being the Tableau serial number (I think)
2.6.5 - changed report writing output and summary writer (add a space in summary if you don't want it to write anything.)
2.6.2 - reportStatus gets colored if it's marked Finalized, Draft or Y
2.6.1 - if you change agency, agencyFull and divisionFull it writes a more customized report
2.6.0 - added -L option to parse a folder full of logs all at once. -I and -O are optional now
2.5.6 - Logs: CellebritePremium DeviceInfo.txt, Berla iVE
2.5.0 - Column Re-order to group like items together (case, description, lab chain of custody, acquisition, notes)
2.1.3 - fixed log parser to populate storageSize, storageMakeModel, storageSerial, storageSize (Tableau)
2.1.2 - Added about a dozen columns for additional info (the columns need to be re-ordered one of these days.)
2.1.1 - Added ISP pdf templates for pdf writing (just change agency = to agency = 'ISP'
2.1.0 - Added CaseNotes.pdf output if you add -c to -r
2.0.3 - Added Recon imager log parsing
2.0.2 - ActivityReport....docx output works best from the template.
2.0.1 - Reorginized column orders, fixed serial #
1.0.1 - Created a Tableau log parser
1.0.0 - Created forensic report writer
0.1.2 - converted tabs to 4 spaces for #pep8

"""

# <<<<<<<<<<<<<<<<<<<<<<<<<< Future Wishlist  >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
fix: forensicsreporter.exe -h blows an unhandled exception in script line 2344
fix serial and storageSerial for TableauImager_21-41803200001_Ex1_Seagate3TBHDD.txt 

add a brother label printer output to xlsx with qrCode
qrCode could be caseNumber_exhibit_serial (it depends on what the evidence staff want displayed on their inventory scanner)

figure out DocX tags or variables to insert data into the first fields
parse: GrayKey, MagentAcuire, MagnetAxiom, SumuriReconImager, TableauTX1 (MS shared samples)

if qrCode = '_': qrCode = ''

can't parse:
CellebriteUFED4PC_log.txt   # UnicodeDecodeError: 'utf-8' codec can't decode byte 0xff in position 0: invalid start byte
MagnetAXIOM_Case Information_001.txt
MagnetAXIOM_Case Information_002.txt


"""

# <<<<<<<<<<<<<<<<<<<<<<<<<<      notes            >>>>>>>>>>>>>>>>>>>>>>>>>>

"""
change the agency, agencyFull, divisionFull to your agencies info

if you have log output (like GrayKey) you want parsed, send it my way. As long as there is a key:value pair on one line, I can do it.
If you want your agencies forms filled, you just need to insert these variables into your pdf.


“black iPhone 11 currently stored as item number x under case number 22-xxxx which was located on master bedroom nightstand….”

"""


# <<<<<<<<<<<<<<<<<<<<<<<<<<      Copyright        >>>>>>>>>>>>>>>>>>>>>>>>>>

# Copyright (C) 2022 LincolnLandForensics
#
# This program is free software; you can redistribute it and/or modify it under
# the terms of the GNU General Public License version 2, as published by the
# Free Software Foundation
#
# This program is distributed in the hope that it will be useful, but WITHOUT
# ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or FITNESS
# FOR A PARTICULAR PURPOSE. See the GNU General Public License for more
# details (http://www.gnu.org/licenses/gpl.txt).


# <<<<<<<<<<<<<<<<<<<<<<<<<<      The End        >>>>>>>>>>>>>>>>>>>>>>>>>>